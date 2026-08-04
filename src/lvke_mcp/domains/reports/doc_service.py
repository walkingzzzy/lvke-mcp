"""报告域文档服务 —— MCP 自有实现（零外部依赖）。

既有报告文档服务的域内复刻（原样保留被引符号，仅改 import 路径与
存储根，不重写业务逻辑）：

- 存储根：MCP 自有 ``runtime.workspace.workspace_root``
- 内部依赖改写：``workspace_migration``(WAL 控制面) 删除、审计
  调用删除（MCP 域内无对应设施，best-effort 块直接移除）、``docx_fonts`` /
  ``env_templates`` / ``finance_model`` / ``run_service`` 指向 lvke_mcp 域内实现

存储落点（MCP 自管，每个工作区一个报告）：
``{LVKE_MCP_DATA_DIR}/workspaces/{workspace_id}/``
  ├─ ``workspace_meta.json``        工作区元信息 + 当前修订指针
  ├─ ``revisions/{rev_id}/report.md`` 各修订正文
  ├─ ``revisions/{rev_id}/meta.json`` 各修订元信息
  ├─ ``agent_proposals/{pid}/proposed_report.md`` / ``diff.html`` / ``meta.json``
  ├─ ``issues/issues.json``         issue_center
  └─ ``finance.json``               (可选)财务摘要,只读
"""

from __future__ import annotations

import difflib
import hashlib
import json
import logging
import re
import secrets
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from filelock import FileLock

from lvke_mcp.runtime import workspace as runtime_workspace

logger = logging.getLogger(__name__)

# ---- B 型 9 章大纲(参考可研 report_template B 型,域内内置) ----------------

REPORT_CHAPTERS: tuple[str, ...] = (
    "总论",
    "项目背景与建设必要性",
    "需求分析与建设规模",
    "总体建设方案",
    "投资估算与资金筹措",
    "财务分析与评价",
    "风险分析与对策",
    "保障措施",
    "结论与建议",
)

DEFAULT_REPORT_TYPE = "gov10"

# ── 文档类型维度（T0.5，Phase 1 仅 feasibility 生成，其余先占位待 Phase 2/3）──
DEFAULT_DOC_KIND = "feasibility"
DOC_KINDS: dict[str, dict[str, str]] = {
    "feasibility": {"label": "可行性研究报告", "hint": "立项论证主力文档（政府 gov10 / 企业 ent9）。"},
    "asset_acquisition": {"label": "资产收购可行性研究报告", "hint": "存量资产、酒店经营、租约、交易和融资一体化收购论证。"},
    "project_apply": {"label": "项目申请报告", "hint": "企业投资核准类（5 章）。Phase 2 支持。"},
    "fund_apply": {"label": "资金申请报告", "hint": "申请中央预算内投资（8-10 节）。Phase 2 支持。"},
    "implementation": {"label": "实施方案", "hint": "污水治理/以工代赈/道路等工程类。Phase 3 支持。"},
}
# Phase 1 已放开生成能力的文档类型
ENABLED_DOC_KINDS = ("feasibility", "asset_acquisition")


def resolve_doc_kind(meta: dict[str, Any]) -> str:
    """从 meta 解析文档类型；缺省/未知回退 feasibility。"""
    dk = str((meta or {}).get("doc_kind") or "").strip()
    return dk if dk in DOC_KINDS else DEFAULT_DOC_KIND


REPORT_STRUCTURES: dict[str, dict[str, Any]] = {
    "gov10": {
        "label": "政府投资项目（发改委2023通用大纲 10 章）",
        "hint": "适用于政府投资、政府核准/备案的公共服务、基础设施、产业园等项目，遵循发改委2023政府投资项目可行性研究报告编写通用大纲。",
        "chapters": [
            {"title": "概述", "subs": ["项目概况", "建设单位概况", "编制依据和研究范围", "主要结论和建议"]},
            {"title": "项目建设背景和必要性", "subs": ["规划政策符合性", "项目建设背景", "项目建设的必要性"]},
            {"title": "项目需求分析与产出方案", "subs": ["需求分析", "项目建设内容、规模及产出方案", "项目商业模式"]},
            {"title": "项目选址与要素保障", "subs": ["项目选址", "项目建设条件", "要素保障分析"]},
            {"title": "项目建设方案", "subs": ["技术方案", "设备方案", "工程方案", "项目招标方案"]},
            {"title": "项目运营方案", "subs": ["运营模式选择", "经营方案", "安全保障方案", "运营管理方案"]},
            {"title": "项目投融资与财务方案", "subs": ["投资估算", "融资方案", "财务效益评价", "财务可持续性分析"]},
            {"title": "项目影响效果分析", "subs": ["经济影响分析", "社会影响分析", "生态环境影响分析", "资源和能源利用效果分析"]},
            {"title": "项目风险管控方案", "subs": ["主要风险识别与评价", "盈亏平衡分析", "敏感性分析", "风险管控方案与应急预案"]},
            {"title": "研究结论及建议", "subs": ["主要研究结论", "问题与建议"]},
        ],
    },
    "gov9": {
        "label": "政府投资项目（9 章兼容结构 · 历史）",
        "hint": "早期工程结构：把发改委2023政府投资通用大纲中的“背景必要性”和“需求产出”合并为一章。仅用于历史工作区兼容，新建不推荐。",
        "chapters": [
            {"title": "概述", "subs": ["项目概况", "建设单位概况", "编制依据和研究范围", "主要结论和建议"]},
            {"title": "项目建设背景、需求分析及产出方案", "subs": ["规划政策符合性", "项目建设背景", "项目建设的必要性", "需求分析", "项目建设内容、规模及产出方案", "项目商业模式"]},
            {"title": "项目选址与要素保障", "subs": ["项目选址", "项目建设条件", "要素保障分析"]},
            {"title": "项目建设方案", "subs": ["技术方案", "设备方案", "工程方案", "项目招标方案"]},
            {"title": "项目运营方案", "subs": ["运营模式选择", "经营方案", "安全保障方案", "运营管理方案"]},
            {"title": "项目投融资与财务方案", "subs": ["投资估算", "融资方案", "财务效益评价", "财务可持续性分析"]},
            {"title": "项目影响效果分析", "subs": ["经济影响分析", "社会影响分析", "生态环境影响分析", "资源和能源利用效果分析"]},
            {"title": "项目风险管控方案", "subs": ["主要风险识别与评价", "盈亏平衡分析", "敏感性分析", "风险管控方案与应急预案"]},
            {"title": "研究结论及建议", "subs": ["主要研究结论", "问题与建议"]},
        ],
    },
    "ent9": {
        "label": "企业投资项目（发改委新版 9 章）",
        "hint": "适用于企业投资的工业、制造、能源等生产性项目，遵循发改委2023企业投资参考大纲（9 章，侧重市场需求/商业模式/产品产能）。",
        "chapters": [
            {"title": "概述", "subs": ["项目概况", "建设单位概况", "研究的依据和内容", "主要结论和建议"]},
            {"title": "项目建设背景、需求分析及产出方案", "subs": ["规划政策符合性", "项目建设的必要性", "项目市场需求分析", "项目建设内容、规模和产出方案", "项目商业模式"]},
            {"title": "项目选址与要素保障", "subs": ["项目选址", "项目建设条件", "要素保障分析"]},
            {"title": "项目建设方案", "subs": ["技术方案", "设备方案", "工程方案", "建设管理方案", "数字化方案"]},
            {"title": "项目运营方案", "subs": ["生产经营方案", "安全保障方案", "运营管理方案"]},
            {"title": "项目投融资与财务方案", "subs": ["投资估算", "融资方案", "财务效益评价", "财务可持续性分析"]},
            {"title": "项目影响效果分析", "subs": ["经济影响分析", "社会影响分析", "生态环境影响分析", "资源和能源利用效果分析"]},
            {"title": "项目风险管控方案", "subs": ["风险识别及评价", "风险管控方案", "风险应急预案"]},
            {"title": "研究结论及建议", "subs": ["主要研究结论", "问题与建议"]},
        ],
    },
    "ent14": {
        "label": "企业投资项目（14 章 · 2006 旧版，仅兼容历史）",
        "hint": "2006 旧版企业投资 14 章结构，绿科现已改用 ent9；仅用于兼容历史工作区，新建不推荐。",
        "chapters": [
            {"title": "总论", "subs": ["项目概要", "项目申报单位概况", "编制依据和研究范围", "建设规模及内容", "主要研究结论"]},
            {"title": "项目建设背景及必要性", "subs": ["项目建设背景", "项目建设的必要性"]},
            {"title": "建设地点及建设条件", "subs": ["项目建设地点", "建设条件"]},
            {"title": "建设方案", "subs": ["总图布局", "技术标准", "主要建设规模及内容", "工程方案"]},
            {"title": "环境影响与保护", "subs": ["环境保护法规及标准", "项目区环境现状", "项目环境影响分析", "环境保护措施", "环境影响评价"]},
            {"title": "节能评价", "subs": ["节能评价依据", "项目能耗情况", "项目节能措施"]},
            {"title": "劳动保护", "subs": ["设计主要依据", "主要危害因素分析", "劳动保护措施", "劳动保护结论"]},
            {"title": "项目管理、实施进度和工程招标", "subs": ["项目管理", "项目实施进度", "招标方案"]},
            {"title": "社会稳定风险分析", "subs": ["编制依据", "风险调查", "风险识别", "风险估计及初始风险等级判断", "项目风险应对措施", "风险分析结论"]},
            {"title": "投资估算与资金筹措", "subs": ["投资估算", "流动资金估算", "总投资估算", "资金来源", "资本金"]},
            {"title": "财务效益分析", "subs": ["财务效益评价的依据", "营业收入、税金测算", "总成本费用测算", "利润测算", "偿债能力分析", "财务可持续性分析"]},
            {"title": "不确定性分析", "subs": ["盈亏平衡分析", "敏感性分析"]},
            {"title": "社会影响分析", "subs": ["社会影响效果分析", "社会适应性分析", "社会风险及对策"]},
            {"title": "结论与建议", "subs": ["结论", "建议"]},
        ],
    },
    "legacy_b9": {
        "label": "内置 B 型 9 章（旧版，兼容历史工作区）",
        "hint": "早期内置结构，仅用于兼容已创建的历史工作区。",
        "chapters": [{"title": c, "subs": []} for c in REPORT_CHAPTERS],
    },
}

REPORT_STRUCTURES["asset_acquisition"] = {
    "label": "资产收购可行性研究报告",
    "hint": "适用于酒店及其他存量资产收购，覆盖主体、资产边界、历史经营、租约、估值、融资和交易条件。",
    "chapters": [
        {"title": "执行摘要", "subs": ["建议报价", "关键条件", "否决事项"]},
        {"title": "交易主体与资产边界", "subs": ["交易主体", "权属与许可", "资产范围"]},
        {"title": "历史财务与经营标准化", "subs": ["历史报表", "异常调整", "三表勾稽"]},
        {"title": "酒店经营与租赁分析", "subs": ["ADR与入住率", "EBITDAR", "租约与租金覆盖"]},
        {"title": "交易方案与融资", "subs": ["收购价格", "交易税费", "融资与偿债"]},
        {"title": "估值、情景与最高可接受价", "subs": ["六档价格情景", "独立敏感性", "最高可接受价"]},
        {"title": "风险、成交条件与结论", "subs": ["风险清单", "成交前提", "结论与建议"]},
    ],
}


def report_structure(report_type: str = "") -> dict[str, Any]:
    """按 report_type 返回结构定义；未知/为空回退默认。"""
    return REPORT_STRUCTURES.get(report_type or "", REPORT_STRUCTURES[DEFAULT_REPORT_TYPE])


def report_chapter_titles(report_type: str = "") -> list[str]:
    return [c["title"] for c in report_structure(report_type)["chapters"]]


def resolve_report_type(meta: dict[str, Any]) -> str:
    """从工作区 meta 解析结构类型。历史工作区(无该字段)按 legacy_b9 兼容。"""
    rt = str((meta or {}).get("report_type") or "").strip()
    if rt in REPORT_STRUCTURES:
        return rt
    return "legacy_b9"


def workspace_report_type(workspace_id: str) -> str:
    return resolve_report_type(_read_meta(workspace_id))


# 待补充占位标记(参考可研 MISSING="【待补充】";域内内置模板用全角括号占位)。
MISSING_MARKER = "（待补充）"

ISSUE_SOURCES = {"preview_gate", "check_issues", "missing_items", "review_comment"}
ISSUE_SEVERITIES = {"info", "low", "medium", "high", "critical"}
ISSUE_STATUSES = {"open", "in_progress", "resolved", "ignored"}
PROPOSAL_STATUSES = {"proposed", "applied", "rejected"}


class DocServiceError(RuntimeError):
    """文档服务统一异常,携带机器可读 code。"""

    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code
        self.message = message


# ---- 时间 ------------------------------------------------------------------


def _now_iso() -> str:
    return datetime.now().astimezone().replace(microsecond=0).isoformat()


def _new_id(prefix: str) -> str:
    return f"{prefix}_{secrets.token_hex(8)}"


# ---- 路径（存储根：MCP runtime workspace，不读 hermes 配置）-----------------


def _workspace_root(workspace_id: str) -> Path:
    return runtime_workspace.workspace_root(workspace_id)


def _meta_path(workspace_id: str) -> Path:
    return _workspace_root(workspace_id) / "workspace_meta.json"


def _revisions_dir(workspace_id: str) -> Path:
    return _workspace_root(workspace_id) / "revisions"


def _revision_dir(workspace_id: str, revision_id: str) -> Path:
    return _revisions_dir(workspace_id) / revision_id


def _proposals_dir(workspace_id: str) -> Path:
    return _workspace_root(workspace_id) / "agent_proposals"


def _proposal_dir(workspace_id: str, proposal_id: str) -> Path:
    return _proposals_dir(workspace_id) / proposal_id


def _issues_path(workspace_id: str) -> Path:
    return _workspace_root(workspace_id) / "issues" / "issues.json"


def _finance_path(workspace_id: str) -> Path:
    return _workspace_root(workspace_id) / "finance.json"


# ---- 原子 JSON 读写 --------------------------------------------------------


def _read_json(path: Path, default: Any) -> Any:
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return default


def _write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(path)


def _write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(text, encoding="utf-8")
    tmp.replace(path)


# ---- 报告模板与结构 --------------------------------------------------------


def default_report_markdown(title: str = "可行性研究报告", report_type: str = "") -> str:
    """按结构类型生成大纲骨架（含三级小节 + 待补充占位）。"""
    struct = report_structure(report_type)
    lines = [f"# {title}", ""]
    for idx, chapter in enumerate(struct["chapters"], start=1):
        lines.append(f"## {idx}. {chapter['title']}")
        lines.append("")
        subs = chapter.get("subs") or []
        if subs:
            for sidx, sub in enumerate(subs, start=1):
                lines.append(f"### {idx}.{sidx} {sub}")
                lines.append("")
                lines.append("（待补充）")
                lines.append("")
        else:
            lines.append("（待补充）")
            lines.append("")
    return "\n".join(lines).rstrip() + "\n"


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*\S)\s*$")


def parse_revision_sections(markdown: str) -> list[dict[str, Any]]:
    """把报告 markdown 解析为章节列表(参考可研 parse_revision_sections)。

    每节: ``{level, title, anchor, line, body}``。``anchor`` 用标题归一化生成,
    便于审查意见定位。仅按 ``##``(2 级)切分章节主体。
    """
    lines = markdown.splitlines()
    sections: list[dict[str, Any]] = []
    current: Optional[dict[str, Any]] = None
    body_lines: list[str] = []

    def _flush() -> None:
        if current is not None:
            current["body"] = "\n".join(body_lines).strip()
            sections.append(current)

    for i, raw in enumerate(lines):
        m = _HEADING_RE.match(raw)
        if m and len(m.group(1)) == 2:
            _flush()
            title = m.group(2).strip()
            current = {
                "level": 2,
                "title": title,
                "anchor": _anchor_for(title),
                "line": i + 1,
                "body": "",
            }
            body_lines = []
        elif current is not None:
            body_lines.append(raw)
    _flush()
    return sections


def _anchor_for(title: str) -> str:
    cleaned = re.sub(r"^[0-9.\s、]+", "", title).strip()
    return cleaned or title.strip()


def _strip_leading_chapter_title(text: str, target_title: str) -> str:
    """剥离 proposed 开头「与目标章同名」的标题行（任意 #/##/### 级别）。"""
    lines = text.splitlines()
    idx = 0
    while idx < len(lines) and not lines[idx].strip():
        idx += 1
    if idx < len(lines):
        m = re.match(r"^#{0,3}\s*(.+?)\s*$", lines[idx])
        if m and _anchor_for(m.group(1)) == _anchor_for(target_title):
            idx += 1
            while idx < len(lines) and not lines[idx].strip():
                idx += 1
            return "\n".join(lines[idx:]).strip()
    return text.strip()


def merge_single_chapter_proposal(
    base_md: str, target_title: str, proposed_content: str
) -> Optional[str]:
    """把「单章草稿」合并进完整 base 文档，返回完整 proposed_report。"""
    base_secs = parse_revision_sections(base_md)
    if not base_secs:
        return None
    tgt_anchor = _anchor_for(target_title)
    if not any(s["anchor"] == tgt_anchor for s in base_secs):
        return None

    base_lines = base_md.splitlines()
    preamble = "\n".join(base_lines[: base_secs[0]["line"] - 1]).rstrip()
    body_text = _strip_leading_chapter_title(proposed_content, target_title)

    out: list[str] = []
    if preamble:
        out.append(preamble)
        out.append("")
    for s in base_secs:
        out.append(f"## {s['title']}")
        body = body_text if s["anchor"] == tgt_anchor else s["body"]
        if body:
            out.append("")
            out.append(body)
        out.append("")
    return "\n".join(out).rstrip() + "\n"


def validate_report_structure(
    markdown: str,
    report_type: str = "",
    *,
    expected_chapters: Optional[list[str]] = None,
) -> dict[str, Any]:
    """按结构类型校验章节完整性。

    返回 ``{ok, missing_chapters, present_chapters, issues}``。缺少任一章节标题
    即 ``ok=False``。匹配按"章节名包含"宽松判定,容忍编号前缀差异。
    """
    chapters = [
        _anchor_for(str(chapter))
        for chapter in (expected_chapters or [])
        if str(chapter).strip()
    ] or report_chapter_titles(report_type)
    if expected_chapters:
        section_titles = [
            match.group(2).strip()
            for raw_line in str(markdown or "").splitlines()
            if (match := _HEADING_RE.match(raw_line))
        ]
    else:
        section_titles = [s["title"] for s in parse_revision_sections(markdown)]
    normalized = [_anchor_for(t) for t in section_titles]
    missing: list[str] = []
    present: list[str] = []
    for chapter in chapters:
        if any(chapter in t or t in chapter for t in normalized):
            present.append(chapter)
        else:
            missing.append(chapter)
    issues = [f"缺少章节：{c}" for c in missing]
    return {
        "ok": not missing,
        "missing_chapters": missing,
        "present_chapters": present,
        "issues": issues,
    }


def consistency_check(
    workspace_id: str = "",
    *,
    report_text: str = "",
    finance: Optional[dict[str, Any]] = None,
) -> dict[str, Any]:
    """报告一致性自检(参考可研 ``consistency_check(params, fin, report_text)``)。

    - **结构**:B 型 9 章是否齐全(缺章 → ``check_issues``)。
    - **财务**:finance.json 的 ``required_markers`` 是否都在正文出现(缺失 → ``missing_items``)。
    - **待补充**:正文是否仍含占位标记 ``（待补充）``(逐章 → ``missing_items``)。

    返回 ``{ok, check_issues, missing_items, structure}``。
    """
    if not report_text and workspace_id:
        _rev, report_text = _current_revision_content(workspace_id)
    if finance is None and workspace_id:
        finance = finance_summary(workspace_id)
    finance = finance if isinstance(finance, dict) else {}

    rt = workspace_report_type(workspace_id) if workspace_id else ""
    structure = validate_report_structure(report_text, rt)
    check_issues: list[str] = list(structure["issues"])
    missing_items: list[str] = []

    markers = finance.get("required_markers")
    if isinstance(markers, (list, tuple)):
        for marker in markers:
            text = str(marker).strip()
            if text and text not in report_text:
                missing_items.append(f"正文缺少关键财务指标：{text}")

    for sec in parse_revision_sections(report_text):
        if MISSING_MARKER in str(sec.get("body") or ""):
            missing_items.append(f"章节待补充：{sec.get('title')}")

    # M5 T5.4：环保方案齐全性校验（环评硬门槛）。
    if workspace_id:
        try:
            meta = _read_meta(workspace_id)
            req = meta.get("requirement") if isinstance(meta, dict) else {}
            req = req if isinstance(req, dict) else {}
            fin_req = req.get("finance") if isinstance(req.get("finance"), dict) else {}
            if fin_req.get("is_operating"):
                from lvke_mcp.domains.finance import env_templates

                prof = env_templates.env_profile(str(req.get("industry") or ""))
                if prof.get("available"):
                    kw_hit = any(
                        k in report_text
                        for k in ("环保", "环境影响", "三废", "废水", "废气", "污染", "达标排放")
                    )
                    if not kw_hit:
                        missing_items.append(
                            "环保方案缺失：经营性项目须在影响效果章体现三废治理措施与环保投入"
                        )
        except Exception:  # noqa: BLE001 - 环保校验失败不阻断自检
            pass

    # M1 T1.2：财务数值勾稽层（资金筹措=总投资、三处 IRR 一致、成本勾稽、建设期利息）。
    finance_reconciliation: list[dict[str, Any]] = []
    if workspace_id:
        try:
            from lvke_mcp.domains.finance import finance_model

            fm_r = workspace_finance_model(workspace_id)
            finance_reconciliation = finance_model.check_consistency(fm_r)
        except Exception:  # noqa: BLE001 - 勾稽失败不阻断自检主流程
            finance_reconciliation = []
    for chk in finance_reconciliation:
        if not chk.get("ok"):
            check_issues.append(f"财务勾稽不一致：{chk.get('rule')}（{chk.get('detail')}）")

    return {
        "ok": not check_issues and not missing_items,
        "check_issues": check_issues,
        "missing_items": missing_items,
        "structure": structure,
        "finance_reconciliation": finance_reconciliation,
    }


# ---- 工作区与修订 ----------------------------------------------------------


def _default_meta(workspace_id: str) -> dict[str, Any]:
    return {
        "schema_version": "keyui_workspace.v1",
        "workspace_id": workspace_id,
        "title": "",
        "current_revision_id": "",
        "created_at": "",
        "updated_at": "",
    }


def _read_meta(workspace_id: str) -> dict[str, Any]:
    # MCP 域内无 hermes WAL 控制面：直接读 JSON，损坏时回退默认 meta。
    value = _read_json(_meta_path(workspace_id), None)
    if isinstance(value, dict):
        return value
    return _default_meta(workspace_id)


def _write_meta(workspace_id: str, metadata: dict[str, Any]) -> None:
    value = dict(metadata)
    value.pop("_metadata_source", None)
    value.pop("_read_only_recovery", None)
    value["schema_version"] = "keyui_workspace.v1"
    _write_json(_meta_path(workspace_id), value)


def _save_revision(workspace_id: str, *, content: str, parent_id: str, summary: str, source: str) -> dict[str, Any]:
    revision_id = _new_id("rev")
    rev_dir = _revision_dir(workspace_id, revision_id)
    _write_text(rev_dir / "report.md", content)
    meta = {
        "schema_version": "keyui_revision.v1",
        "revision_id": revision_id,
        "parent_id": parent_id,
        "summary": summary,
        "source": source,
        "created_at": _now_iso(),
    }
    _write_json(rev_dir / "meta.json", meta)
    return meta


def ensure_workspace(
    workspace_id: str,
    *,
    title: str = "可行性研究报告",
    report_type: str = "",
    doc_kind: str = "",
    requirement: Optional[dict[str, Any]] = None,
    cover: Optional[dict[str, Any]] = None,
) -> dict[str, Any]:
    """确保工作区存在;不存在则按结构类型创建初始修订(默认发改委新版 9 章)。"""
    if not workspace_id or not re.fullmatch(r"[A-Za-z0-9_.-]{1,64}", workspace_id):
        raise DocServiceError("invalid_workspace_id", "工作区 id 不合法。")
    lock_dir = runtime_workspace.data_root() / "workspace_locks"
    lock_dir.mkdir(parents=True, exist_ok=True)
    with FileLock(str(lock_dir / f"{workspace_id}.init.lock"), timeout=30):
        meta = _read_meta(workspace_id)
        if meta.get("current_revision_id"):
            return meta
        rt = report_type if report_type in REPORT_STRUCTURES else DEFAULT_REPORT_TYPE
        dk = doc_kind if doc_kind in DOC_KINDS else DEFAULT_DOC_KIND
        now = _now_iso()
        rev = _save_revision(
            workspace_id,
            content=default_report_markdown(title, rt),
            parent_id="",
            summary="初始化报告大纲",
            source="bootstrap",
        )
        meta = {
            "schema_version": "keyui_workspace.v1",
            "workspace_id": workspace_id,
            "title": title,
            "report_type": rt,
            "doc_kind": dk,
            "requirement": dict(requirement) if isinstance(requirement, dict) else {},
            "cover": dict(cover) if isinstance(cover, dict) else {},
            "current_revision_id": rev["revision_id"],
            "created_at": now,
            "updated_at": now,
        }
        _write_meta(workspace_id, meta)
        return meta


def _current_revision_content(workspace_id: str) -> tuple[str, str]:
    """返回 ``(current_revision_id, markdown)``;工作区未初始化则自动初始化。"""
    meta = ensure_workspace(workspace_id)
    rev_id = str(meta.get("current_revision_id") or "")
    content = ""
    if rev_id:
        path = _revision_dir(workspace_id, rev_id) / "report.md"
        content = path.read_text(encoding="utf-8") if path.exists() else ""
    return rev_id, content


def revision_content(workspace_id: str, revision_id: str) -> Optional[str]:
    path = _revision_dir(workspace_id, revision_id) / "report.md"
    if not path.exists():
        return None
    return path.read_text(encoding="utf-8")


# ---- snapshot / read -------------------------------------------------------


def load_workspace_snapshot(workspace_id: str) -> dict[str, Any]:
    """工作区快照(参考可研 load_workspace_snapshot)。"""
    meta = ensure_workspace(workspace_id)
    rev_id, content = _current_revision_content(workspace_id)
    rt = resolve_report_type(meta)
    struct = report_structure(rt)
    dk = resolve_doc_kind(meta)
    return {
        "workspace_id": workspace_id,
        "title": meta.get("title") or "",
        "report_type": rt,
        "report_type_label": struct.get("label", ""),
        "doc_kind": dk,
        "doc_kind_label": DOC_KINDS.get(dk, {}).get("label", ""),
        "requirement": meta.get("requirement") or {},
        "cover": meta.get("cover") or {},
        "report_outline": [c["title"] for c in struct["chapters"]],
        "current_revision_id": rev_id,
        "updated_at": meta.get("updated_at") or "",
        "sections": parse_revision_sections(content),
        "issue_center": list_issues(workspace_id),
        "finance_summary": finance_summary(workspace_id),
    }


def read_document(workspace_id: str, *, section: str = "", revision_id: str = "") -> dict[str, Any]:
    """读取文档全文或某章节(参考可研 doc_read)。"""
    if revision_id:
        content = revision_content(workspace_id, revision_id)
        if content is None:
            raise DocServiceError("revision_not_found", f"修订不存在：{revision_id}")
    else:
        revision_id, content = _current_revision_content(workspace_id)
    if section:
        for sec in parse_revision_sections(content):
            if section in sec["title"] or section == sec["anchor"]:
                return {
                    "workspace_id": workspace_id,
                    "revision_id": revision_id,
                    "section": sec["title"],
                    "content": sec["body"],
                }
        raise DocServiceError("section_not_found", f"章节不存在：{section}")
    return {
        "workspace_id": workspace_id,
        "revision_id": revision_id,
        "content": content,
        "sections": [s["title"] for s in parse_revision_sections(content)],
    }


# ---- finance / context -----------------------------------------------------


def finance_summary(workspace_id: str) -> dict[str, Any]:
    """只读财务摘要(读 finance.json,缺省返回空结构)。"""
    data = _read_json(_finance_path(workspace_id), {})
    if not isinstance(data, dict):
        return {}
    return data


def workspace_finance_model(workspace_id: str, *, force_flat: bool = False) -> dict[str, Any]:
    """基于工作区 requirement.finance 运行完整财务模型并返回 13 表结果。

    兼容入口：内部委托 ``domains.finance.run_service``。
    """
    try:
        from lvke_mcp.domains.finance import run_service

        return run_service.run_workspace_finance_model(
            workspace_id,
            force_flat=force_flat,
            allow_prepare_llm=not force_flat,
            record_audit=False,  # 审计由 API/报告链路显式登记，避免隐式写副作用
            mode="estimate_preview",
        )
    except Exception as exc:  # noqa: BLE001 - 财务模型失败不阻断生成
        try:
            meta = _read_meta(workspace_id)
            req = meta.get("requirement") or {}
            fin = dict((req.get("finance") if isinstance(req, dict) else {}) or {})
        except Exception:  # noqa: BLE001
            fin = {}
        return {
            "available": False,
            "ok": False,
            "reason": str(exc)[:200],
            "finance_inputs": fin,
        }


# ---- issue_center ----------------------------------------------------------


def list_issues(workspace_id: str, *, status: str = "", source: str = "") -> list[dict[str, Any]]:
    issues = _read_json(_issues_path(workspace_id), [])
    if not isinstance(issues, list):
        issues = []
    result = []
    for issue in issues:
        if status and issue.get("status") != status:
            continue
        if source and issue.get("source") != source:
            continue
        result.append(issue)
    return result


# ---- proposal --------------------------------------------------------------


def _read_proposal(workspace_id: str, proposal_id: str) -> dict[str, Any]:
    meta = _read_json(_proposal_dir(workspace_id, proposal_id) / "meta.json", None)
    if not isinstance(meta, dict):
        raise DocServiceError("proposal_not_found", f"提案不存在：{proposal_id}")
    return meta


def create_agent_proposal(
    workspace_id: str,
    *,
    session_id: str = "",
    summary: str,
    proposed_content: str,
    target_sections: Optional[list[str]] = None,
    basis: str = "",
    expected_outline: Optional[list[str]] = None,
) -> dict[str, Any]:
    """创建文档修改提案(参考可研 create_agent_proposal:写 proposed_report.md + diff.html + meta)。"""
    meta = ensure_workspace(workspace_id)
    base_rev = str(meta.get("current_revision_id") or "")
    base_content = revision_content(workspace_id, base_rev) or ""
    proposed_content = proposed_content if proposed_content is not None else ""
    if not proposed_content.strip():
        raise DocServiceError("empty_proposal", "提案内容为空。")
    rt = resolve_report_type(meta)

    proposal_id = _new_id("prop")
    pdir = _proposal_dir(workspace_id, proposal_id)

    # 单章/少章提案自动合并进完整文档（幂等：agent 若已传完整文档则不进此分支）。
    merged_from_single_chapter = False
    expected_outline = [str(item) for item in (expected_outline or []) if str(item).strip()]
    pre_struct = validate_report_structure(
        proposed_content, rt, expected_chapters=expected_outline,
    )
    if pre_struct["missing_chapters"] and target_sections and len(target_sections) == 1:
        merged = merge_single_chapter_proposal(
            base_content, str(target_sections[0]), proposed_content
        )
        if merged is not None:
            vmerged = validate_report_structure(
                merged, rt, expected_chapters=expected_outline,
            )
            if not vmerged["missing_chapters"]:
                proposed_content = merged
                merged_from_single_chapter = True

    _write_text(pdir / "proposed_report.md", proposed_content)
    html = _html_diff(base_content, proposed_content, base_rev, "proposed")
    _write_text(pdir / "diff.html", html)

    structure = validate_report_structure(
        proposed_content, rt, expected_chapters=expected_outline,
    )
    now = _now_iso()
    record = {
        "id": proposal_id,
        "workspace_id": workspace_id,
        "session_id": session_id,
        "base_revision_id": base_rev,
        "target_sections": list(target_sections or []),
        "edit_summary": summary,
        "basis": basis,
        "expected_outline": expected_outline,
        "proposed_report_path": str(pdir / "proposed_report.md"),
        "diff_path": str(pdir / "diff.html"),
        "structure_ok": structure["ok"],
        "structure_issues": structure["issues"],
        "merged_from_single_chapter": merged_from_single_chapter,
        "status": "proposed",
        "applied_revision_id": "",
        "created_at": now,
        "updated_at": now,
    }
    _write_json(pdir / "meta.json", record)
    return record


def diff_agent_proposal(workspace_id: str, *, proposal_id: str = "", from_revision: str = "", to_revision: str = "") -> dict[str, Any]:
    """生成 diff(参考可研 diff_agent_proposal,difflib.HtmlDiff 章节级)。"""
    if proposal_id:
        proposal = _read_proposal(workspace_id, proposal_id)
        base_rev = str(proposal.get("base_revision_id") or "")
        base_content = revision_content(workspace_id, base_rev) or ""
        proposed_path = Path(proposal["proposed_report_path"])
        proposed_content = proposed_path.read_text(encoding="utf-8") if proposed_path.exists() else ""
        html = _html_diff(
            base_content, proposed_content,
            _revision_label(workspace_id, base_rev) or "当前版本", "拟应用提案",
        )
        return {"workspace_id": workspace_id, "proposal_id": proposal_id, "html_diff": html}
    # revision-to-revision
    left = revision_content(workspace_id, from_revision) if from_revision else _current_revision_content(workspace_id)[1]
    right = revision_content(workspace_id, to_revision) if to_revision else _current_revision_content(workspace_id)[1]
    html = _html_diff(
        left or "", right or "",
        _revision_label(workspace_id, from_revision) or "基准版本",
        _revision_label(workspace_id, to_revision) or "对照版本",
    )
    return {"workspace_id": workspace_id, "html_diff": html}


_REVISION_SOURCE_CN = {
    "bootstrap": "初始化", "fast_draft": "首版生成", "enhance_pass": "专业化增强",
    "apply": "应用提案", "agent": "AI 修改", "manual": "手动编辑",
}


def _revision_label(workspace_id: str, revision_id: str) -> str:
    """把修订编码映射为友好列头「第 N 版 · YYYY-MM-DD HH:mm · 来源」（D-2）。"""
    if not revision_id:
        return ""
    try:
        revs = list_revisions(workspace_id)  # 最新在前
    except Exception:  # noqa: BLE001
        return revision_id[:10]
    total = len(revs)
    for idx, r in enumerate(revs):
        if r.get("revision_id") == revision_id:
            no = total - idx
            tag = "最新版" if r.get("is_current") else ("初版" if no == 1 else f"第 {no} 版")
            when = str(r.get("created_at") or "")[:16].replace("T", " ")
            src = _REVISION_SOURCE_CN.get(str(r.get("source") or ""), "")
            return f"{tag} · {when}{(' · ' + src) if src else ''}"
    return revision_id[:10]


def _html_diff(left: str, right: str, left_label: str, right_label: str) -> str:
    differ = difflib.HtmlDiff(wrapcolumn=80)
    return differ.make_file(
        left.splitlines(),
        right.splitlines(),
        fromdesc=left_label,
        todesc=right_label,
        context=True,
        numlines=3,
    )


def apply_agent_proposal(
    workspace_id: str,
    proposal_id: str,
    *,
    actor: str = "",
    readonly: bool = False,
    enforce_structure: bool = True,
) -> dict[str, Any]:
    """应用提案,多步校验后落新修订(参考可研 apply_agent_proposal 校验顺序)。

    校验顺序:① 只读锁 → ② 提案存在 → ③ status==proposed → ④ 版本新鲜度
    (base_revision_id == 当前修订) → ⑤ 结构校验 → ⑥ 财务一致性(轻量) → ⑦ 落版本。
    任一步失败抛 :class:`DocServiceError` 并阻断。
    """
    # ① 只读锁
    if readonly:
        raise DocServiceError("readonly", "工作区处于只读模式（被其他会话占用），无法应用提案。")
    # ② 提案存在
    proposal = _read_proposal(workspace_id, proposal_id)
    # ②.5 归属:提案必须属于本工作区。
    if str(proposal.get("workspace_id") or workspace_id) != workspace_id:
        raise DocServiceError("proposal_ownership", "提案不属于当前工作区。")
    # ③ 状态
    if proposal.get("status") != "proposed":
        raise DocServiceError(
            "invalid_proposal_status",
            f"提案状态为 {proposal.get('status')}，只能应用 proposed 状态的提案。",
        )
    # ④ 版本新鲜度
    meta = ensure_workspace(workspace_id)
    current_rev = str(meta.get("current_revision_id") or "")
    if str(proposal.get("base_revision_id") or "") != current_rev:
        raise DocServiceError(
            "stale_proposal",
            "提案基于的版本已过期（文档已被更新），请基于最新版本重新提案。",
        )
    proposed_path = Path(proposal["proposed_report_path"])
    if not proposed_path.exists():
        raise DocServiceError("proposal_payload_missing", "提案正文文件缺失。")
    proposed_content = proposed_path.read_text(encoding="utf-8")
    # ⑤ 结构校验
    if enforce_structure:
        structure = validate_report_structure(
            proposed_content,
            resolve_report_type(meta),
            expected_chapters=list(proposal.get("expected_outline") or []),
        )
        if not structure["ok"]:
            raise DocServiceError(
                "structure_invalid",
                "结构校验未通过：" + "；".join(structure["issues"]),
            )
    # ⑥ 财务一致性(轻量:占位通过,有 finance.json 时校验金额标记存在)
    _check_finance_consistency(workspace_id, proposed_content)
    # ⑦ 落版本
    rev = _save_revision(
        workspace_id,
        content=proposed_content,
        parent_id=current_rev,
        summary=str(proposal.get("edit_summary") or "应用 agent 提案"),
        source="agent_proposal",
    )
    meta["current_revision_id"] = rev["revision_id"]
    meta["updated_at"] = _now_iso()
    _write_meta(workspace_id, meta)

    proposal["status"] = "applied"
    proposal["applied_revision_id"] = rev["revision_id"]
    proposal["updated_at"] = _now_iso()
    _write_json(_proposal_dir(workspace_id, proposal_id) / "meta.json", proposal)
    return {
        "workspace_id": workspace_id,
        "proposal_id": proposal_id,
        "applied_revision_id": rev["revision_id"],
        "status": "applied",
    }


def _check_finance_consistency(workspace_id: str, proposed_content: str) -> None:
    """财务一致性校验(委托具名的 ``consistency_check`` 的财务部分)。

    有 finance.json 且声明了 ``required_markers`` 时,要求这些标记词出现在正文,
    否则阻断 apply;占位章节(``（待补充）``)只进入待补充清单,不阻断 apply。
    """
    finance = finance_summary(workspace_id)
    markers = finance.get("required_markers") if isinstance(finance, dict) else None
    if not markers:
        return
    result = consistency_check(workspace_id, report_text=proposed_content, finance=finance)
    missing_markers = [
        item.split("：", 1)[-1]
        for item in result["missing_items"]
        if item.startswith("正文缺少关键财务指标")
    ]
    if missing_markers:
        raise DocServiceError(
            "finance_inconsistent",
            "财务一致性校验未通过，正文缺少关键财务指标：" + "、".join(missing_markers),
        )


# ---- 修订列表（供 diff 版本标签）-------------------------------------------


def list_revisions(workspace_id: str) -> list[dict[str, Any]]:
    """列出工作区所有修订(按创建时间倒序,标注当前修订)。"""
    meta = _read_meta(workspace_id)
    current = str(meta.get("current_revision_id") or "")
    rdir = _revisions_dir(workspace_id)
    if not rdir.is_dir():
        return []
    revs: list[dict[str, Any]] = []
    for child in rdir.iterdir():
        if not child.is_dir():
            continue
        rmeta = _read_json(child / "meta.json", None)
        if not isinstance(rmeta, dict):
            rmeta = _read_json(child / "revision.json", None)
        if isinstance(rmeta, dict) and rmeta.get("revision_id"):
            revs.append(
                {
                    "revision_id": rmeta.get("revision_id"),
                    "parent_id": rmeta.get("parent_id") or "",
                    "summary": rmeta.get("summary") or "",
                    "source": rmeta.get("source") or "",
                    "created_at": rmeta.get("created_at") or "",
                    "is_current": rmeta.get("revision_id") == current,
                }
            )
    revs.sort(key=lambda r: str(r.get("created_at") or ""), reverse=True)
    return revs


# ---- 生成任务落盘（重启不丢进度，可恢复状态与续跑）-------------------------


def _gen_task_path(workspace_id: str) -> Path:
    return _workspace_root(workspace_id) / "gen_task.json"


def _gen_task_snapshot_path(
    workspace_id: str,
    task_id: str,
) -> Path:
    """Return the durable per-task snapshot path for a generated task id."""
    normalized = str(task_id or "").strip()
    if not re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9_.-]{0,127}", normalized):
        raise ValueError("invalid report generation task id")
    return _workspace_root(workspace_id) / "gen_tasks" / f"{normalized}.json"


def _gen_task_is_latest(candidate: dict[str, Any], current: Any) -> bool:
    """Whether ``candidate`` may replace the legacy latest-task snapshot."""
    if not isinstance(current, dict):
        return True
    candidate_id = str(candidate.get("task_id") or "")
    current_id = str(current.get("task_id") or "")
    if candidate_id and candidate_id == current_id:
        return True
    try:
        candidate_created = float(candidate.get("created_at") or 0)
        current_created = float(current.get("created_at") or 0)
    except (TypeError, ValueError):
        return True
    return candidate_created >= current_created


def save_gen_task(workspace_id: str, task: dict[str, Any]) -> None:
    """持久化生成任务快照：按 task_id 落单文件 + 维护 legacy latest 单例。"""
    try:
        task_id = str(task.get("task_id") or "").strip()
        if task_id:
            _write_json(
                _gen_task_snapshot_path(workspace_id, task_id),
                task,
            )
        latest_path = _gen_task_path(workspace_id)
        if not task_id or _gen_task_is_latest(task, _read_json(latest_path, None)):
            _write_json(latest_path, task)
    except Exception:  # noqa: BLE001
        pass


def load_gen_task(
    workspace_id: str,
    task_id: str = "",
) -> Optional[dict[str, Any]]:
    """Load a task by id, or the legacy latest snapshot when id is omitted.

    A task-id lookup falls back to a matching legacy singleton so snapshots
    written before per-task history was introduced remain addressable.
    """
    normalized = str(task_id or "").strip()
    if normalized:
        try:
            data = _read_json(
                _gen_task_snapshot_path(workspace_id, normalized),
                None,
            )
        except ValueError:
            return None
        if isinstance(data, dict):
            return data
        legacy = _read_json(_gen_task_path(workspace_id), None)
        if isinstance(legacy, dict) and str(legacy.get("task_id") or "") == normalized:
            return legacy
        return None
    data = _read_json(_gen_task_path(workspace_id), None)
    if isinstance(data, dict):
        return data
    rows = list_gen_tasks(workspace_id, limit=1)
    return rows[0] if rows else None


def list_gen_tasks(
    workspace_id: str,
    *,
    owner_user_id: str = "",
    limit: int = 100,
) -> list[dict[str, Any]]:
    """List durable report-generation snapshots newest first.

    The legacy singleton is included when it has not yet been migrated into
    ``gen_tasks``.  Optional owner filtering prevents a user's implicit
    "latest" lookup from adopting another user's task in a shared workspace.
    """

    rows: dict[str, dict[str, Any]] = {}
    base_directory = _workspace_root(workspace_id) / "gen_tasks"
    if base_directory.is_dir():
        for path in base_directory.glob("*.json"):
            data = _read_json(path, None)
            if isinstance(data, dict) and data.get("task_id"):
                rows.setdefault(str(data["task_id"]), data)
    legacy = _read_json(_gen_task_path(workspace_id), None)
    if isinstance(legacy, dict) and legacy.get("task_id"):
        rows.setdefault(str(legacy["task_id"]), legacy)
    owner = str(owner_user_id or "").strip()
    values = [
        row for row in rows.values()
        if not owner or str(row.get("owner_user_id") or "").strip() == owner
    ]

    def sort_key(row: dict[str, Any]) -> tuple[float, str]:
        try:
            created = float(row.get("created_at") or 0)
        except (TypeError, ValueError):
            created = 0.0
        return created, str(row.get("task_id") or "")

    values.sort(key=sort_key, reverse=True)
    return values[: max(1, min(int(limit or 100), 500))]


# ---- DOCX 导出(方案 5.4,pandoc 优先 / python-docx 回退) -------------------


def _export_docx_via_pandoc(content: str) -> Optional[bytes]:
    """优先用 pandoc 把 markdown 转 docx(标题/列表/表格保真度更高)。

    pandoc 未安装、调用失败或超时时返回 ``None``,由调用方回退到 python-docx。
    走外部进程,无 pandoc 的环境零影响(纯回退)。
    """
    import shutil
    import subprocess
    import tempfile

    pandoc = shutil.which("pandoc")
    if not pandoc:
        return None
    with tempfile.TemporaryDirectory() as tmp:
        out_path = Path(tmp) / "report.docx"
        try:
            proc = subprocess.run(
                [pandoc, "-f", "markdown", "-t", "docx", "-o", str(out_path)],
                input=content.encode("utf-8"),
                capture_output=True,
                timeout=60,
            )
        except (OSError, subprocess.SubprocessError):
            return None
        if proc.returncode != 0 or not out_path.exists():
            return None
        try:
            data = out_path.read_bytes()
        except OSError:
            return None
        return data or None


_DOCX_IMAGE_LINE_RE = re.compile(
    r"^\s*!\[(?P<alt>[^\]]*)\]\(\s*"
    r"(?P<source><[^>]+>|[^\s)]+)"
    r"(?:\s+(?P<quote>[\"'])(?P<title>.*?)(?P=quote))?\s*\)\s*$"
)
_DOCX_APPENDIX_HEADING_RE = re.compile(r"^(?:附表|附录|附件(?!索引))\s*", re.IGNORECASE)
_DOCX_PAGE_BREAK_MARKERS = {
    "<!-- pagebreak -->",
    "<!-- page-break -->",
    r"\newpage",
    r"\pagebreak",
}
_DOCX_MAX_IMAGE_BYTES = 25 * 1024 * 1024
_DOCX_DATA_IMAGE_TYPES = {
    "image/bmp",
    "image/gif",
    "image/jpeg",
    "image/png",
    "image/tiff",
}


def _docx_image_bytes(
    source: str,
    *,
    image_base_dir: Optional[Path],
) -> bytes:
    """Resolve a Markdown image without network access or path traversal."""

    import base64
    import binascii
    from urllib.parse import unquote, urlparse
    from urllib.request import url2pathname

    value = str(source or "").strip()
    if value.startswith("<") and value.endswith(">"):
        value = value[1:-1].strip()
    if value.lower().startswith("data:"):
        header, separator, encoded = value.partition(",")
        media_type = header[5:].split(";", 1)[0].lower()
        if (
            separator != ","
            or ";base64" not in header.lower()
            or media_type not in _DOCX_DATA_IMAGE_TYPES
        ):
            raise DocServiceError(
                "docx_image_invalid",
                "DOCX 图片 data URI 必须是受支持的 base64 位图。",
            )
        try:
            data = base64.b64decode(encoded, validate=True)
        except (ValueError, binascii.Error) as exc:
            raise DocServiceError(
                "docx_image_invalid",
                "DOCX 图片 data URI 的 base64 数据无效。",
            ) from exc
    else:
        parsed = urlparse(value)
        if parsed.scheme.lower() == "file":
            if parsed.netloc not in {"", "localhost"}:
                raise DocServiceError(
                    "docx_image_invalid",
                    "DOCX 图片不允许使用远程 file URI。",
                )
            path = Path(url2pathname(unquote(parsed.path)))
        elif parsed.scheme and not re.match(r"^[A-Za-z]:[\\/]", value):
            raise DocServiceError(
                "docx_image_invalid",
                "DOCX 图片仅支持本地文件或 data URI，不执行网络下载。",
            )
        else:
            path = Path(unquote(value))

        base = image_base_dir.resolve() if image_base_dir is not None else None
        if not path.is_absolute():
            if base is None:
                raise DocServiceError(
                    "docx_image_invalid",
                    "相对图片路径缺少受控的基准目录。",
                )
            path = base / path
        try:
            resolved = path.resolve(strict=True)
        except OSError as exc:
            raise DocServiceError(
                "docx_image_invalid",
                f"DOCX 图片不存在或不可读取：{path}",
            ) from exc
        if base is not None and not resolved.is_relative_to(base):
            raise DocServiceError(
                "docx_image_invalid",
                "DOCX 图片路径超出工作区基准目录。",
            )
        try:
            if not resolved.is_file() or resolved.stat().st_size > _DOCX_MAX_IMAGE_BYTES:
                raise DocServiceError(
                    "docx_image_invalid",
                    "DOCX 图片不是普通文件或超过 25 MiB 限制。",
                )
            data = resolved.read_bytes()
        except DocServiceError:
            raise
        except OSError as exc:
            raise DocServiceError(
                "docx_image_invalid",
                f"DOCX 图片不可读取：{resolved}",
            ) from exc

    if not data or len(data) > _DOCX_MAX_IMAGE_BYTES:
        raise DocServiceError(
            "docx_image_invalid",
            "DOCX 图片为空或超过 25 MiB 限制。",
        )
    return data


def _export_docx_via_python_docx(
    content: str,
    *,
    image_base_dir: Optional[Path] = None,
) -> bytes:
    """Formal fallback with pagination, sections, images, captions and tables."""
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover - dep declared in pyproject
        raise DocServiceError("docx_unavailable", "python-docx 未安装,无法导出 DOCX。") from exc

    from docx.enum.section import WD_ORIENT, WD_SECTION  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.image.image import Image as DocxImage  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Cm, Emu, Pt  # type: ignore

    import io

    document = Document()
    title = next((m.group(2).strip() for raw in content.splitlines() if (m := _HEADING_RE.match(raw.rstrip())) and len(m.group(1)) == 1), "正式报告")
    document.core_properties.title = title
    section = document.sections[0]

    def configure_section(target: Any, *, landscape: bool) -> None:
        target.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
        target.page_width = Cm(29.7 if landscape else 21)
        target.page_height = Cm(21 if landscape else 29.7)
        target.top_margin = Cm(2.54)
        target.bottom_margin = Cm(2.54)
        target.left_margin = Cm(2.8)
        target.right_margin = Cm(2.6)

    configure_section(section, landscape=False)
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # noqa: SLF001
    for level in range(1, 4):
        style = document.styles[f"Heading {level}"]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")  # noqa: SLF001

    cover_title = document.add_heading(title, level=0)
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()
    toc_title = document.add_paragraph()
    toc_title.add_run("目录").bold = True
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc = document.add_paragraph()
    run = toc.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])  # noqa: SLF001
    document.add_page_break()

    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    document.settings.element.append(update_fields)
    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("第 ")
    page_begin = OxmlElement("w:fldChar")
    page_begin.set(qn("w:fldCharType"), "begin")
    page_instruction = OxmlElement("w:instrText")
    page_instruction.set(qn("xml:space"), "preserve")
    page_instruction.text = "PAGE"
    page_end = OxmlElement("w:fldChar")
    page_end.set(qn("w:fldCharType"), "end")
    footer_run._r.extend([page_begin, page_instruction, page_end])  # noqa: SLF001
    footer.add_run(" 页")

    lines = content.splitlines()
    index = 0
    skipped_title = False
    body_heading_count = 0
    appendix_heading_count = 0
    figure_number = 0
    is_landscape = False
    explicit_landscape = False
    landscape_context_level: Optional[int] = None

    def switch_orientation(*, landscape: bool) -> bool:
        nonlocal is_landscape
        if is_landscape == landscape:
            return False
        next_section = document.add_section(WD_SECTION.NEW_PAGE)
        configure_section(next_section, landscape=landscape)
        is_landscape = landscape
        return True

    def add_picture(image_match: re.Match[str]) -> None:
        nonlocal figure_number
        image_data = _docx_image_bytes(
            image_match.group("source"),
            image_base_dir=image_base_dir,
        )
        stream = io.BytesIO(image_data)
        try:
            image = DocxImage.from_file(stream)
        except Exception as exc:  # noqa: BLE001 - normalize decoder failures
            raise DocServiceError(
                "docx_image_invalid",
                "DOCX 图片内容不是受支持的有效位图。",
            ) from exc

        current_section = document.sections[-1]
        max_width = int(
            current_section.page_width
            - current_section.left_margin
            - current_section.right_margin
        )
        max_height = int(
            current_section.page_height
            - current_section.top_margin
            - current_section.bottom_margin
            - Cm(2)
        )
        scale = min(
            1.0,
            max_width / max(1, int(image.width)),
            max_height / max(1, int(image.height)),
        )
        width = max(1, round(int(image.width) * scale))
        height = max(1, round(int(image.height) * scale))
        stream.seek(0)
        document.add_picture(stream, width=Emu(width), height=Emu(height))
        picture_paragraph = document.paragraphs[-1]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.keep_with_next = True

        figure_number += 1
        caption_text = (
            str(image_match.group("alt") or "").strip()
            or str(image_match.group("title") or "").strip()
            or "图片"
        )
        caption = document.add_paragraph(style="Caption")
        caption.add_run(f"图 {figure_number} {caption_text}")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.keep_together = True

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        lowered = stripped.lower()
        if lowered in _DOCX_PAGE_BREAK_MARKERS:
            document.add_page_break()
            index += 1
            continue
        if lowered == "<!-- section:landscape -->":
            switch_orientation(landscape=True)
            explicit_landscape = True
            landscape_context_level = None
            index += 1
            continue
        if lowered == "<!-- section:portrait -->":
            switch_orientation(landscape=False)
            explicit_landscape = False
            landscape_context_level = None
            index += 1
            continue

        m = _HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            appendix_heading = bool(_DOCX_APPENDIX_HEADING_RE.match(text))
            section_started = False
            if appendix_heading:
                if not is_landscape:
                    section_started = switch_orientation(landscape=True)
                if not explicit_landscape:
                    landscape_context_level = (
                        level
                        if landscape_context_level is None
                        else min(landscape_context_level, level)
                    )
            elif (
                is_landscape
                and not explicit_landscape
                and landscape_context_level is not None
                and level <= landscape_context_level
            ):
                section_started = switch_orientation(landscape=False)
                landscape_context_level = None

            if level == 1:
                if not skipped_title:
                    skipped_title = True
                else:
                    heading = document.add_heading(text, level=1)
                    if body_heading_count and not section_started:
                        heading.paragraph_format.page_break_before = True
                    body_heading_count += 1
            else:
                heading = document.add_heading(text, level=min(level - 1, 9))
                if level == 2:
                    if body_heading_count and not section_started:
                        heading.paragraph_format.page_break_before = True
                    body_heading_count += 1
                if appendix_heading:
                    if appendix_heading_count and not section_started:
                        heading.paragraph_format.page_break_before = True
                    appendix_heading_count += 1
        elif (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and re.match(r"^\s*\|?\s*:?-{3,}", lines[index + 1])
        ):
            table_rows: list[list[str]] = []
            table_rows.append([cell.strip() for cell in stripped.strip("|").split("|")])
            index += 2  # skip markdown separator
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            columns = max(len(row) for row in table_rows)
            temporary_landscape = columns >= 8 and not is_landscape
            if temporary_landscape:
                switch_orientation(landscape=True)
            table = document.add_table(rows=0, cols=columns)
            table.style = "Table Grid"
            for row_index, values in enumerate(table_rows):
                cells = table.add_row().cells
                tr_pr = cells[0]._tc.getparent().get_or_add_trPr()  # noqa: SLF001
                cant_split = OxmlElement("w:cantSplit")
                tr_pr.append(cant_split)
                if row_index == 0:
                    repeat = OxmlElement("w:tblHeader")
                    repeat.set(qn("w:val"), "true")
                    tr_pr.append(repeat)
                for col_index, value in enumerate(values):
                    cells[col_index].text = value
                    if row_index == 0:
                        for paragraph in cells[col_index].paragraphs:
                            for cell_run in paragraph.runs:
                                cell_run.bold = True
            if temporary_landscape:
                switch_orientation(landscape=False)
            continue
        elif image_match := _DOCX_IMAGE_LINE_RE.match(line):
            add_picture(image_match)
        elif re.match(r"^\s*[-*+]\s+", line):
            document.add_paragraph(re.sub(r"^\s*[-*+]\s+", "", line), style="List Bullet")
        elif re.match(r"^\s*\d+[.)]\s+", line):
            document.add_paragraph(re.sub(r"^\s*\d+[.)]\s+", "", line), style="List Number")
        elif line.strip():
            document.add_paragraph(line)
        else:
            document.add_paragraph("")
        index += 1

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def markdown_to_docx(
    content: str,
    *,
    image_base_dir: Optional[Path] = None,
) -> bytes:
    """把任意 markdown 文本导出为 .docx 字节流（供 DR 研究简报等复用导出链路）。

    与 ``export_report_docx`` 同策略：优先 pandoc，回退 python-docx。不依赖工作区修订。
    """
    generated = _export_docx_via_pandoc(content or "")
    if generated is None:
        generated = _export_docx_via_python_docx(
            content or "",
            image_base_dir=image_base_dir,
        )
    from lvke_mcp.domains.reports.docx_fonts import normalize_docx_fonts

    normalized, _audit = normalize_docx_fonts(generated)
    return normalized