"""docx 导出适配：pandoc 路径、图片内联与 python-docx 渲染。"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Optional



from .paths import (
    DocServiceError,
)

from .structure import (
    _HEADING_RE,
)


def _export_docx_via_pandoc(content: str) -> Optional[bytes]:
    """优先用 pandoc 把 markdown 转 docx(标题/列表/表格保真度更高)。

    pandoc 未安装、调用失败或超时时返回 ``None``,由调用方回退到 python-docx。
    走外部进程,无 pandoc 的环境零影响(纯回退)。
    """
    import shutil
    import subprocess
    import tempfile

    pandoc = shutil.which("pandoc")
    if not pandoc:
        return None
    with tempfile.TemporaryDirectory() as tmp:
        out_path = Path(tmp) / "report.docx"
        try:
            proc = subprocess.run(
                [pandoc, "-f", "markdown", "-t", "docx", "-o", str(out_path)],
                input=content.encode("utf-8"),
                capture_output=True,
                timeout=60,
            )
        except (OSError, subprocess.SubprocessError):
            return None
        if proc.returncode != 0 or not out_path.exists():
            return None
        try:
            data = out_path.read_bytes()
        except OSError:
            return None
        return data or None


_DOCX_IMAGE_LINE_RE = re.compile(
    r"^\s*!\[(?P<alt>[^\]]*)\]\(\s*"
    r"(?P<source><[^>]+>|[^\s)]+)"
    r"(?:\s+(?P<quote>[\"'])(?P<title>.*?)(?P=quote))?\s*\)\s*$"
)


_DOCX_APPENDIX_HEADING_RE = re.compile(r"^(?:附表|附录|附件(?!索引))\s*", re.IGNORECASE)


_DOCX_PAGE_BREAK_MARKERS = {
    "<!-- pagebreak -->",
    "<!-- page-break -->",
    r"\newpage",
    r"\pagebreak",
}


_DOCX_MAX_IMAGE_BYTES = 25 * 1024 * 1024


_DOCX_DATA_IMAGE_TYPES = {
    "image/bmp",
    "image/gif",
    "image/jpeg",
    "image/png",
    "image/tiff",
}


def _docx_image_bytes(
    source: str,
    *,
    image_base_dir: Optional[Path],
) -> bytes:
    """Resolve a Markdown image without network access or path traversal."""

    import base64
    import binascii
    from urllib.parse import unquote, urlparse
    from urllib.request import url2pathname

    value = str(source or "").strip()
    if value.startswith("<") and value.endswith(">"):
        value = value[1:-1].strip()
    if value.lower().startswith("data:"):
        header, separator, encoded = value.partition(",")
        media_type = header[5:].split(";", 1)[0].lower()
        if (
            separator != ","
            or ";base64" not in header.lower()
            or media_type not in _DOCX_DATA_IMAGE_TYPES
        ):
            raise DocServiceError(
                "docx_image_invalid",
                "DOCX 图片 data URI 必须是受支持的 base64 位图。",
            )
        try:
            data = base64.b64decode(encoded, validate=True)
        except (ValueError, binascii.Error) as exc:
            raise DocServiceError(
                "docx_image_invalid",
                "DOCX 图片 data URI 的 base64 数据无效。",
            ) from exc
    else:
        parsed = urlparse(value)
        if parsed.scheme.lower() == "file":
            if parsed.netloc not in {"", "localhost"}:
                raise DocServiceError(
                    "docx_image_invalid",
                    "DOCX 图片不允许使用远程 file URI。",
                )
            path = Path(url2pathname(unquote(parsed.path)))
        elif parsed.scheme and not re.match(r"^[A-Za-z]:[\\/]", value):
            raise DocServiceError(
                "docx_image_invalid",
                "DOCX 图片仅支持本地文件或 data URI，不执行网络下载。",
            )
        else:
            path = Path(unquote(value))

        base = image_base_dir.resolve() if image_base_dir is not None else None
        if not path.is_absolute():
            if base is None:
                raise DocServiceError(
                    "docx_image_invalid",
                    "相对图片路径缺少受控的基准目录。",
                )
            path = base / path
        try:
            resolved = path.resolve(strict=True)
        except OSError as exc:
            raise DocServiceError(
                "docx_image_invalid",
                f"DOCX 图片不存在或不可读取：{path}",
            ) from exc
        if base is not None and not resolved.is_relative_to(base):
            raise DocServiceError(
                "docx_image_invalid",
                "DOCX 图片路径超出工作区基准目录。",
            )
        try:
            if not resolved.is_file() or resolved.stat().st_size > _DOCX_MAX_IMAGE_BYTES:
                raise DocServiceError(
                    "docx_image_invalid",
                    "DOCX 图片不是普通文件或超过 25 MiB 限制。",
                )
            data = resolved.read_bytes()
        except DocServiceError:
            raise
        except OSError as exc:
            raise DocServiceError(
                "docx_image_invalid",
                f"DOCX 图片不可读取：{resolved}",
            ) from exc

    if not data or len(data) > _DOCX_MAX_IMAGE_BYTES:
        raise DocServiceError(
            "docx_image_invalid",
            "DOCX 图片为空或超过 25 MiB 限制。",
        )
    return data


def _export_docx_via_python_docx(
    content: str,
    *,
    image_base_dir: Optional[Path] = None,
) -> bytes:
    """Formal fallback with pagination, sections, images, captions and tables."""
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover - dep declared in pyproject
        raise DocServiceError("docx_unavailable", "python-docx 未安装,无法导出 DOCX。") from exc

    from docx.enum.section import WD_ORIENT, WD_SECTION  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.image.image import Image as DocxImage  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Cm, Emu, Pt  # type: ignore

    import io

    document = Document()
    title = next((m.group(2).strip() for raw in content.splitlines() if (m := _HEADING_RE.match(raw.rstrip())) and len(m.group(1)) == 1), "正式报告")
    document.core_properties.title = title
    section = document.sections[0]

    def configure_section(target: Any, *, landscape: bool) -> None:
        target.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
        target.page_width = Cm(29.7 if landscape else 21)
        target.page_height = Cm(21 if landscape else 29.7)
        target.top_margin = Cm(2.54)
        target.bottom_margin = Cm(2.54)
        target.left_margin = Cm(2.8)
        target.right_margin = Cm(2.6)

    configure_section(section, landscape=False)
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")  # noqa: SLF001
    for level in range(1, 4):
        style = document.styles[f"Heading {level}"]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")  # noqa: SLF001

    cover_title = document.add_heading(title, level=0)
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()
    toc_title = document.add_paragraph()
    toc_title.add_run("目录").bold = True
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc = document.add_paragraph()
    run = toc.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])  # noqa: SLF001
    document.add_page_break()

    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    document.settings.element.append(update_fields)
    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("第 ")
    page_begin = OxmlElement("w:fldChar")
    page_begin.set(qn("w:fldCharType"), "begin")
    page_instruction = OxmlElement("w:instrText")
    page_instruction.set(qn("xml:space"), "preserve")
    page_instruction.text = "PAGE"
    page_end = OxmlElement("w:fldChar")
    page_end.set(qn("w:fldCharType"), "end")
    footer_run._r.extend([page_begin, page_instruction, page_end])  # noqa: SLF001
    footer.add_run(" 页")

    lines = content.splitlines()
    index = 0
    skipped_title = False
    body_heading_count = 0
    appendix_heading_count = 0
    figure_number = 0
    is_landscape = False
    explicit_landscape = False
    landscape_context_level: Optional[int] = None

    def switch_orientation(*, landscape: bool) -> bool:
        nonlocal is_landscape
        if is_landscape == landscape:
            return False
        next_section = document.add_section(WD_SECTION.NEW_PAGE)
        configure_section(next_section, landscape=landscape)
        is_landscape = landscape
        return True

    def add_picture(image_match: re.Match[str]) -> None:
        nonlocal figure_number
        image_data = _docx_image_bytes(
            image_match.group("source"),
            image_base_dir=image_base_dir,
        )
        stream = io.BytesIO(image_data)
        try:
            image = DocxImage.from_file(stream)
        except Exception as exc:  # noqa: BLE001 - normalize decoder failures
            raise DocServiceError(
                "docx_image_invalid",
                "DOCX 图片内容不是受支持的有效位图。",
            ) from exc

        current_section = document.sections[-1]
        max_width = int(
            current_section.page_width
            - current_section.left_margin
            - current_section.right_margin
        )
        max_height = int(
            current_section.page_height
            - current_section.top_margin
            - current_section.bottom_margin
            - Cm(2)
        )
        scale = min(
            1.0,
            max_width / max(1, int(image.width)),
            max_height / max(1, int(image.height)),
        )
        width = max(1, round(int(image.width) * scale))
        height = max(1, round(int(image.height) * scale))
        stream.seek(0)
        document.add_picture(stream, width=Emu(width), height=Emu(height))
        picture_paragraph = document.paragraphs[-1]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.keep_with_next = True

        figure_number += 1
        caption_text = (
            str(image_match.group("alt") or "").strip()
            or str(image_match.group("title") or "").strip()
            or "图片"
        )
        caption = document.add_paragraph(style="Caption")
        caption.add_run(f"图 {figure_number} {caption_text}")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.keep_together = True

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        lowered = stripped.lower()
        if lowered in _DOCX_PAGE_BREAK_MARKERS:
            document.add_page_break()
            index += 1
            continue
        if lowered == "<!-- section:landscape -->":
            switch_orientation(landscape=True)
            explicit_landscape = True
            landscape_context_level = None
            index += 1
            continue
        if lowered == "<!-- section:portrait -->":
            switch_orientation(landscape=False)
            explicit_landscape = False
            landscape_context_level = None
            index += 1
            continue

        m = _HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            appendix_heading = bool(_DOCX_APPENDIX_HEADING_RE.match(text))
            section_started = False
            if appendix_heading:
                if not is_landscape:
                    section_started = switch_orientation(landscape=True)
                if not explicit_landscape:
                    landscape_context_level = (
                        level
                        if landscape_context_level is None
                        else min(landscape_context_level, level)
                    )
            elif (
                is_landscape
                and not explicit_landscape
                and landscape_context_level is not None
                and level <= landscape_context_level
            ):
                section_started = switch_orientation(landscape=False)
                landscape_context_level = None

            if level == 1:
                if not skipped_title:
                    skipped_title = True
                else:
                    heading = document.add_heading(text, level=1)
                    if body_heading_count and not section_started:
                        heading.paragraph_format.page_break_before = True
                    body_heading_count += 1
            else:
                heading = document.add_heading(text, level=min(level - 1, 9))
                if level == 2:
                    if body_heading_count and not section_started:
                        heading.paragraph_format.page_break_before = True
                    body_heading_count += 1
                if appendix_heading:
                    if appendix_heading_count and not section_started:
                        heading.paragraph_format.page_break_before = True
                    appendix_heading_count += 1
        elif (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and re.match(r"^\s*\|?\s*:?-{3,}", lines[index + 1])
        ):
            table_rows: list[list[str]] = []
            table_rows.append([cell.strip() for cell in stripped.strip("|").split("|")])
            index += 2  # skip markdown separator
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            columns = max(len(row) for row in table_rows)
            temporary_landscape = columns >= 8 and not is_landscape
            if temporary_landscape:
                switch_orientation(landscape=True)
            table = document.add_table(rows=0, cols=columns)
            table.style = "Table Grid"
            for row_index, values in enumerate(table_rows):
                cells = table.add_row().cells
                tr_pr = cells[0]._tc.getparent().get_or_add_trPr()  # noqa: SLF001
                cant_split = OxmlElement("w:cantSplit")
                tr_pr.append(cant_split)
                if row_index == 0:
                    repeat = OxmlElement("w:tblHeader")
                    repeat.set(qn("w:val"), "true")
                    tr_pr.append(repeat)
                for col_index, value in enumerate(values):
                    cells[col_index].text = value
                    if row_index == 0:
                        for paragraph in cells[col_index].paragraphs:
                            for cell_run in paragraph.runs:
                                cell_run.bold = True
            if temporary_landscape:
                switch_orientation(landscape=False)
            continue
        elif image_match := _DOCX_IMAGE_LINE_RE.match(line):
            add_picture(image_match)
        elif re.match(r"^\s*[-*+]\s+", line):
            document.add_paragraph(re.sub(r"^\s*[-*+]\s+", "", line), style="List Bullet")
        elif re.match(r"^\s*\d+[.)]\s+", line):
            document.add_paragraph(re.sub(r"^\s*\d+[.)]\s+", "", line), style="List Number")
        elif line.strip():
            document.add_paragraph(line)
        else:
            document.add_paragraph("")
        index += 1

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def markdown_to_docx(
    content: str,
    *,
    image_base_dir: Optional[Path] = None,
) -> bytes:
    """把任意 markdown 文本导出为 .docx 字节流（供 DR 研究简报等复用导出链路）。

    与 ``export_report_docx`` 同策略：优先 pandoc，回退 python-docx。不依赖工作区修订。
    """
    generated = _export_docx_via_pandoc(content or "")
    if generated is None:
        generated = _export_docx_via_python_docx(
            content or "",
            image_base_dir=image_base_dir,
        )
    from lvke_mcp.domains.reports.docx_fonts import normalize_docx_fonts

    normalized, _audit = normalize_docx_fonts(generated)
    return normalized
