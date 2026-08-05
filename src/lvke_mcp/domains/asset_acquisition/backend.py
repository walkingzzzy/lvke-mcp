"""Persistence and artifact orchestration for FinanceSpec v3 acquisition runs."""

from __future__ import annotations

import copy
import hashlib
import io
import json
import logging
import math
import mimetypes
import os
import re
import shutil
import threading
import uuid
import zipfile
from concurrent.futures import ThreadPoolExecutor
from contextlib import contextmanager
from collections.abc import Mapping
from datetime import datetime, timezone
from itertools import product
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from filelock import FileLock

from lvke_mcp.runtime.workspace import workspace_root
from lvke_mcp.domains.asset_acquisition.model import (
    INDEPENDENT_SCENARIO_FIELDS,
    AcquisitionModelError,
    apply_scenario,
    run_acquisition_model,
    solve_max_acquisition_price,
)
from lvke_mcp.domains.finance.spec import (
    LATEST_SPEC_VERSION,
    mark_spec_confirmed,
    validate,
    validate_for_formal,
)

_LOCK = threading.RLock()
_LOG = logging.getLogger(__name__)
_RECOVERY_POOL = ThreadPoolExecutor(max_workers=2, thread_name_prefix="acquisition-recovery")
_DEFAULT_IDEMPOTENCY_TTL_SECONDS = 24 * 60 * 60
_UNTRUSTED_EVIDENCE_ASSERTION_KEYS = {
    "binding_hash",
    "source_sha256",
    "source_size_bytes",
    "parse_job",
    "parse_job_id",
    "attempt",
    "evidence_content_hash",
    "integrity_status",
}
_SOURCE_EVIDENCE_FAILURE_MESSAGE = "资料证据状态当前不可验证"
_RUN_VALIDATION_FAILURE_MESSAGE = "资产收购财务输入未通过模型校验"
_RUN_EXECUTION_FAILURE_MESSAGE = "资产收购财务测算执行失败"
_ARTIFACT_GENERATION_FAILURE_MESSAGE = "资产收购正式工件生成失败"


def _is_estimate_preview_spec(spec: Mapping[str, Any]) -> bool:
    if str(spec.get("delivery_mode") or "") != "estimate_preview":
        return False
    assumptions = spec.get("controlled_assumptions")
    if not isinstance(assumptions, list) or not assumptions:
        return False
    required = {
        "field", "value", "unit", "basis", "impact", "sensitivity",
        "validation_condition",
    }
    return all(
        isinstance(item, dict)
        and required <= set(item)
        and all(item.get(key) not in (None, "") for key in required)
        for item in assumptions
    )


# P1-018：两组必填键提成常量，让校验、诊断与 server 层输入 schema 共用一份定义。
RECONSTRUCTION_RECORD_FIELDS = (
    "reconstruction_id", "source_uri", "content_hash", "locator", "source_kind",
    "method", "limitations",
)
PROCESS_ACCEPTANCE_BASIS_FIELDS = (
    "field", "value", "source_ref", "locator", "content_hash", "method", "limitation",
)


def _valid_reconstruction_records(value: Any) -> bool:
    records = value if isinstance(value, list) else []
    required = set(RECONSTRUCTION_RECORD_FIELDS)
    return bool(records) and all(
        isinstance(item, Mapping)
        and required <= set(item)
        and all(item.get(field) not in (None, "") for field in required - {"limitations"})
        and isinstance(item.get("limitations"), list)
        and str(item.get("content_hash") or "").startswith("sha256:")
        for item in records
    )


def _valid_process_acceptance_basis(value: Any) -> bool:
    records = value if isinstance(value, list) else []
    required = set(PROCESS_ACCEPTANCE_BASIS_FIELDS)
    return bool(records) and all(
        isinstance(item, Mapping)
        and required <= set(item)
        and all(item.get(field) not in (None, "") for field in required)
        and str(item.get("content_hash") or "").startswith("sha256:")
        for item in records
    )


def _record_gaps(
    value: Any,
    label: str,
    fields: tuple[str, ...],
    *,
    list_fields: tuple[str, ...] = (),
) -> list[str]:
    """Describe why a reconstruction/basis record array fails its contract.

    ``list_fields`` names keys that must be arrays rather than non-empty scalars,
    mirroring how the matching ``_valid_*`` predicate exempts them from the blank
    check.
    """

    if not isinstance(value, list) or not value:
        return [f"{label} 缺失或为空数组，process_acceptance 要求至少一条记录"]
    gaps: list[str] = []
    for index, item in enumerate(value):
        if not isinstance(item, Mapping):
            gaps.append(f"{label}[{index}] 不是对象")
            continue
        missing = [field for field in fields if field not in item]
        if missing:
            gaps.append(f"{label}[{index}] 缺少必填键：{'、'.join(missing)}")
        blank = [
            field for field in fields
            if field not in list_fields and field in item and item.get(field) in (None, "")
        ]
        if blank:
            gaps.append(f"{label}[{index}] 以下键为空值：{'、'.join(blank)}")
        for field in list_fields:
            if field in item and not isinstance(item.get(field), list):
                gaps.append(f"{label}[{index}].{field} 必须是数组")
        content_hash = str(item.get("content_hash") or "")
        if content_hash and not content_hash.startswith("sha256:"):
            gaps.append(f"{label}[{index}].content_hash 必须以 sha256: 开头")
    return gaps


def process_acceptance_gaps(spec: Mapping[str, Any]) -> list[str]:
    """List the unmet ``process_acceptance`` conditions, most actionable first.

    ``_is_process_acceptance_spec`` is exactly ``not process_acceptance_gaps(spec)``,
    so the gate itself is unchanged — this only names what is missing so the caller
    knows which field to supply instead of reading an opaque error code.
    """

    gaps: list[str] = []
    if str(spec.get("confirmation_scope") or "") != "process_acceptance":
        gaps.append("spec.confirmation_scope 必须为 process_acceptance")
    if str(spec.get("evidence_policy") or "") != "source_reconstructed":
        gaps.append(
            "spec.evidence_policy 必须为 source_reconstructed"
            "（process_acceptance 只用于来源重建资料的流程验收）"
        )
    if spec.get("project_fact_certified") is not False:
        gaps.append("spec.project_fact_certified 必须显式为 false，不认证项目事实")
    if str(spec.get("business_decision_status") or "") != "not_selected":
        gaps.append("spec.business_decision_status 必须为 not_selected")
    gaps.extend(_record_gaps(
        spec.get("reconstruction_records"), "reconstruction_records",
        RECONSTRUCTION_RECORD_FIELDS,
        list_fields=("limitations",),
    ))
    gaps.extend(_record_gaps(
        spec.get("process_acceptance_basis"), "process_acceptance_basis",
        PROCESS_ACCEPTANCE_BASIS_FIELDS,
    ))
    return gaps


def _is_process_acceptance_spec(spec: Mapping[str, Any]) -> bool:
    return not process_acceptance_gaps(spec)




def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _idempotency_ttl_seconds() -> int:
    try:
        value = int(os.environ.get("LVKE_MCP_IDEMPOTENCY_TTL_SECONDS", ""))
    except (TypeError, ValueError):
        value = _DEFAULT_IDEMPOTENCY_TTL_SECONDS
    return max(1, value)


def _root(
    workspace_id: str,
) -> Path:
    return workspace_root(workspace_id) / "finance_acquisition"


def _state_path(
    workspace_id: str,
) -> Path:
    return _root(workspace_id) / "state.json"


@contextmanager
def _state_guard(
    workspace_id: str,
):
    lock_path = _root(workspace_id) / "state.lock"
    lock_path.parent.mkdir(parents=True, exist_ok=True)
    with _LOCK, FileLock(str(lock_path), timeout=30):
        yield


def _load(
    workspace_id: str,
) -> dict[str, Any]:
    try:
        raw = json.loads(
            _state_path(workspace_id).read_text(
                encoding="utf-8"
            )
        )
    except FileNotFoundError:
        raw = {}
    except (json.JSONDecodeError, TypeError) as exc:
        raise RuntimeError(f"acquisition state is corrupt for workspace {workspace_id}") from exc
    return {
        "version": "acquisition_store.v2",
        "specs": dict(raw.get("specs") or {}),
        "runs": dict(raw.get("runs") or {}),
        "idempotency": dict(raw.get("idempotency") or {}),
        "artifacts": dict(raw.get("artifacts") or {}),
        "scenario_matrices": dict(raw.get("scenario_matrices") or {}),
    }


def _save(
    workspace_id: str,
    state: dict[str, Any],
) -> None:
    path = _state_path(workspace_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_name(f"state.{uuid.uuid4().hex}.tmp")
    try:
        with tmp.open("w", encoding="utf-8", newline="\n") as target:
            json.dump(state, target, ensure_ascii=False, indent=2, default=str)
            target.flush()
            os.fsync(target.fileno())
        os.replace(tmp, path)
    finally:
        tmp.unlink(missing_ok=True)


def _hash(value: Any) -> str:
    payload = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"), default=str)
    return "sha256:" + hashlib.sha256(payload.encode("utf-8")).hexdigest()
























def _idempotency_record(scope: str, body_hash: str, **resource: Any) -> dict[str, Any]:
    created = datetime.now(timezone.utc)
    expires = datetime.fromtimestamp(
        created.timestamp() + _idempotency_ttl_seconds(), tz=timezone.utc,
    )
    return {
        "scope": scope,
        "body_hash": body_hash,
        "request_body_hash": body_hash,
        "created_at": created.isoformat(),
        "expires_at": expires.isoformat(),
        **resource,
    }


def _active_idempotency_record(
    records: dict[str, Any], scope: str,
) -> dict[str, Any] | None:
    record = records.get(scope)
    if not isinstance(record, dict):
        return None
    expires_at = str(record.get("expires_at") or "")
    if not expires_at:
        return record
    try:
        expires = datetime.fromisoformat(expires_at.replace("Z", "+00:00"))
    except ValueError:
        records.pop(scope, None)
        return None
    if expires.tzinfo is None:
        expires = expires.replace(tzinfo=timezone.utc)
    if expires <= datetime.now(timezone.utc):
        records.pop(scope, None)
        return None
    return record


def _history_event(status: str, *, request_id: str = "", **details: Any) -> dict[str, Any]:
    return {
        "status": status, "at": _now(),
        "request_id": request_id, "details": details,
    }


def _close_issue(run: dict[str, Any], code: str, *, reason: str) -> None:
    for issue in run.get("issues") or []:
        if issue.get("code") == code and issue.get("status") == "open":
            issue.update({"status": "closed", "closed_at": _now(), "resolution": reason})


def _open_issue(run: dict[str, Any], code: str, detail: str) -> None:
    existing = next(
        (issue for issue in run.get("issues") or [] if issue.get("code") == code and issue.get("status") == "open"),
        None,
    )
    if existing:
        existing.update({"detail": detail, "updated_at": _now()})
        return
    run.setdefault("issues", []).append({
        "code": code, "blocking": True, "status": "open", "detail": detail, "created_at": _now(),
    })


def _migration_binding(spec: dict[str, Any]) -> dict[str, Any]:
    trace = copy.deepcopy(spec.get("migration_trace") or {})
    if not isinstance(trace, dict):
        trace = {}
    return {
        "source_spec_version": trace.get("source_spec_version") or spec.get("version"),
        "migration_trace": trace,
        "migration_steps": copy.deepcopy(trace.get("steps") or []),
    }


def _sanitize_client_evidence_claims(value: Any) -> Any:
    """Remove client-authored integrity facts before hashing or persistence.

    Evidence identifiers and locators remain assertions that the server resolves.
    Pre-built bindings, source hashes, parse attempts and review revisions are
    authority data and are always reconstructed from the workspace source state.
    """

    if isinstance(value, dict):
        cleaned: dict[str, Any] = {}
        for key, item in value.items():
            if key == "evidence_bindings" or key in _UNTRUSTED_EVIDENCE_ASSERTION_KEYS:
                continue
            cleaned[str(key)] = _sanitize_client_evidence_claims(item)
        return cleaned
    if isinstance(value, list):
        return [_sanitize_client_evidence_claims(item) for item in value]
    if isinstance(value, tuple):
        return [_sanitize_client_evidence_claims(item) for item in value]
    return copy.deepcopy(value)


def _bind_spec_evidence(
    workspace_id: str,
    spec: dict[str, Any],
) -> dict[str, Any]:
    from lvke_mcp.domains.finance.evidence_binding import bind_finance_spec_evidence

    return bind_finance_spec_evidence(
        workspace_id,
        spec,
    )


def _evidence_error_strings(binding: dict[str, Any]) -> list[str]:
    rows = [
        *(binding.get("invalid") or []),
        *(binding.get("missing") or []),
        *(binding.get("pending") or []),
    ]
    return [
        "finance_spec.v3 证据绑定未通过: "
        f"{row.get('source_path') or '$'} [{row.get('code') or 'EVIDENCE_BINDING_FAILED'}] "
        f"{row.get('message') or ''}".strip()
        for row in rows
    ]


def _formal_assessment(
    workspace_id: str,
    spec: dict[str, Any],
    *,
    evidence_binding: dict[str, Any] | None = None,
) -> tuple[bool, list[str], dict[str, Any]]:
    schema_ok, schema_errors = validate_for_formal(spec)
    binding = copy.deepcopy(evidence_binding) if evidence_binding is not None else _bind_spec_evidence(
        workspace_id,
        spec,
    )
    errors = [*schema_errors, *_evidence_error_strings(binding)]
    return bool(schema_ok and binding.get("formal_ok") and not errors), errors, binding


def _evidence_blocking_issues(
    binding: dict[str, Any],
    *,
    created_at: str,
) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    for category in ("invalid", "missing", "pending"):
        for item in binding.get(category) or []:
            issues.append({
                "code": str(item.get("code") or "EVIDENCE_BINDING_FAILED"),
                "blocking": True,
                "status": "open",
                "detail": json.dumps(item, ensure_ascii=False, sort_keys=True),
                "source_path": item.get("source_path") or "$",
                "created_at": created_at,
            })
    if not issues and not binding.get("formal_ok"):
        issues.append({
            "code": "EVIDENCE_BINDING_FAILED",
            "blocking": True,
            "status": "open",
            "detail": "FinanceSpec 未形成可发布的服务端证据绑定",
            "created_at": created_at,
        })
    return issues


def _current_evidence_matches_run(
    workspace_id: str,
    run: dict[str, Any],
    spec: dict[str, Any],
) -> tuple[bool, dict[str, Any]]:
    current = _bind_spec_evidence(
        workspace_id,
        spec,
    )
    return bool(
        current.get("formal_ok")
        and current.get("binding_hash") == run.get("evidence_binding_hash")
        and current.get("binding_version") == run.get("evidence_binding_version")
    ), current


def sanitize_spec_input(spec: dict[str, Any]) -> dict[str, Any]:
    """Public contract helper returning the server-trusted spec assertion body."""

    cleaned = _sanitize_client_evidence_claims(spec)
    return cleaned if isinstance(cleaned, dict) else {}


def assess_spec_evidence(
    workspace_id: str,
    spec: dict[str, Any],
) -> dict[str, Any]:
    """Resolve a spec's evidence using only workspace-owned server state."""

    return _bind_spec_evidence(
        workspace_id,
        sanitize_spec_input(spec),
    )


def _decision_thresholds(spec: dict[str, Any]) -> tuple[dict[str, float | None], str]:
    raw = spec.get("decision_thresholds") or {}
    if not isinstance(raw, dict):
        return {}, "DECISION_THRESHOLDS_INVALID"
    try:
        target = float(raw["target_project_irr"])
        minimum_dscr = (
            float(raw["minimum_dscr"])
            if raw.get("minimum_dscr") is not None else None
        )
    except (KeyError, TypeError, ValueError):
        return {}, "DECISION_THRESHOLDS_INVALID"
    if not math.isfinite(target) or not 0 < target <= 5:
        return {}, "DECISION_THRESHOLDS_INVALID"
    if minimum_dscr is not None and (
        not math.isfinite(minimum_dscr) or minimum_dscr <= 0
    ):
        return {}, "DECISION_THRESHOLDS_INVALID"
    return {
        "target_project_irr": target,
        "minimum_dscr": minimum_dscr,
    }, ""


def _same_optional_number(left: Any, right: Any) -> bool:
    if left is None or right is None:
        return left is None and right is None
    try:
        return math.isclose(float(left), float(right), rel_tol=1e-12, abs_tol=1e-12)
    except (TypeError, ValueError):
        return False


def _max_price_validation(
    run: dict[str, Any], spec: dict[str, Any],
) -> tuple[bool, str, dict[str, Any]]:
    thresholds, threshold_error = _decision_thresholds(spec)
    if threshold_error:
        return False, threshold_error, {"decision_thresholds": spec.get("decision_thresholds")}
    analysis = run.get("max_acquisition_price_analysis") or {}
    if not isinstance(analysis, dict) or not analysis:
        return False, "MAX_PRICE_REQUIRED", {"decision_thresholds": thresholds}
    parameters = analysis.get("parameters") or {}
    if (
        not _same_optional_number(
            parameters.get("target_irr"), thresholds["target_project_irr"],
        )
        or not _same_optional_number(
            parameters.get("min_dscr"), thresholds["minimum_dscr"],
        )
    ):
        return False, "MAX_PRICE_THRESHOLD_MISMATCH", {
            "decision_thresholds": thresholds, "parameters": parameters,
        }
    result = analysis.get("result") or {}
    if not result.get("feasible") or not result.get("converged"):
        return False, "MAX_PRICE_NOT_FEASIBLE", {
            "decision_thresholds": thresholds,
            "feasible": bool(result.get("feasible")),
            "converged": bool(result.get("converged")),
            "reason": result.get("reason"),
        }
    expected_hash = _hash({
        key: value for key, value in analysis.items() if key != "analysis_hash"
    })
    if analysis.get("analysis_hash") != expected_hash:
        return False, "MAX_PRICE_HASH_MISMATCH", {
            "expected_analysis_hash": expected_hash,
            "actual_analysis_hash": analysis.get("analysis_hash"),
        }
    return True, "", {"decision_thresholds": thresholds}


def save_spec(
    workspace_id: str, spec: dict[str, Any], *, idempotency_key: str = "", request_id: str = "",
    trusted_confirmation: bool = False,
) -> dict[str, Any]:
    spec = _sanitize_client_evidence_claims(spec)
    if not trusted_confirmation:
        spec["confirmation_status"] = "candidate"
        spec.pop("confirmed_by", None)
        spec.pop("confirmed_at", None)
    body_hash = _hash(spec)
    evidence_binding = _bind_spec_evidence(
        workspace_id,
        spec,
    )
    evidence_binding_hash = str(evidence_binding.get("binding_hash") or "")
    snapshot_hash = _hash({
        "spec_hash": body_hash,
        "evidence_binding_hash": evidence_binding_hash,
        "evidence_binding_version": evidence_binding.get("binding_version"),
    })
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        scope = f"spec:{idempotency_key}" if idempotency_key else ""
        prior = _active_idempotency_record(state["idempotency"], scope) if scope else None
        if prior:
            if prior["body_hash"] != body_hash:
                return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior["spec_id"]}
            return {**state["specs"][prior["spec_id"]], "idempotent_replay": True}
        matching = [
            row for row in state["specs"].values()
            if row.get("spec_hash") == body_hash
            and row.get("evidence_binding_hash") == evidence_binding_hash
        ]
        if matching:
            existing = matching[0]
            if scope:
                state["idempotency"][scope] = _idempotency_record(
                    scope, body_hash, spec_id=existing["spec_id"],
                )
                _save(workspace_id, state)
            return {**existing, "idempotent_replay": True}
        spec_id = f"spec_{uuid.uuid4().hex}"
        revision = max((int(row.get("revision") or 0) for row in state["specs"].values()), default=0) + 1
        row = {
            "ok": True, "spec_id": spec_id, "version": spec.get("version"),
            "spec_hash": body_hash, "revision": revision, "spec": spec,
            "snapshot_hash": snapshot_hash,
            "evidence_binding_version": evidence_binding.get("binding_version"),
            "evidence_binding_hash": evidence_binding_hash,
            "evidence_status": evidence_binding.get("status"),
            "evidence_formal_ok": bool(evidence_binding.get("formal_ok")),
            "evidence_binding": evidence_binding,
            "confirmation_status": spec.get("confirmation_status") or "candidate",
            "request_id": request_id, "created_at": _now(),
        }
        state["specs"][spec_id] = row
        if scope:
            state["idempotency"][scope] = _idempotency_record(
                scope, body_hash, spec_id=spec_id,
            )
        _save(workspace_id, state)
        return row


def confirm_saved_spec(
    workspace_id: str,
    spec_id: str,
    *,
    note: str = "",
    confirmation_scope: str = "project_candidate",
    idempotency_key: str = "",
    request_id: str = "",
) -> dict[str, Any]:
    """Create an immutable confirmed revision from a saved candidate spec."""

    now = _now()
    body_hash = _hash({
        "action": "confirm_spec", "spec_id": spec_id, "note": str(note or ""),
        "confirmation_scope": confirmation_scope,
    })
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        source = state["specs"].get(str(spec_id or ""))
        if not source:
            return {"ok": False, "error": "SPEC_NOT_FOUND"}
        candidate = copy.deepcopy(source.get("spec") or {})
        if confirmation_scope not in {"project_candidate", "process_acceptance"}:
            return {"ok": False, "error": "CONFIRMATION_SCOPE_INVALID"}
        if (
            confirmation_scope == "project_candidate"
            and str(candidate.get("evidence_policy") or "") == "source_reconstructed"
        ):
            return {"ok": False, "error": "PROJECT_FACT_EVIDENCE_MISSING"}
        if confirmation_scope == "process_acceptance":
            candidate.update({
                "confirmation_scope": "process_acceptance",
                "project_fact_certified": False,
                "business_decision_status": "not_selected",
            })
            # P1-018：原先只回一个不透明的错误码，调用方不知道 6 项条件差哪一项。
            # 判定不变（gaps 为空等价于原 _is_process_acceptance_spec），只把缺项列出来。
            gaps = process_acceptance_gaps(candidate)
            if gaps:
                return {
                    "ok": False,
                    "error": "PROCESS_ACCEPTANCE_BASIS_INCOMPLETE",
                    "message": "process_acceptance 确认所需的重建依据不完整：" + "；".join(gaps),
                    "details": {
                        "gaps": gaps,
                        "required_reconstruction_record_fields": list(RECONSTRUCTION_RECORD_FIELDS),
                        "required_process_acceptance_basis_fields": list(PROCESS_ACCEPTANCE_BASIS_FIELDS),
                        "next_action": (
                            "在 acquisition_save_spec 的 spec 里补全 reconstruction_records "
                            "与 process_acceptance_basis，再调用 acquisition_confirm_spec "
                            "并传 confirmation_scope=process_acceptance"
                        ),
                    },
                }
        formal_candidate = mark_spec_confirmed(candidate)
        estimate_preview = _is_estimate_preview_spec(candidate)
        process_acceptance = _is_process_acceptance_spec(candidate)
        schema_ok, schema_errors = (
            validate(formal_candidate)
            if estimate_preview or process_acceptance
            else validate_for_formal(formal_candidate)
        )
        evidence_binding = _bind_spec_evidence(
            workspace_id,
            candidate,
        )
        if not schema_ok:
            return {
                "ok": False,
                "error": "SPEC_VALIDATION_FAILED",
                "details": list(schema_errors),
            }
        if not estimate_preview and not process_acceptance and not evidence_binding.get("formal_ok"):
            return {
                "ok": False,
                "error": "EVIDENCE_REVIEW_REQUIRED",
                "evidence_status": evidence_binding.get("status"),
            }
        scope = f"confirm-spec:{idempotency_key}" if idempotency_key else ""
        prior = _active_idempotency_record(state["idempotency"], scope) if scope else None
        if prior:
            if prior.get("body_hash") != body_hash:
                return {
                    "ok": False, "error": "IDEMPOTENCY_CONFLICT",
                    "resource_id": prior.get("spec_id"),
                }
            existing = state["specs"].get(str(prior.get("spec_id") or ""))
            if not existing:
                raise RuntimeError("spec confirmation idempotency record points to missing spec")
            return {**existing, "idempotent_replay": True}
        existing = next((
            row for row in state["specs"].values()
            if row.get("parent_spec_id") == spec_id
            and row.get("confirmation_status") == "confirmed"
            and str((row.get("confirmation") or {}).get("note") or "") == str(note or "")
        ), None)
        if existing:
            if scope:
                state["idempotency"][scope] = _idempotency_record(
                    scope, body_hash, spec_id=existing["spec_id"],
                )
                _save(workspace_id, state)
            return {**existing, "idempotent_replay": True}
        confirmed = formal_candidate
        if estimate_preview:
            confirmed["confirmation_scope"] = "estimate_preview"
        elif process_acceptance:
            confirmed["confirmation_scope"] = "process_acceptance"
            confirmed["project_fact_certified"] = False
            confirmed["business_decision_status"] = "not_selected"
        evidence_binding = _bind_spec_evidence(
            workspace_id,
            confirmed,
        )
        spec_hash = _hash(confirmed)
        snapshot_hash = _hash({
            "spec_hash": spec_hash,
            "evidence_binding_hash": evidence_binding.get("binding_hash"),
            "evidence_binding_version": evidence_binding.get("binding_version"),
        })
        confirmed_id = f"spec_{uuid.uuid4().hex}"
        revision = max(
            (int(row.get("revision") or 0) for row in state["specs"].values()),
            default=0,
        ) + 1
        row = {
            "ok": True,
            "spec_id": confirmed_id,
            "parent_spec_id": spec_id,
            "version": confirmed.get("version"),
            "spec_hash": spec_hash,
            "revision": revision,
            "spec": confirmed,
            "snapshot_hash": snapshot_hash,
            "evidence_binding_version": evidence_binding.get("binding_version"),
            "evidence_binding_hash": evidence_binding.get("binding_hash"),
            "evidence_status": evidence_binding.get("status"),
            "evidence_formal_ok": bool(evidence_binding.get("formal_ok")),
            "evidence_binding": evidence_binding,
            "confirmation_status": "confirmed",
            "confirmation_scope": (
                "estimate_preview" if estimate_preview else (
                    "process_acceptance" if process_acceptance else "formal_input"
                )
            ),
            "confirmation": {
                "note": str(note or ""),
                "confirmed_at": now, "request_id": request_id,
            },
            "request_id": request_id,
            "created_at": now,
        }
        state["specs"][confirmed_id] = row
        if scope:
            state["idempotency"][scope] = _idempotency_record(
                scope, body_hash, spec_id=confirmed_id,
            )
        _save(workspace_id, state)
        return row


def list_specs(
    workspace_id: str,
    *,
    limit: int = 50,
) -> list[dict[str, Any]]:
    """Return saved acquisition specs newest-first for workbench recovery."""

    rows = [
        copy.deepcopy(row)
        for row in _load(workspace_id)["specs"].values()
    ]
    rows.sort(
        key=lambda row: (int(row.get("revision") or 0), str(row.get("created_at") or "")),
        reverse=True,
    )
    return rows[: max(1, min(int(limit or 50), 100))]


def get_spec(
    workspace_id: str,
    spec_id: str,
) -> dict[str, Any]:
    return copy.deepcopy(
        _load(workspace_id)["specs"].get(spec_id) or {}
    )


def _is_selected_scenario(spec: Mapping[str, Any], scenario_id: str) -> bool:
    """Only run the scenario whose assumptions are already selected in the Spec."""

    selected = str(spec.get("selected_scenario_id") or "base").strip()
    requested = str(scenario_id or "base").strip()
    return bool(selected and requested and requested == selected)


def create_run(
    workspace_id: str, spec: dict[str, Any], *, discount_rate: float = 0.08,
    scenario_id: str = "base", idempotency_key: str = "", request_id: str = "",
    scenario_change_ledger: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    if not _is_selected_scenario(spec, scenario_id):
        return {"ok": False, "error": "SCENARIO_NOT_FOUND"}
    estimate_preview = _is_estimate_preview_spec(spec)
    process_acceptance = _is_process_acceptance_spec(spec)
    schema_ok, schema_errors = (
        validate(spec) if estimate_preview or process_acceptance else validate_for_formal(spec)
    )
    evidence_binding = _bind_spec_evidence(
        workspace_id,
        spec,
    )
    if not schema_ok:
        return {"ok": False, "error": "SPEC_VALIDATION_FAILED", "details": list(schema_errors)}
    if not estimate_preview and not process_acceptance and not evidence_binding.get("formal_ok"):
        return {
            "ok": False,
            "error": "EVIDENCE_REVIEW_REQUIRED",
            "evidence_status": evidence_binding.get("status"),
        }
    request_id = request_id or f"req_{uuid.uuid4().hex}"
    saved_spec = save_spec(
        workspace_id, spec, request_id=request_id, trusted_confirmation=True,
    )
    if not saved_spec.get("ok"):
        return saved_spec
    spec = copy.deepcopy(saved_spec.get("spec") or {})
    evidence_binding = copy.deepcopy(saved_spec.get("evidence_binding") or {})
    body = {"spec": spec, "discount_rate": discount_rate, "scenario_id": scenario_id}
    body_hash = _hash(body)
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        scope = f"run:{idempotency_key}" if idempotency_key else ""
        prior = _active_idempotency_record(state["idempotency"], scope) if scope else None
        if prior:
            if prior["body_hash"] != body_hash:
                return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior["run_id"]}
            existing = state["runs"].get(prior["run_id"])
            if not existing:
                raise RuntimeError("run idempotency record points to a missing run")
            return {**existing, "idempotent_replay": True}
        result = run_acquisition_model(spec, discount_rate=discount_rate, scenario_id=scenario_id)
        run_id = f"acqrun_{uuid.uuid4().hex}"
        schema_formal_ok, schema_errors = validate_for_formal(spec)
        formal_ok, formal_errors, evidence_binding = _formal_assessment(
            workspace_id,
            spec,
            evidence_binding=evidence_binding,
        )
        created_at = _now()
        issues = [
            {"code": "SPEC_VALIDATION_FAILED", "blocking": True, "status": "open", "detail": item, "created_at": created_at}
            for item in schema_errors
        ]
        issues.extend(_evidence_blocking_issues(evidence_binding, created_at=created_at))
        scenario_ledger = list(scenario_change_ledger or [])
        invalid_scenario_sources = [
            index for index, item in enumerate(scenario_ledger)
            if not isinstance(item, Mapping)
            or not str(item.get("source") or "").strip()
        ]
        if invalid_scenario_sources:
            issues.append({
                "code": "SCENARIO_SOURCE_REQUIRED", "blocking": True, "status": "open",
                "detail": "情景调整缺少可追溯来源", "rows": invalid_scenario_sources,
                "created_at": created_at,
            })
        row = {
            "ok": True, "available": True, "run_id": run_id, "workspace_id": workspace_id,
            "status": "succeeded",
            "lifecycle_status": "validated" if not issues else "validation_failed",
            "delivery_mode": (
                "estimate_preview" if estimate_preview else (
                    "process_acceptance" if process_acceptance else "formal_candidate"
                )
            ),
            "model_version": result["model_version"],
            "spec_version": LATEST_SPEC_VERSION, "spec_hash": _hash(spec), "input_hash": body_hash,
            "spec_id": saved_spec.get("spec_id"),
            "spec_snapshot_hash": saved_spec.get("snapshot_hash"),
            "evidence_binding_version": evidence_binding.get("binding_version"),
            "evidence_binding_hash": evidence_binding.get("binding_hash"),
            "evidence_status": evidence_binding.get("status"),
            "evidence_formal_ok": bool(evidence_binding.get("formal_ok")),
            "evidence_binding": evidence_binding,
            "scenario_id": scenario_id, "scenario_change_ledger": scenario_ledger,
            "discount_rate": discount_rate, "result": result, "consistency_ok": True,
            "validation_status": "passed" if not issues else "failed",
            "formal_spec_valid": bool(schema_formal_ok and formal_ok),
            "process_acceptance_valid": bool(process_acceptance and schema_ok),
            "formal_spec_errors": formal_errors, "request_id": request_id,
            "created_at": created_at,
            "evidence_policy": str(spec.get("evidence_policy") or "formal_evidence"),
            "project_fact_certified": bool(spec.get(
                "project_fact_certified",
                str(spec.get("evidence_policy") or "") != "source_reconstructed",
            )),
            "reconstruction_records": copy.deepcopy(spec.get("reconstruction_records") or []),
            "reconstructed_source_ids": copy.deepcopy(spec.get("reconstructed_source_ids") or []),
            "unresolved_inputs": copy.deepcopy(spec.get("unresolved_inputs") or []),
            "release_limitations": copy.deepcopy(spec.get("release_limitations") or []),
            "business_decision_status": str(spec.get("business_decision_status") or "not_selected"),
            **_migration_binding(spec),
            "issues": issues,
            "state_history": [
                _history_event("validated_spec", request_id=request_id),
                _history_event("running", request_id=request_id),
                _history_event("calculated", request_id=request_id),
                _history_event("internally_consistent", request_id=request_id),
                _history_event(
                    "validated" if not issues else "validation_failed",
                    request_id=request_id,
                ),
            ],
        }
        state["runs"][run_id] = row
        if scope:
            state["idempotency"][scope] = _idempotency_record(
                scope, body_hash, run_id=run_id,
            )
        _save(workspace_id, state)
    return row


def enqueue_run(
    workspace_id: str, spec: dict[str, Any], *, discount_rate: float = 0.08,
    scenario_id: str = "base", idempotency_key: str = "", request_id: str = "",
    scenario_change_ledger: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """Persist a pollable queued run before any model calculation starts."""

    if not _is_selected_scenario(spec, scenario_id):
        return {"ok": False, "error": "SCENARIO_NOT_FOUND"}
    schema_ok, schema_errors = validate_for_formal(spec)
    evidence_binding = _bind_spec_evidence(
        workspace_id,
        spec,
    )
    if not schema_ok:
        return {"ok": False, "error": "SPEC_VALIDATION_FAILED", "details": list(schema_errors)}
    if not evidence_binding.get("formal_ok"):
        return {
            "ok": False,
            "error": "EVIDENCE_REVIEW_REQUIRED",
            "evidence_status": evidence_binding.get("status"),
        }
    request_id = request_id or f"req_{uuid.uuid4().hex}"
    saved_spec = save_spec(
        workspace_id, spec, request_id=request_id, trusted_confirmation=True,
    )
    if not saved_spec.get("ok"):
        return saved_spec
    spec = copy.deepcopy(saved_spec.get("spec") or {})
    evidence_binding = copy.deepcopy(saved_spec.get("evidence_binding") or {})
    body = {"spec": spec, "discount_rate": discount_rate, "scenario_id": scenario_id}
    body_hash = _hash(body)
    schema_formal_ok, _schema_errors = validate_for_formal(spec)
    formal_ok, formal_errors, evidence_binding = _formal_assessment(
        workspace_id,
        spec,
        evidence_binding=evidence_binding,
    )
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        scope = f"run:{idempotency_key}" if idempotency_key else ""
        prior = _active_idempotency_record(state["idempotency"], scope) if scope else None
        if prior:
            if prior.get("body_hash") != body_hash:
                return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior.get("run_id", "")}
            existing = state["runs"].get(prior.get("run_id"))
            if not existing:
                raise RuntimeError("run idempotency record points to a missing run")
            return {**existing, "idempotent_replay": True}
        run_id = f"acqrun_{uuid.uuid4().hex}"
        created_at = _now()
        row = {
            "ok": True, "available": False, "run_id": run_id, "workspace_id": workspace_id,
            "status": "queued", "progress": 0, "lifecycle_status": "validated_spec",
            "spec_version": LATEST_SPEC_VERSION, "spec_hash": _hash(spec), "input_hash": body_hash,
            "spec_id": saved_spec.get("spec_id"), "scenario_id": scenario_id,
            "spec_snapshot_hash": saved_spec.get("snapshot_hash"),
            "evidence_binding_version": evidence_binding.get("binding_version"),
            "evidence_binding_hash": evidence_binding.get("binding_hash"),
            "evidence_status": evidence_binding.get("status"),
            "evidence_formal_ok": bool(evidence_binding.get("formal_ok")),
            "evidence_binding": evidence_binding,
            "scenario_change_ledger": list(scenario_change_ledger or []), "discount_rate": discount_rate,
            "validation_status": "pending",
            "formal_spec_valid": bool(schema_formal_ok and formal_ok),
            "formal_spec_errors": formal_errors,
            "request_id": request_id, "created_at": created_at, "updated_at": created_at,
            "issues": [],
            "evidence_policy": str(spec.get("evidence_policy") or "formal_evidence"),
            "project_fact_certified": bool(spec.get("project_fact_certified", str(spec.get("evidence_policy") or "") != "source_reconstructed")),
            "reconstruction_records": copy.deepcopy(spec.get("reconstruction_records") or []),
            "reconstructed_source_ids": copy.deepcopy(spec.get("reconstructed_source_ids") or []),
            "unresolved_inputs": copy.deepcopy(spec.get("unresolved_inputs") or []),
            "release_limitations": copy.deepcopy(spec.get("release_limitations") or []),
            "business_decision_status": str(spec.get("business_decision_status") or "not_selected"),
            **_migration_binding(spec),
            "state_history": [
                _history_event("validated_spec", request_id=request_id),
                _history_event("queued", request_id=request_id),
            ],
        }
        state["runs"][run_id] = row
        if scope:
            state["idempotency"][scope] = _idempotency_record(
                scope, body_hash, run_id=run_id,
            )
        _save(workspace_id, state)
        return row


def execute_queued_run(
    workspace_id: str,
    run_id: str,
) -> None:
    """Execute one durable queued run and converge it to succeeded/failed."""

    with _state_guard(workspace_id):
        state = _load(workspace_id)
        run = state["runs"].get(run_id)
        if not run or run.get("status") in {"succeeded", "failed", "cancelled"}:
            return
        spec_row = state["specs"].get(str(run.get("spec_id") or "")) or {}
        spec = spec_row.get("spec")
        if not isinstance(spec, dict):
            run.update({
                "status": "failed", "progress": 100, "updated_at": _now(),
                "error": {"code": "SPEC_VALIDATION_FAILED", "message": "spec snapshot missing", "retryable": False},
            })
            _save(workspace_id, state)
            return
        run.update({"status": "running", "progress": 10, "lifecycle_status": "running", "updated_at": _now()})
        run.setdefault("state_history", []).append(_history_event(
            "running", request_id=str(run.get("request_id") or ""),
        ))
        discount_rate = float(run.get("discount_rate") or 0.08)
        scenario_id = str(run.get("scenario_id") or "base")
        _save(workspace_id, state)
    try:
        current_evidence = _bind_spec_evidence(
            workspace_id,
            spec,
        )
    except Exception as exc:  # noqa: BLE001
        _LOG.warning(
            "asset acquisition evidence binding failed; error_type=%s",
            type(exc).__name__,
        )
        current_evidence = {
            "formal_ok": False,
            "status": "invalid",
            "binding_version": "",
            "binding_hash": "",
            "bindings": [],
            "missing": [],
            "pending": [],
            "invalid": [{
                "source_path": "$",
                "code": "SOURCE_EVIDENCE_STATE_INVALID",
                "message": _SOURCE_EVIDENCE_FAILURE_MESSAGE,
            }],
        }
    evidence_hash_matches = bool(
        current_evidence.get("binding_hash") == run.get("evidence_binding_hash")
        and current_evidence.get("binding_version") == run.get("evidence_binding_version")
    )
    evidence_snapshot_matches = bool(evidence_hash_matches and current_evidence.get("formal_ok"))
    try:
        result = run_acquisition_model(spec, discount_rate=discount_rate, scenario_id=scenario_id)
    except Exception as exc:  # noqa: BLE001
        code = "SPEC_VALIDATION_FAILED" if isinstance(exc, AcquisitionModelError) else "RUN_FAILED"
        safe_message = (
            _RUN_VALIDATION_FAILURE_MESSAGE
            if code == "SPEC_VALIDATION_FAILED"
            else _RUN_EXECUTION_FAILURE_MESSAGE
        )
        _LOG.warning(
            "asset acquisition run failed; run_id=%s error_type=%s code=%s",
            run_id,
            type(exc).__name__,
            code,
        )
        with _state_guard(workspace_id):
            state = _load(workspace_id)
            run = state["runs"].get(run_id)
            if run and run.get("status") != "cancelled":
                run.update({
                    "status": "failed", "progress": 100, "lifecycle_status": "failed", "updated_at": _now(),
                    "error": {"code": code, "message": safe_message, "retryable": False},
                })
                run.setdefault("state_history", []).append(_history_event(
                    "failed", request_id=str(run.get("request_id") or ""), error_code=code,
                ))
                _save(workspace_id, state)
        return

    with _state_guard(workspace_id):
        state = _load(workspace_id)
        run = state["runs"].get(run_id)
        if not run or run.get("status") == "cancelled":
            return
        created_at = str(run.get("created_at") or _now())
        schema_formal_ok, schema_errors = validate_for_formal(spec)
        formal_errors = [*schema_errors, *_evidence_error_strings(current_evidence)]
        if not evidence_hash_matches:
            formal_errors.append(
                "finance_spec.v3 证据绑定快照已变化或不再满足正式复核条件"
            )
        issues = [
            {"code": "SPEC_VALIDATION_FAILED", "blocking": True, "status": "open", "detail": item, "created_at": created_at}
            for item in schema_errors
        ]
        issues.extend(_evidence_blocking_issues(current_evidence, created_at=created_at))
        if not evidence_hash_matches:
            issues.append({
                "code": "EVIDENCE_BINDING_STALE",
                "blocking": True,
                "status": "open",
                "detail": (
                    "运行排队后证据绑定状态发生变化；必须保存新Spec修订并重新运行。"
                    f" snapshot={run.get('evidence_binding_hash')} current={current_evidence.get('binding_hash')}"
                ),
                "created_at": created_at,
            })
        scenario_ledger = list(run.get("scenario_change_ledger") or [])
        invalid_scenario_sources = [
            index for index, item in enumerate(scenario_ledger)
            if not isinstance(item, Mapping)
            or not str(item.get("source") or "").strip()
        ]
        if invalid_scenario_sources:
            issues.append({
                "code": "SCENARIO_SOURCE_REQUIRED", "blocking": True, "status": "open",
                "detail": "情景调整缺少可追溯来源", "rows": invalid_scenario_sources,
                "created_at": created_at,
            })
        run.update({
            "available": True, "status": "succeeded", "progress": 100,
            "lifecycle_status": "validated" if not issues else "validation_failed",
            "model_version": result["model_version"],
            "result": result, "consistency_ok": True, "issues": issues, "updated_at": _now(),
            "validation_status": "passed" if not issues else "failed",
            "formal_spec_valid": bool(schema_formal_ok and evidence_snapshot_matches),
            "formal_spec_errors": formal_errors,
            "evidence_revalidation": {
                "checked_at": _now(),
                "matches_snapshot": evidence_hash_matches,
                "binding_version": current_evidence.get("binding_version"),
                "binding_hash": current_evidence.get("binding_hash"),
                "status": current_evidence.get("status"),
                "formal_ok": bool(current_evidence.get("formal_ok")),
            },
        })
        run.setdefault("state_history", []).extend([
            _history_event("calculated", request_id=str(run.get("request_id") or "")),
            _history_event("internally_consistent", request_id=str(run.get("request_id") or "")),
            _history_event(
                "validated" if not issues else "validation_failed",
                request_id=str(run.get("request_id") or ""),
            ),
        ])
        _save(workspace_id, state)


def get_run(
    workspace_id: str,
    run_id: str,
) -> dict[str, Any]:
    return dict(
        _load(workspace_id)["runs"].get(run_id) or {}
    )


def list_runs(
    workspace_id: str,
    *,
    limit: int = 50,
) -> list[dict[str, Any]]:
    """Return lightweight acquisition run summaries newest-first."""

    rows: list[dict[str, Any]] = []
    summary_keys = (
        "run_id", "workspace_id", "status", "lifecycle_status",
        "model_version", "spec_version", "spec_hash", "input_hash",
        "spec_id", "scenario_id", "discount_rate", "consistency_ok",
        "spec_snapshot_hash", "evidence_binding_version", "evidence_binding_hash",
        "evidence_status", "evidence_formal_ok", "evidence_revalidation",
        "validation_status", "formal_spec_valid", "formal_spec_errors",
        "request_id", "created_at", "issues",
    )
    for run in _load(workspace_id)["runs"].values():
        indicators = (run.get("result") or {}).get("indicators") or {}
        row = {key: copy.deepcopy(run.get(key)) for key in summary_keys}
        row.update({
            "invest_type": "asset_acquisition",
            "industry": "资产收购",
            "available": bool(run.get("available")),
            "started_at": run.get("created_at"),
            "finished_at": run.get("finished_at") or run.get("created_at"),
            "indicators": {
                key: copy.deepcopy(indicators.get(key))
                for key in ("project_irr_pct", "equity_irr_pct", "npv_wan", "minimum_dscr")
            },
        })
        rows.append(row)
    rows.sort(
        key=lambda row: (str(row.get("created_at") or ""), str(row.get("run_id") or "")),
        reverse=True,
    )
    return rows[: max(1, min(int(limit or 50), 100))]


MAX_SCENARIO_MATRIX_COMBINATIONS = 64


def create_scenario_matrix(
    workspace_id: str, run_id: str, dimensions: dict[str, Any], *,
    idempotency_key: str = "", request_id: str = "",
) -> dict[str, Any]:
    """Evaluate a bounded Cartesian product without mutating the bound run/spec."""

    request_id = request_id or f"req_{uuid.uuid4().hex}"
    if not isinstance(dimensions, dict) or not dimensions:
        return {"ok": False, "error": "SCENARIO_MATRIX_INVALID", "details": {"reason": "dimensions_required"}}
    unknown = sorted(set(dimensions) - INDEPENDENT_SCENARIO_FIELDS)
    if unknown:
        return {"ok": False, "error": "SCENARIO_FIELD_UNSUPPORTED", "details": {"fields": unknown}}
    normalized: dict[str, list[Any]] = {}
    for field in sorted(dimensions):
        values = dimensions[field]
        if not isinstance(values, (list, tuple)) or not values:
            return {
                "ok": False, "error": "SCENARIO_MATRIX_INVALID",
                "details": {"field": field, "reason": "non_empty_array_required"},
            }
        rows = list(values)
        fingerprints = [_hash(value) for value in rows]
        if len(fingerprints) != len(set(fingerprints)):
            return {
                "ok": False, "error": "SCENARIO_MATRIX_INVALID",
                "details": {"field": field, "reason": "duplicate_values"},
            }
        normalized[field] = rows
    combination_count = math.prod(len(values) for values in normalized.values())
    if combination_count > MAX_SCENARIO_MATRIX_COMBINATIONS:
        return {
            "ok": False, "error": "SCENARIO_MATRIX_TOO_LARGE",
            "details": {"combination_count": combination_count, "max_combinations": MAX_SCENARIO_MATRIX_COMBINATIONS},
        }
    body_hash = _hash({"run_id": run_id, "dimensions": normalized})
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        run = state["runs"].get(run_id)
        if not run:
            return {"ok": False, "error": "RUN_NOT_FOUND"}
        if run.get("status") != "succeeded" or not run.get("available"):
            return {"ok": False, "error": "RUN_NOT_READY", "details": {"status": run.get("status")}}
        spec_row = state["specs"].get(str(run.get("spec_id") or "")) or {}
        spec = spec_row.get("spec")
        if not isinstance(spec, dict) or _hash(spec) != run.get("spec_hash"):
            return {"ok": False, "error": "RUN_SPEC_MISMATCH"}
        scope = f"scenario_matrix:{run_id}:{idempotency_key}" if idempotency_key else ""
        prior = _active_idempotency_record(state["idempotency"], scope) if scope else None
        if prior:
            if prior.get("body_hash") != body_hash:
                return {
                    "ok": False, "error": "IDEMPOTENCY_CONFLICT",
                    "resource_id": prior.get("matrix_id", ""),
                }
            existing = state["scenario_matrices"].get(prior.get("matrix_id"))
            if not existing:
                raise RuntimeError("scenario matrix idempotency record points to a missing matrix")
            return {**existing, "idempotent_replay": True}
        bound = {
            "run_id": run_id, "spec_id": run.get("spec_id"), "spec_hash": run.get("spec_hash"),
            "input_hash": run.get("input_hash"), "model_version": run.get("model_version"),
            "scenario_id": run.get("scenario_id"),
        }

    fields = list(normalized)
    rows: list[dict[str, Any]] = []
    matrix_id = f"scenario_matrix_{uuid.uuid4().hex}"
    asset_type = str(spec.get("asset_type") or "hotel_lease")
    base_market_rent = copy.deepcopy((spec.get("lease_portfolio") or {}).get("market_rent"))
    for index, values in enumerate(product(*(normalized[field] for field in fields)), 1):
        scenario_id = f"{matrix_id}:row_{index:03d}"
        changes = {field: copy.deepcopy(value) for field, value in zip(fields, values)}
        try:
            scenario_spec, ledger = apply_scenario(spec, changes)
        except AcquisitionModelError as exc:
            return {
                "ok": False, "error": "SCENARIO_MATRIX_INVALID",
                "details": {"scenario_index": index, "reason": str(exc)},
            }
        expected_market_rent = copy.deepcopy(changes.get("lease_portfolio.market_rent", base_market_rent))
        actual_market_rent = copy.deepcopy((scenario_spec.get("lease_portfolio") or {}).get("market_rent"))
        if asset_type == "hotel_lease" and actual_market_rent != expected_market_rent:
            return {
                "ok": False, "error": "SCENARIO_INDEPENDENCE_VIOLATION",
                "details": {"field": "lease_portfolio.market_rent", "expected": expected_market_rent, "actual": actual_market_rent},
            }
        try:
            result = run_acquisition_model(
                scenario_spec, discount_rate=float(run.get("discount_rate") or 0.08),
                scenario_id=scenario_id,
            )
        except AcquisitionModelError as exc:
            return {
                "ok": False, "error": "SCENARIO_MATRIX_INVALID",
                "details": {"scenario_index": index, "changes": changes, "reason": str(exc)},
            }
        solar = dict(scenario_spec.get("solar_operation") or {})
        row = {
            "scenario_id": scenario_id, "changes": changes,
            "scenario_change_ledger": ledger, "scenario_spec_hash": _hash(scenario_spec),
            "result_hash": _hash(result),
            "purchase_price_wan": result.get("purchase_price_wan"),
            "financing_ratio": (scenario_spec.get("transaction") or {}).get("financing_ratio"),
            "indicators": copy.deepcopy(result.get("indicators") or {}),
        }
        if asset_type == "solar_power":
            row.update({
                "tariff_yuan_per_kwh": solar.get("tariff_yuan_per_kwh"),
                "annual_generation_mwh": solar.get("annual_generation_mwh"),
                "utilization_hours": solar.get("utilization_hours"),
                "annual_opex_wan": solar.get("annual_opex_wan"),
            })
        else:
            row.update({
                "market_rent": actual_market_rent,
                "occupancy": copy.deepcopy(
                    (scenario_spec.get("hotel_operation") or {}).get("occupancy")
                ),
            })
        rows.append(row)
    created_at = _now()
    matrix = {
        "ok": True, "matrix_id": matrix_id, "status": "succeeded", "read_only": True,
        "workspace_id": workspace_id, **bound, "dimensions": normalized,
        "combination_count": combination_count, "max_combinations": MAX_SCENARIO_MATRIX_COMBINATIONS,
        "rows": rows, "request_id": request_id, "created_at": created_at,
    }
    matrix["matrix_hash"] = _hash({
        "bound": bound, "dimensions": normalized,
        "rows": [{"changes": row["changes"], "scenario_spec_hash": row["scenario_spec_hash"], "result_hash": row["result_hash"]} for row in rows],
    })
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        current = state["runs"].get(run_id)
        if not current or current.get("spec_hash") != bound["spec_hash"] or current.get("input_hash") != bound["input_hash"]:
            return {"ok": False, "error": "RUN_SPEC_MISMATCH"}
        scope = f"scenario_matrix:{run_id}:{idempotency_key}" if idempotency_key else ""
        if scope and (prior := _active_idempotency_record(state["idempotency"], scope)):
            if prior.get("body_hash") != body_hash:
                return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior.get("matrix_id", "")}
            return {**state["scenario_matrices"][prior["matrix_id"]], "idempotent_replay": True}
        state["scenario_matrices"][matrix_id] = matrix
        if scope:
            state["idempotency"][scope] = _idempotency_record(
                scope, body_hash, matrix_id=matrix_id,
            )
        _save(workspace_id, state)
    return matrix


def get_scenario_matrix(
    workspace_id: str,
    run_id: str,
    matrix_id: str,
) -> dict[str, Any]:
    matrix = dict(
        _load(workspace_id)["scenario_matrices"].get(
            matrix_id
        )
        or {}
    )
    if not matrix or matrix.get("run_id") != run_id:
        return {}
    return matrix


def list_scenario_matrices(
    workspace_id: str,
    run_id: str,
) -> list[dict[str, Any]]:
    rows = [
        value
        for value in _load(workspace_id)[
            "scenario_matrices"
        ].values()
        if value.get("run_id") == run_id
    ]
    rows.sort(key=lambda value: str(value.get("created_at") or ""), reverse=True)
    return [
        {key: value.get(key) for key in (
            "matrix_id", "matrix_hash", "run_id", "spec_id", "spec_hash", "input_hash",
            "combination_count", "status", "created_at",
        )}
        for value in rows
    ]


def _diff_is_blocking(row: dict[str, Any], tolerance: float) -> bool:
    if bool(row.get("blocking")) or row.get("within_tolerance") is False:
        return True
    relative = row.get("relative_difference")
    if relative is None:
        try:
            reference = float(row["reference_value"])
            calculated = float(row["calculated_value"])
            relative = abs(calculated - reference) / max(abs(reference), 1e-12)
        except (KeyError, TypeError, ValueError):
            return False
    try:
        exceeds = abs(float(relative)) > max(float(tolerance), 0.0)
    except (TypeError, ValueError):
        return True
    if not exceeds:
        return False
    validation = str(row.get("validation_status") or "").lower()
    has_reason = bool(str(row.get("validation_reason") or row.get("reason") or "").strip())
    has_evidence = bool(row.get("evidence_ids") or row.get("source_locators"))
    return validation not in {"passed", "resolved", "within_tolerance"} or not has_reason or not has_evidence










def max_price(
    workspace_id: str, run_id: str, *, target_irr: float | None = None,
    min_dscr: float | None = None,
    lower: float = 0.0, upper: float | None = None,
    request_id: str = "",
) -> dict[str, Any]:
    run = get_run(workspace_id, run_id)
    if not run:
        return {"ok": False, "error": "run_not_found"}
    spec = (
        _load(workspace_id)["specs"].get(
            str(run.get("spec_id") or "")
        )
        or {}
    ).get("spec")
    if not isinstance(spec, dict):
        return {"ok": False, "error": "spec_snapshot_missing"}
    thresholds, threshold_error = _decision_thresholds(spec)
    if threshold_error:
        return {
            "ok": False, "error": threshold_error,
            "decision_thresholds": spec.get("decision_thresholds"),
        }
    expected_target = float(thresholds["target_project_irr"] or 0.0)
    expected_dscr = thresholds["minimum_dscr"]
    try:
        effective_target = expected_target if target_irr is None else float(target_irr)
        effective_dscr = expected_dscr if min_dscr is None else float(min_dscr)
        effective_lower = float(lower)
        effective_upper = float(upper) if upper is not None else None
    except (TypeError, ValueError):
        return {
            "ok": False, "error": "MAX_PRICE_PARAMETERS_INVALID",
            "parameters": {
                "target_irr": target_irr, "min_dscr": min_dscr,
                "lower": lower, "upper": upper,
            },
        }
    parameters = {
        "target_irr": effective_target, "min_dscr": effective_dscr,
        "lower": effective_lower, "upper": effective_upper,
    }
    if (
        not _same_optional_number(effective_target, expected_target)
        or not _same_optional_number(effective_dscr, expected_dscr)
    ):
        return {
            "ok": False, "error": "MAX_PRICE_THRESHOLD_MISMATCH",
            "parameters": parameters, "decision_thresholds": thresholds,
        }
    if (
        not 0 < effective_target <= 5
        or (effective_dscr is not None and effective_dscr <= 0)
        or effective_lower < 0
        or (effective_upper is not None and effective_upper <= effective_lower)
    ):
        return {"ok": False, "error": "MAX_PRICE_PARAMETERS_INVALID", "parameters": parameters}
    existing = run.get("max_acquisition_price_analysis") or {}
    if existing.get("parameters") == parameters and isinstance(existing.get("result"), dict):
        expected_hash = _hash({
            key: value for key, value in existing.items() if key != "analysis_hash"
        })
        if existing.get("analysis_hash") != expected_hash:
            return {"ok": False, "error": "MAX_PRICE_HASH_MISMATCH"}
        return {
            "ok": True,
            "run_id": run_id,
            "analysis_hash": existing.get("analysis_hash"),
            "validation_status": existing.get("validation_status") or "passed",
            "idempotent_replay": True,
            **copy.deepcopy(existing.get("result") or {}),
        }
    baseline_analysis_hash = str(existing.get("analysis_hash") or "")
    solved = solve_max_acquisition_price(
        spec, target_irr=effective_target, min_dscr=effective_dscr,
        lower=effective_lower, upper=effective_upper,
    )
    analysis = {
        "status": "calculated",
        "validation_status": (
            "passed" if solved.get("feasible") and solved.get("converged") else "failed"
        ),
        "parameters": parameters,
        "result": copy.deepcopy(solved),
        "calculated_at": _now(),
        "engine_version": run.get("model_version"),
        "request_id": request_id,
    }
    analysis["analysis_hash"] = _hash(analysis)
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        current = state["runs"].get(run_id)
        if not current:
            return {"ok": False, "error": "run_not_found"}
        if current.get("spec_hash") != run.get("spec_hash") or current.get("input_hash") != run.get("input_hash"):
            return {"ok": False, "error": "RUN_SPEC_MISMATCH"}
        current_analysis = current.get("max_acquisition_price_analysis") or {}
        if (
            current_analysis.get("parameters") == parameters
            and isinstance(current_analysis.get("result"), dict)
        ):
            return {
                "ok": True,
                "run_id": run_id,
                "analysis_hash": current_analysis.get("analysis_hash"),
                "validation_status": current_analysis.get("validation_status") or "passed",
                "idempotent_replay": True,
                **copy.deepcopy(current_analysis.get("result") or {}),
            }
        if str(current_analysis.get("analysis_hash") or "") != baseline_analysis_hash:
            return {
                "ok": False,
                "error": "WORKSPACE_VERSION_CONFLICT",
                "resource_id": run_id,
            }
        current["max_acquisition_price_analysis"] = analysis
        if solved.get("feasible") and solved.get("converged"):
            _close_issue(current, "MAX_PRICE_NOT_FEASIBLE", reason="最高可接受价求解可行且收敛")
        else:
            _open_issue(current, "MAX_PRICE_NOT_FEASIBLE", "最高可接受价求解不可行或未收敛")
        current.setdefault("max_acquisition_price_history", []).append(copy.deepcopy(analysis))
        current.setdefault("state_history", []).append(_history_event(
            "max_acquisition_price_calculated",
            request_id=request_id,
            analysis_hash=analysis["analysis_hash"],
            validation_status=analysis["validation_status"],
        ))
        _save(workspace_id, state)
    return {
        "ok": True,
        "run_id": run_id,
        "analysis_hash": analysis["analysis_hash"],
        "validation_status": analysis["validation_status"],
        **solved,
    }


def build_acquisition_report_data(
    workspace_id: str,
    run: dict[str, Any],
) -> dict[str, Any]:
    """Project a bound, validated run/spec into the acquisition report contract."""

    state = _load(workspace_id)
    spec_row = state["specs"].get(str(run.get("spec_id") or "")) or {}
    spec = copy.deepcopy(spec_row.get("spec") or {})
    if not isinstance(spec, dict) or _hash(spec) != run.get("spec_hash"):
        raise RuntimeError("run spec snapshot is missing or does not match spec_hash")
    evidence_ok, current_evidence = _current_evidence_matches_run(
        workspace_id,
        run,
        spec,
    )
    if not evidence_ok:
        raise RuntimeError(
            "run evidence binding is stale or no longer valid: "
            f"snapshot={run.get('evidence_binding_hash')} current={current_evidence.get('binding_hash')}"
        )
    transaction = copy.deepcopy(spec.get("transaction") or {})
    source_ledger = [
        {
            "field": binding.get("source_path"),
            "evidence_ids": [binding.get("evidence_id")],
            "source": {
                "file_id": binding.get("file_id"),
                "locator": binding.get("locator"),
                "source_sha256": binding.get("source_sha256"),
                "source_size_bytes": binding.get("source_size_bytes"),
                "parse_job": binding.get("parse_job"),
                "attempt": binding.get("attempt"),
                "evidence_content_hash": binding.get("evidence_content_hash"),
                "binding_hash": binding.get("binding_hash"),
            },
            "validation_status": "bound",
        }
        for binding in current_evidence.get("bindings") or []
    ]
    parties = copy.deepcopy(spec.get("project_parties") or [])
    licenses = copy.deepcopy(transaction.get("licenses") or [])
    for party in parties:
        if isinstance(party, dict) and "license_holder" in (party.get("roles") or []):
            licenses.append({
                "license_type": "license_holder_role",
                "holder_id": party.get("entity_id"),
                "holder_name": party.get("name"),
                "status": party.get("status"),
                "evidence_ids": copy.deepcopy(party.get("evidence_ids") or []),
            })
    matrices = [
        copy.deepcopy(value)
        for value in state["scenario_matrices"].values()
        if value.get("run_id") == run.get("run_id")
        and value.get("spec_hash") == run.get("spec_hash")
        and value.get("input_hash") == run.get("input_hash")
    ]
    matrices.sort(key=lambda value: str(value.get("created_at") or ""))
    result = run.get("result") or {}
    asset_type = str(result.get("asset_type") or spec.get("asset_type") or "hotel_lease")
    valuation = float(transaction.get("valuation_value") or 0.0)
    purchase = float(result.get("purchase_price_wan") or transaction.get("purchase_price") or 0.0)
    report = {
        "schema_version": "asset_acquisition_report_data.v1",
        "asset_type": asset_type,
        "bindings": {
            "workspace_id": workspace_id,
            "run_id": run.get("run_id"),
            "spec_id": run.get("spec_id"),
            "spec_hash": run.get("spec_hash"),
            "input_hash": run.get("input_hash"),
            "model_version": run.get("model_version"),
            "spec_snapshot_hash": run.get("spec_snapshot_hash"),
            "evidence_binding_version": run.get("evidence_binding_version"),
            "evidence_binding_hash": run.get("evidence_binding_hash"),
            "validation_status": run.get("validation_status"),
            "consistency_ok": bool(run.get("consistency_ok")),
        },
        "source_processing_ledger": source_ledger,
        "party_relationships": parties,
        "asset_boundary": copy.deepcopy(transaction.get("asset_scope") or []),
        "license_ledger": licenses,
        "historical_financial_comparison": copy.deepcopy(spec.get("historical_statements") or []),
        "valuation_transaction_bridge": {
            "valuation_value_wan": valuation,
            "purchase_price_wan": purchase,
            "purchase_price_vs_valuation_wan": purchase - valuation,
            "transaction_tax_wan": result.get("transaction_tax_wan"),
            "total_acquisition_cost_wan": result.get("total_acquisition_cost_wan"),
            "valuation_date": transaction.get("valuation_date"),
            "closing_date": transaction.get("closing_date"),
        },
        "maximum_acceptable_price": copy.deepcopy(
            run.get("max_acquisition_price_analysis")
            or {"status": "not_calculated", "validation_status": "not_run"}
        ),
        "scenario_matrices": matrices,
        "red_flags": copy.deepcopy(transaction.get("red_flags") or []),
        "closing_conditions": copy.deepcopy(transaction.get("closing_conditions") or []),
        "veto_items": copy.deepcopy(transaction.get("veto_items") or []),
        "validation_summary": {
            "validation_status": run.get("validation_status"),
            "consistency_ok": bool(run.get("consistency_ok")),
            "formal_spec_valid": bool(run.get("formal_spec_valid")),
            "evidence_status": current_evidence.get("status"),
            "evidence_formal_ok": bool(current_evidence.get("formal_ok")),
            "evidence_binding_hash": current_evidence.get("binding_hash"),
            "open_blocking_issues": [
                copy.deepcopy(issue) for issue in (run.get("issues") or [])
                if issue.get("blocking") and issue.get("status") == "open"
            ],
        },
    }
    if asset_type == "solar_power":
        report["solar_operation"] = copy.deepcopy(result.get("solar_operation") or {})
        report["solar_operating_ledger"] = copy.deepcopy(result.get("annual_summary") or [])
    else:
        report["lease_ledger"] = copy.deepcopy(
            (spec.get("lease_portfolio") or {}).get("units") or []
        )
    report["report_data_hash"] = _hash(report)
    return report


def render_markdown(run: dict[str, Any], report_data: dict[str, Any] | None = None) -> str:
    result = run.get("result") or {}
    ind = result.get("indicators") or {}
    report_data = report_data or {}
    asset_type = str(result.get("asset_type") or report_data.get("asset_type") or "hotel_lease")
    is_solar = asset_type == "solar_power"

    def text(value: Any) -> str:
        if isinstance(value, (dict, list)):
            value = json.dumps(value, ensure_ascii=False, sort_keys=True, default=str)
        return str(value if value not in (None, "") else "—").replace("|", "¦").replace("\n", " ")

    lines = [
        "# 资产收购可行性研究报告",
        "",
        f"> 财务运行：`{run.get('run_id')}`；场景：`{run.get('scenario_id')}`；模型：`{run.get('model_version')}`",
        f"> Spec：`{run.get('spec_hash')}`；事实版本：`{run.get('spec_id')}`",
        f"> 证据绑定：`{run.get('evidence_binding_hash')}`；版本：`{run.get('evidence_binding_version')}`",
        "",
        "## 一、交易概览",
        "",
        f"- 收购价格：{result.get('purchase_price_wan', 0):,.2f} 万元",
        f"- 交易税费：{result.get('transaction_tax_wan', 0):,.2f} 万元",
        f"- 总收购成本：{result.get('total_acquisition_cost_wan', 0):,.2f} 万元",
        "",
        "## 二、核心财务指标",
        "",
        "| 指标 | 结果 |",
        "|---|---:|",
        f"| 项目 IRR | {_pct(ind.get('project_irr_pct'))} |",
        f"| 资本金 IRR | {_pct(ind.get('equity_irr_pct'))} |",
        f"| NPV | {_num(ind.get('npv_wan'))} 万元 |",
        f"| 最低 DSCR | {_num(ind.get('minimum_dscr'))} |",
        f"| 最低 ICR | {_num(ind.get('minimum_icr'))} |",
    ]
    if is_solar:
        solar = result.get("solar_operation") or {}
        lines.extend([
            "",
            "## 三、光伏电站运营",
            "",
            f"- 装机容量：{_num(solar.get('installed_capacity_mw'))} MW",
            f"- 基准发电量：{_num(solar.get('base_generation_mwh'))} MWh",
            f"- 上网电价：{_num(solar.get('tariff_yuan_per_kwh'))} 元/kWh",
            f"- 限电率：{_pct_ratio(solar.get('curtailment_rate'))}",
            f"- 年衰减率：{_pct_ratio(solar.get('degradation_rate'))}",
            "",
            "| 年度 | 理论发电量(MWh) | 上网电量(MWh) | 售电收入(万元) | 运维费(万元) | 维护性资本开支(万元) | 所得税(万元) | 债务服务(万元) | 项目现金流(万元) | 资本金现金流(万元) |",
            "|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|",
        ])
        for row in result.get("annual_summary") or []:
            lines.append(
                f"| {text(row.get('year'))} | {_num(row.get('gross_generation_mwh'))} | "
                f"{_num(row.get('sold_generation_mwh'))} | {_num(row.get('revenue_wan'))} | "
                f"{_num(row.get('operating_cost_wan'))} | {_num(row.get('maintenance_capex_wan'))} | "
                f"{_num(row.get('income_tax_wan'))} | {_num(row.get('debt_service_wan'))} | "
                f"{_num(row.get('project_cf_wan'))} | {_num(row.get('equity_cf_wan'))} |"
            )
    else:
        lines.extend([
            f"| 最低租金覆盖率 | {_num(ind.get('minimum_tenant_rent_coverage'))} |",
            f"| 租约覆盖年限 | {_num(ind.get('lease_coverage_years'))} 年 |",
            f"| 合同收入占比 | {_pct_ratio(ind.get('contract_income_ratio'))} |",
            f"| 未锁定收入占比 | {_pct_ratio(ind.get('unlocked_income_ratio'))} |",
            "",
            "## 三、酒店经营与租约",
            "",
            "| 年度 | ADR(元) | 入住率 | RevPAR(元) | EBITDAR(万元) | 可支付租金(万元) |",
            "|---:|---:|---:|---:|---:|---:|",
        ])
        for row in (result.get("hotel_operation") or {}).get("years") or []:
            lines.append(
                f"| {row['year']} | {row['adr_yuan']:.2f} | {row['occupancy']:.2%} | "
                f"{row['revpar_yuan']:.2f} | {row['ebitdar_wan']:.2f} | {row['affordable_rent_wan']:.2f} |"
            )
    lines.extend(["", "## 四、主体关系", ""])
    parties = report_data.get("party_relationships") or []
    if parties:
        lines.extend(["| 主体ID | 名称 | 角色 | 状态 | 证据 |", "|---|---|---|---|---|"])
        for row in parties:
            lines.append(
                f"| {text(row.get('entity_id'))} | {text(row.get('name'))} | "
                f"{text(row.get('roles') or [])} | {text(row.get('status'))} | "
                f"{text(row.get('evidence_ids') or [])} |"
            )
    else:
        lines.append("- 无已绑定主体。")
    lines.extend(["", "## 五、资产边界与权证许可", ""])
    assets = report_data.get("asset_boundary") or []
    if assets:
        lines.extend(["| 资产ID | 类型 | 是否纳入 | 面积(㎡) | 状态 | 冲突/裁决 |", "|---|---|---:|---:|---|---|"])
        for row in assets:
            lines.append(
                f"| {text(row.get('scope_id'))} | {text(row.get('type'))} | "
                f"{text(row.get('included'))} | {text(row.get('area_sqm'))} | "
                f"{text(row.get('status'))} | {text(row.get('conflicts') or row.get('resolution'))} |"
            )
    licenses = report_data.get("license_ledger") or []
    lines.extend(["", "### 5.1 权证/许可台账", ""])
    if licenses:
        lines.extend(["| 许可/角色 | 持有人 | 状态 | 证据 |", "|---|---|---|---|"])
        for row in licenses:
            lines.append(
                f"| {text(row.get('license_type') or row.get('type'))} | "
                f"{text(row.get('holder_name') or row.get('holder_id'))} | "
                f"{text(row.get('status'))} | {text(row.get('evidence_ids') or [])} |"
            )
    else:
        lines.append("- 无已绑定许可记录。")
    if is_solar:
        lines.extend(["", "## 六、光伏运营关键参数", ""])
        solar = report_data.get("solar_operation") or {}
        lines.extend([
            "| 参数 | 值 |", "|---|---:|",
            f"| 装机容量(MW) | {text(solar.get('installed_capacity_mw'))} |",
            f"| 基准发电量(MWh) | {text(solar.get('base_generation_mwh'))} |",
            f"| 上网电价(元/kWh) | {text(solar.get('tariff_yuan_per_kwh'))} |",
            f"| 限电率 | {text(solar.get('curtailment_rate'))} |",
            f"| 年衰减率 | {text(solar.get('degradation_rate'))} |",
        ])
    else:
        lines.extend(["", "## 六、租约台账", ""])
        leases = report_data.get("lease_ledger") or []
        if leases:
            lines.extend(["| 单元 | 位置 | 面积(㎡) | 出租人/承租人 | 起止日 | 基础租金(万元) | 证据 |", "|---|---|---:|---|---|---:|---|"])
            for row in leases:
                lines.append(
                    f"| {text(row.get('unit_id'))} | {text(row.get('asset_location'))} | "
                    f"{text(row.get('area_sqm'))} | {text(row.get('lessor_id'))}/{text(row.get('lessee_id'))} | "
                    f"{text(row.get('start_date'))}—{text(row.get('end_date'))} | "
                    f"{text(row.get('base_rent_wan'))} | {text(row.get('evidence_ids') or [])} |"
                )
        else:
            lines.append("- 无已绑定租约。")
    lines.extend(["", "## 七、历史财务对比与勾稽", ""])
    statements = report_data.get("historical_financial_comparison") or []
    if statements:
        lines.extend(["| 主体 | 期间 | 报表类型 | 数值勾稽 | 异常 | 来源定位 |", "|---|---|---|---|---|---|"])
        for row in statements:
            lines.append(
                f"| {text(row.get('entity_id'))} | {text(row.get('period_start'))}—{text(row.get('period_end'))} | "
                f"{text(row.get('statement_type'))} | {text(row.get('reconciliation'))} | "
                f"{text(row.get('anomalies') or [])} | {text(row.get('source_locators') or [])} |"
            )
    else:
        lines.append("- 无已绑定历史报表。")
    bridge = report_data.get("valuation_transaction_bridge") or {}
    max_price_analysis = report_data.get("maximum_acceptable_price") or {}
    max_price_result = max_price_analysis.get("result") or {}
    lines.extend([
        "", "## 八、估值—成交价桥接", "",
        "| 项目 | 金额/日期 |", "|---|---:|",
        f"| 评估值(万元) | {text(bridge.get('valuation_value_wan'))} |",
        f"| 收购价(万元) | {text(bridge.get('purchase_price_wan'))} |",
        f"| 收购价较评估值差额(万元) | {text(bridge.get('purchase_price_vs_valuation_wan'))} |",
        f"| 交易税费(万元) | {text(bridge.get('transaction_tax_wan'))} |",
        f"| 总收购成本(万元) | {text(bridge.get('total_acquisition_cost_wan'))} |",
        f"| 评估日/交割日 | {text(bridge.get('valuation_date'))}/{text(bridge.get('closing_date'))} |",
        "", "### 8.1 最高可接受收购价", "",
        f"- 计算状态：{text(max_price_analysis.get('status'))}；验证状态：{text(max_price_analysis.get('validation_status'))}。",
        f"- 最高可接受价：{text(max_price_result.get('max_acquisition_price_wan'))} 万元。",
        f"- 目标IRR/最低DSCR：{text((max_price_analysis.get('parameters') or {}).get('target_irr'))}/"
        f"{text((max_price_analysis.get('parameters') or {}).get('min_dscr'))}。",
        f"- 求解哈希：`{text(max_price_analysis.get('analysis_hash'))}`；引擎版本：{text(max_price_analysis.get('engine_version'))}。",
        "", "## 九、独立情景矩阵", "",
    ])
    matrices = report_data.get("scenario_matrices") or []
    if matrices:
        if is_solar:
            lines.extend(["| 矩阵/场景 | 变更 | 收购价(万元) | 上网电价(元/kWh) | 年发电量(MWh) | 利用小时 | 年运维费(万元) | 项目IRR | 最低DSCR |", "|---|---|---:|---:|---:|---:|---:|---:|---:|"])
        else:
            lines.extend(["| 矩阵/场景 | 变更 | 收购价(万元) | 市场租金 | 入住率 | 项目IRR | 最低DSCR |", "|---|---|---:|---:|---|---:|---:|"])
        for matrix in matrices:
            for row in matrix.get("rows") or []:
                indicators = row.get("indicators") or {}
                prefix = (
                    f"| {text(matrix.get('matrix_id'))}/{text(row.get('scenario_id'))} | "
                    f"{text(row.get('changes'))} | {text(row.get('purchase_price_wan'))} | "
                )
                if is_solar:
                    lines.append(
                        prefix + f"{text(row.get('tariff_yuan_per_kwh'))} | "
                        f"{text(row.get('annual_generation_mwh'))} | {text(row.get('utilization_hours'))} | "
                        f"{text(row.get('annual_opex_wan'))} | {text(indicators.get('project_irr_pct'))} | "
                        f"{text(indicators.get('minimum_dscr'))} |"
                    )
                else:
                    lines.append(
                        prefix + f"{text(row.get('market_rent'))} | {text(row.get('occupancy'))} | "
                        f"{text(indicators.get('project_irr_pct'))} | {text(indicators.get('minimum_dscr'))} |"
                    )
    else:
        lines.append("- 当前 run 未绑定独立情景矩阵。")
    lines.extend(["", "## 十、红旗、成交条件与否决事项", ""])
    for row in report_data.get("red_flags") or []:
        lines.append(f"- 红旗 `{text(row.get('code'))}`：{text(row.get('status'))}；{text(row.get('resolution'))}")
    for item in report_data.get("closing_conditions") or []:
        lines.append(f"- 成交条件：{text(item)}")
    for item in report_data.get("veto_items") or []:
        lines.append(f"- 否决事项：{text(item)}")
    lines.extend(["", "## 十一、资料处理与字段证据台账", ""])
    ledger = report_data.get("source_processing_ledger") or []
    if ledger:
        lines.extend(["| 字段 | 证据ID | 来源 | 绑定状态 |", "|---|---|---|---|"])
        for row in ledger:
            lines.append(
                f"| {text(row.get('field'))} | {text(row.get('evidence_ids'))} | "
                f"{text(row.get('source'))} | {text(row.get('validation_status'))} |"
            )
    lines.extend([
        "", "## 十二、确定性校验与工件绑定", "",
        f"- 运行校验：{text(run.get('validation_status'))}",
        f"- Spec 校验：{text(run.get('formal_spec_valid'))}",
        f"- 证据绑定：{text(run.get('evidence_status'))}；`{text(run.get('evidence_binding_hash'))}`",
        f"- 数值一致性：{text(run.get('consistency_ok'))}",
        "- 本报告数字仅绑定上述不可变 run、输入快照与内容哈希。",
    ])
    return "\n".join(lines) + "\n"


def enqueue_artifact(
    workspace_id: str, run_id: str, *, idempotency_key: str = "", request_id: str = "",
) -> dict[str, Any]:
    """Create a durable artifact job that can be polled before rendering."""

    request_id = request_id or f"req_{uuid.uuid4().hex}"
    run = get_run(workspace_id, run_id)
    if not run:
        return {"ok": False, "error": "run_not_found"}
    if run.get("status") != "succeeded" or run.get("consistency_ok") is not True:
        return {
            "ok": False,
            "error": "RUN_INCONSISTENT",
            "reason": "succeeded_consistent_run_required",
        }
    required_bindings = (
        "run_id", "spec_hash", "input_hash", "spec_snapshot_hash",
        "evidence_binding_hash", "model_version",
    )
    if any(not str(run.get(field) or "").strip() for field in required_bindings):
        return {
            "ok": False,
            "error": "RUN_BINDING_INCOMPLETE",
            "missing": [field for field in required_bindings if not str(run.get(field) or "").strip()],
        }
    body_hash = _hash({
        "run_id": run_id, "spec_hash": run.get("spec_hash"), "fact_revision": run.get("spec_id"),
        "spec_snapshot_hash": run.get("spec_snapshot_hash"),
        "evidence_binding_hash": run.get("evidence_binding_hash"),
        "template_version": "asset_acquisition.v2",
    })
    scope = f"artifact:{idempotency_key}" if idempotency_key else ""
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        prior = _active_idempotency_record(state["idempotency"], scope) if scope else None
        if prior:
            if prior.get("body_hash") != body_hash:
                return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior.get("artifact_id", "")}
            existing = state["artifacts"].get(prior.get("artifact_id"))
            if not existing:
                raise RuntimeError("artifact idempotency record points to a missing artifact")
            return {**existing, "idempotent_replay": True}
        artifact_id = f"artifact_{uuid.uuid4().hex}"
        job_id = f"artifact_job_{uuid.uuid4().hex}"
        created_at = _now()
        row = {
            "ok": True, "artifact_id": artifact_id, "artifact_job_id": job_id,
            "status": "queued", "progress": 0, "type": "asset_acquisition", "run_id": run_id,
            "spec_hash": run.get("spec_hash"), "fact_revision": run.get("spec_id"),
            "spec_snapshot_hash": run.get("spec_snapshot_hash"),
            "evidence_binding_version": run.get("evidence_binding_version"),
            "evidence_binding_hash": run.get("evidence_binding_hash"),
            "template_version": "asset_acquisition.v2", "created_at": created_at,
            "updated_at": created_at, "request_id": request_id, "files": [],
            "numeric_consistency": "pending", "integrity_status": "pending",
            "state_history": [
                _history_event("queued", request_id=request_id, run_id=run_id),
            ],
        }
        state["artifacts"][artifact_id] = row
        if scope:
            state["idempotency"][scope] = _idempotency_record(
                scope, body_hash, artifact_id=artifact_id,
            )
        _save(workspace_id, state)
        return row


def _bind_succeeded_artifact(
    workspace_id: str,
    run: dict[str, Any],
    artifact: dict[str, Any],
) -> dict[str, Any]:
    """Bind a successful formal pack to the report-side finance revision."""

    from lvke_mcp.domains.reports import artifacts as report_artifacts

    expected = {
        "finance_run_id": str(run.get("run_id") or ""),
        "input_hash": run.get("input_hash"),
        "spec_hash": run.get("spec_hash"),
        "spec_id": run.get("spec_id"),
        "fact_revision": run.get("spec_id"),
        "spec_snapshot_hash": run.get("spec_snapshot_hash"),
        "evidence_binding_version": run.get("evidence_binding_version"),
        "evidence_binding_hash": run.get("evidence_binding_hash"),
        "model_version": run.get("model_version"),
        "template_version": artifact.get("template_version"),
        "artifact_id": artifact.get("artifact_id"),
        "artifact_job_id": artifact.get("artifact_job_id"),
        "artifact_status": "succeeded",
        "report_data_hash": artifact.get("report_data_hash"),
        "binding_kind": "asset_acquisition",
    }
    report_artifacts.bind_finance_run(
        workspace_id,
        expected["finance_run_id"],
        section="asset_acquisition_artifact",
        fin={
            key: value for key, value in expected.items()
            if key != "finance_run_id"
        } | {"validation_level": "complete"},
    )
    actual = report_artifacts.load(
        workspace_id,
        "finance_binding",
        {},
    ) or {}
    mismatches = [
        {"field": key, "expected": value, "actual": actual.get(key)}
        for key, value in expected.items()
        if actual.get(key) != value
    ]
    if mismatches:
        return {
            "ok": False,
            "error": "ARTIFACT_BINDING_FAILED",
            "reason": "finance_binding_mismatch",
            "mismatches": mismatches,
        }
    return {"ok": True, "binding": actual}


def execute_queued_artifact(
    workspace_id: str,
    artifact_id: str,
) -> None:
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        row = state["artifacts"].get(artifact_id)
        if not row or row.get("status") in {"succeeded", "failed", "cancelled"}:
            return
        row.update({"status": "running", "progress": 10, "updated_at": _now()})
        run_id = str(row.get("run_id") or "")
        job_id = str(row.get("artifact_job_id") or "")
        request_id = str(row.get("request_id") or "")
        _save(workspace_id, state)
    try:
        result = generate_artifacts(
            workspace_id, run_id, request_id=request_id,
            artifact_id=artifact_id, artifact_job_id=job_id,
        )
        if result.get("ok"):
            return
        raw_code = str(result.get("error") or "ARTIFACT_MISMATCH")
        code = (
            raw_code
            if re.fullmatch(r"[A-Z][A-Z0-9_]{1,79}", raw_code)
            else "ARTIFACT_MISMATCH"
        )
        error = {
            "code": code,
            "message": _ARTIFACT_GENERATION_FAILURE_MESSAGE,
            "retryable": False,
        }
    except Exception as exc:  # noqa: BLE001
        _LOG.warning(
            "asset acquisition artifact generation failed; artifact_id=%s "
            "error_type=%s",
            artifact_id,
            type(exc).__name__,
        )
        error = {
            "code": "ARTIFACT_GENERATION_FAILED",
            "message": _ARTIFACT_GENERATION_FAILURE_MESSAGE,
            "retryable": False,
        }
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        row = state["artifacts"].get(artifact_id)
        if row and row.get("status") != "cancelled":
            row.update({
                "ok": False, "status": "failed", "progress": 100,
                "numeric_consistency": "failed", "error": error, "updated_at": _now(),
            })
            _save(workspace_id, state)


def generate_artifacts(
    workspace_id: str, run_id: str, *, idempotency_key: str = "", request_id: str = "",
    artifact_id: str = "", artifact_job_id: str = "",
) -> dict[str, Any]:
    """Generate an atomically published, consistency-checked artifact pack."""

    run = get_run(workspace_id, run_id)
    if not run:
        return {"ok": False, "error": "run_not_found"}
    if run.get("status") != "succeeded" or not run.get("consistency_ok"):
        return {"ok": False, "error": "RUN_INCONSISTENT"}
    body_hash = _hash({
        "run_id": run_id, "spec_hash": run.get("spec_hash"), "fact_revision": run.get("spec_id"),
        "spec_snapshot_hash": run.get("spec_snapshot_hash"),
        "evidence_binding_hash": run.get("evidence_binding_hash"),
        "template_version": "asset_acquisition.v2",
    })
    scope = f"artifact:{idempotency_key}" if idempotency_key else ""
    if scope:
        with _state_guard(workspace_id):
            state = _load(workspace_id)
            prior = _active_idempotency_record(state["idempotency"], scope)
            if prior:
                if prior.get("body_hash") != body_hash:
                    return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior.get("artifact_id", "")}
                existing = state["artifacts"].get(prior.get("artifact_id"))
                if not existing:
                    raise RuntimeError("artifact idempotency record points to a missing artifact")
                return {**existing, "idempotent_replay": True}

    artifact_id = artifact_id or f"artifact_{uuid.uuid4().hex}"
    artifacts_root = _root(workspace_id) / "artifacts"
    artifacts_root.mkdir(parents=True, exist_ok=True)
    staging = artifacts_root / f".{artifact_id}.staging-{uuid.uuid4().hex}"
    final_root = artifacts_root / artifact_id
    staging.mkdir(parents=True, exist_ok=False)
    report_data = build_acquisition_report_data(
        workspace_id,
        run,
    )
    markdown = render_markdown(run, report_data)
    md_path = staging / "资产收购可行性研究报告.md"
    docx_path = staging / "资产收购可行性研究报告.docx"
    xlsx_path = staging / "资产收购财务模型.xlsx"
    report_data_path = staging / "资产收购报告数据.json"
    index_path = staging / "附件索引.json"
    try:
        md_path.write_text(markdown, encoding="utf-8")
        from lvke_mcp.domains.reports import doc_service as _doc_svc

        docx_path.write_bytes(_doc_svc.markdown_to_docx(markdown))
        report_data_path.write_text(json.dumps(report_data, ensure_ascii=False, indent=2), encoding="utf-8")
        _write_minimal_xlsx(xlsx_path, run, report_data=report_data)
        consistency = _check_artifact_consistency(
            run, markdown, docx_path, xlsx_path, report_data_path=report_data_path,
        )
        if consistency["status"] != "passed":
            return {
                "ok": False, "error": "ARTIFACT_MISMATCH", "reason": "numeric_or_binding_mismatch",
                "consistency": consistency,
            }
        files = [md_path, docx_path, xlsx_path, report_data_path]
        index = {
            "artifact_id": artifact_id, "run_id": run_id,
            "spec_hash": run.get("spec_hash"), "fact_revision": run.get("spec_id"),
            "spec_snapshot_hash": run.get("spec_snapshot_hash"),
            "evidence_binding_version": run.get("evidence_binding_version"),
            "evidence_binding_hash": run.get("evidence_binding_hash"),
            "model_version": run.get("model_version"), "generated_at": _now(),
            "report_data_hash": report_data.get("report_data_hash"),
            "numeric_consistency": consistency,
            "files": [
                {"name": path.name, "size_bytes": path.stat().st_size, "sha256": _file_hash(path)}
                for path in files
            ],
        }
        index_path.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
        index["files"].append({"name": index_path.name, "size_bytes": index_path.stat().st_size, "sha256": _file_hash(index_path)})
        os.replace(staging, final_root)
    finally:
        if staging.exists():
            shutil.rmtree(staging, ignore_errors=True)

    row = {
        "ok": True, "artifact_id": artifact_id,
        "artifact_job_id": artifact_job_id or f"artifact_job_{uuid.uuid4().hex}",
        "status": "succeeded", "progress": 100,
        "type": "asset_acquisition", "run_id": run_id,
        "spec_hash": run.get("spec_hash"), "fact_revision": run.get("spec_id"),
        "spec_snapshot_hash": run.get("spec_snapshot_hash"),
        "evidence_binding_version": run.get("evidence_binding_version"),
        "evidence_binding_hash": run.get("evidence_binding_hash"),
        "template_version": "asset_acquisition.v2", "created_at": _now(),
        "updated_at": _now(), "request_id": request_id,
        "files": index["files"], "directory": str(final_root),
        "report_data_hash": report_data.get("report_data_hash"),
        "numeric_consistency": "passed", "consistency_checks": consistency["checks"],
        "state_history": [
            _history_event(
                "artifact_generated",
                request_id=request_id, run_id=run_id, artifact_id=artifact_id,
            ),
        ],
    }
    try:
        with _state_guard(workspace_id):
            state = _load(workspace_id)
            if scope and (prior := _active_idempotency_record(state["idempotency"], scope)):
                # Another worker completed the same request while this one rendered.
                if prior.get("body_hash") != body_hash:
                    shutil.rmtree(final_root, ignore_errors=True)
                    return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior.get("artifact_id", "")}
                shutil.rmtree(final_root, ignore_errors=True)
                return {**state["artifacts"][prior["artifact_id"]], "idempotent_replay": True}
            state["artifacts"][artifact_id] = row
            if scope:
                state["idempotency"][scope] = _idempotency_record(
                    scope, body_hash, artifact_id=artifact_id,
                )
            stored_run = state["runs"].get(run_id)
            if stored_run:
                stored_run["lifecycle_status"] = "artifact_generated"
                stored_run.setdefault("state_history", []).append(_history_event(
                    "artifact_generated", request_id=request_id,
                    artifact_id=artifact_id,
                ))
            _save(workspace_id, state)
    except BaseException:
        # State publication and the visible directory are one logical commit.
        # If metadata cannot be committed, do not leave an untracked formal pack.
        shutil.rmtree(final_root, ignore_errors=True)
        raise

    binding_result = _bind_succeeded_artifact(
        workspace_id,
        run,
        row,
    )
    if not binding_result.get("ok"):
        # The artifact is successful only when the report-side binding proves
        # the same run/spec/fact revision.
        with _state_guard(workspace_id):
            state = _load(workspace_id)
            stored = state["artifacts"].get(artifact_id)
            if stored:
                stored.update({
                    "ok": False, "status": "failed",
                    "error": {
                        "code": "ARTIFACT_BINDING_FAILED",
                        "message": "formal artifact finance binding failed",
                        "retryable": False,
                        "details": binding_result.get("mismatches") or [],
                    },
                    "updated_at": _now(),
                })
            stored_run = state["runs"].get(run_id)
            if stored_run:
                stored_run["lifecycle_status"] = "artifact_binding_failed"
                stored_run.setdefault("state_history", []).append(_history_event(
                    "artifact_binding_failed", request_id=request_id,
                    artifact_id=artifact_id,
                ))
            _save(workspace_id, state)
        shutil.rmtree(final_root, ignore_errors=True)
        return {
            "ok": False, "error": "ARTIFACT_BINDING_FAILED",
            "reason": binding_result.get("reason") or "finance_binding_mismatch",
            "mismatches": binding_result.get("mismatches") or [],
        }

    return dict(_load(workspace_id)["artifacts"].get(artifact_id) or row)


def _recover_published_artifact(
    workspace_id: str, artifact_id: str, row: dict[str, Any],
) -> tuple[bool, str]:
    """Recover a pack atomically published just before a process crash."""

    directory = _root(workspace_id) / "artifacts" / artifact_id
    index_path = directory / "附件索引.json"
    if not index_path.is_file():
        return False, "not_published"
    try:
        index = json.loads(index_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError, TypeError):
        return False, "invalid_index"
    if (
        index.get("artifact_id") != artifact_id
        or index.get("run_id") != row.get("run_id")
    ):
        return False, "binding_mismatch"
    consistency = index.get("numeric_consistency") or {}
    if not isinstance(consistency, dict) or consistency.get("status") != "passed":
        return False, "numeric_consistency_failed"
    files = list(index.get("files") or [])
    for item in files:
        path = directory / str(item.get("name") or "")
        if not path.is_file() or _file_hash(path) != item.get("sha256"):
            return False, "file_hash_mismatch"
    files.append({
        "name": index_path.name,
        "size_bytes": index_path.stat().st_size,
        "sha256": _file_hash(index_path),
    })
    row.update({
        "ok": True, "status": "succeeded", "progress": 100,
        "directory": str(directory), "files": files,
        "numeric_consistency": "passed",
        "report_data_hash": index.get("report_data_hash"),
        "consistency_checks": list(consistency.get("checks") or []),
        "recovered_from_published_artifact": True,
        "updated_at": _now(),
    })
    row.setdefault("state_history", []).append(_history_event(
        "artifact_generated",
        request_id=str(row.get("request_id") or ""),
        artifact_id=artifact_id,
        recovery=True,
    ))
    return True, "recovered"


def _workspace_ids() -> list[str]:
    root = data_root() / "workspaces"
    if not root.is_dir():
        return []
    return sorted(
        path.name
        for path in root.iterdir()
        if path.is_dir() and str(path.name).strip()
    )


def recover_incomplete_acquisition_tasks(
    submit_run: Any | None = None, submit_artifact: Any | None = None,
) -> dict[str, Any]:
    """Requeue durable acquisition runs/artifacts left by a process restart."""

    run_submitter = submit_run or (
        lambda workspace_id, run_id: _RECOVERY_POOL.submit(
            execute_queued_run,
            workspace_id,
            run_id,
        )
    )
    artifact_submitter = submit_artifact or (
        lambda workspace_id, artifact_id: _RECOVERY_POOL.submit(
            execute_queued_artifact,
            workspace_id,
            artifact_id,
        )
    )
    runs: list[tuple[str, str]] = []
    artifacts: list[tuple[str, str]] = []
    recovered_bindings: list[tuple[str, str]] = []
    recovered_artifacts = 0
    failed_artifacts = 0
    for workspace_id in _workspace_ids():
        if not _state_path(workspace_id).is_file():
            continue
        with _state_guard(workspace_id):
            state = _load(workspace_id)
            changed = False
            for run_id, row in state["runs"].items():
                if row.get("status") not in {"queued", "running"}:
                    continue
                row.update({
                    "status": "queued", "progress": 0, "updated_at": _now(),
                    "recovered_at": _now(),
                    "recovery_count": int(row.get("recovery_count") or 0) + 1,
                })
                row.setdefault("state_history", []).append(_history_event(
                    "queued", request_id=str(row.get("request_id") or ""), recovery=True,
                ))
                runs.append((workspace_id, run_id))
                changed = True
            for artifact_id, row in state["artifacts"].items():
                if row.get("status") not in {"queued", "running"}:
                    continue
                recovered, reason = _recover_published_artifact(
                    workspace_id,
                    artifact_id,
                    row,
                )
                if recovered:
                    recovered_artifacts += 1
                    recovered_bindings.append((workspace_id, artifact_id))
                    changed = True
                    continue
                if reason != "not_published":
                    row.update({
                        "ok": False, "status": "failed", "progress": 100,
                        "numeric_consistency": "failed", "updated_at": _now(),
                        "error": {
                            "code": "ARTIFACT_RECOVERY_FAILED", "message": reason,
                            "retryable": False,
                        },
                    })
                    failed_artifacts += 1
                    changed = True
                    continue
                row.update({
                    "status": "queued", "progress": 0, "updated_at": _now(),
                    "recovered_at": _now(),
                    "recovery_count": int(row.get("recovery_count") or 0) + 1,
                })
                artifacts.append((workspace_id, artifact_id))
                changed = True
            if changed:
                _save(workspace_id, state)
    for workspace_id, artifact_id in recovered_bindings:
        artifact = dict(
            _load(workspace_id)["artifacts"].get(artifact_id)
            or {}
        )
        run = get_run(
            workspace_id,
            str(artifact.get("run_id") or ""),
        )
        if run:
            with _state_guard(workspace_id):
                state = _load(workspace_id)
                stored_run = state["runs"].get(str(artifact.get("run_id") or ""))
                if stored_run:
                    stored_run["lifecycle_status"] = "artifact_generated"
                    stored_run.setdefault("state_history", []).append(_history_event(
                        "artifact_generated",
                        request_id=str(artifact.get("request_id") or ""),
                        artifact_id=artifact_id,
                        recovery=True,
                    ))
                    _save(workspace_id, state)
                    run = copy.deepcopy(stored_run)
        bound = _bind_succeeded_artifact(
            workspace_id,
            run,
            artifact,
        ) if run else {
            "ok": False, "error": "ARTIFACT_BINDING_FAILED",
        }
        if not bound.get("ok"):
            with _state_guard(workspace_id):
                state = _load(workspace_id)
                stored = state["artifacts"].get(artifact_id)
                if stored:
                    stored.update({
                        "ok": False, "status": "failed",
                        "error": {
                            "code": "ARTIFACT_BINDING_FAILED",
                            "message": "recovered artifact finance binding failed",
                            "retryable": False,
                        },
                        "updated_at": _now(),
                    })
                    recovered_artifacts -= 1
                    failed_artifacts += 1
                    _save(workspace_id, state)
    for workspace_id, run_id in runs:
        run_submitter(workspace_id, run_id)
    for workspace_id, artifact_id in artifacts:
        artifact_submitter(workspace_id, artifact_id)
    return {
        "runs_requeued": len(runs),
        "artifacts_requeued": len(artifacts),
        "artifacts_recovered": recovered_artifacts,
        "artifacts_failed": failed_artifacts,
    }


def _artifact_filename_is_safe(filename: str) -> bool:
    value = str(filename or "")
    return bool(
        value
        and value not in {".", ".."}
        and "\x00" not in value
        and "/" not in value
        and "\\" not in value
        and Path(value).name == value
        and not Path(value).is_absolute()
    )


def _artifact_media_type(filename: str) -> str:
    stable_types = {
        ".md": "text/markdown; charset=utf-8",
        ".json": "application/json; charset=utf-8",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }
    suffix = Path(str(filename or "")).suffix.lower()
    return stable_types.get(
        suffix,
        mimetypes.guess_type(str(filename or ""))[0]
        or "application/octet-stream",
    )


def get_artifact(
    workspace_id: str,
    artifact_id: str,
) -> dict[str, Any]:
    """Return an acquisition artifact only after re-verifying its formal pack."""

    row = dict(
        _load(workspace_id)["artifacts"].get(artifact_id)
        or {}
    )
    if not row:
        return {}
    if row.get("status") != "succeeded":
        return {
            **row,
            "integrity_status": (
                "pending" if row.get("status") in {"queued", "running"} else "failed"
            ),
        }
    run = get_run(
        workspace_id,
        str(row.get("run_id") or ""),
    )

    artifacts_root = (
        _root(workspace_id) / "artifacts"
    ).resolve()
    expected_directory = (artifacts_root / artifact_id).resolve()
    directory = Path(str(row.get("directory") or "")).resolve()
    try:
        expected_directory.relative_to(artifacts_root)
        directory.relative_to(artifacts_root)
    except ValueError:
        return {
            **row, "ok": False, "error": "ARTIFACT_MISMATCH",
            "integrity_status": "invalid_directory",
            "failures": [{"reason": "directory_outside_artifact_root"}],
        }
    if directory != expected_directory or not directory.is_dir():
        return {
            **row, "ok": False, "error": "ARTIFACT_MISMATCH",
            "integrity_status": "invalid_directory",
            "failures": [{"reason": "artifact_directory_binding_mismatch"}],
        }

    failures: list[dict[str, Any]] = []
    spec_row = get_spec(
        workspace_id,
        str(run.get("spec_id") or ""),
    )
    bound_spec = spec_row.get("spec") if isinstance(spec_row, dict) else None
    if not isinstance(bound_spec, dict) or _hash(bound_spec) != run.get("spec_hash"):
        failures.append({"reason": "run_spec_snapshot_mismatch"})
    else:
        evidence_ok, current_evidence = _current_evidence_matches_run(
            workspace_id,
            run,
            bound_spec,
        )
        if not evidence_ok:
            failures.append({
                "reason": "evidence_binding_stale",
                "expected": run.get("evidence_binding_hash"),
                "actual": current_evidence.get("binding_hash"),
                "status": current_evidence.get("status"),
            })
    file_rows: dict[str, dict[str, Any]] = {}
    for item in row.get("files") or []:
        if not isinstance(item, dict):
            failures.append({"reason": "invalid_file_manifest_entry"})
            continue
        name = str(item.get("name") or "")
        if not _artifact_filename_is_safe(name):
            failures.append({"name": name, "reason": "invalid_filename_path"})
            continue
        if name in file_rows:
            failures.append({"name": name, "reason": "duplicate_file_manifest_entry"})
            continue
        file_rows[name] = item
        path = directory / name
        try:
            path.resolve().relative_to(directory)
        except ValueError:
            failures.append({"name": name, "reason": "path_escape"})
            continue
        if path.is_symlink():
            failures.append({"name": name, "reason": "symlink_not_allowed"})
            continue
        if not path.is_file():
            failures.append({"name": name, "reason": "missing"})
            continue
        actual_size = path.stat().st_size
        if actual_size != item.get("size_bytes"):
            failures.append({
                "name": name, "reason": "size_mismatch",
                "expected": item.get("size_bytes"), "actual": actual_size,
            })
        actual_hash = _file_hash(path)
        if actual_hash != item.get("sha256"):
            failures.append({
                "name": name, "reason": "sha256_mismatch", "actual": actual_hash,
            })

    required_names = {
        "资产收购可行性研究报告.md",
        "资产收购可行性研究报告.docx",
        "资产收购财务模型.xlsx",
        "资产收购报告数据.json",
        "附件索引.json",
    }
    for missing_name in sorted(required_names - set(file_rows)):
        failures.append({"name": missing_name, "reason": "manifest_entry_missing"})

    index: dict[str, Any] = {}
    index_path = directory / "附件索引.json"
    if index_path.is_file():
        try:
            loaded_index = json.loads(index_path.read_text(encoding="utf-8"))
            index = loaded_index if isinstance(loaded_index, dict) else {}
        except (OSError, json.JSONDecodeError, TypeError):
            failures.append({"name": index_path.name, "reason": "invalid_json"})
    if not index:
        failures.append({"name": index_path.name, "reason": "index_missing_or_invalid"})
    else:
        for field, expected in (
            ("artifact_id", artifact_id),
            ("run_id", row.get("run_id")),
            ("spec_hash", row.get("spec_hash")),
            ("fact_revision", row.get("fact_revision")),
            ("spec_snapshot_hash", row.get("spec_snapshot_hash")),
            ("evidence_binding_version", row.get("evidence_binding_version")),
            ("evidence_binding_hash", row.get("evidence_binding_hash")),
            ("report_data_hash", row.get("report_data_hash")),
        ):
            if index.get(field) != expected:
                failures.append({
                    "name": index_path.name, "reason": f"index_{field}_mismatch",
                    "expected": expected, "actual": index.get(field),
                })
        indexed_files = {
            str(item.get("name") or ""): item
            for item in (index.get("files") or [])
            if isinstance(item, dict)
        }
        state_files = {
            name: item for name, item in file_rows.items()
            if name != index_path.name
        }
        if indexed_files != state_files:
            failures.append({"name": index_path.name, "reason": "file_manifest_mismatch"})
        index_consistency = index.get("numeric_consistency") or {}
        index_checks = index_consistency.get("checks") or [] if isinstance(index_consistency, dict) else []
        if (
            not isinstance(index_consistency, dict)
            or index_consistency.get("status") != "passed"
            or not index_checks
            or any(not check.get("passed") for check in index_checks if isinstance(check, dict))
            or any(not isinstance(check, dict) for check in index_checks)
        ):
            failures.append({"name": index_path.name, "reason": "numeric_consistency_incomplete"})
        if row.get("numeric_consistency") != "passed" or row.get("consistency_checks") != index_checks:
            failures.append({"name": index_path.name, "reason": "numeric_consistency_state_mismatch"})

    report_data: dict[str, Any] | None = None
    report_data_path = directory / "资产收购报告数据.json"
    if report_data_path.is_file():
        try:
            loaded = json.loads(report_data_path.read_text(encoding="utf-8"))
            report_data = loaded if isinstance(loaded, dict) else None
        except (OSError, json.JSONDecodeError, TypeError):
            report_data = None
    if not isinstance(report_data, dict):
        failures.append({"name": report_data_path.name, "reason": "invalid_json"})
    else:
        embedded_hash = report_data.get("report_data_hash")
        payload = {key: value for key, value in report_data.items() if key != "report_data_hash"}
        calculated_hash = _hash(payload)
        if embedded_hash != calculated_hash or embedded_hash != row.get("report_data_hash"):
            failures.append({
                "name": report_data_path.name, "reason": "report_data_hash_mismatch",
                "expected": row.get("report_data_hash"), "actual": calculated_hash,
            })
        bindings = report_data.get("bindings") or {}
        for field in (
            "run_id", "spec_hash", "input_hash", "model_version",
            "spec_snapshot_hash", "evidence_binding_version", "evidence_binding_hash",
        ):
            if bindings.get(field) != run.get(field):
                failures.append({
                    "name": report_data_path.name,
                    "reason": f"report_binding_{field}_mismatch",
                })
        if str(bindings.get("spec_id") or "") != str(row.get("fact_revision") or ""):
            failures.append({
                "name": report_data_path.name,
                "reason": "report_binding_fact_revision_mismatch",
            })
        if report_data.get("maximum_acceptable_price") != run.get("max_acquisition_price_analysis"):
            failures.append({
                "name": report_data_path.name,
                "reason": "report_maximum_price_mismatch",
            })

    # Re-run the numeric/binding verifier against the current immutable run;
    # trusting only the checks serialized at generation time would miss a
    # stale artifact after an out-of-band run mutation.
    if isinstance(report_data, dict):
        try:
            current_consistency = _check_artifact_consistency(
                run,
                (directory / "资产收购可行性研究报告.md").read_text(encoding="utf-8"),
                directory / "资产收购可行性研究报告.docx",
                directory / "资产收购财务模型.xlsx",
                report_data_path=report_data_path,
            )
        except Exception as exc:  # noqa: BLE001
            failures.append({
                "reason": "numeric_consistency_unverifiable",
                "error_type": type(exc).__name__,
            })
        else:
            if current_consistency.get("status") != "passed":
                failures.append({
                    "reason": "numeric_consistency_failed",
                    "checks": [
                        check for check in (current_consistency.get("checks") or [])
                        if not check.get("passed")
                    ],
                })
            if current_consistency.get("checks") != row.get("consistency_checks"):
                failures.append({"reason": "numeric_consistency_snapshot_mismatch"})

    if failures:
        return {
            **row, "ok": False, "error": "ARTIFACT_MISMATCH",
            "integrity_status": "failed", "failures": failures,
        }
    return {
        **row, "ok": True, "integrity_status": "passed",
        "report_data": report_data,
    }


def list_artifacts(
    workspace_id: str,
    *,
    limit: int = 50,
) -> list[dict[str, Any]]:
    """Return artifact jobs newest-first, including their integrity projection."""

    state = _load(workspace_id)
    ids = sorted(
        state["artifacts"],
        key=lambda artifact_id: (
            str(state["artifacts"][artifact_id].get("created_at") or ""),
            artifact_id,
        ),
        reverse=True,
    )
    rows = [
        get_artifact(workspace_id, artifact_id)
        for artifact_id in ids
    ]
    return [
        row for row in rows if row
    ][: max(1, min(int(limit or 50), 100))]


def _resolve_artifact_download(
    workspace_id: str,
    artifact_id: str,
    filename: str,
) -> dict[str, Any]:
    """Resolve one consistent artifact file after path and hash validation."""

    if not _artifact_filename_is_safe(filename):
        return {"ok": False, "error": "INVALID_FILENAME_PATH"}
    artifact = get_artifact(workspace_id, artifact_id)
    if not artifact:
        return {"ok": False, "error": "ARTIFACT_NOT_FOUND"}
    if artifact.get("status") != "succeeded":
        return {
            "ok": False, "error": "ARTIFACT_NOT_READY",
            "status": artifact.get("status"),
        }
    if artifact.get("integrity_status") != "passed" or not artifact.get("ok", True):
        return {
            "ok": False, "error": "ARTIFACT_MISMATCH",
            "failures": artifact.get("failures") or [],
        }
    run = get_run(
        workspace_id,
        str(artifact.get("run_id") or ""),
    )
    if run.get("status") != "succeeded" or not run.get("consistency_ok"):
        return {"ok": False, "error": "RUN_INCONSISTENT"}

    file_row = next(
        (
            item for item in (artifact.get("files") or [])
            if str(item.get("name") or "") == filename
        ),
        None,
    )
    if not file_row:
        return {"ok": False, "error": "ARTIFACT_FILE_NOT_FOUND"}
    directory = Path(str(artifact.get("directory") or "")).resolve()
    path = (directory / filename).resolve()
    try:
        path.relative_to(directory)
    except ValueError:
        return {"ok": False, "error": "INVALID_FILENAME_PATH"}
    if not path.is_file() or path.is_symlink():
        return {"ok": False, "error": "ARTIFACT_FILE_NOT_FOUND"}
    actual_hash = _file_hash(path)
    if actual_hash != file_row.get("sha256"):
        return {
            "ok": False, "error": "ARTIFACT_MISMATCH",
            "failures": [{
                "name": filename, "reason": "sha256_mismatch",
                "expected": file_row.get("sha256"), "actual": actual_hash,
            }],
        }
    return {
        "ok": True,
        "artifact_id": artifact_id,
        "run_id": artifact.get("run_id"),
        "filename": filename,
        "path": path,
        "size_bytes": path.stat().st_size,
        "sha256": actual_hash,
        "media_type": _artifact_media_type(filename),
        "integrity_status": artifact.get("integrity_status"),
    }


def resolve_artifact_download(
    workspace_id: str,
    artifact_id: str,
    filename: str,
) -> dict[str, Any]:
    """Resolve a complete acquisition artifact after integrity validation."""

    return _resolve_artifact_download(
        workspace_id,
        artifact_id,
        filename,
    )


def resolve_artifact_candidate_download(
    workspace_id: str,
    artifact_id: str,
    filename: str,
) -> dict[str, Any]:
    """Compatibility alias for integrity-validated artifact resolution."""

    return _resolve_artifact_download(
        workspace_id,
        artifact_id,
        filename,
    )


def _read_resolved_artifact(
    workspace_id: str,
    artifact_id: str,
    resolved: dict[str, Any],
) -> dict[str, Any]:
    if not resolved.get("ok"):
        return resolved
    path = resolved.get("path")
    try:
        content = path.read_bytes() if isinstance(path, Path) else b""
    except OSError:
        return {"ok": False, "error": "ARTIFACT_FILE_NOT_FOUND"}
    digest = hashlib.sha256(content).hexdigest()
    if digest != resolved.get("sha256") or len(content) != resolved.get("size_bytes"):
        get_artifact(workspace_id, artifact_id)
        return {
            "ok": False,
            "error": "ARTIFACT_MISMATCH",
            "failures": [{
                "name": resolved.get("filename"),
                "reason": "content_changed_after_validation",
            }],
        }
    return {**resolved, "content": content}


def read_artifact_download(
    workspace_id: str,
    artifact_id: str,
    filename: str,
) -> dict[str, Any]:
    return _read_resolved_artifact(
        workspace_id,
        artifact_id,
        resolve_artifact_download(
            workspace_id,
            artifact_id,
            filename,
        ),
    )


def read_artifact_candidate_download(
    workspace_id: str,
    artifact_id: str,
    filename: str,
) -> dict[str, Any]:
    return _read_resolved_artifact(
        workspace_id,
        artifact_id,
        resolve_artifact_candidate_download(
            workspace_id,
            artifact_id,
            filename,
        ),
    )
















def _xlsx_summary_values(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("xl/worksheets/sheet1.xml"))
    values: dict[str, str] = {}
    for row in root.iter():
        if not row.tag.endswith("}row"):
            continue
        cells: list[str] = []
        for cell in row:
            if not cell.tag.endswith("}c"):
                continue
            inline = "".join(node.text or "" for node in cell.iter() if node.tag.endswith("}t"))
            raw = next((node.text or "" for node in cell if node.tag.endswith("}v")), "")
            cells.append(inline if inline else raw)
        if len(cells) >= 2:
            values[cells[0]] = cells[1]
    return values


def _same_number(actual: str, expected: Any) -> bool:
    try:
        return math.isclose(float(actual), float(expected), rel_tol=1e-12, abs_tol=1e-9)
    except (TypeError, ValueError):
        return False


def _check_artifact_consistency(
    run: dict[str, Any], markdown: str, docx_path: Path, xlsx_path: Path,
    *, report_data_path: Path | None = None,
) -> dict[str, Any]:
    from docx import Document  # type: ignore

    document = Document(io.BytesIO(docx_path.read_bytes()))
    docx_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    result = run.get("result") or {}
    indicators = result.get("indicators") or {}
    asset_type = str(result.get("asset_type") or "hotel_lease")
    is_solar = asset_type == "solar_power"
    max_price_analysis = run.get("max_acquisition_price_analysis") or {}
    max_price_result = max_price_analysis.get("result") or {}
    checks: list[dict[str, Any]] = []

    binding_tokens = [
        str(run.get("run_id") or ""), str(run.get("spec_hash") or ""),
        str(run.get("model_version") or ""), str(run.get("evidence_binding_hash") or ""),
    ]
    numeric_tokens = [
        f"{float(result.get('purchase_price_wan') or 0):,.2f}",
        f"{float(result.get('total_acquisition_cost_wan') or 0):,.2f}",
    ]
    numeric_fields = [
        (indicators.get("project_irr_pct"), _pct),
        (indicators.get("equity_irr_pct"), _pct),
        (indicators.get("npv_wan"), _num),
        (indicators.get("minimum_dscr"), _num),
    ]
    if not is_solar:
        numeric_fields.append((indicators.get("minimum_tenant_rent_coverage"), _num))
    for value, formatter in numeric_fields:
        if value is not None:
            numeric_tokens.append(formatter(value))
    for token in binding_tokens + numeric_tokens:
        checks.append({
            "artifact": "markdown", "field": token, "passed": bool(token and token in markdown),
        })
        checks.append({
            "artifact": "docx", "field": token, "passed": bool(token and token in docx_text),
        })

    summary = _xlsx_summary_values(xlsx_path)
    string_expectations = {
        "资产收购财务模型": run.get("run_id"),
        "模型版本": run.get("model_version"),
        "Spec哈希": run.get("spec_hash"),
        "证据绑定哈希": run.get("evidence_binding_hash"),
        "证据绑定版本": run.get("evidence_binding_version"),
        "最高价验证状态": max_price_analysis.get("validation_status") or "not_run",
    }
    number_expectations = {
        "收购价格(万元)": result.get("purchase_price_wan"),
        "总收购成本(万元)": result.get("total_acquisition_cost_wan"),
        "项目IRR(%)": indicators.get("project_irr_pct"),
        "资本金IRR(%)": indicators.get("equity_irr_pct"),
        "NPV(万元)": indicators.get("npv_wan"),
        "最低DSCR": indicators.get("minimum_dscr"),
        "最低ICR": indicators.get("minimum_icr"),
        "最高可接受收购价(万元)": max_price_result.get("max_acquisition_price_wan"),
        "最高价目标IRR": (max_price_analysis.get("parameters") or {}).get("target_irr"),
        "最高价最低DSCR": (max_price_analysis.get("parameters") or {}).get("min_dscr"),
    }
    if is_solar:
        solar = result.get("solar_operation") or {}
        number_expectations.update({
            "装机容量(MW)": solar.get("installed_capacity_mw"),
            "基准发电量(MWh)": solar.get("base_generation_mwh"),
            "上网电价(元/kWh)": solar.get("tariff_yuan_per_kwh"),
            "限电率": solar.get("curtailment_rate"),
            "年衰减率": solar.get("degradation_rate"),
        })
    else:
        number_expectations.update({
            "最低租金覆盖率": indicators.get("minimum_tenant_rent_coverage"),
            "租约覆盖年限": indicators.get("lease_coverage_years"),
            "合同收入占比": indicators.get("contract_income_ratio"),
            "未锁定收入占比": indicators.get("unlocked_income_ratio"),
        })
    for label, expected in string_expectations.items():
        checks.append({
            "artifact": "xlsx", "field": label, "expected": expected,
            "actual": summary.get(label), "passed": summary.get(label) == str(expected or ""),
        })
    for label, expected in number_expectations.items():
        passed = summary.get(label, "") == "" if expected is None else _same_number(summary.get(label, ""), expected)
        checks.append({
            "artifact": "xlsx", "field": label, "expected": expected,
            "actual": summary.get(label), "passed": passed,
        })
    if report_data_path is not None:
        try:
            report_data = json.loads(report_data_path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError, TypeError):
            report_data = {}
        bindings = report_data.get("bindings") or {}
        for field in (
            "run_id", "spec_hash", "input_hash", "model_version",
            "spec_snapshot_hash", "evidence_binding_hash", "evidence_binding_version",
        ):
            checks.append({
                "artifact": "report_data", "field": field,
                "expected": run.get(field), "actual": bindings.get(field),
                "passed": bindings.get(field) == run.get(field),
            })
        expected_hash = report_data.get("report_data_hash")
        payload = {key: value for key, value in report_data.items() if key != "report_data_hash"}
        checks.append({
            "artifact": "report_data", "field": "report_data_hash",
            "expected": _hash(payload), "actual": expected_hash,
            "passed": expected_hash == _hash(payload),
        })
    return {"status": "passed" if all(row["passed"] for row in checks) else "failed", "checks": checks}


def _file_hash(path: Path) -> str:
    hasher = hashlib.sha256()
    with path.open("rb") as source:
        for chunk in iter(lambda: source.read(1024 * 1024), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


def _xlsx_col(index: int) -> str:
    out = ""
    while index:
        index, rem = divmod(index - 1, 26)
        out = chr(65 + rem) + out
    return out


def _xml_cell(row: int, column: int, value: Any) -> str:
    from xml.sax.saxutils import escape

    ref = f"{_xlsx_col(column)}{row}"
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return f'<c r="{ref}"><v>{value}</v></c>'
    return f'<c r="{ref}" t="inlineStr"><is><t>{escape(str(value or ""))}</t></is></c>'


def _sheet_xml(rows: list[list[Any]]) -> str:
    body = []
    for row_index, values in enumerate(rows, 1):
        cells = "".join(_xml_cell(row_index, col, value) for col, value in enumerate(values, 1))
        body.append(f'<row r="{row_index}">{cells}</row>')
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        f'<sheetData>{"".join(body)}</sheetData></worksheet>'
    )


def _write_minimal_xlsx(
    path: Path, run: dict[str, Any], *, report_data: dict[str, Any] | None = None,
) -> None:
    result = run.get("result") or {}
    indicators = result.get("indicators") or {}
    report_data = report_data or {}
    asset_type = str(result.get("asset_type") or report_data.get("asset_type") or "hotel_lease")
    is_solar = asset_type == "solar_power"
    max_price_analysis = report_data.get("maximum_acceptable_price") or {}
    max_price_result = max_price_analysis.get("result") or {}
    summary = [
        ["资产收购财务模型", run.get("run_id")],
        ["模型版本", run.get("model_version")],
        ["Spec哈希", run.get("spec_hash")],
        ["证据绑定哈希", run.get("evidence_binding_hash")],
        ["证据绑定版本", run.get("evidence_binding_version")],
        ["收购价格(万元)", result.get("purchase_price_wan")],
        ["总收购成本(万元)", result.get("total_acquisition_cost_wan")],
        ["项目IRR(%)", indicators.get("project_irr_pct")],
        ["资本金IRR(%)", indicators.get("equity_irr_pct")],
        ["NPV(万元)", indicators.get("npv_wan")],
        ["最低DSCR", indicators.get("minimum_dscr")],
        ["最低ICR", indicators.get("minimum_icr")],
        ["维修资本开支覆盖", indicators.get("maintenance_capex_coverage")],
        ["退出价值NPV占比", indicators.get("exit_value_npv_ratio")],
        ["最高可接受收购价(万元)", max_price_result.get("max_acquisition_price_wan")],
        ["最高价目标IRR", (max_price_analysis.get("parameters") or {}).get("target_irr")],
        ["最高价最低DSCR", (max_price_analysis.get("parameters") or {}).get("min_dscr")],
        ["最高价求解哈希", max_price_analysis.get("analysis_hash")],
        ["最高价验证状态", max_price_analysis.get("validation_status")],
    ]
    if is_solar:
        solar = result.get("solar_operation") or {}
        summary[13:13] = [
            ["装机容量(MW)", solar.get("installed_capacity_mw")],
            ["基准发电量(MWh)", solar.get("base_generation_mwh")],
            ["上网电价(元/kWh)", solar.get("tariff_yuan_per_kwh")],
            ["限电率", solar.get("curtailment_rate")],
            ["年衰减率", solar.get("degradation_rate")],
        ]
    else:
        summary[13:13] = [
            ["最低租金覆盖率", indicators.get("minimum_tenant_rent_coverage")],
            ["租约覆盖年限", indicators.get("lease_coverage_years")],
            ["合同收入占比", indicators.get("contract_income_ratio")],
            ["未锁定收入占比", indicators.get("unlocked_income_ratio")],
        ]
    project = result.get("project_cashflows_wan") or []
    equity = result.get("equity_cashflows_wan") or []
    cashflows = [["年度", "项目现金流(万元)", "资本金现金流(万元)"]]
    for index in range(max(len(project), len(equity))):
        cashflows.append([index, project[index] if index < len(project) else "", equity[index] if index < len(equity) else ""])
    def json_text(value: Any) -> str:
        if isinstance(value, (dict, list)):
            return json.dumps(value, ensure_ascii=False, sort_keys=True, default=str)
        return str(value or "")

    parties = [["主体ID", "名称", "角色", "状态", "证据"]] + [
        [row.get("entity_id"), row.get("name"), json_text(row.get("roles")), row.get("status"), json_text(row.get("evidence_ids"))]
        for row in report_data.get("party_relationships") or []
    ]
    assets = [["资产ID", "类型", "是否纳入", "面积(㎡)", "状态", "冲突", "裁决", "证据"]] + [
        [row.get("scope_id"), row.get("type"), row.get("included"), row.get("area_sqm"), row.get("status"),
         json_text(row.get("conflicts")), row.get("resolution"), json_text(row.get("evidence_ids"))]
        for row in report_data.get("asset_boundary") or []
    ]
    if is_solar:
        operations = [[
            "年度", "理论发电量(MWh)", "上网电量(MWh)", "上网电价(元/kWh)",
            "售电收入(万元)", "运维费(万元)", "维护性资本开支(万元)",
            "所得税(万元)", "债务服务(万元)", "项目现金流(万元)", "资本金现金流(万元)",
        ]] + [
            [
                row.get("year"), row.get("gross_generation_mwh"), row.get("sold_generation_mwh"),
                row.get("tariff_yuan_per_kwh"), row.get("revenue_wan"), row.get("operating_cost_wan"),
                row.get("maintenance_capex_wan"), row.get("income_tax_wan"), row.get("debt_service_wan"),
                row.get("project_cf_wan"), row.get("equity_cf_wan"),
            ]
            for row in report_data.get("solar_operating_ledger") or []
        ]
    else:
        leases = [["单元ID", "位置", "面积(㎡)", "出租人", "承租人", "起始日", "终止日", "基础租金(万元)", "证据"]] + [
            [row.get("unit_id"), row.get("asset_location"), row.get("area_sqm"), row.get("lessor_id"), row.get("lessee_id"),
             row.get("start_date"), row.get("end_date"), row.get("base_rent_wan"), json_text(row.get("evidence_ids"))]
            for row in report_data.get("lease_ledger") or []
        ]
    history = [["主体", "开始日", "结束日", "报表类型", "来源格式", "数值勾稽", "异常", "来源定位"]] + [
        [row.get("entity_id"), row.get("period_start"), row.get("period_end"), row.get("statement_type"), row.get("source_format"),
         json_text(row.get("reconciliation")), json_text(row.get("anomalies")), json_text(row.get("source_locators"))]
        for row in report_data.get("historical_financial_comparison") or []
    ]
    scenarios = ([
        ["矩阵ID", "场景ID", "变更", "收购价(万元)", "上网电价(元/kWh)", "年发电量(MWh)", "利用小时", "年运维费(万元)", "融资比例", "指标"]
    ] if is_solar else [
        ["矩阵ID", "场景ID", "变更", "收购价(万元)", "市场租金", "入住率", "融资比例", "指标"]
    ])
    for matrix in report_data.get("scenario_matrices") or []:
        for row in matrix.get("rows") or []:
            common = [
                matrix.get("matrix_id"), row.get("scenario_id"), json_text(row.get("changes")),
                row.get("purchase_price_wan"),
            ]
            scenarios.append(common + ([
                row.get("tariff_yuan_per_kwh"), row.get("annual_generation_mwh"),
                row.get("utilization_hours"), row.get("annual_opex_wan"),
                row.get("financing_ratio"), json_text(row.get("indicators")),
            ] if is_solar else [
                json_text(row.get("market_rent")), json_text(row.get("occupancy")),
                row.get("financing_ratio"), json_text(row.get("indicators")),
            ]))
    risks = [["类型", "编码/内容", "状态", "裁决/说明", "证据"]]
    risks.extend([
        ["red_flag", row.get("code"), row.get("status"), row.get("resolution"), json_text(row.get("evidence_ids"))]
        for row in report_data.get("red_flags") or []
    ])
    risks.extend([["closing_condition", item, "", "", ""] for item in report_data.get("closing_conditions") or []])
    risks.extend([["veto_item", item, "", "", ""] for item in report_data.get("veto_items") or []])
    evidence = [["字段", "证据ID", "来源", "绑定状态"]] + [
        [row.get("field"), json_text(row.get("evidence_ids")), json_text(row.get("source")), row.get("validation_status")]
        for row in report_data.get("source_processing_ledger") or []
    ]
    sheets: list[tuple[str, list[list[Any]]]] = [("收购摘要", summary), ("现金流", cashflows)]
    if report_data:
        sheets.extend([("主体关系", parties), ("资产边界", assets)])
        sheets.append(("光伏运营", operations) if is_solar else ("租约台账", leases))
        sheets.extend([
            ("历史财务", history), ("情景矩阵", scenarios),
            ("风险与条件", risks), ("证据台账", evidence),
        ])
    content_types = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
        + ''.join(
            f'<Override PartName="/xl/worksheets/sheet{index}.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
            for index in range(1, len(sheets) + 1)
        )
        + '</Types>'
    )
    from xml.sax.saxutils import escape

    workbook = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<sheets>'
        + ''.join(
            f'<sheet name="{escape(name)}" sheetId="{index}" r:id="rId{index}"/>'
            for index, (name, _rows) in enumerate(sheets, 1)
        )
        + '</sheets></workbook>'
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        + ''.join(
            f'<Relationship Id="rId{index}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet{index}.xml"/>'
            for index in range(1, len(sheets) + 1)
        )
        + '</Relationships>'
    )
    root_rels = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>'
        '</Relationships>'
    )
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", root_rels)
        archive.writestr("xl/workbook.xml", workbook)
        archive.writestr("xl/_rels/workbook.xml.rels", rels)
        for index, (_name, rows) in enumerate(sheets, 1):
            archive.writestr(f"xl/worksheets/sheet{index}.xml", _sheet_xml(rows))


def _num(value: Any) -> str:
    try:
        return f"{float(value):,.2f}"
    except (TypeError, ValueError):
        return "—"


def _pct(value: Any) -> str:
    try:
        return f"{float(value):.2f}%"
    except (TypeError, ValueError):
        return "—"


def _pct_ratio(value: Any) -> str:
    try:
        return f"{float(value) * 100:.2f}%"
    except (TypeError, ValueError):
        return "—"
