"""工件生成与一致性核对：入队、执行、绑定与产出。"""

from __future__ import annotations

import io
import json
import os
import re
import shutil
import uuid
from pathlib import Path
from typing import Any



from .base import (
    _ARTIFACT_GENERATION_FAILURE_MESSAGE,
    _LOG,
    _hash,
    _now,
    _num,
    _pct,
    _same_number,
)

from .report_data import (
    build_acquisition_report_data,
    render_markdown,
)

from .evidence import (
    _bind_spec_evidence,
)

from .runs import (
    get_run,
)

from .store import (
    _active_idempotency_record,
    _artifacts_root,
    _history_event,
    _idempotency_record,
    _load,
    _save,
    _state_guard,
)

from .xlsx import (
    _file_hash,
    _write_minimal_xlsx,
    _xlsx_summary_values,
)


_FORMAL_ARTIFACT_REQUIRED = "FORMAL_ARTIFACT_QUALIFICATION_REQUIRED"
_EVIDENCE_BINDING_STALE = "EVIDENCE_BINDING_STALE"


def _artifact_blocked(
    code: str,
    message: str,
    *,
    details: dict[str, Any] | None = None,
    next_actions: list[str] | None = None,
) -> dict[str, Any]:
    return {
        "ok": False,
        "error": code,
        "message": message,
        "details": dict(details or {}),
        "blockers": [code],
        "next_actions": list(next_actions or []),
    }


def _preflight_formal_artifact(
    workspace_id: str,
    run_id: str,
) -> tuple[dict[str, Any] | None, dict[str, Any] | None]:
    """Check whether a calculable run exists and collect non-blocking quality diagnostics."""

    run = get_run(workspace_id, run_id)
    if not run:
        return None, _artifact_blocked(
            "RUN_NOT_FOUND",
            "未找到资产收购运行",
            details={"run_id": run_id},
            next_actions=["读取当前工作区的收购运行并使用有效 run_id 重试"],
        )
    if run.get("status") != "succeeded" or not isinstance(run.get("result"), dict):
        return None, _artifact_blocked(
            "RUN_UNAVAILABLE",
            "资产收购运行尚未成功产生可读取结果",
            details={
                "run_id": run_id,
                "run_status": run.get("status"),
                "has_result": isinstance(run.get("result"), dict),
            },
            next_actions=["完成或重新执行该运行后再生成工件"],
        )

    quality_issues: list[dict[str, Any]] = []
    if run.get("consistency_ok") is not True:
        quality_issues.append({
            "code": "RUN_INCONSISTENT",
            "message": "运行一致性检查未通过；工件仍会生成并披露该限制。",
        })
    state = _load(workspace_id)
    spec_row = state["specs"].get(str(run.get("spec_id") or "")) or {}
    spec = spec_row.get("spec") or {}
    if not isinstance(spec, dict):
        return None, _artifact_blocked(
            "SPEC_SNAPSHOT_MISSING",
            "运行绑定的 Spec 快照实体缺失，无法构造工件",
            details={"run_id": run_id},
            next_actions=["恢复运行绑定的 Spec 快照后重试"],
        )
    if _hash(spec) != run.get("spec_hash"):
        quality_issues.append({
            "code": "SPEC_SNAPSHOT_MISMATCH",
            "message": "当前 Spec 快照哈希与运行绑定值不一致；工件仍基于运行结果生成。",
            "expected": run.get("spec_hash"),
            "actual": _hash(spec),
        })

    current_evidence = _bind_spec_evidence(workspace_id, spec)
    evidence_hash_matches = (
        current_evidence.get("binding_hash") == run.get("evidence_binding_hash")
    )
    evidence_version_matches = (
        current_evidence.get("binding_version") == run.get("evidence_binding_version")
    )
    if not evidence_hash_matches or not evidence_version_matches:
        quality_issues.append({
            "code": _EVIDENCE_BINDING_STALE,
            "message": "运行绑定的证据快照已变化；工件仍会生成并披露该限制。",
            "snapshot_binding_hash": run.get("evidence_binding_hash"),
            "current_binding_hash": current_evidence.get("binding_hash"),
            "snapshot_binding_version": run.get("evidence_binding_version"),
            "current_binding_version": current_evidence.get("binding_version"),
        })

    open_blockers = [
        str(issue.get("code") or "FORMAL_QUALIFICATION_BLOCKED")
        for issue in (run.get("issues") or [])
        if isinstance(issue, dict)
        and issue.get("blocking") is True
        and issue.get("status") == "open"
    ]
    qualification_failures: list[str] = []
    if run.get("delivery_mode") != "formal_candidate":
        qualification_failures.append("delivery_mode_not_formal_candidate")
    if run.get("validation_status") != "passed":
        qualification_failures.append("validation_not_passed")
    if run.get("formal_spec_valid") is not True:
        qualification_failures.append("formal_spec_invalid")
    if run.get("evidence_formal_ok") is not True or current_evidence.get("formal_ok") is not True:
        qualification_failures.append("formal_evidence_not_qualified")
    if open_blockers:
        qualification_failures.append("open_blocking_issues")
    if qualification_failures:
        quality_issues.append({
            "code": _FORMAL_ARTIFACT_REQUIRED,
            "message": "正式资格未满足；工件仍会生成并携带限制说明。",
            "delivery_mode": run.get("delivery_mode"),
            "qualification_failures": qualification_failures,
            "open_blockers": open_blockers,
        })
    run = dict(run)
    run["artifact_quality_issues"] = quality_issues
    return run, None


def enqueue_artifact(
    workspace_id: str, run_id: str, *, idempotency_key: str = "", request_id: str = "",
) -> dict[str, Any]:
    """Create a durable artifact job that can be polled before rendering."""

    request_id = request_id or f"req_{uuid.uuid4().hex}"
    run, blocked = _preflight_formal_artifact(workspace_id, run_id)
    if blocked:
        return blocked
    assert run is not None
    required_bindings = (
        "run_id", "spec_hash", "input_hash", "spec_snapshot_hash",
        "evidence_binding_hash", "model_version",
    )
    if any(not str(run.get(field) or "").strip() for field in required_bindings):
        return {
            "ok": False,
            "error": "RUN_BINDING_INCOMPLETE",
            "missing": [field for field in required_bindings if not str(run.get(field) or "").strip()],
        }
    body_hash = _hash({
        "run_id": run_id, "spec_hash": run.get("spec_hash"), "fact_revision": run.get("spec_id"),
        "spec_snapshot_hash": run.get("spec_snapshot_hash"),
        "evidence_binding_hash": run.get("evidence_binding_hash"),
        "template_version": "asset_acquisition.v2",
    })
    scope = f"artifact:{idempotency_key}" if idempotency_key else ""
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        prior = _active_idempotency_record(state["idempotency"], scope) if scope else None
        if prior:
            if prior.get("body_hash") != body_hash:
                return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior.get("artifact_id", "")}
            existing = state["artifacts"].get(prior.get("artifact_id"))
            if not existing:
                raise RuntimeError("artifact idempotency record points to a missing artifact")
            return {**existing, "idempotent_replay": True}
        artifact_id = f"artifact_{uuid.uuid4().hex}"
        job_id = f"artifact_job_{uuid.uuid4().hex}"
        created_at = _now()
        row = {
            "ok": True, "artifact_id": artifact_id, "artifact_job_id": job_id,
            "status": "queued", "progress": 0, "type": "asset_acquisition", "run_id": run_id,
            "spec_hash": run.get("spec_hash"), "fact_revision": run.get("spec_id"),
            "spec_snapshot_hash": run.get("spec_snapshot_hash"),
            "evidence_binding_version": run.get("evidence_binding_version"),
            "evidence_binding_hash": run.get("evidence_binding_hash"),
            "template_version": "asset_acquisition.v2", "created_at": created_at,
            "updated_at": created_at, "request_id": request_id, "files": [],
            "numeric_consistency": "pending", "integrity_status": "pending",
            "state_history": [
                _history_event("queued", request_id=request_id, run_id=run_id),
            ],
        }
        state["artifacts"][artifact_id] = row
        if scope:
            state["idempotency"][scope] = _idempotency_record(
                scope, body_hash, artifact_id=artifact_id,
            )
        _save(workspace_id, state)
        return row


def _bind_succeeded_artifact(
    workspace_id: str,
    run: dict[str, Any],
    artifact: dict[str, Any],
) -> dict[str, Any]:
    """Bind a successful formal pack to the report-side finance revision."""

    from lvke_mcp.domains.reports import artifacts as report_artifacts

    expected = {
        "finance_run_id": str(run.get("run_id") or ""),
        "input_hash": run.get("input_hash"),
        "spec_hash": run.get("spec_hash"),
        "spec_id": run.get("spec_id"),
        "fact_revision": run.get("spec_id"),
        "spec_snapshot_hash": run.get("spec_snapshot_hash"),
        "evidence_binding_version": run.get("evidence_binding_version"),
        "evidence_binding_hash": run.get("evidence_binding_hash"),
        "model_version": run.get("model_version"),
        "template_version": artifact.get("template_version"),
        "artifact_id": artifact.get("artifact_id"),
        "artifact_job_id": artifact.get("artifact_job_id"),
        "artifact_status": "succeeded",
        "report_data_hash": artifact.get("report_data_hash"),
        "binding_kind": "asset_acquisition",
    }
    report_artifacts.bind_finance_run(
        workspace_id,
        expected["finance_run_id"],
        section="asset_acquisition_artifact",
        fin={
            key: value for key, value in expected.items()
            if key != "finance_run_id"
        } | {"validation_level": "complete"},
    )
    actual = report_artifacts.load(
        workspace_id,
        "finance_binding",
        {},
    ) or {}
    mismatches = [
        {"field": key, "expected": value, "actual": actual.get(key)}
        for key, value in expected.items()
        if actual.get(key) != value
    ]
    if mismatches:
        return {
            "ok": False,
            "error": "ARTIFACT_BINDING_FAILED",
            "reason": "finance_binding_mismatch",
            "mismatches": mismatches,
        }
    return {"ok": True, "binding": actual}


def _mark_artifact_binding_failed(
    workspace_id: str,
    run_id: str,
    artifact_id: str,
    *,
    request_id: str,
    details: list[dict[str, Any]] | None = None,
) -> None:
    """Persist a failed terminal state when report-side binding does not commit."""

    with _state_guard(workspace_id):
        state = _load(workspace_id)
        stored = state["artifacts"].get(artifact_id)
        if stored:
            stored.update({
                "ok": False,
                "status": "failed",
                "integrity_status": "failed",
                "error": {
                    "code": "ARTIFACT_BINDING_FAILED",
                    "message": "formal artifact finance binding failed",
                    "retryable": False,
                    "details": list(details or []),
                },
                "updated_at": _now(),
            })
        stored_run = state["runs"].get(run_id)
        if stored_run:
            stored_run["lifecycle_status"] = "artifact_binding_failed"
            stored_run.setdefault("state_history", []).append(_history_event(
                "artifact_binding_failed",
                request_id=request_id,
                artifact_id=artifact_id,
            ))
        _save(workspace_id, state)


def execute_queued_artifact(
    workspace_id: str,
    artifact_id: str,
) -> None:
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        row = state["artifacts"].get(artifact_id)
        if not row or row.get("status") in {"succeeded", "failed", "cancelled"}:
            return
        row.update({"status": "running", "progress": 10, "updated_at": _now()})
        run_id = str(row.get("run_id") or "")
        job_id = str(row.get("artifact_job_id") or "")
        request_id = str(row.get("request_id") or "")
        _save(workspace_id, state)
    try:
        result = generate_artifacts(
            workspace_id, run_id, request_id=request_id,
            artifact_id=artifact_id, artifact_job_id=job_id,
        )
        if result.get("ok"):
            return
        raw_code = str(result.get("error") or "ARTIFACT_MISMATCH")
        code = (
            raw_code
            if re.fullmatch(r"[A-Z][A-Z0-9_]{1,79}", raw_code)
            else "ARTIFACT_MISMATCH"
        )
        error = {
            "code": code,
            "message": _ARTIFACT_GENERATION_FAILURE_MESSAGE,
            "retryable": False,
        }
    except Exception as exc:  # noqa: BLE001
        _LOG.warning(
            "asset acquisition artifact generation failed; artifact_id=%s "
            "error_type=%s",
            artifact_id,
            type(exc).__name__,
        )
        error = {
            "code": "ARTIFACT_GENERATION_FAILED",
            "message": _ARTIFACT_GENERATION_FAILURE_MESSAGE,
            "retryable": False,
        }
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        row = state["artifacts"].get(artifact_id)
        if row and row.get("status") != "cancelled":
            row.update({
                "ok": False, "status": "failed", "progress": 100,
                "numeric_consistency": "failed", "error": error, "updated_at": _now(),
            })
            _save(workspace_id, state)


def generate_artifacts(
    workspace_id: str, run_id: str, *, idempotency_key: str = "", request_id: str = "",
    artifact_id: str = "", artifact_job_id: str = "",
) -> dict[str, Any]:
    """Generate an atomically published, consistency-checked artifact pack."""

    run, blocked = _preflight_formal_artifact(workspace_id, run_id)
    if blocked:
        return blocked
    assert run is not None
    body_hash = _hash({
        "run_id": run_id, "spec_hash": run.get("spec_hash"), "fact_revision": run.get("spec_id"),
        "spec_snapshot_hash": run.get("spec_snapshot_hash"),
        "evidence_binding_hash": run.get("evidence_binding_hash"),
        "template_version": "asset_acquisition.v2",
    })
    scope = f"artifact:{idempotency_key}" if idempotency_key else ""
    if scope:
        with _state_guard(workspace_id):
            state = _load(workspace_id)
            prior = _active_idempotency_record(state["idempotency"], scope)
            if prior:
                if prior.get("body_hash") != body_hash:
                    return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior.get("artifact_id", "")}
                existing = state["artifacts"].get(prior.get("artifact_id"))
                if not existing:
                    raise RuntimeError("artifact idempotency record points to a missing artifact")
                return {**existing, "idempotent_replay": True}

    artifact_id = artifact_id or f"artifact_{uuid.uuid4().hex}"
    artifacts_root = _artifacts_root(workspace_id)
    created_artifacts_root = not artifacts_root.exists()
    artifacts_root.mkdir(parents=True, exist_ok=True)
    staging = artifacts_root / f".{artifact_id}.staging-{uuid.uuid4().hex}"
    final_root = artifacts_root / artifact_id
    committed = False
    try:
        staging.mkdir(parents=True, exist_ok=False)
        report_data = build_acquisition_report_data(
            workspace_id,
            run,
        )
        markdown = render_markdown(run, report_data)
        md_path = staging / "资产收购可行性研究报告.md"
        docx_path = staging / "资产收购可行性研究报告.docx"
        xlsx_path = staging / "资产收购财务模型.xlsx"
        report_data_path = staging / "资产收购报告数据.json"
        index_path = staging / "附件索引.json"
        md_path.write_text(markdown, encoding="utf-8")
        from lvke_mcp.domains.reports import doc_service as _doc_svc

        docx_path.write_bytes(_doc_svc.markdown_to_docx(markdown))
        report_data_path.write_text(json.dumps(report_data, ensure_ascii=False, indent=2), encoding="utf-8")
        _write_minimal_xlsx(xlsx_path, run, report_data=report_data)
        consistency = _check_artifact_consistency(
            run, markdown, docx_path, xlsx_path, report_data_path=report_data_path,
        )
        if consistency["status"] != "passed":
            run.setdefault("artifact_quality_issues", []).append({
                "code": "ARTIFACT_MISMATCH",
                "message": "工件数值或绑定一致性检查未通过；文件仍会发布并保留检查结果。",
                "consistency": consistency,
            })
        files = [md_path, docx_path, xlsx_path, report_data_path]
        index = {
            "artifact_id": artifact_id, "run_id": run_id,
            "spec_hash": run.get("spec_hash"), "fact_revision": run.get("spec_id"),
            "spec_snapshot_hash": run.get("spec_snapshot_hash"),
            "evidence_binding_version": run.get("evidence_binding_version"),
            "evidence_binding_hash": run.get("evidence_binding_hash"),
            "model_version": run.get("model_version"), "generated_at": _now(),
            "report_data_hash": report_data.get("report_data_hash"),
            "numeric_consistency": consistency,
            "quality_issues": run.get("artifact_quality_issues") or [],
            "files": [
                {"name": path.name, "size_bytes": path.stat().st_size, "sha256": _file_hash(path)}
                for path in files
            ],
        }
        index_path.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
        index["files"].append({"name": index_path.name, "size_bytes": index_path.stat().st_size, "sha256": _file_hash(index_path)})
        os.replace(staging, final_root)
        row = {
            "ok": True, "artifact_id": artifact_id,
            "artifact_job_id": artifact_job_id or f"artifact_job_{uuid.uuid4().hex}",
            "status": "succeeded", "progress": 100,
            "type": "asset_acquisition", "run_id": run_id,
            "spec_hash": run.get("spec_hash"), "fact_revision": run.get("spec_id"),
            "spec_snapshot_hash": run.get("spec_snapshot_hash"),
            "evidence_binding_version": run.get("evidence_binding_version"),
            "evidence_binding_hash": run.get("evidence_binding_hash"),
            "template_version": "asset_acquisition.v2", "created_at": _now(),
            "updated_at": _now(), "request_id": request_id,
            "files": index["files"], "directory": str(final_root),
            "report_data_hash": report_data.get("report_data_hash"),
            "numeric_consistency": consistency["status"],
            "consistency_checks": consistency["checks"],
            "quality_issues": run.get("artifact_quality_issues") or [],
            "state_history": [
                _history_event(
                    "artifact_generated",
                    request_id=request_id, run_id=run_id, artifact_id=artifact_id,
                ),
            ],
        }
        with _state_guard(workspace_id):
            state = _load(workspace_id)
            if scope and (prior := _active_idempotency_record(state["idempotency"], scope)):
                # Another worker completed the same request while this one rendered.
                if prior.get("body_hash") != body_hash:
                    shutil.rmtree(final_root, ignore_errors=True)
                    return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior.get("artifact_id", "")}
                shutil.rmtree(final_root, ignore_errors=True)
                return {**state["artifacts"][prior["artifact_id"]], "idempotent_replay": True}
            state["artifacts"][artifact_id] = row
            if scope:
                state["idempotency"][scope] = _idempotency_record(
                    scope, body_hash, artifact_id=artifact_id,
                )
            stored_run = state["runs"].get(run_id)
            if stored_run:
                stored_run["lifecycle_status"] = "artifact_generated"
                stored_run.setdefault("state_history", []).append(_history_event(
                    "artifact_generated", request_id=request_id,
                    artifact_id=artifact_id,
                ))
            _save(workspace_id, state)
        try:
            binding_result = _bind_succeeded_artifact(
                workspace_id,
                run,
                row,
            )
        except BaseException:
            try:
                _mark_artifact_binding_failed(
                    workspace_id,
                    run_id,
                    artifact_id,
                    request_id=request_id,
                )
            except Exception:  # noqa: BLE001 - preserve original binding error
                _LOG.exception(
                    "failed to persist artifact binding failure; artifact_id=%s",
                    artifact_id,
                )
            raise
        if not binding_result.get("ok"):
            mismatches = binding_result.get("mismatches") or []
            _mark_artifact_binding_failed(
                workspace_id,
                run_id,
                artifact_id,
                request_id=request_id,
                details=mismatches,
            )
            return {
                "ok": False, "error": "ARTIFACT_BINDING_FAILED",
                "reason": binding_result.get("reason") or "finance_binding_mismatch",
                "mismatches": mismatches,
            }

        committed = True
        return dict(_load(workspace_id)["artifacts"].get(artifact_id) or row)
    finally:
        if staging.exists():
            shutil.rmtree(staging, ignore_errors=True)
        if not committed and final_root.exists():
            shutil.rmtree(final_root, ignore_errors=True)
        if created_artifacts_root and artifacts_root.exists():
            try:
                artifacts_root.rmdir()
            except OSError:
                pass


def _check_artifact_consistency(
    run: dict[str, Any], markdown: str, docx_path: Path, xlsx_path: Path,
    *, report_data_path: Path | None = None,
) -> dict[str, Any]:
    from docx import Document  # type: ignore

    document = Document(io.BytesIO(docx_path.read_bytes()))
    docx_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    result = run.get("result") or {}
    indicators = result.get("indicators") or {}
    asset_type = str(result.get("asset_type") or "hotel_lease")
    is_solar = asset_type == "solar_power"
    max_price_analysis = run.get("max_acquisition_price_analysis") or {}
    max_price_result = max_price_analysis.get("result") or {}
    checks: list[dict[str, Any]] = []

    binding_tokens = [
        str(run.get("run_id") or ""), str(run.get("spec_hash") or ""),
        str(run.get("model_version") or ""), str(run.get("evidence_binding_hash") or ""),
    ]
    numeric_tokens = [
        f"{float(result.get('purchase_price_wan') or 0):,.2f}",
        f"{float(result.get('total_acquisition_cost_wan') or 0):,.2f}",
    ]
    numeric_fields = [
        (indicators.get("project_irr_pct"), _pct),
        (indicators.get("equity_irr_pct"), _pct),
        (indicators.get("npv_wan"), _num),
        (indicators.get("minimum_dscr"), _num),
    ]
    if not is_solar:
        numeric_fields.append((indicators.get("minimum_tenant_rent_coverage"), _num))
    for value, formatter in numeric_fields:
        if value is not None:
            numeric_tokens.append(formatter(value))
    for token in binding_tokens + numeric_tokens:
        checks.append({
            "artifact": "markdown", "field": token, "passed": bool(token and token in markdown),
        })
        checks.append({
            "artifact": "docx", "field": token, "passed": bool(token and token in docx_text),
        })

    summary = _xlsx_summary_values(xlsx_path)
    string_expectations = {
        "资产收购财务模型": run.get("run_id"),
        "模型版本": run.get("model_version"),
        "Spec哈希": run.get("spec_hash"),
        "证据绑定哈希": run.get("evidence_binding_hash"),
        "证据绑定版本": run.get("evidence_binding_version"),
        "最高价验证状态": max_price_analysis.get("validation_status") or "not_run",
    }
    number_expectations = {
        "收购价格(万元)": result.get("purchase_price_wan"),
        "总收购成本(万元)": result.get("total_acquisition_cost_wan"),
        "项目IRR(%)": indicators.get("project_irr_pct"),
        "资本金IRR(%)": indicators.get("equity_irr_pct"),
        "NPV(万元)": indicators.get("npv_wan"),
        "最低DSCR": indicators.get("minimum_dscr"),
        "最低ICR": indicators.get("minimum_icr"),
        "最高可接受收购价(万元)": max_price_result.get("max_acquisition_price_wan"),
        "最高价目标IRR": (max_price_analysis.get("parameters") or {}).get("target_irr"),
        "最高价最低DSCR": (max_price_analysis.get("parameters") or {}).get("min_dscr"),
    }
    if is_solar:
        solar = result.get("solar_operation") or {}
        number_expectations.update({
            "装机容量(MW)": solar.get("installed_capacity_mw"),
            "基准发电量(MWh)": solar.get("base_generation_mwh"),
            "上网电价(元/kWh)": solar.get("tariff_yuan_per_kwh"),
            "限电率": solar.get("curtailment_rate"),
            "年衰减率": solar.get("degradation_rate"),
        })
    else:
        number_expectations.update({
            "最低租金覆盖率": indicators.get("minimum_tenant_rent_coverage"),
            "租约覆盖年限": indicators.get("lease_coverage_years"),
            "合同收入占比": indicators.get("contract_income_ratio"),
            "未锁定收入占比": indicators.get("unlocked_income_ratio"),
        })
    for label, expected in string_expectations.items():
        checks.append({
            "artifact": "xlsx", "field": label, "expected": expected,
            "actual": summary.get(label), "passed": summary.get(label) == str(expected or ""),
        })
    for label, expected in number_expectations.items():
        passed = summary.get(label, "") == "" if expected is None else _same_number(summary.get(label, ""), expected)
        checks.append({
            "artifact": "xlsx", "field": label, "expected": expected,
            "actual": summary.get(label), "passed": passed,
        })
    if report_data_path is not None:
        try:
            report_data = json.loads(report_data_path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError, TypeError):
            report_data = {}
        bindings = report_data.get("bindings") or {}
        for field in (
            "run_id", "spec_hash", "input_hash", "model_version",
            "spec_snapshot_hash", "evidence_binding_hash", "evidence_binding_version",
        ):
            checks.append({
                "artifact": "report_data", "field": field,
                "expected": run.get(field), "actual": bindings.get(field),
                "passed": bindings.get(field) == run.get(field),
            })
        expected_hash = report_data.get("report_data_hash")
        payload = {key: value for key, value in report_data.items() if key != "report_data_hash"}
        checks.append({
            "artifact": "report_data", "field": "report_data_hash",
            "expected": _hash(payload), "actual": expected_hash,
            "passed": expected_hash == _hash(payload),
        })
    return {"status": "passed" if all(row["passed"] for row in checks) else "failed", "checks": checks}
