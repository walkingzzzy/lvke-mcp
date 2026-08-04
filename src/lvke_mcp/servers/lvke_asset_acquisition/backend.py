"""Persistence and artifact orchestration for FinanceSpec v3 acquisition runs."""

from __future__ import annotations

import copy
import hashlib
import io
import json
import logging
import math
import mimetypes
import os
import re
import shutil
import threading
import uuid
import zipfile
from concurrent.futures import ThreadPoolExecutor
from contextlib import contextmanager
from collections.abc import Mapping
from datetime import datetime, timezone
from itertools import product
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from filelock import FileLock

from lvke_mcp.runtime.workspace import workspace_root
from lvke_mcp.servers.lvke_asset_acquisition.model import (
    INDEPENDENT_SCENARIO_FIELDS,
    AcquisitionModelError,
    apply_scenario,
    run_acquisition_model,
    solve_max_acquisition_price,
)
from lvke_mcp.servers.lvke_asset_acquisition.spec import (
    LATEST_SPEC_VERSION,
    mark_spec_confirmed,
    validate,
    validate_for_formal,
)

_LOCK = threading.RLock()
_LOG = logging.getLogger(__name__)
_RECOVERY_POOL = ThreadPoolExecutor(max_workers=2, thread_name_prefix="acquisition-recovery")
_DEFAULT_IDEMPOTENCY_TTL_SECONDS = 24 * 60 * 60
_SAFE_PROFESSIONAL_SIGNOFF_ID = re.compile(r"^signoff_[0-9a-f]{32}$")
_SHA256_EXTERNAL_EVIDENCE_HASH = re.compile(r"^sha256:[0-9a-f]{64}$")
_SAFE_SIGNOFF_OPERATION_ID = re.compile(r"^[A-Za-z0-9_.:-]{1,200}$")
_UNTRUSTED_EVIDENCE_ASSERTION_KEYS = {
    "binding_hash",
    "source_sha256",
    "source_size_bytes",
    "parse_job",
    "parse_job_id",
    "attempt",
    "review_revision",
    "integrity_status",
}
_SOURCE_EVIDENCE_FAILURE_MESSAGE = "资料证据状态当前不可验证"
_RUN_VALIDATION_FAILURE_MESSAGE = "资产收购财务输入未通过模型校验"
_RUN_EXECUTION_FAILURE_MESSAGE = "资产收购财务测算执行失败"
_ARTIFACT_GENERATION_FAILURE_MESSAGE = "资产收购正式工件生成失败"


def _is_estimate_preview_spec(spec: Mapping[str, Any]) -> bool:
    if str(spec.get("delivery_mode") or "") != "estimate_preview":
        return False
    assumptions = spec.get("controlled_assumptions")
    if not isinstance(assumptions, list) or not assumptions:
        return False
    required = {
        "field", "value", "unit", "basis", "impact", "sensitivity",
        "release_condition",
    }
    return all(
        isinstance(item, dict)
        and required <= set(item)
        and all(item.get(key) not in (None, "") for key in required)
        for item in assumptions
    )


class AcquisitionSignoffError(RuntimeError):
    """Machine-readable failure for external professional-signoff evidence."""

    def __init__(
        self,
        code: str,
        message: str,
        *,
        details: dict[str, Any] | None = None,
    ) -> None:
        super().__init__(message)
        self.code = code
        self.message = message
        self.details = details or {}


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _idempotency_ttl_seconds() -> int:
    try:
        value = int(os.environ.get("LVKE_MCP_IDEMPOTENCY_TTL_SECONDS", ""))
    except (TypeError, ValueError):
        value = _DEFAULT_IDEMPOTENCY_TTL_SECONDS
    return max(1, value)


def _root(
    workspace_id: str,
) -> Path:
    return workspace_root(workspace_id) / "finance_acquisition"


def _state_path(
    workspace_id: str,
) -> Path:
    return _root(workspace_id) / "state.json"


@contextmanager
def _state_guard(
    workspace_id: str,
):
    lock_path = _root(workspace_id) / "state.lock"
    lock_path.parent.mkdir(parents=True, exist_ok=True)
    with _LOCK, FileLock(str(lock_path), timeout=30):
        yield


def _load(
    workspace_id: str,
) -> dict[str, Any]:
    try:
        raw = json.loads(
            _state_path(workspace_id).read_text(
                encoding="utf-8"
            )
        )
    except FileNotFoundError:
        raw = {}
    except (json.JSONDecodeError, TypeError) as exc:
        raise RuntimeError(f"acquisition state is corrupt for workspace {workspace_id}") from exc
    return {
        "version": "acquisition_store.v2",
        "specs": dict(raw.get("specs") or {}),
        "runs": dict(raw.get("runs") or {}),
        "idempotency": dict(raw.get("idempotency") or {}),
        "artifacts": dict(raw.get("artifacts") or {}),
        "scenario_matrices": dict(raw.get("scenario_matrices") or {}),
    }


def _save(
    workspace_id: str,
    state: dict[str, Any],
) -> None:
    path = _state_path(workspace_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_name(f"state.{uuid.uuid4().hex}.tmp")
    try:
        with tmp.open("w", encoding="utf-8", newline="\n") as target:
            json.dump(state, target, ensure_ascii=False, indent=2, default=str)
            target.flush()
            os.fsync(target.fileno())
        os.replace(tmp, path)
    finally:
        tmp.unlink(missing_ok=True)


def _hash(value: Any) -> str:
    payload = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"), default=str)
    return "sha256:" + hashlib.sha256(payload.encode("utf-8")).hexdigest()


def _required_signoff_text(value: Any, *, field: str, maximum: int) -> str:
    text = str(value or "").strip()
    if (
        not text
        or len(text) > maximum
        or any(ord(char) < 32 for char in text)
    ):
        raise AcquisitionSignoffError(
            "INVALID_SIGNOFF_EVIDENCE", f"{field} 不合法",
        )
    return text


def _normalize_signoff_evidence_hash(value: Any) -> str:
    normalized = str(value or "").strip().lower()
    if not _SHA256_EXTERNAL_EVIDENCE_HASH.fullmatch(normalized):
        raise AcquisitionSignoffError(
            "INVALID_SIGNOFF_EVIDENCE",
            "evidence_hash 必须是 sha256:<64位十六进制> 格式",
        )
    return normalized


def _normalize_signoff_time(value: Any, *, field: str) -> str:
    raw = _required_signoff_text(value, field=field, maximum=80)
    try:
        parsed = datetime.fromisoformat(raw.replace("Z", "+00:00"))
    except ValueError as exc:
        raise AcquisitionSignoffError(
            "INVALID_SIGNOFF_EVIDENCE", f"{field} 必须是 ISO 8601 时间",
        ) from exc
    if parsed.tzinfo is None or parsed.utcoffset() is None:
        raise AcquisitionSignoffError(
            "INVALID_SIGNOFF_EVIDENCE", f"{field} 必须包含时区",
        )
    return parsed.astimezone(timezone.utc).isoformat()


def _validated_signoff_actor(actor: Mapping[str, Any]) -> dict[str, Any]:
    from lvke_mcp.domains.reports.professional_review import (
        ProfessionalReviewError,
        validate_authenticated_actor,
    )

    try:
        validated = validate_authenticated_actor(actor)
    except ProfessionalReviewError as exc:
        raise AcquisitionSignoffError(
            exc.code, exc.message, details=exc.details,
        ) from exc
    return validated


def _stored_signoff_actor_is_valid(actor: Any) -> bool:
    if not isinstance(actor, Mapping):
        return False
    from lvke_mcp.domains.reports.professional_review import (
        ProfessionalReviewError,
        validate_authenticated_actor,
    )

    try:
        normalized = validate_authenticated_actor(actor)
    except ProfessionalReviewError:
        return False
    return normalized == dict(actor)


def _asset_release_integrity_material(
    release: Mapping[str, Any],
) -> dict[str, Any]:
    return {
        str(key): copy.deepcopy(value)
        for key, value in release.items()
        if str(key) not in {"release_hash", "professional_signoff"}
    }


def _asset_signoff_registration_integrity_material(
    signoff: Mapping[str, Any],
) -> dict[str, Any]:
    fields = (
        "schema_version",
        "signoff_id",
        "artifact_id",
        "release_id",
        "run_id",
        "spec_hash",
        "fact_revision",
        "spec_snapshot_hash",
        "evidence_binding_hash",
        "report_data_hash",
        "artifact_files_hash",
        "internal_release_hash",
        "evidence_hash",
        "evidence_reference",
        "signer",
        "signer_credential",
        "signed_at",
        "note",
        "registration_hash",
        "registration_operation_id",
        "registered_by",
        "registered_at",
        "manual_external_evidence",
        "evidence_hash_source",
        "system_signature_performed",
    )
    return {field: copy.deepcopy(signoff.get(field)) for field in fields}


def _asset_signoff_verification_integrity_material(
    signoff: Mapping[str, Any],
) -> dict[str, Any]:
    fields = (
        "signoff_id",
        "artifact_id",
        "release_id",
        "run_id",
        "evidence_hash",
        "registration_integrity_hash",
        "verified",
        "verified_by",
        "verified_at",
        "verification_reference",
        "verification_note",
        "verification_hash",
        "verification_operation_id",
        "verification_kind",
        "evidence_hash_verification",
        "system_signature_performed",
    )
    return {field: copy.deepcopy(signoff.get(field)) for field in fields}


def _asset_signoff_integrity_failures(
    artifact: Mapping[str, Any],
) -> list[dict[str, Any]]:
    """Recompute asset signoff envelopes; never infer a real signature."""

    release = artifact.get("release")
    if not isinstance(release, Mapping):
        return []
    failures: list[dict[str, Any]] = []

    def fail(reason: str, **details: Any) -> None:
        failures.append({"reason": reason, **details})

    actual_release_hash = _hash(_asset_release_integrity_material(release))
    if release.get("release_hash") != actual_release_hash:
        fail("internal_release_hash_mismatch")
    if (
        str(release.get("release_id") or "") == ""
        or release.get("artifact_id") != artifact.get("artifact_id")
        or release.get("run_id") != artifact.get("run_id")
        or not str(release.get("actor") or "").strip()
    ):
        fail("invalid_internal_release_envelope")

    signoff = release.get("professional_signoff")
    if not isinstance(signoff, Mapping) or not signoff:
        return failures
    if signoff.get("schema_version") != "asset_external_professional_signoff.v1":
        fail("unsupported_signoff_schema")
    signoff_id = str(signoff.get("signoff_id") or "")
    if not _SAFE_PROFESSIONAL_SIGNOFF_ID.fullmatch(signoff_id):
        fail("invalid_signoff_id")
    expected_bindings = {
        "artifact_id": artifact.get("artifact_id"),
        "release_id": release.get("release_id"),
        "run_id": artifact.get("run_id"),
        "spec_hash": artifact.get("spec_hash"),
        "fact_revision": artifact.get("fact_revision"),
        "spec_snapshot_hash": artifact.get("spec_snapshot_hash"),
        "evidence_binding_hash": artifact.get("evidence_binding_hash"),
        "report_data_hash": artifact.get("report_data_hash"),
        "artifact_files_hash": _hash(artifact.get("files") or []),
        "internal_release_hash": actual_release_hash,
    }
    for field, expected in expected_bindings.items():
        if signoff.get(field) != expected:
            fail(
                "signoff_binding_mismatch",
                field=field,
                expected=expected,
                actual=signoff.get(field),
            )
    if (
        signoff.get("manual_external_evidence") is not True
        or signoff.get("evidence_hash_source")
        != "human_supplied_external_evidence_digest"
        or signoff.get("system_signature_performed") is not False
    ):
        fail("signoff_boundary_flags_invalid")
    try:
        evidence_hash = _normalize_signoff_evidence_hash(signoff.get("evidence_hash"))
        signed_at = _normalize_signoff_time(signoff.get("signed_at"), field="signed_at")
        evidence_reference = _required_signoff_text(
            signoff.get("evidence_reference"),
            field="evidence_reference",
            maximum=1000,
        )
        signer = _required_signoff_text(
            signoff.get("signer"), field="signer", maximum=300,
        )
        signer_credential = _required_signoff_text(
            signoff.get("signer_credential"),
            field="signer_credential",
            maximum=500,
        )
    except AcquisitionSignoffError:
        fail("registration_fields_invalid")
        evidence_hash = str(signoff.get("evidence_hash") or "")
        signed_at = str(signoff.get("signed_at") or "")
        evidence_reference = str(signoff.get("evidence_reference") or "")
        signer = str(signoff.get("signer") or "")
        signer_credential = str(signoff.get("signer_credential") or "")
    if (
        evidence_hash != signoff.get("evidence_hash")
        or signed_at != signoff.get("signed_at")
        or evidence_reference != signoff.get("evidence_reference")
        or signer != signoff.get("signer")
        or signer_credential != signoff.get("signer_credential")
    ):
        fail("registration_fields_not_normalized")
    registered_by = signoff.get("registered_by")
    if not _stored_signoff_actor_is_valid(registered_by):
        fail("registrar_identity_invalid")
    registration_body = {
        "artifact_id": signoff.get("artifact_id"),
        "evidence_hash": signoff.get("evidence_hash"),
        "evidence_reference": signoff.get("evidence_reference"),
        "signer": signoff.get("signer"),
        "signer_credential": signoff.get("signer_credential"),
        "signed_at": signoff.get("signed_at"),
        "note": signoff.get("note"),
    }
    if signoff.get("registration_hash") != _hash(registration_body):
        fail("registration_hash_mismatch")
    if signoff.get("registration_integrity_hash") != _hash(
        _asset_signoff_registration_integrity_material(signoff)
    ):
        fail("registration_integrity_hash_mismatch")

    verified = signoff.get("verified")
    verification_fields = (
        "verified_by",
        "verified_at",
        "verification_reference",
        "verification_hash",
        "verification_integrity_hash",
    )
    if verified is False:
        if any(signoff.get(field) not in (None, "") for field in verification_fields):
            fail("unverified_signoff_contains_verification_state")
    elif verified is True:
        verified_by = signoff.get("verified_by")
        if not _stored_signoff_actor_is_valid(verified_by):
            fail("verifier_identity_invalid")
            verifier_id = ""
        else:
            verifier_id = str(verified_by.get("actor_id") or "")
        registrar_id = str(
            (registered_by.get("actor_id") or "")
            if isinstance(registered_by, Mapping)
            else ""
        )
        release_actor_id = str(release.get("actor") or "")
        if not verifier_id or verifier_id in {registrar_id, release_actor_id}:
            fail("verification_duty_separation_invalid")
        try:
            verified_at = _normalize_signoff_time(
                signoff.get("verified_at"), field="verified_at",
            )
            verification_reference = _required_signoff_text(
                signoff.get("verification_reference"),
                field="verification_reference",
                maximum=1000,
            )
        except AcquisitionSignoffError:
            fail("verification_fields_invalid")
            verified_at = str(signoff.get("verified_at") or "")
            verification_reference = str(
                signoff.get("verification_reference") or ""
            )
        if (
            verified_at != signoff.get("verified_at")
            or verification_reference != signoff.get("verification_reference")
        ):
            fail("verification_fields_not_normalized")
        verification_body = {
            "signoff_id": signoff.get("signoff_id"),
            "evidence_hash": signoff.get("evidence_hash"),
            "verification_reference": signoff.get("verification_reference"),
            "note": signoff.get("verification_note"),
            "verified_by": verifier_id,
        }
        if signoff.get("verification_hash") != _hash(verification_body):
            fail("verification_hash_mismatch")
        if (
            signoff.get("verification_kind")
            != "independent_human_evidence_check"
            or signoff.get("evidence_hash_verification")
            != "independent_human_digest_match"
            or signoff.get("system_signature_performed") is not False
        ):
            fail("verification_boundary_flags_invalid")
        if signoff.get("verification_integrity_hash") != _hash(
            _asset_signoff_verification_integrity_material(signoff)
        ):
            fail("verification_integrity_hash_mismatch")
    else:
        fail("verified_flag_invalid")

    matching_releases = [
        item
        for item in (artifact.get("release_history") or [])
        if isinstance(item, Mapping)
        and str(item.get("release_id") or "") == str(release.get("release_id") or "")
    ]
    if len(matching_releases) != 1:
        fail("release_history_binding_missing")
    elif matching_releases[0].get("professional_signoff") != dict(signoff):
        fail("release_history_signoff_mismatch")
    signoff_history = [
        item
        for item in (artifact.get("professional_signoff_history") or [])
        if isinstance(item, Mapping)
        and isinstance(item.get("signoff"), Mapping)
        and str((item.get("signoff") or {}).get("signoff_id") or "") == signoff_id
    ]
    expected_event = "verified" if verified is True else "registered"
    if not signoff_history:
        fail("professional_signoff_history_missing")
    elif (
        signoff_history[-1].get("event") != expected_event
        or signoff_history[-1].get("signoff") != dict(signoff)
    ):
        fail("professional_signoff_history_mismatch")
    return failures


def asset_external_professional_signoff_verified(
    artifact: Mapping[str, Any],
) -> bool:
    release = artifact.get("release")
    signoff = (
        release.get("professional_signoff")
        if isinstance(release, Mapping)
        else None
    )
    return bool(
        isinstance(signoff, Mapping)
        and signoff.get("verified") is True
        and not _asset_signoff_integrity_failures(
            artifact,
        )
    )


def asset_external_professional_signoff_registered(
    artifact: Mapping[str, Any],
) -> bool:
    release = artifact.get("release")
    signoff = (
        release.get("professional_signoff")
        if isinstance(release, Mapping)
        else None
    )
    return bool(
        isinstance(signoff, Mapping)
        and signoff.get("signoff_id")
        and not _asset_signoff_integrity_failures(
            artifact,
        )
    )


def _idempotency_record(scope: str, body_hash: str, **resource: Any) -> dict[str, Any]:
    created = datetime.now(timezone.utc)
    expires = datetime.fromtimestamp(
        created.timestamp() + _idempotency_ttl_seconds(), tz=timezone.utc,
    )
    return {
        "scope": scope,
        "body_hash": body_hash,
        "request_body_hash": body_hash,
        "created_at": created.isoformat(),
        "expires_at": expires.isoformat(),
        **resource,
    }


def _active_idempotency_record(
    records: dict[str, Any], scope: str,
) -> dict[str, Any] | None:
    record = records.get(scope)
    if not isinstance(record, dict):
        return None
    expires_at = str(record.get("expires_at") or "")
    if not expires_at:
        return record
    try:
        expires = datetime.fromisoformat(expires_at.replace("Z", "+00:00"))
    except ValueError:
        records.pop(scope, None)
        return None
    if expires.tzinfo is None:
        expires = expires.replace(tzinfo=timezone.utc)
    if expires <= datetime.now(timezone.utc):
        records.pop(scope, None)
        return None
    return record


def _history_event(status: str, *, actor: str = "system", request_id: str = "", **details: Any) -> dict[str, Any]:
    return {
        "status": status, "at": _now(), "actor": actor,
        "request_id": request_id, "details": details,
    }


def _close_issue(run: dict[str, Any], code: str, *, actor: str, reason: str) -> None:
    for issue in run.get("issues") or []:
        if issue.get("code") == code and issue.get("status") == "open":
            issue.update({"status": "closed", "closed_at": _now(), "closed_by": actor, "resolution": reason})


def _open_issue(run: dict[str, Any], code: str, detail: str) -> None:
    existing = next(
        (issue for issue in run.get("issues") or [] if issue.get("code") == code and issue.get("status") == "open"),
        None,
    )
    if existing:
        existing.update({"detail": detail, "updated_at": _now()})
        return
    run.setdefault("issues", []).append({
        "code": code, "blocking": True, "status": "open", "detail": detail, "created_at": _now(),
    })


def _migration_binding(spec: dict[str, Any]) -> dict[str, Any]:
    trace = copy.deepcopy(spec.get("migration_trace") or {})
    if not isinstance(trace, dict):
        trace = {}
    return {
        "source_spec_version": trace.get("source_spec_version") or spec.get("version"),
        "migration_trace": trace,
        "migration_steps": copy.deepcopy(trace.get("steps") or []),
    }


def _sanitize_client_evidence_claims(value: Any) -> Any:
    """Remove client-authored integrity facts before hashing or persistence.

    Evidence identifiers and locators remain assertions that the server resolves.
    Pre-built bindings, source hashes, parse attempts and review revisions are
    authority data and are always reconstructed from the workspace source state.
    """

    if isinstance(value, dict):
        cleaned: dict[str, Any] = {}
        for key, item in value.items():
            if key == "evidence_bindings" or key in _UNTRUSTED_EVIDENCE_ASSERTION_KEYS:
                continue
            cleaned[str(key)] = _sanitize_client_evidence_claims(item)
        return cleaned
    if isinstance(value, list):
        return [_sanitize_client_evidence_claims(item) for item in value]
    if isinstance(value, tuple):
        return [_sanitize_client_evidence_claims(item) for item in value]
    return copy.deepcopy(value)


def _bind_spec_evidence(
    workspace_id: str,
    spec: dict[str, Any],
) -> dict[str, Any]:
    from lvke_mcp.domains.finance.evidence_binding import bind_finance_spec_evidence

    return bind_finance_spec_evidence(
        workspace_id,
        spec,
    )


def _evidence_error_strings(binding: dict[str, Any]) -> list[str]:
    rows = [
        *(binding.get("invalid") or []),
        *(binding.get("missing") or []),
        *(binding.get("pending") or []),
    ]
    return [
        "finance_spec.v3 证据绑定未通过: "
        f"{row.get('source_path') or '$'} [{row.get('code') or 'EVIDENCE_BINDING_FAILED'}] "
        f"{row.get('message') or ''}".strip()
        for row in rows
    ]


def _formal_assessment(
    workspace_id: str,
    spec: dict[str, Any],
    *,
    evidence_binding: dict[str, Any] | None = None,
) -> tuple[bool, list[str], dict[str, Any]]:
    schema_ok, schema_errors = validate_for_formal(spec)
    binding = copy.deepcopy(evidence_binding) if evidence_binding is not None else _bind_spec_evidence(
        workspace_id,
        spec,
    )
    errors = [*schema_errors, *_evidence_error_strings(binding)]
    return bool(schema_ok and binding.get("formal_ok") and not errors), errors, binding


def _evidence_blocking_issues(
    binding: dict[str, Any],
    *,
    created_at: str,
) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    for category in ("invalid", "missing", "pending"):
        for item in binding.get(category) or []:
            issues.append({
                "code": str(item.get("code") or "EVIDENCE_BINDING_FAILED"),
                "blocking": True,
                "status": "open",
                "detail": json.dumps(item, ensure_ascii=False, sort_keys=True),
                "source_path": item.get("source_path") or "$",
                "created_at": created_at,
            })
    if not issues and not binding.get("formal_ok"):
        issues.append({
            "code": "EVIDENCE_BINDING_FAILED",
            "blocking": True,
            "status": "open",
            "detail": "FinanceSpec 未形成可发布的服务端证据绑定",
            "created_at": created_at,
        })
    return issues


def _current_evidence_matches_run(
    workspace_id: str,
    run: dict[str, Any],
    spec: dict[str, Any],
) -> tuple[bool, dict[str, Any]]:
    current = _bind_spec_evidence(
        workspace_id,
        spec,
    )
    return bool(
        current.get("formal_ok")
        and current.get("binding_hash") == run.get("evidence_binding_hash")
        and current.get("binding_version") == run.get("evidence_binding_version")
    ), current


def sanitize_spec_input(spec: dict[str, Any]) -> dict[str, Any]:
    """Public contract helper returning the server-trusted spec assertion body."""

    cleaned = _sanitize_client_evidence_claims(spec)
    return cleaned if isinstance(cleaned, dict) else {}


def assess_spec_evidence(
    workspace_id: str,
    spec: dict[str, Any],
) -> dict[str, Any]:
    """Resolve a spec's evidence using only workspace-owned server state."""

    return _bind_spec_evidence(
        workspace_id,
        sanitize_spec_input(spec),
    )


def _decision_thresholds(spec: dict[str, Any]) -> tuple[dict[str, float | None], str]:
    raw = spec.get("decision_thresholds") or {}
    if not isinstance(raw, dict):
        return {}, "DECISION_THRESHOLDS_INVALID"
    try:
        target = float(raw["target_project_irr"])
        minimum_dscr = (
            float(raw["minimum_dscr"])
            if raw.get("minimum_dscr") is not None else None
        )
    except (KeyError, TypeError, ValueError):
        return {}, "DECISION_THRESHOLDS_INVALID"
    if not math.isfinite(target) or not 0 < target <= 5:
        return {}, "DECISION_THRESHOLDS_INVALID"
    if minimum_dscr is not None and (
        not math.isfinite(minimum_dscr) or minimum_dscr <= 0
    ):
        return {}, "DECISION_THRESHOLDS_INVALID"
    return {
        "target_project_irr": target,
        "minimum_dscr": minimum_dscr,
    }, ""


def _same_optional_number(left: Any, right: Any) -> bool:
    if left is None or right is None:
        return left is None and right is None
    try:
        return math.isclose(float(left), float(right), rel_tol=1e-12, abs_tol=1e-12)
    except (TypeError, ValueError):
        return False


def _max_price_gate(
    run: dict[str, Any], spec: dict[str, Any], *, require_business_decision: bool,
) -> tuple[bool, str, dict[str, Any]]:
    thresholds, threshold_error = _decision_thresholds(spec)
    if threshold_error:
        return False, threshold_error, {"decision_thresholds": spec.get("decision_thresholds")}
    analysis = run.get("max_acquisition_price_analysis") or {}
    if not isinstance(analysis, dict) or not analysis:
        return False, "MAX_PRICE_REQUIRED", {"decision_thresholds": thresholds}
    parameters = analysis.get("parameters") or {}
    if (
        not _same_optional_number(
            parameters.get("target_irr"), thresholds["target_project_irr"],
        )
        or not _same_optional_number(
            parameters.get("min_dscr"), thresholds["minimum_dscr"],
        )
    ):
        return False, "MAX_PRICE_THRESHOLD_MISMATCH", {
            "decision_thresholds": thresholds, "parameters": parameters,
        }
    result = analysis.get("result") or {}
    if not result.get("feasible") or not result.get("converged"):
        return False, "MAX_PRICE_NOT_FEASIBLE", {
            "decision_thresholds": thresholds,
            "feasible": bool(result.get("feasible")),
            "converged": bool(result.get("converged")),
            "reason": result.get("reason"),
        }
    if require_business_decision and analysis.get("decision_status") != "business_approved":
        return False, "MAX_PRICE_REVIEW_REQUIRED", {
            "decision_status": analysis.get("decision_status") or "candidate",
        }
    return True, "", {"decision_thresholds": thresholds}


def save_spec(
    workspace_id: str, spec: dict[str, Any], *, idempotency_key: str = "", request_id: str = "",
    trusted_confirmation: bool = False,
) -> dict[str, Any]:
    spec = _sanitize_client_evidence_claims(spec)
    if not trusted_confirmation:
        # Saving and confirming are separate authority-bearing operations.
        # Client-authored confirmation fields can never turn a candidate into
        # a formal spec through the generic save endpoint.
        spec["confirmation_status"] = "candidate"
        spec["confirmed_by"] = ""
        spec["confirmed_at"] = ""
    body_hash = _hash(spec)
    evidence_binding = _bind_spec_evidence(
        workspace_id,
        spec,
    )
    evidence_binding_hash = str(evidence_binding.get("binding_hash") or "")
    snapshot_hash = _hash({
        "spec_hash": body_hash,
        "evidence_binding_hash": evidence_binding_hash,
        "evidence_binding_version": evidence_binding.get("binding_version"),
    })
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        scope = f"spec:{idempotency_key}" if idempotency_key else ""
        prior = _active_idempotency_record(state["idempotency"], scope) if scope else None
        if prior:
            if prior["body_hash"] != body_hash:
                return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior["spec_id"]}
            return {**state["specs"][prior["spec_id"]], "idempotent_replay": True}
        matching = [
            row for row in state["specs"].values()
            if row.get("spec_hash") == body_hash
            and row.get("evidence_binding_hash") == evidence_binding_hash
        ]
        if matching:
            existing = matching[0]
            if scope:
                state["idempotency"][scope] = _idempotency_record(
                    scope, body_hash, spec_id=existing["spec_id"],
                )
                _save(workspace_id, state)
            return {**existing, "idempotent_replay": True}
        spec_id = f"spec_{uuid.uuid4().hex}"
        revision = max((int(row.get("revision") or 0) for row in state["specs"].values()), default=0) + 1
        row = {
            "ok": True, "spec_id": spec_id, "version": spec.get("version"),
            "spec_hash": body_hash, "revision": revision, "spec": spec,
            "snapshot_hash": snapshot_hash,
            "evidence_binding_version": evidence_binding.get("binding_version"),
            "evidence_binding_hash": evidence_binding_hash,
            "evidence_status": evidence_binding.get("status"),
            "evidence_formal_ok": bool(evidence_binding.get("formal_ok")),
            "evidence_binding": evidence_binding,
            "confirmation_status": spec.get("confirmation_status") or "candidate",
            "request_id": request_id, "created_at": _now(),
        }
        state["specs"][spec_id] = row
        if scope:
            state["idempotency"][scope] = _idempotency_record(
                scope, body_hash, spec_id=spec_id,
            )
        _save(workspace_id, state)
        return row


def confirm_saved_spec(
    workspace_id: str,
    spec_id: str,
    *,
    actor: str,
    note: str = "",
    idempotency_key: str = "",
    request_id: str = "",
) -> dict[str, Any]:
    """Create an immutable confirmed revision from a saved candidate spec."""

    actor = str(actor or "").strip()
    if not actor:
        return {"ok": False, "error": "AUTHENTICATION_REQUIRED"}
    now = _now()
    body_hash = _hash({
        "action": "confirm_spec", "spec_id": spec_id,
        "actor": actor, "note": str(note or ""),
    })
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        source = state["specs"].get(str(spec_id or ""))
        if not source:
            return {"ok": False, "error": "SPEC_NOT_FOUND"}
        candidate = copy.deepcopy(source.get("spec") or {})
        formal_candidate = mark_spec_confirmed(
            candidate, confirmed_by=actor, confirmed_at=now,
        )
        estimate_preview = _is_estimate_preview_spec(candidate)
        schema_ok, schema_errors = (
            validate(formal_candidate)
            if estimate_preview
            else validate_for_formal(formal_candidate)
        )
        evidence_binding = _bind_spec_evidence(
            workspace_id,
            candidate,
        )
        if not schema_ok:
            return {
                "ok": False,
                "error": "SPEC_VALIDATION_FAILED",
                "details": list(schema_errors),
            }
        if not estimate_preview and not evidence_binding.get("formal_ok"):
            return {
                "ok": False,
                "error": "EVIDENCE_REVIEW_REQUIRED",
                "evidence_status": evidence_binding.get("status"),
            }
        scope = f"confirm-spec:{idempotency_key}" if idempotency_key else ""
        prior = _active_idempotency_record(state["idempotency"], scope) if scope else None
        if prior:
            if prior.get("body_hash") != body_hash:
                return {
                    "ok": False, "error": "IDEMPOTENCY_CONFLICT",
                    "resource_id": prior.get("spec_id"),
                }
            existing = state["specs"].get(str(prior.get("spec_id") or ""))
            if not existing:
                raise RuntimeError("spec confirmation idempotency record points to missing spec")
            return {**existing, "idempotent_replay": True}
        existing = next((
            row for row in state["specs"].values()
            if row.get("parent_spec_id") == spec_id
            and row.get("confirmation_status") == "confirmed"
            and str((row.get("confirmation") or {}).get("actor") or "") == actor
            and str((row.get("confirmation") or {}).get("note") or "") == str(note or "")
        ), None)
        if existing:
            if scope:
                state["idempotency"][scope] = _idempotency_record(
                    scope, body_hash, spec_id=existing["spec_id"],
                )
                _save(workspace_id, state)
            return {**existing, "idempotent_replay": True}
        confirmed = formal_candidate
        if estimate_preview:
            confirmed["confirmation_scope"] = "estimate_preview"
        evidence_binding = _bind_spec_evidence(
            workspace_id,
            confirmed,
        )
        spec_hash = _hash(confirmed)
        snapshot_hash = _hash({
            "spec_hash": spec_hash,
            "evidence_binding_hash": evidence_binding.get("binding_hash"),
            "evidence_binding_version": evidence_binding.get("binding_version"),
        })
        confirmed_id = f"spec_{uuid.uuid4().hex}"
        revision = max(
            (int(row.get("revision") or 0) for row in state["specs"].values()),
            default=0,
        ) + 1
        row = {
            "ok": True,
            "spec_id": confirmed_id,
            "parent_spec_id": spec_id,
            "version": confirmed.get("version"),
            "spec_hash": spec_hash,
            "revision": revision,
            "spec": confirmed,
            "snapshot_hash": snapshot_hash,
            "evidence_binding_version": evidence_binding.get("binding_version"),
            "evidence_binding_hash": evidence_binding.get("binding_hash"),
            "evidence_status": evidence_binding.get("status"),
            "evidence_formal_ok": bool(evidence_binding.get("formal_ok")),
            "evidence_binding": evidence_binding,
            "confirmation_status": "confirmed",
            "confirmation_scope": (
                "estimate_preview" if estimate_preview else "formal_input"
            ),
            "confirmation": {
                "actor": actor, "note": str(note or ""),
                "confirmed_at": now, "request_id": request_id,
            },
            "request_id": request_id,
            "created_at": now,
        }
        state["specs"][confirmed_id] = row
        if scope:
            state["idempotency"][scope] = _idempotency_record(
                scope, body_hash, spec_id=confirmed_id,
            )
        _save(workspace_id, state)
        return row


def list_specs(
    workspace_id: str,
    *,
    limit: int = 50,
) -> list[dict[str, Any]]:
    """Return saved acquisition specs newest-first for workbench recovery."""

    rows = [
        copy.deepcopy(row)
        for row in _load(workspace_id)["specs"].values()
    ]
    rows.sort(
        key=lambda row: (int(row.get("revision") or 0), str(row.get("created_at") or "")),
        reverse=True,
    )
    return rows[: max(1, min(int(limit or 50), 100))]


def get_spec(
    workspace_id: str,
    spec_id: str,
) -> dict[str, Any]:
    return copy.deepcopy(
        _load(workspace_id)["specs"].get(spec_id) or {}
    )


def _is_selected_scenario(spec: Mapping[str, Any], scenario_id: str) -> bool:
    """Only run the scenario whose assumptions are already selected in the Spec."""

    selected = str(spec.get("selected_scenario_id") or "base").strip()
    requested = str(scenario_id or "base").strip()
    return bool(selected and requested and requested == selected)


def create_run(
    workspace_id: str, spec: dict[str, Any], *, discount_rate: float = 0.08,
    scenario_id: str = "base", idempotency_key: str = "", request_id: str = "",
    scenario_change_ledger: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    if not _is_selected_scenario(spec, scenario_id):
        return {"ok": False, "error": "SCENARIO_NOT_FOUND"}
    estimate_preview = _is_estimate_preview_spec(spec)
    schema_ok, schema_errors = (
        validate(spec) if estimate_preview else validate_for_formal(spec)
    )
    evidence_binding = _bind_spec_evidence(
        workspace_id,
        spec,
    )
    if not schema_ok:
        return {"ok": False, "error": "SPEC_VALIDATION_FAILED", "details": list(schema_errors)}
    if not estimate_preview and not evidence_binding.get("formal_ok"):
        return {
            "ok": False,
            "error": "EVIDENCE_REVIEW_REQUIRED",
            "evidence_status": evidence_binding.get("status"),
        }
    request_id = request_id or f"req_{uuid.uuid4().hex}"
    saved_spec = save_spec(
        workspace_id, spec, request_id=request_id, trusted_confirmation=True,
    )
    if not saved_spec.get("ok"):
        return saved_spec
    spec = copy.deepcopy(saved_spec.get("spec") or {})
    evidence_binding = copy.deepcopy(saved_spec.get("evidence_binding") or {})
    body = {"spec": spec, "discount_rate": discount_rate, "scenario_id": scenario_id}
    body_hash = _hash(body)
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        scope = f"run:{idempotency_key}" if idempotency_key else ""
        prior = _active_idempotency_record(state["idempotency"], scope) if scope else None
        if prior:
            if prior["body_hash"] != body_hash:
                return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior["run_id"]}
            existing = state["runs"].get(prior["run_id"])
            if not existing:
                raise RuntimeError("run idempotency record points to a missing run")
            return {**existing, "idempotent_replay": True}
        result = run_acquisition_model(spec, discount_rate=discount_rate, scenario_id=scenario_id)
        run_id = f"acqrun_{uuid.uuid4().hex}"
        schema_formal_ok, schema_errors = validate_for_formal(spec)
        formal_ok, formal_errors, evidence_binding = _formal_assessment(
            workspace_id,
            spec,
            evidence_binding=evidence_binding,
        )
        created_at = _now()
        issues = [
            {"code": "SPEC_VALIDATION_FAILED", "blocking": True, "status": "open", "detail": item, "created_at": created_at}
            for item in schema_errors
        ]
        issues.extend(_evidence_blocking_issues(evidence_binding, created_at=created_at))
        issues.extend([
            {"code": "REFERENCE_DIFF_BLOCKING", "blocking": True, "status": "open", "detail": "参考轨尚未复核", "created_at": created_at},
            {"code": "BUSINESS_REVIEW_REQUIRED", "blocking": True, "status": "open", "detail": "业务复核尚未完成", "created_at": created_at},
            {"code": "MAX_PRICE_REQUIRED", "blocking": True, "status": "open", "detail": "最高可接受价尚未按决策阈值求解", "created_at": created_at},
        ])
        scenario_ledger = list(scenario_change_ledger or [])
        invalid_scenario_approvals = [
            item for item in scenario_ledger
            if not all(str(item.get(key) or "").strip() for key in ("source", "proposed_by", "approved_by", "approval_reason"))
            or item.get("proposed_by") == item.get("approved_by")
        ]
        if invalid_scenario_approvals:
            issues.append({
                "code": "SCENARIO_APPROVAL_REQUIRED", "blocking": True, "status": "open",
                "detail": "情景调整缺来源/提议人/独立批准人/批准理由", "created_at": created_at,
            })
        row = {
            "ok": True, "available": True, "run_id": run_id, "workspace_id": workspace_id,
            "status": "succeeded", "lifecycle_status": "reference_review_pending",
            "delivery_mode": (
                "estimate_preview" if estimate_preview else "formal_candidate"
            ),
            "model_version": result["model_version"],
            "spec_version": LATEST_SPEC_VERSION, "spec_hash": _hash(spec), "input_hash": body_hash,
            "spec_id": saved_spec.get("spec_id"),
            "spec_snapshot_hash": saved_spec.get("snapshot_hash"),
            "evidence_binding_version": evidence_binding.get("binding_version"),
            "evidence_binding_hash": evidence_binding.get("binding_hash"),
            "evidence_status": evidence_binding.get("status"),
            "evidence_formal_ok": bool(evidence_binding.get("formal_ok")),
            "evidence_binding": evidence_binding,
            "scenario_id": scenario_id, "scenario_change_ledger": scenario_ledger,
            "discount_rate": discount_rate, "result": result, "consistency_ok": True,
            "reference_review_status": "pending", "business_review_status": "pending",
            "review_status": "draft", "formal_spec_valid": bool(schema_formal_ok and formal_ok),
            "formal_spec_errors": formal_errors, "request_id": request_id,
            "created_at": created_at, "approved_at": "", "approved_by": "",
            **_migration_binding(spec),
            "issues": issues,
            "state_history": [
                _history_event("validated_spec", request_id=request_id),
                _history_event("running", request_id=request_id),
                _history_event("calculated", request_id=request_id),
                _history_event("internally_consistent", request_id=request_id),
                _history_event("reference_review_pending", request_id=request_id),
            ],
        }
        state["runs"][run_id] = row
        if scope:
            state["idempotency"][scope] = _idempotency_record(
                scope, body_hash, run_id=run_id,
            )
        _save(workspace_id, state)
    return row


def enqueue_run(
    workspace_id: str, spec: dict[str, Any], *, discount_rate: float = 0.08,
    scenario_id: str = "base", idempotency_key: str = "", request_id: str = "",
    scenario_change_ledger: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """Persist a pollable queued run before any model calculation starts."""

    if not _is_selected_scenario(spec, scenario_id):
        return {"ok": False, "error": "SCENARIO_NOT_FOUND"}
    schema_ok, schema_errors = validate_for_formal(spec)
    evidence_binding = _bind_spec_evidence(
        workspace_id,
        spec,
    )
    if not schema_ok:
        return {"ok": False, "error": "SPEC_VALIDATION_FAILED", "details": list(schema_errors)}
    if not evidence_binding.get("formal_ok"):
        return {
            "ok": False,
            "error": "EVIDENCE_REVIEW_REQUIRED",
            "evidence_status": evidence_binding.get("status"),
        }
    request_id = request_id or f"req_{uuid.uuid4().hex}"
    saved_spec = save_spec(
        workspace_id, spec, request_id=request_id, trusted_confirmation=True,
    )
    if not saved_spec.get("ok"):
        return saved_spec
    spec = copy.deepcopy(saved_spec.get("spec") or {})
    evidence_binding = copy.deepcopy(saved_spec.get("evidence_binding") or {})
    body = {"spec": spec, "discount_rate": discount_rate, "scenario_id": scenario_id}
    body_hash = _hash(body)
    schema_formal_ok, _schema_errors = validate_for_formal(spec)
    formal_ok, formal_errors, evidence_binding = _formal_assessment(
        workspace_id,
        spec,
        evidence_binding=evidence_binding,
    )
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        scope = f"run:{idempotency_key}" if idempotency_key else ""
        prior = _active_idempotency_record(state["idempotency"], scope) if scope else None
        if prior:
            if prior.get("body_hash") != body_hash:
                return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior.get("run_id", "")}
            existing = state["runs"].get(prior.get("run_id"))
            if not existing:
                raise RuntimeError("run idempotency record points to a missing run")
            return {**existing, "idempotent_replay": True}
        run_id = f"acqrun_{uuid.uuid4().hex}"
        created_at = _now()
        row = {
            "ok": True, "available": False, "run_id": run_id, "workspace_id": workspace_id,
            "status": "queued", "progress": 0, "lifecycle_status": "validated_spec",
            "spec_version": LATEST_SPEC_VERSION, "spec_hash": _hash(spec), "input_hash": body_hash,
            "spec_id": saved_spec.get("spec_id"), "scenario_id": scenario_id,
            "spec_snapshot_hash": saved_spec.get("snapshot_hash"),
            "evidence_binding_version": evidence_binding.get("binding_version"),
            "evidence_binding_hash": evidence_binding.get("binding_hash"),
            "evidence_status": evidence_binding.get("status"),
            "evidence_formal_ok": bool(evidence_binding.get("formal_ok")),
            "evidence_binding": evidence_binding,
            "scenario_change_ledger": list(scenario_change_ledger or []), "discount_rate": discount_rate,
            "reference_review_status": "pending", "business_review_status": "pending",
            "review_status": "draft", "formal_spec_valid": bool(schema_formal_ok and formal_ok),
            "formal_spec_errors": formal_errors,
            "request_id": request_id, "created_at": created_at, "updated_at": created_at,
            "approved_at": "", "approved_by": "", "issues": [],
            **_migration_binding(spec),
            "state_history": [
                _history_event("validated_spec", request_id=request_id),
                _history_event("queued", request_id=request_id),
            ],
        }
        state["runs"][run_id] = row
        if scope:
            state["idempotency"][scope] = _idempotency_record(
                scope, body_hash, run_id=run_id,
            )
        _save(workspace_id, state)
        return row


def execute_queued_run(
    workspace_id: str,
    run_id: str,
) -> None:
    """Execute one durable queued run and converge it to succeeded/failed."""

    with _state_guard(workspace_id):
        state = _load(workspace_id)
        run = state["runs"].get(run_id)
        if not run or run.get("status") in {"succeeded", "failed", "cancelled"}:
            return
        spec_row = state["specs"].get(str(run.get("spec_id") or "")) or {}
        spec = spec_row.get("spec")
        if not isinstance(spec, dict):
            run.update({
                "status": "failed", "progress": 100, "updated_at": _now(),
                "error": {"code": "SPEC_VALIDATION_FAILED", "message": "spec snapshot missing", "retryable": False},
            })
            _save(workspace_id, state)
            return
        run.update({"status": "running", "progress": 10, "lifecycle_status": "running", "updated_at": _now()})
        run.setdefault("state_history", []).append(_history_event(
            "running", request_id=str(run.get("request_id") or ""),
        ))
        discount_rate = float(run.get("discount_rate") or 0.08)
        scenario_id = str(run.get("scenario_id") or "base")
        _save(workspace_id, state)
    try:
        current_evidence = _bind_spec_evidence(
            workspace_id,
            spec,
        )
    except Exception as exc:  # noqa: BLE001
        _LOG.warning(
            "asset acquisition evidence binding failed; error_type=%s",
            type(exc).__name__,
        )
        current_evidence = {
            "formal_ok": False,
            "status": "invalid",
            "binding_version": "",
            "binding_hash": "",
            "bindings": [],
            "missing": [],
            "pending": [],
            "invalid": [{
                "source_path": "$",
                "code": "SOURCE_EVIDENCE_STATE_INVALID",
                "message": _SOURCE_EVIDENCE_FAILURE_MESSAGE,
            }],
        }
    evidence_hash_matches = bool(
        current_evidence.get("binding_hash") == run.get("evidence_binding_hash")
        and current_evidence.get("binding_version") == run.get("evidence_binding_version")
    )
    evidence_snapshot_matches = bool(evidence_hash_matches and current_evidence.get("formal_ok"))
    try:
        result = run_acquisition_model(spec, discount_rate=discount_rate, scenario_id=scenario_id)
    except Exception as exc:  # noqa: BLE001
        code = "SPEC_VALIDATION_FAILED" if isinstance(exc, AcquisitionModelError) else "RUN_FAILED"
        safe_message = (
            _RUN_VALIDATION_FAILURE_MESSAGE
            if code == "SPEC_VALIDATION_FAILED"
            else _RUN_EXECUTION_FAILURE_MESSAGE
        )
        _LOG.warning(
            "asset acquisition run failed; run_id=%s error_type=%s code=%s",
            run_id,
            type(exc).__name__,
            code,
        )
        with _state_guard(workspace_id):
            state = _load(workspace_id)
            run = state["runs"].get(run_id)
            if run and run.get("status") != "cancelled":
                run.update({
                    "status": "failed", "progress": 100, "lifecycle_status": "failed", "updated_at": _now(),
                    "error": {"code": code, "message": safe_message, "retryable": False},
                })
                run.setdefault("state_history", []).append(_history_event(
                    "failed", request_id=str(run.get("request_id") or ""), error_code=code,
                ))
                _save(workspace_id, state)
        return

    with _state_guard(workspace_id):
        state = _load(workspace_id)
        run = state["runs"].get(run_id)
        if not run or run.get("status") == "cancelled":
            return
        created_at = str(run.get("created_at") or _now())
        schema_formal_ok, schema_errors = validate_for_formal(spec)
        formal_errors = [*schema_errors, *_evidence_error_strings(current_evidence)]
        if not evidence_hash_matches:
            formal_errors.append(
                "finance_spec.v3 证据绑定快照已变化或不再满足正式复核条件"
            )
        issues = [
            {"code": "SPEC_VALIDATION_FAILED", "blocking": True, "status": "open", "detail": item, "created_at": created_at}
            for item in schema_errors
        ]
        issues.extend(_evidence_blocking_issues(current_evidence, created_at=created_at))
        if not evidence_hash_matches:
            issues.append({
                "code": "EVIDENCE_BINDING_STALE",
                "blocking": True,
                "status": "open",
                "detail": (
                    "运行排队后证据绑定状态发生变化；必须保存新Spec修订并重新运行。"
                    f" snapshot={run.get('evidence_binding_hash')} current={current_evidence.get('binding_hash')}"
                ),
                "created_at": created_at,
            })
        issues.extend([
            {"code": "REFERENCE_DIFF_BLOCKING", "blocking": True, "status": "open", "detail": "参考轨尚未复核", "created_at": created_at},
            {"code": "BUSINESS_REVIEW_REQUIRED", "blocking": True, "status": "open", "detail": "业务复核尚未完成", "created_at": created_at},
            {"code": "MAX_PRICE_REQUIRED", "blocking": True, "status": "open", "detail": "最高可接受价尚未按决策阈值求解", "created_at": created_at},
        ])
        scenario_ledger = list(run.get("scenario_change_ledger") or [])
        if any(
            not all(str(item.get(key) or "").strip() for key in ("source", "proposed_by", "approved_by", "approval_reason"))
            or item.get("proposed_by") == item.get("approved_by")
            for item in scenario_ledger
        ):
            issues.append({
                "code": "SCENARIO_APPROVAL_REQUIRED", "blocking": True, "status": "open",
                "detail": "情景调整缺来源/提议人/独立批准人/批准理由", "created_at": created_at,
            })
        run.update({
            "available": True, "status": "succeeded", "progress": 100,
            "lifecycle_status": "reference_review_pending", "model_version": result["model_version"],
            "result": result, "consistency_ok": True, "issues": issues, "updated_at": _now(),
            "formal_spec_valid": bool(schema_formal_ok and evidence_snapshot_matches),
            "formal_spec_errors": formal_errors,
            "evidence_revalidation": {
                "checked_at": _now(),
                "matches_snapshot": evidence_hash_matches,
                "binding_version": current_evidence.get("binding_version"),
                "binding_hash": current_evidence.get("binding_hash"),
                "status": current_evidence.get("status"),
                "formal_ok": bool(current_evidence.get("formal_ok")),
            },
        })
        run.setdefault("state_history", []).extend([
            _history_event("calculated", request_id=str(run.get("request_id") or "")),
            _history_event("internally_consistent", request_id=str(run.get("request_id") or "")),
            _history_event("reference_review_pending", request_id=str(run.get("request_id") or "")),
        ])
        _save(workspace_id, state)


def get_run(
    workspace_id: str,
    run_id: str,
) -> dict[str, Any]:
    return dict(
        _load(workspace_id)["runs"].get(run_id) or {}
    )


def list_runs(
    workspace_id: str,
    *,
    limit: int = 50,
) -> list[dict[str, Any]]:
    """Return lightweight acquisition run summaries newest-first."""

    rows: list[dict[str, Any]] = []
    summary_keys = (
        "run_id", "workspace_id", "status", "lifecycle_status",
        "model_version", "spec_version", "spec_hash", "input_hash",
        "spec_id", "scenario_id", "discount_rate", "consistency_ok",
        "spec_snapshot_hash", "evidence_binding_version", "evidence_binding_hash",
        "evidence_status", "evidence_formal_ok", "evidence_revalidation",
        "reference_review_status", "business_review_status",
        "review_status", "formal_spec_valid", "formal_spec_errors",
        "request_id", "created_at", "approved_at", "approved_by", "issues",
    )
    for run in _load(workspace_id)["runs"].values():
        indicators = (run.get("result") or {}).get("indicators") or {}
        row = {key: copy.deepcopy(run.get(key)) for key in summary_keys}
        row.update({
            "invest_type": "asset_acquisition",
            "industry": "资产收购",
            "available": bool(run.get("available")),
            "started_at": run.get("created_at"),
            "finished_at": run.get("finished_at") or run.get("created_at"),
            "indicators": {
                key: copy.deepcopy(indicators.get(key))
                for key in ("project_irr_pct", "equity_irr_pct", "npv_wan", "minimum_dscr")
            },
        })
        rows.append(row)
    rows.sort(
        key=lambda row: (str(row.get("created_at") or ""), str(row.get("run_id") or "")),
        reverse=True,
    )
    return rows[: max(1, min(int(limit or 50), 100))]


MAX_SCENARIO_MATRIX_COMBINATIONS = 64


def create_scenario_matrix(
    workspace_id: str, run_id: str, dimensions: dict[str, Any], *,
    idempotency_key: str = "", request_id: str = "",
) -> dict[str, Any]:
    """Evaluate a bounded Cartesian product without mutating the bound run/spec."""

    request_id = request_id or f"req_{uuid.uuid4().hex}"
    if not isinstance(dimensions, dict) or not dimensions:
        return {"ok": False, "error": "SCENARIO_MATRIX_INVALID", "details": {"reason": "dimensions_required"}}
    unknown = sorted(set(dimensions) - INDEPENDENT_SCENARIO_FIELDS)
    if unknown:
        return {"ok": False, "error": "SCENARIO_FIELD_UNSUPPORTED", "details": {"fields": unknown}}
    normalized: dict[str, list[Any]] = {}
    for field in sorted(dimensions):
        values = dimensions[field]
        if not isinstance(values, (list, tuple)) or not values:
            return {
                "ok": False, "error": "SCENARIO_MATRIX_INVALID",
                "details": {"field": field, "reason": "non_empty_array_required"},
            }
        rows = list(values)
        fingerprints = [_hash(value) for value in rows]
        if len(fingerprints) != len(set(fingerprints)):
            return {
                "ok": False, "error": "SCENARIO_MATRIX_INVALID",
                "details": {"field": field, "reason": "duplicate_values"},
            }
        normalized[field] = rows
    combination_count = math.prod(len(values) for values in normalized.values())
    if combination_count > MAX_SCENARIO_MATRIX_COMBINATIONS:
        return {
            "ok": False, "error": "SCENARIO_MATRIX_TOO_LARGE",
            "details": {"combination_count": combination_count, "max_combinations": MAX_SCENARIO_MATRIX_COMBINATIONS},
        }
    body_hash = _hash({"run_id": run_id, "dimensions": normalized})
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        run = state["runs"].get(run_id)
        if not run:
            return {"ok": False, "error": "RUN_NOT_FOUND"}
        if run.get("review_status") == "approved":
            return {"ok": False, "error": "APPROVED_RUN_IMMUTABLE", "resource_id": run_id}
        if run.get("status") != "succeeded" or not run.get("available"):
            return {"ok": False, "error": "RUN_NOT_READY", "details": {"status": run.get("status")}}
        spec_row = state["specs"].get(str(run.get("spec_id") or "")) or {}
        spec = spec_row.get("spec")
        if not isinstance(spec, dict) or _hash(spec) != run.get("spec_hash"):
            return {"ok": False, "error": "RUN_SPEC_MISMATCH"}
        scope = f"scenario_matrix:{run_id}:{idempotency_key}" if idempotency_key else ""
        prior = _active_idempotency_record(state["idempotency"], scope) if scope else None
        if prior:
            if prior.get("body_hash") != body_hash:
                return {
                    "ok": False, "error": "IDEMPOTENCY_CONFLICT",
                    "resource_id": prior.get("matrix_id", ""),
                }
            existing = state["scenario_matrices"].get(prior.get("matrix_id"))
            if not existing:
                raise RuntimeError("scenario matrix idempotency record points to a missing matrix")
            return {**existing, "idempotent_replay": True}
        bound = {
            "run_id": run_id, "spec_id": run.get("spec_id"), "spec_hash": run.get("spec_hash"),
            "input_hash": run.get("input_hash"), "model_version": run.get("model_version"),
            "scenario_id": run.get("scenario_id"),
        }

    fields = list(normalized)
    rows: list[dict[str, Any]] = []
    matrix_id = f"scenario_matrix_{uuid.uuid4().hex}"
    asset_type = str(spec.get("asset_type") or "hotel_lease")
    base_market_rent = copy.deepcopy((spec.get("lease_portfolio") or {}).get("market_rent"))
    for index, values in enumerate(product(*(normalized[field] for field in fields)), 1):
        scenario_id = f"{matrix_id}:row_{index:03d}"
        changes = {field: copy.deepcopy(value) for field, value in zip(fields, values)}
        try:
            scenario_spec, ledger = apply_scenario(spec, changes)
        except AcquisitionModelError as exc:
            return {
                "ok": False, "error": "SCENARIO_MATRIX_INVALID",
                "details": {"scenario_index": index, "reason": str(exc)},
            }
        expected_market_rent = copy.deepcopy(changes.get("lease_portfolio.market_rent", base_market_rent))
        actual_market_rent = copy.deepcopy((scenario_spec.get("lease_portfolio") or {}).get("market_rent"))
        if asset_type == "hotel_lease" and actual_market_rent != expected_market_rent:
            return {
                "ok": False, "error": "SCENARIO_INDEPENDENCE_VIOLATION",
                "details": {"field": "lease_portfolio.market_rent", "expected": expected_market_rent, "actual": actual_market_rent},
            }
        try:
            result = run_acquisition_model(
                scenario_spec, discount_rate=float(run.get("discount_rate") or 0.08),
                scenario_id=scenario_id,
            )
        except AcquisitionModelError as exc:
            return {
                "ok": False, "error": "SCENARIO_MATRIX_INVALID",
                "details": {"scenario_index": index, "changes": changes, "reason": str(exc)},
            }
        solar = dict(scenario_spec.get("solar_operation") or {})
        row = {
            "scenario_id": scenario_id, "changes": changes,
            "scenario_change_ledger": ledger, "scenario_spec_hash": _hash(scenario_spec),
            "result_hash": _hash(result),
            "purchase_price_wan": result.get("purchase_price_wan"),
            "financing_ratio": (scenario_spec.get("transaction") or {}).get("financing_ratio"),
            "indicators": copy.deepcopy(result.get("indicators") or {}),
        }
        if asset_type == "solar_power":
            row.update({
                "tariff_yuan_per_kwh": solar.get("tariff_yuan_per_kwh"),
                "annual_generation_mwh": solar.get("annual_generation_mwh"),
                "utilization_hours": solar.get("utilization_hours"),
                "annual_opex_wan": solar.get("annual_opex_wan"),
            })
        else:
            row.update({
                "market_rent": actual_market_rent,
                "occupancy": copy.deepcopy(
                    (scenario_spec.get("hotel_operation") or {}).get("occupancy")
                ),
            })
        rows.append(row)
    created_at = _now()
    matrix = {
        "ok": True, "matrix_id": matrix_id, "status": "succeeded", "read_only": True,
        "workspace_id": workspace_id, **bound, "dimensions": normalized,
        "combination_count": combination_count, "max_combinations": MAX_SCENARIO_MATRIX_COMBINATIONS,
        "rows": rows, "request_id": request_id, "created_at": created_at,
    }
    matrix["matrix_hash"] = _hash({
        "bound": bound, "dimensions": normalized,
        "rows": [{"changes": row["changes"], "scenario_spec_hash": row["scenario_spec_hash"], "result_hash": row["result_hash"]} for row in rows],
    })
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        current = state["runs"].get(run_id)
        if not current or current.get("spec_hash") != bound["spec_hash"] or current.get("input_hash") != bound["input_hash"]:
            return {"ok": False, "error": "RUN_SPEC_MISMATCH"}
        if current.get("review_status") == "approved":
            return {"ok": False, "error": "APPROVED_RUN_IMMUTABLE", "resource_id": run_id}
        scope = f"scenario_matrix:{run_id}:{idempotency_key}" if idempotency_key else ""
        if scope and (prior := _active_idempotency_record(state["idempotency"], scope)):
            if prior.get("body_hash") != body_hash:
                return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior.get("matrix_id", "")}
            return {**state["scenario_matrices"][prior["matrix_id"]], "idempotent_replay": True}
        state["scenario_matrices"][matrix_id] = matrix
        if scope:
            state["idempotency"][scope] = _idempotency_record(
                scope, body_hash, matrix_id=matrix_id,
            )
        _save(workspace_id, state)
    return matrix


def get_scenario_matrix(
    workspace_id: str,
    run_id: str,
    matrix_id: str,
) -> dict[str, Any]:
    matrix = dict(
        _load(workspace_id)["scenario_matrices"].get(
            matrix_id
        )
        or {}
    )
    if not matrix or matrix.get("run_id") != run_id:
        return {}
    return matrix


def list_scenario_matrices(
    workspace_id: str,
    run_id: str,
) -> list[dict[str, Any]]:
    rows = [
        value
        for value in _load(workspace_id)[
            "scenario_matrices"
        ].values()
        if value.get("run_id") == run_id
    ]
    rows.sort(key=lambda value: str(value.get("created_at") or ""), reverse=True)
    return [
        {key: value.get(key) for key in (
            "matrix_id", "matrix_hash", "run_id", "spec_id", "spec_hash", "input_hash",
            "combination_count", "status", "created_at",
        )}
        for value in rows
    ]


def _diff_is_blocking(row: dict[str, Any], tolerance: float) -> bool:
    if bool(row.get("blocking")) or row.get("within_tolerance") is False:
        return True
    relative = row.get("relative_difference")
    if relative is None:
        try:
            reference = float(row["reference_value"])
            calculated = float(row["calculated_value"])
            relative = abs(calculated - reference) / max(abs(reference), 1e-12)
        except (KeyError, TypeError, ValueError):
            return False
    try:
        exceeds = abs(float(relative)) > max(float(tolerance), 0.0)
    except (TypeError, ValueError):
        return True
    if not exceeds:
        return False
    decision = str(row.get("decision_status") or "").lower()
    has_reason = bool(str(row.get("decision_reason") or row.get("reason") or "").strip())
    has_evidence = bool(row.get("evidence_ids") or row.get("source_locators"))
    return decision not in {"approved", "resolved", "accepted"} or not has_reason or not has_evidence


def review_reference(
    workspace_id: str, run_id: str, *, status: str, diffs: list[dict[str, Any]],
    tolerance: float, actor: str, note: str = "", reference_hash: str = "",
    reference_kind: str = "generic", request_id: str = "",
) -> dict[str, Any]:
    if status not in {"approved", "pending", "rejected", "out_of_tolerance"}:
        return {"ok": False, "error": "invalid_reference_review_status"}
    actor = str(actor or "").strip()
    reference_hash = str(reference_hash or "").strip()
    reference_kind = str(reference_kind or "generic").strip().lower()
    if reference_kind not in {"generic", "hengli"}:
        return {"ok": False, "error": "REFERENCE_KIND_UNSUPPORTED"}
    if not actor:
        return {"ok": False, "error": "authenticated_actor_required"}
    if status == "approved" and not reference_hash:
        return {"ok": False, "error": "reference_hash_required"}
    if status == "approved" and reference_kind == "hengli":
        from lvke_mcp.domains.finance.hengli_reference import scenario_matrix as hengli_scenario_matrix

        hengli = hengli_scenario_matrix()
        expected_hash = str(hengli.get("reference_hash") or "")
        if reference_hash != expected_hash:
            return {
                "ok": False, "error": "HENGLI_REFERENCE_HASH_MISMATCH",
                "expected_reference_hash": expected_hash, "provided_reference_hash": reference_hash,
            }
        blocking_issues = list((hengli.get("replay") or {}).get("blocking_issues") or [])
        if not hengli.get("valid") or blocking_issues:
            return {
                "ok": False, "error": "HENGLI_REFERENCE_BLOCKING",
                "reference_hash": expected_hash, "issues": blocking_issues,
                "validation_errors": list(hengli.get("errors") or []),
            }
    invalid_rows = [index for index, row in enumerate(diffs) if not isinstance(row, dict)]
    if invalid_rows:
        return {"ok": False, "error": "invalid_reference_diff", "rows": invalid_rows}
    blocking = status != "approved" or any(_diff_is_blocking(row, tolerance) for row in diffs)
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        run = state["runs"].get(run_id)
        if not run:
            return {"ok": False, "error": "run_not_found"}
        if run.get("review_status") == "approved":
            return {"ok": False, "error": "approved_run_immutable"}
        if status == "approved":
            spec_row = state["specs"].get(str(run.get("spec_id") or "")) or {}
            bound_spec = spec_row.get("spec")
            if not isinstance(bound_spec, dict) or _hash(bound_spec) != run.get("spec_hash"):
                return {"ok": False, "error": "RUN_SPEC_MISMATCH"}
            evidence_ok, current_evidence = _current_evidence_matches_run(
                workspace_id,
                run,
                bound_spec,
            )
            if not evidence_ok:
                return {
                    "ok": False,
                    "error": "EVIDENCE_REVIEW_REQUIRED",
                    "evidence_status": current_evidence.get("status"),
                    "expected_binding_hash": run.get("evidence_binding_hash"),
                    "current_binding_hash": current_evidence.get("binding_hash"),
                    "issues": [
                        *(current_evidence.get("invalid") or []),
                        *(current_evidence.get("missing") or []),
                        *(current_evidence.get("pending") or []),
                    ],
                }
        review = {
            "status": status, "diffs": diffs, "tolerance": tolerance,
            "reference_hash": reference_hash, "reference_kind": reference_kind,
            "actor": actor, "note": note,
            "request_id": request_id, "reviewed_at": _now(),
        }
        previous = run.get("reference_review") or {}
        if previous and _hash(previous | {"reviewed_at": None}) == _hash(review | {"reviewed_at": None}):
            return {
                "ok": True, "run_id": run_id,
                "reference_review_status": run.get("reference_review_status"),
                "reference_kind": previous.get("reference_kind") or "generic",
                "reference_hash": previous.get("reference_hash"),
                "blocking": run.get("reference_review_status") != "approved", "idempotent_replay": True,
            }
        run["reference_review_status"] = status if blocking else "approved"
        run["reference_review"] = review
        run.setdefault("reference_review_history", []).append(review)
        if blocking:
            _open_issue(run, "REFERENCE_DIFF_BLOCKING", "参考轨超容差、缺少裁决证据或尚未批准")
            run["lifecycle_status"] = "reference_review_pending"
        else:
            _close_issue(run, "REFERENCE_DIFF_BLOCKING", actor=actor, reason="参考轨已在容差内复现或差异已完成证据化裁决")
            run["lifecycle_status"] = "business_review_pending"
            run.setdefault("state_history", []).append(_history_event(
                "business_review_pending", actor=actor, request_id=request_id, reference_hash=reference_hash,
            ))
        _save(workspace_id, state)
        return {
            "ok": True, "run_id": run_id,
            "reference_review_status": run["reference_review_status"],
            "reference_kind": reference_kind, "reference_hash": reference_hash,
            "blocking": blocking,
        }


def review_business(
    workspace_id: str, run_id: str, *, status: str, actor: str, note: str = "", request_id: str = "",
) -> dict[str, Any]:
    if status not in {"approved", "pending", "rejected"}:
        return {"ok": False, "error": "invalid_business_review_status"}
    actor = str(actor or "").strip()
    if not actor:
        return {"ok": False, "error": "authenticated_actor_required"}
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        run = state["runs"].get(run_id)
        if not run:
            return {"ok": False, "error": "run_not_found"}
        if run.get("review_status") == "approved":
            return {"ok": False, "error": "approved_run_immutable"}
        reference_actor = str((run.get("reference_review") or {}).get("actor") or "")
        if status == "approved" and run.get("reference_review_status") != "approved":
            return {"ok": False, "error": "reference_review_required"}
        if status == "approved" and not run.get("formal_spec_valid"):
            return {"ok": False, "error": "formal_spec_invalid", "issues": run.get("formal_spec_errors") or []}
        if status == "approved" and actor == reference_actor:
            return {"ok": False, "error": "reviewer_separation_required"}
        if status == "approved":
            spec_row = state["specs"].get(str(run.get("spec_id") or "")) or {}
            spec = spec_row.get("spec")
            if not isinstance(spec, dict) or _hash(spec) != run.get("spec_hash"):
                return {"ok": False, "error": "RUN_SPEC_MISMATCH"}
            evidence_ok, current_evidence = _current_evidence_matches_run(
                workspace_id,
                run,
                spec,
            )
            if not evidence_ok:
                run["formal_spec_valid"] = False
                run["reference_review_status"] = "pending"
                run["evidence_revalidation"] = {
                    "checked_at": _now(),
                    "matches_snapshot": current_evidence.get("binding_hash") == run.get("evidence_binding_hash"),
                    "binding_version": current_evidence.get("binding_version"),
                    "binding_hash": current_evidence.get("binding_hash"),
                    "status": current_evidence.get("status"),
                    "formal_ok": bool(current_evidence.get("formal_ok")),
                }
                _open_issue(
                    run,
                    "EVIDENCE_BINDING_STALE",
                    "证据状态未达到正式条件或与运行快照不一致；须保存新Spec修订并重新运行。",
                )
                _save(workspace_id, state)
                return {
                    "ok": False,
                    "error": "EVIDENCE_REVIEW_REQUIRED",
                    "evidence_status": current_evidence.get("status"),
                    "expected_binding_hash": run.get("evidence_binding_hash"),
                    "current_binding_hash": current_evidence.get("binding_hash"),
                    "issues": [
                        *(current_evidence.get("invalid") or []),
                        *(current_evidence.get("missing") or []),
                        *(current_evidence.get("pending") or []),
                    ],
                }
            max_price_ok, max_price_error, max_price_details = _max_price_gate(
                run, spec, require_business_decision=False,
            )
            if not max_price_ok:
                return {
                    "ok": False, "error": max_price_error,
                    "reason": max_price_error, "details": max_price_details,
                }
        review = {"status": status, "actor": actor, "note": note, "request_id": request_id, "reviewed_at": _now()}
        previous = run.get("business_review") or {}
        if (
            previous
            and run.get("business_review_status") == status
            and _hash(previous | {"reviewed_at": None}) == _hash(review | {"reviewed_at": None})
        ):
            return {"ok": True, "run_id": run_id, "business_review_status": status, "idempotent_replay": True}
        run["business_review_status"] = status
        run["business_review"] = review
        run.setdefault("business_review_history", []).append(review)
        if status == "approved":
            _close_issue(run, "BUSINESS_REVIEW_REQUIRED", actor=actor, reason="业务复核批准")
            max_price_analysis = run.get("max_acquisition_price_analysis") or {}
            if max_price_analysis:
                max_price_analysis.update({
                    "decision_status": "business_approved",
                    "reviewed_by": actor,
                    "reviewed_at": review["reviewed_at"],
                    "review_note": note,
                })
                max_price_analysis["decision_hash"] = _hash({
                    key: value for key, value in max_price_analysis.items() if key != "decision_hash"
                })
                _close_issue(run, "MAX_PRICE_REVIEW_REQUIRED", actor=actor, reason="最高可接受价已通过业务复核")
                _close_issue(run, "MAX_PRICE_REQUIRED", actor=actor, reason="最高可接受价已完成求解与业务复核")
                _close_issue(run, "MAX_PRICE_NOT_FEASIBLE", actor=actor, reason="最高可接受价求解已达到可行与收敛条件")
            run["lifecycle_status"] = "business_review_approved"
            run.setdefault("state_history", []).append(_history_event(
                "business_review_approved", actor=actor, request_id=request_id,
            ))
        else:
            _open_issue(run, "BUSINESS_REVIEW_REQUIRED", "业务复核未批准")
            run["lifecycle_status"] = "business_review_pending"
        _save(workspace_id, state)
        return {"ok": True, "run_id": run_id, "business_review_status": status}


def approve_run(
    workspace_id: str, run_id: str, *, actor: str, note: str = "", request_id: str = "",
) -> dict[str, Any]:
    actor = str(actor or "").strip()
    if not actor:
        return {"ok": False, "error": "authenticated_actor_required"}
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        run = state["runs"].get(run_id)
        if not run:
            return {"ok": False, "error": "run_not_found"}
        if run.get("review_status") == "approved":
            if run.get("approved_by") == actor and str(run.get("approval_note") or "") == str(note or ""):
                return {
                    "ok": True, "run_id": run_id, "review_status": "approved",
                    "approved_at": run.get("approved_at"), "idempotent_replay": True,
                }
            return {"ok": False, "error": "approved_run_immutable"}
        if run.get("status") != "succeeded" or not run.get("consistency_ok"):
            return {"ok": False, "error": "RUN_NOT_APPROVABLE", "reason": "INTERNAL_CONSISTENCY_REQUIRED"}
        if not run.get("evidence_formal_ok"):
            return {
                "ok": False,
                "error": "RUN_NOT_APPROVABLE",
                "reason": "EVIDENCE_REVIEW_REQUIRED",
                "details": {
                    "evidence_status": run.get("evidence_status"),
                    "evidence_binding_hash": run.get("evidence_binding_hash"),
                    "issues": [
                        *((run.get("evidence_binding") or {}).get("invalid") or []),
                        *((run.get("evidence_binding") or {}).get("missing") or []),
                        *((run.get("evidence_binding") or {}).get("pending") or []),
                    ],
                },
            }
        if not run.get("formal_spec_valid"):
            return {"ok": False, "error": "RUN_NOT_APPROVABLE", "reason": "SPEC_VALIDATION_FAILED"}
        spec_row = state["specs"].get(str(run.get("spec_id") or "")) or {}
        spec = spec_row.get("spec")
        if not isinstance(spec, dict) or _hash(spec) != run.get("spec_hash"):
            return {"ok": False, "error": "RUN_NOT_APPROVABLE", "reason": "RUN_SPEC_MISMATCH"}
        evidence_ok, current_evidence = _current_evidence_matches_run(
            workspace_id,
            run,
            spec,
        )
        if not evidence_ok:
            run["formal_spec_valid"] = False
            run["evidence_formal_ok"] = False
            run["reference_review_status"] = "pending"
            run["business_review_status"] = "pending"
            _open_issue(
                run,
                "EVIDENCE_BINDING_STALE",
                "批准前复核发现证据绑定不再满足正式条件或已偏离运行快照。",
            )
            _save(workspace_id, state)
            return {
                "ok": False,
                "error": "RUN_NOT_APPROVABLE",
                "reason": "EVIDENCE_REVIEW_REQUIRED",
                "details": {
                    "evidence_status": current_evidence.get("status"),
                    "expected_binding_hash": run.get("evidence_binding_hash"),
                    "current_binding_hash": current_evidence.get("binding_hash"),
                    "issues": [
                        *(current_evidence.get("invalid") or []),
                        *(current_evidence.get("missing") or []),
                        *(current_evidence.get("pending") or []),
                    ],
                },
            }
        max_price_ok, max_price_error, max_price_details = _max_price_gate(
            run, spec, require_business_decision=True,
        )
        if not max_price_ok:
            return {
                "ok": False, "error": "RUN_NOT_APPROVABLE",
                "reason": max_price_error, "details": max_price_details,
            }
        reference_review = run.get("reference_review") or {}
        if reference_review.get("reference_kind") == "hengli":
            from lvke_mcp.domains.finance.hengli_reference import scenario_matrix as hengli_scenario_matrix

            hengli = hengli_scenario_matrix()
            expected_hash = str(hengli.get("reference_hash") or "")
            if reference_review.get("reference_hash") != expected_hash:
                return {
                    "ok": False, "error": "RUN_NOT_APPROVABLE",
                    "reason": "HENGLI_REFERENCE_HASH_MISMATCH",
                }
            hengli_blockers = list((hengli.get("replay") or {}).get("blocking_issues") or [])
            if not hengli.get("valid") or hengli_blockers:
                return {
                    "ok": False, "error": "RUN_NOT_APPROVABLE",
                    "reason": "HENGLI_REFERENCE_BLOCKING", "issues": hengli_blockers,
                }
        blockers = [issue for issue in run.get("issues") or [] if issue.get("blocking") and issue.get("status") == "open"]
        if run.get("reference_review_status") != "approved":
            return {"ok": False, "error": "RUN_NOT_APPROVABLE", "reason": "REFERENCE_DIFF_BLOCKING"}
        if run.get("business_review_status") != "approved" or blockers:
            return {"ok": False, "error": "RUN_NOT_APPROVABLE", "reason": "BUSINESS_REVIEW_REQUIRED", "issues": blockers}
        review_actors = {
            str((run.get("reference_review") or {}).get("actor") or ""),
            str((run.get("business_review") or {}).get("actor") or ""),
        }
        if actor in review_actors:
            return {"ok": False, "error": "RUN_NOT_APPROVABLE", "reason": "APPROVER_SEPARATION_REQUIRED"}
        approved_at = _now()
        for other_id, other in state["runs"].items():
            if other_id == run_id or other.get("review_status") != "approved":
                continue
            other.update({
                "review_status": "superseded",
                "lifecycle_status": "superseded",
                "superseded_at": approved_at,
                "superseded_by_run_id": run_id,
            })
            other.setdefault("state_history", []).append(_history_event(
                "superseded", actor=actor, request_id=request_id,
                superseded_by_run_id=run_id,
            ))
            for artifact in state["artifacts"].values():
                if str(artifact.get("run_id") or "") != other_id:
                    continue
                artifact["current_status"] = "superseded"
                artifact["superseded_at"] = approved_at
                artifact["superseded_by_run_id"] = run_id
                artifact.setdefault("state_history", []).append(_history_event(
                    "superseded", actor=actor, request_id=request_id,
                    superseded_by_run_id=run_id,
                ))
        run.update({
            "review_status": "approved", "lifecycle_status": "approved",
            "approved_at": approved_at, "approved_by": actor, "approval_note": note,
        })
        approval = {"actor": actor, "note": note, "request_id": request_id, "approved_at": approved_at}
        run.setdefault("approval_history", []).append(approval)
        run.setdefault("state_history", []).append(_history_event("approved", actor=actor, request_id=request_id))
        _save(workspace_id, state)
        return {"ok": True, "run_id": run_id, "review_status": "approved", "approved_at": run["approved_at"]}


def reject_run(
    workspace_id: str,
    run_id: str,
    *,
    actor: str,
    note: str = "",
    request_id: str = "",
) -> dict[str, Any]:
    """Record a final rejection without allowing an approved run to regress."""

    actor = str(actor or "").strip()
    note = str(note or "")
    if not actor:
        return {"ok": False, "error": "authenticated_actor_required"}
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        run = state["runs"].get(run_id)
        if not run:
            return {"ok": False, "error": "run_not_found"}
        if run.get("review_status") == "approved":
            return {"ok": False, "error": "approved_run_immutable"}
        if run.get("review_status") == "rejected":
            if (
                str(run.get("rejected_by") or "") == actor
                and str(run.get("rejection_note") or "") == note
            ):
                return {
                    "ok": True,
                    "run_id": run_id,
                    "review_status": "rejected",
                    "rejected_at": run.get("rejected_at"),
                    "idempotent_replay": True,
                }
            return {"ok": False, "error": "rejected_run_immutable"}
        rejected_at = _now()
        rejection = {
            "actor": actor,
            "note": note,
            "request_id": request_id,
            "rejected_at": rejected_at,
        }
        run.update({
            "review_status": "rejected",
            "lifecycle_status": "rejected",
            "rejected_at": rejected_at,
            "rejected_by": actor,
            "rejection_note": note,
        })
        run.setdefault("rejection_history", []).append(rejection)
        run.setdefault("state_history", []).append(_history_event(
            "rejected", actor=actor, request_id=request_id,
        ))
        _open_issue(run, "RUN_REJECTED", "财务运行已被最终驳回，须创建新修订后重新运行")
        _save(workspace_id, state)
        return {
            "ok": True,
            "run_id": run_id,
            "review_status": "rejected",
            "rejected_at": rejected_at,
        }


def max_price(
    workspace_id: str, run_id: str, *, target_irr: float | None = None,
    min_dscr: float | None = None,
    lower: float = 0.0, upper: float | None = None,
    actor: str = "system:finance-model", request_id: str = "",
) -> dict[str, Any]:
    actor = str(actor or "system:finance-model").strip() or "system:finance-model"
    run = get_run(workspace_id, run_id)
    if not run:
        return {"ok": False, "error": "run_not_found"}
    spec = (
        _load(workspace_id)["specs"].get(
            str(run.get("spec_id") or "")
        )
        or {}
    ).get("spec")
    if not isinstance(spec, dict):
        return {"ok": False, "error": "spec_snapshot_missing"}
    thresholds, threshold_error = _decision_thresholds(spec)
    if threshold_error:
        return {
            "ok": False, "error": threshold_error,
            "decision_thresholds": spec.get("decision_thresholds"),
        }
    expected_target = float(thresholds["target_project_irr"] or 0.0)
    expected_dscr = thresholds["minimum_dscr"]
    try:
        effective_target = expected_target if target_irr is None else float(target_irr)
        effective_dscr = expected_dscr if min_dscr is None else float(min_dscr)
        effective_lower = float(lower)
        effective_upper = float(upper) if upper is not None else None
    except (TypeError, ValueError):
        return {
            "ok": False, "error": "MAX_PRICE_PARAMETERS_INVALID",
            "parameters": {
                "target_irr": target_irr, "min_dscr": min_dscr,
                "lower": lower, "upper": upper,
            },
        }
    parameters = {
        "target_irr": effective_target, "min_dscr": effective_dscr,
        "lower": effective_lower, "upper": effective_upper,
    }
    if (
        not _same_optional_number(effective_target, expected_target)
        or not _same_optional_number(effective_dscr, expected_dscr)
    ):
        return {
            "ok": False, "error": "MAX_PRICE_THRESHOLD_MISMATCH",
            "parameters": parameters, "decision_thresholds": thresholds,
        }
    if (
        not 0 < effective_target <= 5
        or (effective_dscr is not None and effective_dscr <= 0)
        or effective_lower < 0
        or (effective_upper is not None and effective_upper <= effective_lower)
    ):
        return {"ok": False, "error": "MAX_PRICE_PARAMETERS_INVALID", "parameters": parameters}
    existing = run.get("max_acquisition_price_analysis") or {}
    if existing.get("parameters") == parameters and isinstance(existing.get("result"), dict):
        return {
            "ok": True,
            "run_id": run_id,
            "analysis_hash": existing.get("analysis_hash"),
            "decision_status": existing.get("decision_status") or "candidate",
            "idempotent_replay": True,
            **copy.deepcopy(existing.get("result") or {}),
        }
    if run.get("review_status") == "approved":
        return {"ok": False, "error": "APPROVED_RUN_IMMUTABLE"}
    baseline_analysis_hash = str(existing.get("analysis_hash") or "")
    solved = solve_max_acquisition_price(
        spec, target_irr=effective_target, min_dscr=effective_dscr,
        lower=effective_lower, upper=effective_upper,
    )
    analysis = {
        "status": "calculated",
        "decision_status": "candidate",
        "parameters": parameters,
        "result": copy.deepcopy(solved),
        "calculated_at": _now(),
        "calculation_timing": "before_approval",
        "calculated_by": actor,
        "request_id": request_id,
    }
    analysis["analysis_hash"] = _hash(analysis)
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        current = state["runs"].get(run_id)
        if not current:
            return {"ok": False, "error": "run_not_found"}
        if current.get("spec_hash") != run.get("spec_hash") or current.get("input_hash") != run.get("input_hash"):
            return {"ok": False, "error": "RUN_SPEC_MISMATCH"}
        if current.get("review_status") == "approved":
            return {"ok": False, "error": "APPROVED_RUN_IMMUTABLE"}
        current_analysis = current.get("max_acquisition_price_analysis") or {}
        if (
            current_analysis.get("parameters") == parameters
            and isinstance(current_analysis.get("result"), dict)
        ):
            return {
                "ok": True,
                "run_id": run_id,
                "analysis_hash": current_analysis.get("analysis_hash"),
                "decision_status": current_analysis.get("decision_status") or "candidate",
                "idempotent_replay": True,
                **copy.deepcopy(current_analysis.get("result") or {}),
            }
        if str(current_analysis.get("analysis_hash") or "") != baseline_analysis_hash:
            return {
                "ok": False,
                "error": "WORKSPACE_VERSION_CONFLICT",
                "resource_id": run_id,
            }
        current["max_acquisition_price_analysis"] = analysis
        _close_issue(current, "MAX_PRICE_REQUIRED", actor="system", reason="最高可接受价已按决策阈值完成求解")
        if solved.get("feasible") and solved.get("converged"):
            _close_issue(current, "MAX_PRICE_NOT_FEASIBLE", actor="system", reason="最高可接受价求解可行且收敛")
            _open_issue(current, "MAX_PRICE_REVIEW_REQUIRED", "最高可接受价尚未完成业务复核")
        else:
            _close_issue(current, "MAX_PRICE_REVIEW_REQUIRED", actor="system", reason="最高价求解结果已变更为不可行或未收敛")
            _open_issue(current, "MAX_PRICE_NOT_FEASIBLE", "最高可接受价求解不可行或未收敛")
        if current.get("business_review_status") == "approved":
            current["business_review_status"] = "pending"
            _open_issue(current, "BUSINESS_REVIEW_REQUIRED", "最高可接受价新增或变更后须重新业务复核")
            current["lifecycle_status"] = "business_review_pending"
            current.setdefault("state_history", []).append(_history_event(
                "business_review_pending", actor=actor, request_id=request_id,
                reason="max_price_changed",
            ))
        current.setdefault("max_acquisition_price_history", []).append(copy.deepcopy(analysis))
        current.setdefault("state_history", []).append(_history_event(
            "max_acquisition_price_calculated", actor=actor,
            request_id=request_id,
            analysis_hash=analysis["analysis_hash"],
            calculation_timing=analysis["calculation_timing"],
        ))
        _save(workspace_id, state)
    return {"ok": True, "run_id": run_id, "analysis_hash": analysis["analysis_hash"], **solved}


def build_acquisition_report_data(
    workspace_id: str,
    run: dict[str, Any],
) -> dict[str, Any]:
    """Project an approved run/spec into the formal acquisition report contract."""

    state = _load(workspace_id)
    spec_row = state["specs"].get(str(run.get("spec_id") or "")) or {}
    spec = copy.deepcopy(spec_row.get("spec") or {})
    if not isinstance(spec, dict) or _hash(spec) != run.get("spec_hash"):
        raise RuntimeError("approved run spec snapshot is missing or does not match spec_hash")
    evidence_ok, current_evidence = _current_evidence_matches_run(
        workspace_id,
        run,
        spec,
    )
    if not evidence_ok:
        raise RuntimeError(
            "approved run evidence binding is stale or no longer formally approved: "
            f"snapshot={run.get('evidence_binding_hash')} current={current_evidence.get('binding_hash')}"
        )
    transaction = copy.deepcopy(spec.get("transaction") or {})
    source_ledger = [
        {
            "field": binding.get("source_path"),
            "evidence_ids": [binding.get("evidence_id")],
            "source": {
                "file_id": binding.get("file_id"),
                "locator": binding.get("locator"),
                "source_sha256": binding.get("source_sha256"),
                "source_size_bytes": binding.get("source_size_bytes"),
                "parse_job": binding.get("parse_job"),
                "attempt": binding.get("attempt"),
                "review_revision": copy.deepcopy(binding.get("review_revision")),
                "binding_hash": binding.get("binding_hash"),
            },
            "review_status": "server_bound_approved",
        }
        for binding in current_evidence.get("bindings") or []
    ]
    parties = copy.deepcopy(spec.get("project_parties") or [])
    licenses = copy.deepcopy(transaction.get("licenses") or [])
    for party in parties:
        if isinstance(party, dict) and "license_holder" in (party.get("roles") or []):
            licenses.append({
                "license_type": "license_holder_role",
                "holder_id": party.get("entity_id"),
                "holder_name": party.get("name"),
                "status": party.get("status"),
                "evidence_ids": copy.deepcopy(party.get("evidence_ids") or []),
            })
    matrices = [
        copy.deepcopy(value)
        for value in state["scenario_matrices"].values()
        if value.get("run_id") == run.get("run_id")
        and value.get("spec_hash") == run.get("spec_hash")
        and value.get("input_hash") == run.get("input_hash")
    ]
    matrices.sort(key=lambda value: str(value.get("created_at") or ""))
    result = run.get("result") or {}
    asset_type = str(result.get("asset_type") or spec.get("asset_type") or "hotel_lease")
    valuation = float(transaction.get("valuation_value") or 0.0)
    purchase = float(result.get("purchase_price_wan") or transaction.get("purchase_price") or 0.0)
    report = {
        "schema_version": "asset_acquisition_report_data.v1",
        "asset_type": asset_type,
        "bindings": {
            "workspace_id": workspace_id,
            "run_id": run.get("run_id"),
            "spec_id": run.get("spec_id"),
            "spec_hash": run.get("spec_hash"),
            "input_hash": run.get("input_hash"),
            "model_version": run.get("model_version"),
            "spec_snapshot_hash": run.get("spec_snapshot_hash"),
            "evidence_binding_version": run.get("evidence_binding_version"),
            "evidence_binding_hash": run.get("evidence_binding_hash"),
            "approved_by": run.get("approved_by"),
            "approved_at": run.get("approved_at"),
        },
        "source_processing_ledger": source_ledger,
        "party_relationships": parties,
        "asset_boundary": copy.deepcopy(transaction.get("asset_scope") or []),
        "license_ledger": licenses,
        "historical_financial_comparison": copy.deepcopy(spec.get("historical_statements") or []),
        "valuation_transaction_bridge": {
            "valuation_value_wan": valuation,
            "purchase_price_wan": purchase,
            "purchase_price_vs_valuation_wan": purchase - valuation,
            "transaction_tax_wan": result.get("transaction_tax_wan"),
            "total_acquisition_cost_wan": result.get("total_acquisition_cost_wan"),
            "valuation_date": transaction.get("valuation_date"),
            "closing_date": transaction.get("closing_date"),
        },
        "maximum_acceptable_price": copy.deepcopy(
            run.get("max_acquisition_price_analysis")
            or {"status": "not_calculated", "decision_status": "pending_business_thresholds"}
        ),
        "scenario_matrices": matrices,
        "red_flags": copy.deepcopy(transaction.get("red_flags") or []),
        "closing_conditions": copy.deepcopy(transaction.get("closing_conditions") or []),
        "veto_items": copy.deepcopy(transaction.get("veto_items") or []),
        "golden_regression_status": {
            "reference_review_status": run.get("reference_review_status"),
            "reference_kind": (run.get("reference_review") or {}).get("reference_kind"),
            "reference_hash": (run.get("reference_review") or {}).get("reference_hash"),
            "business_review_status": run.get("business_review_status"),
            "approval_status": run.get("review_status"),
            "evidence_status": current_evidence.get("status"),
            "evidence_formal_ok": bool(current_evidence.get("formal_ok")),
            "evidence_binding_hash": current_evidence.get("binding_hash"),
            "open_blocking_issues": [
                copy.deepcopy(issue) for issue in (run.get("issues") or [])
                if issue.get("blocking") and issue.get("status") == "open"
            ],
        },
    }
    if asset_type == "solar_power":
        report["solar_operation"] = copy.deepcopy(result.get("solar_operation") or {})
        report["solar_operating_ledger"] = copy.deepcopy(result.get("annual_summary") or [])
    else:
        report["lease_ledger"] = copy.deepcopy(
            (spec.get("lease_portfolio") or {}).get("units") or []
        )
    report["report_data_hash"] = _hash(report)
    return report


def render_markdown(run: dict[str, Any], report_data: dict[str, Any] | None = None) -> str:
    result = run.get("result") or {}
    ind = result.get("indicators") or {}
    report_data = report_data or {}
    asset_type = str(result.get("asset_type") or report_data.get("asset_type") or "hotel_lease")
    is_solar = asset_type == "solar_power"

    def text(value: Any) -> str:
        if isinstance(value, (dict, list)):
            value = json.dumps(value, ensure_ascii=False, sort_keys=True, default=str)
        return str(value if value not in (None, "") else "—").replace("|", "¦").replace("\n", " ")

    lines = [
        "# 资产收购可行性研究报告",
        "",
        f"> 财务运行：`{run.get('run_id')}`；场景：`{run.get('scenario_id')}`；模型：`{run.get('model_version')}`",
        f"> Spec：`{run.get('spec_hash')}`；事实版本：`{run.get('spec_id')}`",
        f"> 证据绑定：`{run.get('evidence_binding_hash')}`；版本：`{run.get('evidence_binding_version')}`",
        "",
        "## 一、交易概览",
        "",
        f"- 收购价格：{result.get('purchase_price_wan', 0):,.2f} 万元",
        f"- 交易税费：{result.get('transaction_tax_wan', 0):,.2f} 万元",
        f"- 总收购成本：{result.get('total_acquisition_cost_wan', 0):,.2f} 万元",
        "",
        "## 二、核心财务指标",
        "",
        "| 指标 | 结果 |",
        "|---|---:|",
        f"| 项目 IRR | {_pct(ind.get('project_irr_pct'))} |",
        f"| 资本金 IRR | {_pct(ind.get('equity_irr_pct'))} |",
        f"| NPV | {_num(ind.get('npv_wan'))} 万元 |",
        f"| 最低 DSCR | {_num(ind.get('minimum_dscr'))} |",
        f"| 最低 ICR | {_num(ind.get('minimum_icr'))} |",
    ]
    if is_solar:
        solar = result.get("solar_operation") or {}
        lines.extend([
            "",
            "## 三、光伏电站运营",
            "",
            f"- 装机容量：{_num(solar.get('installed_capacity_mw'))} MW",
            f"- 基准发电量：{_num(solar.get('base_generation_mwh'))} MWh",
            f"- 上网电价：{_num(solar.get('tariff_yuan_per_kwh'))} 元/kWh",
            f"- 限电率：{_pct_ratio(solar.get('curtailment_rate'))}",
            f"- 年衰减率：{_pct_ratio(solar.get('degradation_rate'))}",
            "",
            "| 年度 | 理论发电量(MWh) | 上网电量(MWh) | 售电收入(万元) | 运维费(万元) | 维护性资本开支(万元) | 所得税(万元) | 债务服务(万元) | 项目现金流(万元) | 资本金现金流(万元) |",
            "|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|",
        ])
        for row in result.get("annual_summary") or []:
            lines.append(
                f"| {text(row.get('year'))} | {_num(row.get('gross_generation_mwh'))} | "
                f"{_num(row.get('sold_generation_mwh'))} | {_num(row.get('revenue_wan'))} | "
                f"{_num(row.get('operating_cost_wan'))} | {_num(row.get('maintenance_capex_wan'))} | "
                f"{_num(row.get('income_tax_wan'))} | {_num(row.get('debt_service_wan'))} | "
                f"{_num(row.get('project_cf_wan'))} | {_num(row.get('equity_cf_wan'))} |"
            )
    else:
        lines.extend([
            f"| 最低租金覆盖率 | {_num(ind.get('minimum_tenant_rent_coverage'))} |",
            f"| 租约覆盖年限 | {_num(ind.get('lease_coverage_years'))} 年 |",
            f"| 合同收入占比 | {_pct_ratio(ind.get('contract_income_ratio'))} |",
            f"| 未锁定收入占比 | {_pct_ratio(ind.get('unlocked_income_ratio'))} |",
            "",
            "## 三、酒店经营与租约",
            "",
            "| 年度 | ADR(元) | 入住率 | RevPAR(元) | EBITDAR(万元) | 可支付租金(万元) |",
            "|---:|---:|---:|---:|---:|---:|",
        ])
        for row in (result.get("hotel_operation") or {}).get("years") or []:
            lines.append(
                f"| {row['year']} | {row['adr_yuan']:.2f} | {row['occupancy']:.2%} | "
                f"{row['revpar_yuan']:.2f} | {row['ebitdar_wan']:.2f} | {row['affordable_rent_wan']:.2f} |"
            )
    lines.extend(["", "## 四、主体关系", ""])
    parties = report_data.get("party_relationships") or []
    if parties:
        lines.extend(["| 主体ID | 名称 | 角色 | 状态 | 证据 |", "|---|---|---|---|---|"])
        for row in parties:
            lines.append(
                f"| {text(row.get('entity_id'))} | {text(row.get('name'))} | "
                f"{text(row.get('roles') or [])} | {text(row.get('status'))} | "
                f"{text(row.get('evidence_ids') or [])} |"
            )
    else:
        lines.append("- 无已绑定主体。")
    lines.extend(["", "## 五、资产边界与权证许可", ""])
    assets = report_data.get("asset_boundary") or []
    if assets:
        lines.extend(["| 资产ID | 类型 | 是否纳入 | 面积(㎡) | 状态 | 冲突/裁决 |", "|---|---|---:|---:|---|---|"])
        for row in assets:
            lines.append(
                f"| {text(row.get('scope_id'))} | {text(row.get('type'))} | "
                f"{text(row.get('included'))} | {text(row.get('area_sqm'))} | "
                f"{text(row.get('status'))} | {text(row.get('conflicts') or row.get('resolution'))} |"
            )
    licenses = report_data.get("license_ledger") or []
    lines.extend(["", "### 5.1 权证/许可台账", ""])
    if licenses:
        lines.extend(["| 许可/角色 | 持有人 | 状态 | 证据 |", "|---|---|---|---|"])
        for row in licenses:
            lines.append(
                f"| {text(row.get('license_type') or row.get('type'))} | "
                f"{text(row.get('holder_name') or row.get('holder_id'))} | "
                f"{text(row.get('status'))} | {text(row.get('evidence_ids') or [])} |"
            )
    else:
        lines.append("- 无已绑定许可记录。")
    if is_solar:
        lines.extend(["", "## 六、光伏运营关键参数", ""])
        solar = report_data.get("solar_operation") or {}
        lines.extend([
            "| 参数 | 值 |", "|---|---:|",
            f"| 装机容量(MW) | {text(solar.get('installed_capacity_mw'))} |",
            f"| 基准发电量(MWh) | {text(solar.get('base_generation_mwh'))} |",
            f"| 上网电价(元/kWh) | {text(solar.get('tariff_yuan_per_kwh'))} |",
            f"| 限电率 | {text(solar.get('curtailment_rate'))} |",
            f"| 年衰减率 | {text(solar.get('degradation_rate'))} |",
        ])
    else:
        lines.extend(["", "## 六、租约台账", ""])
        leases = report_data.get("lease_ledger") or []
        if leases:
            lines.extend(["| 单元 | 位置 | 面积(㎡) | 出租人/承租人 | 起止日 | 基础租金(万元) | 证据 |", "|---|---|---:|---|---|---:|---|"])
            for row in leases:
                lines.append(
                    f"| {text(row.get('unit_id'))} | {text(row.get('asset_location'))} | "
                    f"{text(row.get('area_sqm'))} | {text(row.get('lessor_id'))}/{text(row.get('lessee_id'))} | "
                    f"{text(row.get('start_date'))}—{text(row.get('end_date'))} | "
                    f"{text(row.get('base_rent_wan'))} | {text(row.get('evidence_ids') or [])} |"
                )
        else:
            lines.append("- 无已绑定租约。")
    lines.extend(["", "## 七、历史财务对比与勾稽", ""])
    statements = report_data.get("historical_financial_comparison") or []
    if statements:
        lines.extend(["| 主体 | 期间 | 报表类型 | 人工复核 | 勾稽 | 来源定位 |", "|---|---|---|---|---|---|"])
        for row in statements:
            lines.append(
                f"| {text(row.get('entity_id'))} | {text(row.get('period_start'))}—{text(row.get('period_end'))} | "
                f"{text(row.get('statement_type'))} | {text(row.get('manual_review_status'))} | "
                f"{text(row.get('reconciliation'))} | {text(row.get('source_locators') or [])} |"
            )
    else:
        lines.append("- 无已绑定历史报表。")
    bridge = report_data.get("valuation_transaction_bridge") or {}
    max_price_analysis = report_data.get("maximum_acceptable_price") or {}
    max_price_result = max_price_analysis.get("result") or {}
    lines.extend([
        "", "## 八、估值—成交价桥接", "",
        "| 项目 | 金额/日期 |", "|---|---:|",
        f"| 评估值(万元) | {text(bridge.get('valuation_value_wan'))} |",
        f"| 收购价(万元) | {text(bridge.get('purchase_price_wan'))} |",
        f"| 收购价较评估值差额(万元) | {text(bridge.get('purchase_price_vs_valuation_wan'))} |",
        f"| 交易税费(万元) | {text(bridge.get('transaction_tax_wan'))} |",
        f"| 总收购成本(万元) | {text(bridge.get('total_acquisition_cost_wan'))} |",
        f"| 评估日/交割日 | {text(bridge.get('valuation_date'))}/{text(bridge.get('closing_date'))} |",
        "", "### 8.1 最高可接受收购价", "",
        f"- 计算状态：{text(max_price_analysis.get('status'))}；决策状态：{text(max_price_analysis.get('decision_status'))}。",
        f"- 最高可接受价：{text(max_price_result.get('max_acquisition_price_wan'))} 万元。",
        f"- 目标IRR/最低DSCR：{text((max_price_analysis.get('parameters') or {}).get('target_irr'))}/"
        f"{text((max_price_analysis.get('parameters') or {}).get('min_dscr'))}。",
        f"- 求解哈希：`{text(max_price_analysis.get('analysis_hash'))}`；计算时点：{text(max_price_analysis.get('calculation_timing'))}。",
        "", "## 九、独立情景矩阵", "",
    ])
    matrices = report_data.get("scenario_matrices") or []
    if matrices:
        if is_solar:
            lines.extend(["| 矩阵/场景 | 变更 | 收购价(万元) | 上网电价(元/kWh) | 年发电量(MWh) | 利用小时 | 年运维费(万元) | 项目IRR | 最低DSCR |", "|---|---|---:|---:|---:|---:|---:|---:|---:|"])
        else:
            lines.extend(["| 矩阵/场景 | 变更 | 收购价(万元) | 市场租金 | 入住率 | 项目IRR | 最低DSCR |", "|---|---|---:|---:|---|---:|---:|"])
        for matrix in matrices:
            for row in matrix.get("rows") or []:
                indicators = row.get("indicators") or {}
                prefix = (
                    f"| {text(matrix.get('matrix_id'))}/{text(row.get('scenario_id'))} | "
                    f"{text(row.get('changes'))} | {text(row.get('purchase_price_wan'))} | "
                )
                if is_solar:
                    lines.append(
                        prefix + f"{text(row.get('tariff_yuan_per_kwh'))} | "
                        f"{text(row.get('annual_generation_mwh'))} | {text(row.get('utilization_hours'))} | "
                        f"{text(row.get('annual_opex_wan'))} | {text(indicators.get('project_irr_pct'))} | "
                        f"{text(indicators.get('minimum_dscr'))} |"
                    )
                else:
                    lines.append(
                        prefix + f"{text(row.get('market_rent'))} | {text(row.get('occupancy'))} | "
                        f"{text(indicators.get('project_irr_pct'))} | {text(indicators.get('minimum_dscr'))} |"
                    )
    else:
        lines.append("- 本批准 run 未绑定独立情景矩阵。")
    lines.extend(["", "## 十、红旗、成交条件与否决事项", ""])
    for row in report_data.get("red_flags") or []:
        lines.append(f"- 红旗 `{text(row.get('code'))}`：{text(row.get('status'))}；{text(row.get('resolution'))}")
    for item in report_data.get("closing_conditions") or []:
        lines.append(f"- 成交条件：{text(item)}")
    for item in report_data.get("veto_items") or []:
        lines.append(f"- 否决事项：{text(item)}")
    lines.extend(["", "## 十一、资料处理与字段证据台账", ""])
    ledger = report_data.get("source_processing_ledger") or []
    if ledger:
        lines.extend(["| 字段 | 证据ID | 来源 | 复核状态 |", "|---|---|---|---|"])
        for row in ledger:
            lines.append(
                f"| {text(row.get('field'))} | {text(row.get('evidence_ids'))} | "
                f"{text(row.get('source'))} | {text(row.get('review_status'))} |"
            )
    lines.extend([
        "", "## 十二、金标回归、复核与发布边界", "",
        f"- 参考轨复核：{run.get('reference_review_status')}",
        f"- 业务复核：{run.get('business_review_status')}",
        f"- 批准状态：{run.get('review_status')}",
        f"- 批准人：{run.get('approved_by')}；批准时间：{run.get('approved_at')}",
        "- 本报告数字仅绑定上述 run；未批准 run 不得正式发布。",
    ])
    return "\n".join(lines) + "\n"


def enqueue_artifact(
    workspace_id: str, run_id: str, *, idempotency_key: str = "", request_id: str = "",
) -> dict[str, Any]:
    """Create a durable artifact job that can be polled before rendering."""

    request_id = request_id or f"req_{uuid.uuid4().hex}"
    run = get_run(workspace_id, run_id)
    if not run:
        return {"ok": False, "error": "run_not_found"}
    if run.get("review_status") != "approved":
        return {"ok": False, "error": "RUN_NOT_APPROVABLE", "reason": "approved_run_required"}
    body_hash = _hash({
        "run_id": run_id, "spec_hash": run.get("spec_hash"), "fact_revision": run.get("spec_id"),
        "spec_snapshot_hash": run.get("spec_snapshot_hash"),
        "evidence_binding_hash": run.get("evidence_binding_hash"),
        "template_version": "asset_acquisition.v2",
    })
    scope = f"artifact:{idempotency_key}" if idempotency_key else ""
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        prior = _active_idempotency_record(state["idempotency"], scope) if scope else None
        if prior:
            if prior.get("body_hash") != body_hash:
                return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior.get("artifact_id", "")}
            existing = state["artifacts"].get(prior.get("artifact_id"))
            if not existing:
                raise RuntimeError("artifact idempotency record points to a missing artifact")
            return {**existing, "idempotent_replay": True}
        artifact_id = f"artifact_{uuid.uuid4().hex}"
        job_id = f"artifact_job_{uuid.uuid4().hex}"
        created_at = _now()
        row = {
            "ok": True, "artifact_id": artifact_id, "artifact_job_id": job_id,
            "status": "queued", "progress": 0, "type": "asset_acquisition", "run_id": run_id,
            "spec_hash": run.get("spec_hash"), "fact_revision": run.get("spec_id"),
            "spec_snapshot_hash": run.get("spec_snapshot_hash"),
            "evidence_binding_version": run.get("evidence_binding_version"),
            "evidence_binding_hash": run.get("evidence_binding_hash"),
            "template_version": "asset_acquisition.v2", "created_at": created_at,
            "updated_at": created_at, "request_id": request_id, "files": [],
            "numeric_consistency": "pending", "release_status": "pending",
            "state_history": [
                _history_event("queued", request_id=request_id, run_id=run_id),
            ],
        }
        state["artifacts"][artifact_id] = row
        if scope:
            state["idempotency"][scope] = _idempotency_record(
                scope, body_hash, artifact_id=artifact_id,
            )
        _save(workspace_id, state)
        return row


def _bind_succeeded_artifact(
    workspace_id: str,
    run: dict[str, Any],
    artifact: dict[str, Any],
) -> dict[str, Any]:
    """Bind a successful formal pack to the report-side finance revision."""

    from lvke_mcp.domains.reports import artifacts as report_artifacts

    expected = {
        "finance_run_id": str(run.get("run_id") or ""),
        "input_hash": run.get("input_hash"),
        "spec_hash": run.get("spec_hash"),
        "spec_id": run.get("spec_id"),
        "fact_revision": run.get("spec_id"),
        "spec_snapshot_hash": run.get("spec_snapshot_hash"),
        "evidence_binding_version": run.get("evidence_binding_version"),
        "evidence_binding_hash": run.get("evidence_binding_hash"),
        "model_version": run.get("model_version"),
        "template_version": artifact.get("template_version"),
        "review_status": "approved",
        "artifact_id": artifact.get("artifact_id"),
        "artifact_job_id": artifact.get("artifact_job_id"),
        "artifact_status": "succeeded",
        "report_data_hash": artifact.get("report_data_hash"),
        "binding_kind": "asset_acquisition",
    }
    report_artifacts.bind_finance_run(
        workspace_id,
        expected["finance_run_id"],
        section="asset_acquisition_artifact",
        fin={
            key: value for key, value in expected.items()
            if key != "finance_run_id"
        } | {"assurance_level": "review_grade"},
    )
    actual = report_artifacts.load(
        workspace_id,
        "finance_binding",
        {},
    ) or {}
    mismatches = [
        {"field": key, "expected": value, "actual": actual.get(key)}
        for key, value in expected.items()
        if actual.get(key) != value
    ]
    if mismatches:
        return {
            "ok": False,
            "error": "ARTIFACT_BINDING_FAILED",
            "reason": "finance_binding_mismatch",
            "mismatches": mismatches,
        }
    return {"ok": True, "binding": actual}


def mark_artifact_release_ready(
    workspace_id: str,
    artifact_id: str,
    *,
    actor: str = "system",
    request_id: str = "",
) -> dict[str, Any]:
    """Run the unified publish gate and audit the ``release_ready`` state."""

    from lvke_mcp.domains.finance import gate as finance_gate

    artifact = get_artifact(workspace_id, artifact_id)
    if not artifact:
        return {"ok": False, "error": "ARTIFACT_NOT_FOUND"}
    if artifact.get("status") != "succeeded":
        return {
            "ok": False, "error": "ARTIFACT_NOT_READY",
            "status": artifact.get("status"),
        }
    gate_result = finance_gate.assert_publish_finance_binding(
        workspace_id,
        strict=True,
    )
    if (
        not gate_result.get("ok")
        or str(gate_result.get("artifact_id") or "") != artifact_id
    ):
        blockers = list(gate_result.get("blockers") or [])
        if gate_result.get("ok") and str(gate_result.get("artifact_id") or "") != artifact_id:
            blockers.append({
                "code": "finance_artifact_binding_stale",
                "message": "待签发工件不是 finance_binding 当前绑定的正式工件",
            })
        with _state_guard(workspace_id):
            state = _load(workspace_id)
            stored = state["artifacts"].get(artifact_id)
            if stored and stored.get("release_status") != "released":
                stored.update({
                    "release_status": "blocked",
                    "release_gate": {
                        "ok": False,
                        "checked_at": _now(),
                        "blockers": blockers,
                        "warnings": list(gate_result.get("warnings") or []),
                    },
                    "updated_at": _now(),
                })
                _save(workspace_id, state)
        return {
            "ok": False,
            "error": "PUBLISH_GATE_BLOCKED",
            "blockers": blockers,
            "warnings": list(gate_result.get("warnings") or []),
        }

    checked_at = _now()
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        stored = state["artifacts"].get(artifact_id)
        if not stored:
            return {"ok": False, "error": "ARTIFACT_NOT_FOUND"}
        if stored.get("release_status") == "released":
            return {
                "ok": True, "artifact_id": artifact_id,
                "run_id": stored.get("run_id"), "release_status": "released",
                "idempotent_replay": True,
            }
        if stored.get("release_status") == "release_ready":
            return {
                "ok": True, "artifact_id": artifact_id,
                "run_id": stored.get("run_id"), "release_status": "release_ready",
                "release_ready_at": stored.get("release_ready_at"),
                "idempotent_replay": True,
            }
        stored.update({
            "release_status": "release_ready",
            "release_ready_at": checked_at,
            "release_ready_by": actor or "system",
            "release_gate": {
                "ok": True,
                "checked_at": checked_at,
                "warnings": list(gate_result.get("warnings") or []),
                "approved_run_id": gate_result.get("approved_run_id"),
                "artifact_id": artifact_id,
            },
            "updated_at": checked_at,
        })
        stored.setdefault("state_history", []).append(_history_event(
            "release_ready", actor=actor or "system", request_id=request_id,
            run_id=stored.get("run_id"), artifact_id=artifact_id,
        ))
        run = state["runs"].get(str(stored.get("run_id") or ""))
        if run and run.get("lifecycle_status") != "released":
            run["lifecycle_status"] = "release_ready"
            run.setdefault("state_history", []).append(_history_event(
                "release_ready", actor=actor or "system", request_id=request_id,
                artifact_id=artifact_id,
            ))
        _save(workspace_id, state)
    return {
        "ok": True, "artifact_id": artifact_id,
        "run_id": artifact.get("run_id"), "release_status": "release_ready",
        "release_ready_at": checked_at,
    }


def execute_queued_artifact(
    workspace_id: str,
    artifact_id: str,
) -> None:
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        row = state["artifacts"].get(artifact_id)
        if not row or row.get("status") in {"succeeded", "failed", "cancelled"}:
            return
        row.update({"status": "running", "progress": 10, "updated_at": _now()})
        run_id = str(row.get("run_id") or "")
        job_id = str(row.get("artifact_job_id") or "")
        request_id = str(row.get("request_id") or "")
        _save(workspace_id, state)
    try:
        result = generate_artifacts(
            workspace_id, run_id, request_id=request_id,
            artifact_id=artifact_id, artifact_job_id=job_id,
        )
        if result.get("ok"):
            return
        raw_code = str(result.get("error") or "ARTIFACT_MISMATCH")
        code = (
            raw_code
            if re.fullmatch(r"[A-Z][A-Z0-9_]{1,79}", raw_code)
            else "ARTIFACT_MISMATCH"
        )
        error = {
            "code": code,
            "message": _ARTIFACT_GENERATION_FAILURE_MESSAGE,
            "retryable": False,
        }
    except Exception as exc:  # noqa: BLE001
        _LOG.warning(
            "asset acquisition artifact generation failed; artifact_id=%s "
            "error_type=%s",
            artifact_id,
            type(exc).__name__,
        )
        error = {
            "code": "ARTIFACT_GENERATION_FAILED",
            "message": _ARTIFACT_GENERATION_FAILURE_MESSAGE,
            "retryable": False,
        }
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        row = state["artifacts"].get(artifact_id)
        if row and row.get("status") != "cancelled":
            row.update({
                "ok": False, "status": "failed", "progress": 100,
                "numeric_consistency": "failed", "error": error, "updated_at": _now(),
            })
            _save(workspace_id, state)


def generate_artifacts(
    workspace_id: str, run_id: str, *, idempotency_key: str = "", request_id: str = "",
    artifact_id: str = "", artifact_job_id: str = "",
) -> dict[str, Any]:
    """Generate an atomically published, consistency-checked approved-run pack."""

    run = get_run(workspace_id, run_id)
    if not run:
        return {"ok": False, "error": "run_not_found"}
    if run.get("review_status") != "approved":
        return {"ok": False, "error": "RUN_NOT_APPROVABLE", "reason": "approved_run_required"}
    body_hash = _hash({
        "run_id": run_id, "spec_hash": run.get("spec_hash"), "fact_revision": run.get("spec_id"),
        "spec_snapshot_hash": run.get("spec_snapshot_hash"),
        "evidence_binding_hash": run.get("evidence_binding_hash"),
        "template_version": "asset_acquisition.v2",
    })
    scope = f"artifact:{idempotency_key}" if idempotency_key else ""
    if scope:
        with _state_guard(workspace_id):
            state = _load(workspace_id)
            prior = _active_idempotency_record(state["idempotency"], scope)
            if prior:
                if prior.get("body_hash") != body_hash:
                    return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior.get("artifact_id", "")}
                existing = state["artifacts"].get(prior.get("artifact_id"))
                if not existing:
                    raise RuntimeError("artifact idempotency record points to a missing artifact")
                return {**existing, "idempotent_replay": True}

    artifact_id = artifact_id or f"artifact_{uuid.uuid4().hex}"
    artifacts_root = _root(workspace_id) / "artifacts"
    artifacts_root.mkdir(parents=True, exist_ok=True)
    staging = artifacts_root / f".{artifact_id}.staging-{uuid.uuid4().hex}"
    final_root = artifacts_root / artifact_id
    staging.mkdir(parents=True, exist_ok=False)
    report_data = build_acquisition_report_data(
        workspace_id,
        run,
    )
    markdown = render_markdown(run, report_data)
    md_path = staging / "资产收购可行性研究报告.md"
    docx_path = staging / "资产收购可行性研究报告.docx"
    xlsx_path = staging / "资产收购财务模型.xlsx"
    report_data_path = staging / "资产收购报告数据.json"
    index_path = staging / "附件索引.json"
    try:
        md_path.write_text(markdown, encoding="utf-8")
        from lvke_mcp.domains.reports import doc_service as _doc_svc

        docx_path.write_bytes(_doc_svc.markdown_to_docx(markdown))
        report_data_path.write_text(json.dumps(report_data, ensure_ascii=False, indent=2), encoding="utf-8")
        _write_minimal_xlsx(xlsx_path, run, report_data=report_data)
        consistency = _check_artifact_consistency(
            run, markdown, docx_path, xlsx_path, report_data_path=report_data_path,
        )
        if consistency["status"] != "passed":
            return {
                "ok": False, "error": "ARTIFACT_MISMATCH", "reason": "numeric_or_binding_mismatch",
                "consistency": consistency,
            }
        files = [md_path, docx_path, xlsx_path, report_data_path]
        index = {
            "artifact_id": artifact_id, "run_id": run_id,
            "spec_hash": run.get("spec_hash"), "fact_revision": run.get("spec_id"),
            "spec_snapshot_hash": run.get("spec_snapshot_hash"),
            "evidence_binding_version": run.get("evidence_binding_version"),
            "evidence_binding_hash": run.get("evidence_binding_hash"),
            "model_version": run.get("model_version"), "generated_at": _now(),
            "report_data_hash": report_data.get("report_data_hash"),
            "numeric_consistency": consistency,
            "files": [
                {"name": path.name, "size_bytes": path.stat().st_size, "sha256": _file_hash(path)}
                for path in files
            ],
        }
        index_path.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
        index["files"].append({"name": index_path.name, "size_bytes": index_path.stat().st_size, "sha256": _file_hash(index_path)})
        os.replace(staging, final_root)
    finally:
        if staging.exists():
            shutil.rmtree(staging, ignore_errors=True)

    row = {
        "ok": True, "artifact_id": artifact_id,
        "artifact_job_id": artifact_job_id or f"artifact_job_{uuid.uuid4().hex}",
        "status": "succeeded", "progress": 100,
        "type": "asset_acquisition", "run_id": run_id,
        "spec_hash": run.get("spec_hash"), "fact_revision": run.get("spec_id"),
        "spec_snapshot_hash": run.get("spec_snapshot_hash"),
        "evidence_binding_version": run.get("evidence_binding_version"),
        "evidence_binding_hash": run.get("evidence_binding_hash"),
        "template_version": "asset_acquisition.v2", "created_at": _now(),
        "updated_at": _now(), "request_id": request_id,
        "files": index["files"], "directory": str(final_root),
        "report_data_hash": report_data.get("report_data_hash"),
        "numeric_consistency": "passed", "consistency_checks": consistency["checks"],
        "release_status": "pending",
        "state_history": [
            _history_event(
                "artifact_generated", actor=run.get("approved_by") or "system",
                request_id=request_id, run_id=run_id, artifact_id=artifact_id,
            ),
        ],
    }
    try:
        with _state_guard(workspace_id):
            state = _load(workspace_id)
            if scope and (prior := _active_idempotency_record(state["idempotency"], scope)):
                # Another worker completed the same request while this one rendered.
                if prior.get("body_hash") != body_hash:
                    shutil.rmtree(final_root, ignore_errors=True)
                    return {"ok": False, "error": "IDEMPOTENCY_CONFLICT", "resource_id": prior.get("artifact_id", "")}
                shutil.rmtree(final_root, ignore_errors=True)
                return {**state["artifacts"][prior["artifact_id"]], "idempotent_replay": True}
            state["artifacts"][artifact_id] = row
            if scope:
                state["idempotency"][scope] = _idempotency_record(
                    scope, body_hash, artifact_id=artifact_id,
                )
            stored_run = state["runs"].get(run_id)
            if stored_run:
                stored_run["lifecycle_status"] = "artifact_generated"
                stored_run.setdefault("state_history", []).append(_history_event(
                    "artifact_generated", actor=run.get("approved_by") or "system", request_id=request_id,
                    artifact_id=artifact_id,
                ))
            _save(workspace_id, state)
    except BaseException:
        # State publication and the visible directory are one logical commit.
        # If metadata cannot be committed, do not leave an untracked formal pack.
        shutil.rmtree(final_root, ignore_errors=True)
        raise

    binding_result = _bind_succeeded_artifact(
        workspace_id,
        run,
        row,
    )
    if not binding_result.get("ok"):
        # A formal artifact is not successful unless the report-side binding
        # can prove the same approved run/spec/fact revision.  Compensate the
        # state publication and remove the visible pack on a binding failure.
        with _state_guard(workspace_id):
            state = _load(workspace_id)
            stored = state["artifacts"].get(artifact_id)
            if stored:
                stored.update({
                    "ok": False, "status": "failed", "release_status": "blocked",
                    "error": {
                        "code": "ARTIFACT_BINDING_FAILED",
                        "message": "formal artifact finance binding failed",
                        "retryable": False,
                        "details": binding_result.get("mismatches") or [],
                    },
                    "updated_at": _now(),
                })
            stored_run = state["runs"].get(run_id)
            if stored_run:
                stored_run["lifecycle_status"] = "approved"
                stored_run.setdefault("state_history", []).append(_history_event(
                    "artifact_binding_failed", actor="system", request_id=request_id,
                    artifact_id=artifact_id,
                ))
            _save(workspace_id, state)
        shutil.rmtree(final_root, ignore_errors=True)
        return {
            "ok": False, "error": "ARTIFACT_BINDING_FAILED",
            "reason": binding_result.get("reason") or "finance_binding_mismatch",
            "mismatches": binding_result.get("mismatches") or [],
        }

    ready = mark_artifact_release_ready(
        workspace_id,
        artifact_id,
        actor=str(run.get("approved_by") or "system"),
        request_id=request_id,
    )
    latest = dict(
        _load(workspace_id)["artifacts"].get(artifact_id)
        or row
    )
    if not ready.get("ok"):
        latest["release_gate"] = {
            "ok": False,
            "blockers": ready.get("blockers") or [],
            "warnings": ready.get("warnings") or [],
        }
    return latest


def _recover_published_artifact(
    workspace_id: str, artifact_id: str, row: dict[str, Any],
) -> tuple[bool, str]:
    """Recover a pack atomically published just before a process crash."""

    directory = _root(workspace_id) / "artifacts" / artifact_id
    index_path = directory / "附件索引.json"
    if not index_path.is_file():
        return False, "not_published"
    try:
        index = json.loads(index_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError, TypeError):
        return False, "invalid_index"
    if (
        index.get("artifact_id") != artifact_id
        or index.get("run_id") != row.get("run_id")
    ):
        return False, "binding_mismatch"
    consistency = index.get("numeric_consistency") or {}
    if not isinstance(consistency, dict) or consistency.get("status") != "passed":
        return False, "numeric_consistency_failed"
    files = list(index.get("files") or [])
    for item in files:
        path = directory / str(item.get("name") or "")
        if not path.is_file() or _file_hash(path) != item.get("sha256"):
            return False, "file_hash_mismatch"
    files.append({
        "name": index_path.name,
        "size_bytes": index_path.stat().st_size,
        "sha256": _file_hash(index_path),
    })
    row.update({
        "ok": True, "status": "succeeded", "progress": 100,
        "directory": str(directory), "files": files,
        "numeric_consistency": "passed",
        "report_data_hash": index.get("report_data_hash"),
        "consistency_checks": list(consistency.get("checks") or []),
        "recovered_from_published_artifact": True,
        "updated_at": _now(),
    })
    row.setdefault("state_history", []).append(_history_event(
        "artifact_generated",
        actor="system",
        request_id=str(row.get("request_id") or ""),
        artifact_id=artifact_id,
        recovery=True,
    ))
    return True, "recovered"


def _workspace_ids() -> list[str]:
    root = data_root() / "workspaces"
    if not root.is_dir():
        return []
    return sorted(
        path.name
        for path in root.iterdir()
        if path.is_dir() and str(path.name).strip()
    )


def recover_incomplete_acquisition_tasks(
    submit_run: Any | None = None, submit_artifact: Any | None = None,
) -> dict[str, Any]:
    """Requeue durable acquisition runs/artifacts left by a process restart."""

    run_submitter = submit_run or (
        lambda workspace_id, run_id: _RECOVERY_POOL.submit(
            execute_queued_run,
            workspace_id,
            run_id,
        )
    )
    artifact_submitter = submit_artifact or (
        lambda workspace_id, artifact_id: _RECOVERY_POOL.submit(
            execute_queued_artifact,
            workspace_id,
            artifact_id,
        )
    )
    runs: list[tuple[str, str]] = []
    artifacts: list[tuple[str, str]] = []
    recovered_bindings: list[tuple[str, str]] = []
    recovered_artifacts = 0
    failed_artifacts = 0
    for workspace_id in _workspace_ids():
        if not _state_path(workspace_id).is_file():
            continue
        with _state_guard(workspace_id):
            state = _load(workspace_id)
            changed = False
            for run_id, row in state["runs"].items():
                if row.get("status") not in {"queued", "running"}:
                    continue
                row.update({
                    "status": "queued", "progress": 0, "updated_at": _now(),
                    "recovered_at": _now(),
                    "recovery_count": int(row.get("recovery_count") or 0) + 1,
                })
                row.setdefault("state_history", []).append(_history_event(
                    "queued", request_id=str(row.get("request_id") or ""), recovery=True,
                ))
                runs.append((workspace_id, run_id))
                changed = True
            for artifact_id, row in state["artifacts"].items():
                if row.get("status") not in {"queued", "running"}:
                    continue
                recovered, reason = _recover_published_artifact(
                    workspace_id,
                    artifact_id,
                    row,
                )
                if recovered:
                    recovered_artifacts += 1
                    recovered_bindings.append((workspace_id, artifact_id))
                    changed = True
                    continue
                if reason != "not_published":
                    row.update({
                        "ok": False, "status": "failed", "progress": 100,
                        "numeric_consistency": "failed", "updated_at": _now(),
                        "error": {
                            "code": "ARTIFACT_RECOVERY_FAILED", "message": reason,
                            "retryable": False,
                        },
                    })
                    failed_artifacts += 1
                    changed = True
                    continue
                row.update({
                    "status": "queued", "progress": 0, "updated_at": _now(),
                    "recovered_at": _now(),
                    "recovery_count": int(row.get("recovery_count") or 0) + 1,
                })
                artifacts.append((workspace_id, artifact_id))
                changed = True
            if changed:
                _save(workspace_id, state)
    for workspace_id, artifact_id in recovered_bindings:
        artifact = dict(
            _load(workspace_id)["artifacts"].get(artifact_id)
            or {}
        )
        run = get_run(
            workspace_id,
            str(artifact.get("run_id") or ""),
        )
        if run:
            with _state_guard(workspace_id):
                state = _load(workspace_id)
                stored_run = state["runs"].get(str(artifact.get("run_id") or ""))
                if stored_run:
                    stored_run["lifecycle_status"] = "artifact_generated"
                    stored_run.setdefault("state_history", []).append(_history_event(
                        "artifact_generated",
                        actor=str(stored_run.get("approved_by") or "system"),
                        request_id=str(artifact.get("request_id") or ""),
                        artifact_id=artifact_id,
                        recovery=True,
                    ))
                    _save(workspace_id, state)
                    run = copy.deepcopy(stored_run)
        bound = _bind_succeeded_artifact(
            workspace_id,
            run,
            artifact,
        ) if run else {
            "ok": False, "error": "ARTIFACT_BINDING_FAILED",
        }
        if bound.get("ok"):
            mark_artifact_release_ready(
                workspace_id,
                artifact_id,
                actor=str(run.get("approved_by") or "system"),
                request_id=str(artifact.get("request_id") or ""),
            )
        else:
            with _state_guard(workspace_id):
                state = _load(workspace_id)
                stored = state["artifacts"].get(artifact_id)
                if stored:
                    stored.update({
                        "ok": False, "status": "failed", "release_status": "blocked",
                        "error": {
                            "code": "ARTIFACT_BINDING_FAILED",
                            "message": "recovered artifact finance binding failed",
                            "retryable": False,
                        },
                        "updated_at": _now(),
                    })
                    recovered_artifacts -= 1
                    failed_artifacts += 1
                    _save(workspace_id, state)
    for workspace_id, run_id in runs:
        run_submitter(workspace_id, run_id)
    for workspace_id, artifact_id in artifacts:
        artifact_submitter(workspace_id, artifact_id)
    return {
        "runs_requeued": len(runs),
        "artifacts_requeued": len(artifacts),
        "artifacts_recovered": recovered_artifacts,
        "artifacts_failed": failed_artifacts,
    }


def _artifact_filename_is_safe(filename: str) -> bool:
    value = str(filename or "")
    return bool(
        value
        and value not in {".", ".."}
        and "\x00" not in value
        and "/" not in value
        and "\\" not in value
        and Path(value).name == value
        and not Path(value).is_absolute()
    )


def _artifact_media_type(filename: str) -> str:
    stable_types = {
        ".md": "text/markdown; charset=utf-8",
        ".json": "application/json; charset=utf-8",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }
    suffix = Path(str(filename or "")).suffix.lower()
    return stable_types.get(
        suffix,
        mimetypes.guess_type(str(filename or ""))[0]
        or "application/octet-stream",
    )


def get_artifact(
    workspace_id: str,
    artifact_id: str,
) -> dict[str, Any]:
    """Return an acquisition artifact only after re-verifying its formal pack."""

    row = dict(
        _load(workspace_id)["artifacts"].get(artifact_id)
        or {}
    )
    if not row:
        return {}
    signoff = ((row.get("release") or {}).get("professional_signoff") or {})
    if row.get("status") != "succeeded":
        return {
            **row,
            "integrity_status": (
                "pending" if row.get("status") in {"queued", "running"} else "failed"
            ),
        }
    run = get_run(
        workspace_id,
        str(row.get("run_id") or ""),
    )

    artifacts_root = (
        _root(workspace_id) / "artifacts"
    ).resolve()
    expected_directory = (artifacts_root / artifact_id).resolve()
    directory = Path(str(row.get("directory") or "")).resolve()
    try:
        expected_directory.relative_to(artifacts_root)
        directory.relative_to(artifacts_root)
    except ValueError:
        return {
            **row, "ok": False, "error": "ARTIFACT_MISMATCH",
            "integrity_status": "invalid_directory",
            "failures": [{"reason": "directory_outside_artifact_root"}],
        }
    if directory != expected_directory or not directory.is_dir():
        return {
            **row, "ok": False, "error": "ARTIFACT_MISMATCH",
            "integrity_status": "invalid_directory",
            "failures": [{"reason": "artifact_directory_binding_mismatch"}],
        }

    failures: list[dict[str, Any]] = []
    spec_row = get_spec(
        workspace_id,
        str(run.get("spec_id") or ""),
    )
    bound_spec = spec_row.get("spec") if isinstance(spec_row, dict) else None
    if not isinstance(bound_spec, dict) or _hash(bound_spec) != run.get("spec_hash"):
        failures.append({"reason": "run_spec_snapshot_mismatch"})
    else:
        evidence_ok, current_evidence = _current_evidence_matches_run(
            workspace_id,
            run,
            bound_spec,
        )
        if not evidence_ok:
            failures.append({
                "reason": "evidence_binding_stale",
                "expected": run.get("evidence_binding_hash"),
                "actual": current_evidence.get("binding_hash"),
                "status": current_evidence.get("status"),
            })
    file_rows: dict[str, dict[str, Any]] = {}
    for item in row.get("files") or []:
        if not isinstance(item, dict):
            failures.append({"reason": "invalid_file_manifest_entry"})
            continue
        name = str(item.get("name") or "")
        if not _artifact_filename_is_safe(name):
            failures.append({"name": name, "reason": "invalid_filename_path"})
            continue
        if name in file_rows:
            failures.append({"name": name, "reason": "duplicate_file_manifest_entry"})
            continue
        file_rows[name] = item
        path = directory / name
        try:
            path.resolve().relative_to(directory)
        except ValueError:
            failures.append({"name": name, "reason": "path_escape"})
            continue
        if path.is_symlink():
            failures.append({"name": name, "reason": "symlink_not_allowed"})
            continue
        if not path.is_file():
            failures.append({"name": name, "reason": "missing"})
            continue
        actual_size = path.stat().st_size
        if actual_size != item.get("size_bytes"):
            failures.append({
                "name": name, "reason": "size_mismatch",
                "expected": item.get("size_bytes"), "actual": actual_size,
            })
        actual_hash = _file_hash(path)
        if actual_hash != item.get("sha256"):
            failures.append({
                "name": name, "reason": "sha256_mismatch", "actual": actual_hash,
            })

    required_names = {
        "资产收购可行性研究报告.md",
        "资产收购可行性研究报告.docx",
        "资产收购财务模型.xlsx",
        "资产收购报告数据.json",
        "附件索引.json",
    }
    for missing_name in sorted(required_names - set(file_rows)):
        failures.append({"name": missing_name, "reason": "manifest_entry_missing"})

    index: dict[str, Any] = {}
    index_path = directory / "附件索引.json"
    if index_path.is_file():
        try:
            loaded_index = json.loads(index_path.read_text(encoding="utf-8"))
            index = loaded_index if isinstance(loaded_index, dict) else {}
        except (OSError, json.JSONDecodeError, TypeError):
            failures.append({"name": index_path.name, "reason": "invalid_json"})
    if not index:
        failures.append({"name": index_path.name, "reason": "index_missing_or_invalid"})
    else:
        for field, expected in (
            ("artifact_id", artifact_id),
            ("run_id", row.get("run_id")),
            ("spec_hash", row.get("spec_hash")),
            ("fact_revision", row.get("fact_revision")),
            ("spec_snapshot_hash", row.get("spec_snapshot_hash")),
            ("evidence_binding_version", row.get("evidence_binding_version")),
            ("evidence_binding_hash", row.get("evidence_binding_hash")),
            ("report_data_hash", row.get("report_data_hash")),
        ):
            if index.get(field) != expected:
                failures.append({
                    "name": index_path.name, "reason": f"index_{field}_mismatch",
                    "expected": expected, "actual": index.get(field),
                })
        indexed_files = {
            str(item.get("name") or ""): item
            for item in (index.get("files") or [])
            if isinstance(item, dict)
        }
        state_files = {
            name: item for name, item in file_rows.items()
            if name != index_path.name
        }
        if indexed_files != state_files:
            failures.append({"name": index_path.name, "reason": "file_manifest_mismatch"})
        index_consistency = index.get("numeric_consistency") or {}
        index_checks = index_consistency.get("checks") or [] if isinstance(index_consistency, dict) else []
        if (
            not isinstance(index_consistency, dict)
            or index_consistency.get("status") != "passed"
            or not index_checks
            or any(not check.get("passed") for check in index_checks if isinstance(check, dict))
            or any(not isinstance(check, dict) for check in index_checks)
        ):
            failures.append({"name": index_path.name, "reason": "numeric_consistency_incomplete"})
        if row.get("numeric_consistency") != "passed" or row.get("consistency_checks") != index_checks:
            failures.append({"name": index_path.name, "reason": "numeric_consistency_state_mismatch"})

    report_data: dict[str, Any] | None = None
    report_data_path = directory / "资产收购报告数据.json"
    if report_data_path.is_file():
        try:
            loaded = json.loads(report_data_path.read_text(encoding="utf-8"))
            report_data = loaded if isinstance(loaded, dict) else None
        except (OSError, json.JSONDecodeError, TypeError):
            report_data = None
    if not isinstance(report_data, dict):
        failures.append({"name": report_data_path.name, "reason": "invalid_json"})
    else:
        embedded_hash = report_data.get("report_data_hash")
        payload = {key: value for key, value in report_data.items() if key != "report_data_hash"}
        calculated_hash = _hash(payload)
        if embedded_hash != calculated_hash or embedded_hash != row.get("report_data_hash"):
            failures.append({
                "name": report_data_path.name, "reason": "report_data_hash_mismatch",
                "expected": row.get("report_data_hash"), "actual": calculated_hash,
            })
        bindings = report_data.get("bindings") or {}
        for field in (
            "run_id", "spec_hash", "input_hash", "model_version",
            "spec_snapshot_hash", "evidence_binding_version", "evidence_binding_hash",
        ):
            if bindings.get(field) != run.get(field):
                failures.append({
                    "name": report_data_path.name,
                    "reason": f"report_binding_{field}_mismatch",
                })
        if str(bindings.get("spec_id") or "") != str(row.get("fact_revision") or ""):
            failures.append({
                "name": report_data_path.name,
                "reason": "report_binding_fact_revision_mismatch",
            })
        if report_data.get("maximum_acceptable_price") != run.get("max_acquisition_price_analysis"):
            failures.append({
                "name": report_data_path.name,
                "reason": "report_maximum_price_mismatch",
            })

    # Re-run the numeric/binding verifier against the *current approved run*;
    # trusting only the checks serialized at generation time would miss a
    # stale artifact after an out-of-band run mutation.
    if isinstance(report_data, dict):
        try:
            current_consistency = _check_artifact_consistency(
                run,
                (directory / "资产收购可行性研究报告.md").read_text(encoding="utf-8"),
                directory / "资产收购可行性研究报告.docx",
                directory / "资产收购财务模型.xlsx",
                report_data_path=report_data_path,
            )
        except Exception as exc:  # noqa: BLE001
            failures.append({
                "reason": "numeric_consistency_unverifiable",
                "error_type": type(exc).__name__,
            })
        else:
            if current_consistency.get("status") != "passed":
                failures.append({
                    "reason": "numeric_consistency_failed",
                    "checks": [
                        check for check in (current_consistency.get("checks") or [])
                        if not check.get("passed")
                    ],
                })
            if current_consistency.get("checks") != row.get("consistency_checks"):
                failures.append({"reason": "numeric_consistency_snapshot_mismatch"})

    for signoff_failure in _asset_signoff_integrity_failures(
        row,
    ):
        failures.append({
            "reason": "professional_signoff_integrity_failed",
            "details": copy.deepcopy(signoff_failure),
        })
    if failures:
        return {
            **row, "ok": False, "error": "ARTIFACT_MISMATCH",
            "integrity_status": "failed", "failures": failures,
        }
    return {
        **row, "ok": True, "integrity_status": "passed",
        "report_data": report_data,
    }


def list_artifacts(
    workspace_id: str,
    *,
    limit: int = 50,
) -> list[dict[str, Any]]:
    """Return artifact jobs newest-first, including their integrity projection."""

    state = _load(workspace_id)
    ids = sorted(
        state["artifacts"],
        key=lambda artifact_id: (
            str(state["artifacts"][artifact_id].get("created_at") or ""),
            artifact_id,
        ),
        reverse=True,
    )
    rows = [
        get_artifact(workspace_id, artifact_id)
        for artifact_id in ids
    ]
    return [
        row for row in rows if row
    ][: max(1, min(int(limit or 50), 100))]


def _require_unified_asset_review(
    workspace_id: str,
    artifact_id: str,
    review_id: str,
) -> dict[str, Any]:
    normalized_review_id = str(review_id or "").strip()
    if not normalized_review_id:
        return {
            "ok": False,
            "error": "REVIEW_REQUIRED",
            "reason": "formal release requires an explicit review_id",
        }
    from lvke_mcp.servers.lvke_deliverable_review import service as review_service

    binding = review_service.require_released_review_for_target(
        workspace_id,
        normalized_review_id,
        "report_artifact",
        artifact_id,
        artifact_domain="asset_acquisition",
    )
    target = binding.get("target")
    target = target if isinstance(target, Mapping) else {}
    required_binding = {
        "review_id": binding.get("review_id"),
        "review_release_id": binding.get("review_release_id"),
        "review_release_hash": binding.get("review_release_hash"),
        "review_release_basis_hash": binding.get("review_release_basis_hash"),
        "review_event_chain_hash": binding.get("review_event_chain_hash"),
        "review_target_sha256": target.get("target_sha256"),
    }
    binding_valid = (
        binding.get("status") == "ok"
        and str(binding.get("review_id") or "") == normalized_review_id
        and target.get("target_type") == "report_artifact"
        and str(target.get("target_id") or "") == artifact_id
        and all(str(value or "").strip() for value in required_binding.values())
    )
    if not binding_valid:
        return {
            "ok": False,
            "error": "REVIEW_REQUIRED",
            "reason": str(binding.get("message") or "review gate blocked"),
            "review_id": normalized_review_id,
            "review_error": binding.get("code"),
            "blockers": list(binding.get("blockers") or []),
            "binding_complete": all(
                str(value or "").strip()
                for value in required_binding.values()
            ),
        }
    return {"ok": True, "binding": binding}


def _require_current_asset_release_review(
    workspace_id: str,
    artifact: Mapping[str, Any],
) -> dict[str, Any]:
    release = artifact.get("release")
    if artifact.get("release_status") != "released" or not isinstance(
        release, Mapping
    ):
        return {
            "ok": False,
            "error": "REVIEW_REQUIRED",
            "reason": "formal artifact has not completed controlled release",
            "status": artifact.get("release_status"),
        }
    checked = _require_unified_asset_review(
        workspace_id,
        str(artifact.get("artifact_id") or ""),
        str(release.get("review_id") or ""),
    )
    if not checked.get("ok"):
        return checked
    binding = checked["binding"]
    expected = {
        "review_id": str(release.get("review_id") or ""),
        "review_release_id": str(release.get("review_release_id") or ""),
        "review_release_hash": str(release.get("review_release_hash") or ""),
        "review_release_basis_hash": str(
            release.get("review_release_basis_hash") or ""
        ),
        "review_event_chain_hash": str(
            release.get("review_event_chain_hash") or ""
        ),
        "review_target_sha256": str(release.get("review_target_sha256") or ""),
    }
    actual = {
        "review_id": str(binding.get("review_id") or ""),
        "review_release_id": str(binding.get("review_release_id") or ""),
        "review_release_hash": str(binding.get("review_release_hash") or ""),
        "review_release_basis_hash": str(
            binding.get("review_release_basis_hash") or ""
        ),
        "review_event_chain_hash": str(
            binding.get("review_event_chain_hash") or ""
        ),
        "review_target_sha256": str(
            (binding.get("target") or {}).get("target_sha256") or ""
        ),
    }
    if any(not value for value in expected.values()) or actual != expected:
        return {
            "ok": False,
            "error": "REVIEW_REQUIRED",
            "reason": "stored review release binding is missing or changed",
            "expected_review_binding": expected,
            "actual_review_binding": actual,
        }
    return {"ok": True, "binding": binding}


def _resolve_artifact_download(
    workspace_id: str,
    artifact_id: str,
    filename: str,
    *,
    allow_unreleased_candidate: bool = False,
) -> dict[str, Any]:
    """Resolve one approved-run artifact file after path and hash validation."""

    if not _artifact_filename_is_safe(filename):
        return {"ok": False, "error": "INVALID_FILENAME_PATH"}
    artifact = get_artifact(workspace_id, artifact_id)
    if not artifact:
        return {"ok": False, "error": "ARTIFACT_NOT_FOUND"}
    if artifact.get("status") != "succeeded":
        return {
            "ok": False, "error": "ARTIFACT_NOT_READY",
            "status": artifact.get("status"),
        }
    if artifact.get("integrity_status") != "passed" or not artifact.get("ok", True):
        return {
            "ok": False, "error": "ARTIFACT_MISMATCH",
            "failures": artifact.get("failures") or [],
        }
    run = get_run(
        workspace_id,
        str(artifact.get("run_id") or ""),
    )
    if run.get("review_status") != "approved":
        return {"ok": False, "error": "RUN_NOT_APPROVABLE", "reason": "approved_run_required"}
    if not allow_unreleased_candidate:
        review_check = _require_current_asset_release_review(
            workspace_id,
            artifact,
        )
        if not review_check.get("ok"):
            return review_check

    file_row = next(
        (
            item for item in (artifact.get("files") or [])
            if str(item.get("name") or "") == filename
        ),
        None,
    )
    if not file_row:
        return {"ok": False, "error": "ARTIFACT_FILE_NOT_FOUND"}
    directory = Path(str(artifact.get("directory") or "")).resolve()
    path = (directory / filename).resolve()
    try:
        path.relative_to(directory)
    except ValueError:
        return {"ok": False, "error": "INVALID_FILENAME_PATH"}
    if not path.is_file() or path.is_symlink():
        return {"ok": False, "error": "ARTIFACT_FILE_NOT_FOUND"}
    actual_hash = _file_hash(path)
    if actual_hash != file_row.get("sha256"):
        return {
            "ok": False, "error": "ARTIFACT_MISMATCH",
            "failures": [{
                "name": filename, "reason": "sha256_mismatch",
                "expected": file_row.get("sha256"), "actual": actual_hash,
            }],
        }
    return {
        "ok": True,
        "artifact_id": artifact_id,
        "run_id": artifact.get("run_id"),
        "filename": filename,
        "path": path,
        "size_bytes": path.stat().st_size,
        "sha256": actual_hash,
        "media_type": _artifact_media_type(filename),
        "release_status": artifact.get("release_status"),
    }


def resolve_artifact_download(
    workspace_id: str,
    artifact_id: str,
    filename: str,
) -> dict[str, Any]:
    """Resolve an externally downloadable, released acquisition artifact."""

    return _resolve_artifact_download(
        workspace_id,
        artifact_id,
        filename,
    )


def resolve_artifact_candidate_download(
    workspace_id: str,
    artifact_id: str,
    filename: str,
) -> dict[str, Any]:
    """Resolve an unreleased candidate for the unified-review engine only."""

    return _resolve_artifact_download(
        workspace_id,
        artifact_id,
        filename,
        allow_unreleased_candidate=True,
    )


def _read_resolved_artifact(
    workspace_id: str,
    artifact_id: str,
    resolved: dict[str, Any],
) -> dict[str, Any]:
    if not resolved.get("ok"):
        return resolved
    path = resolved.get("path")
    try:
        content = path.read_bytes() if isinstance(path, Path) else b""
    except OSError:
        return {"ok": False, "error": "ARTIFACT_FILE_NOT_FOUND"}
    digest = hashlib.sha256(content).hexdigest()
    if digest != resolved.get("sha256") or len(content) != resolved.get("size_bytes"):
        get_artifact(workspace_id, artifact_id)
        return {
            "ok": False,
            "error": "ARTIFACT_MISMATCH",
            "failures": [{
                "name": resolved.get("filename"),
                "reason": "content_changed_after_validation",
            }],
        }
    return {**resolved, "content": content}


def read_artifact_download(
    workspace_id: str,
    artifact_id: str,
    filename: str,
) -> dict[str, Any]:
    return _read_resolved_artifact(
        workspace_id,
        artifact_id,
        resolve_artifact_download(
            workspace_id,
            artifact_id,
            filename,
        ),
    )


def read_artifact_candidate_download(
    workspace_id: str,
    artifact_id: str,
    filename: str,
) -> dict[str, Any]:
    return _read_resolved_artifact(
        workspace_id,
        artifact_id,
        resolve_artifact_candidate_download(
            workspace_id,
            artifact_id,
            filename,
        ),
    )


def release_artifact(
    workspace_id: str,
    artifact_id: str,
    *,
    actor: str,
    review_id: str,
    note: str = "",
    idempotency_key: str = "",
    request_id: str = "",
) -> dict[str, Any]:
    """Record an authorized internal release after the unified publish gate.

    This operation is not a legal or professional signature.  External human
    signoff evidence is registered and independently verified through separate
    endpoints.
    """

    actor = str(actor or "").strip()
    if not actor:
        return {"ok": False, "error": "AUTHENTICATION_REQUIRED"}
    request_id = request_id or f"req_{uuid.uuid4().hex}"
    artifact = get_artifact(workspace_id, artifact_id)
    if not artifact:
        return {"ok": False, "error": "ARTIFACT_NOT_FOUND"}
    normalized_review_id = str(review_id or "").strip()
    review_check = _require_unified_asset_review(
        workspace_id,
        artifact_id,
        normalized_review_id,
    )
    if not review_check.get("ok"):
        return review_check
    review_binding = review_check["binding"]
    body_hash = _hash({
        "action": "release",
        "artifact_id": artifact_id,
        "run_id": artifact.get("run_id"),
        "spec_hash": artifact.get("spec_hash"),
        "fact_revision": artifact.get("fact_revision"),
        "spec_snapshot_hash": artifact.get("spec_snapshot_hash"),
        "evidence_binding_hash": artifact.get("evidence_binding_hash"),
        "review_id": normalized_review_id,
        "review_release_id": review_binding.get("review_release_id"),
        "review_release_hash": review_binding.get("review_release_hash"),
        "review_release_basis_hash": review_binding.get(
            "review_release_basis_hash"
        ),
        "review_event_chain_hash": review_binding.get(
            "review_event_chain_hash"
        ),
        "review_target_sha256": (
            review_binding.get("target") or {}
        ).get("target_sha256"),
        "actor": actor,
        "note": str(note or ""),
    })
    scope = f"release:{idempotency_key}" if idempotency_key else ""
    if scope:
        with _state_guard(workspace_id):
            state = _load(workspace_id)
            prior = _active_idempotency_record(state["idempotency"], scope)
            if prior:
                if prior.get("body_hash") != body_hash:
                    return {
                        "ok": False, "error": "IDEMPOTENCY_CONFLICT",
                        "resource_id": prior.get("release_id") or prior.get("artifact_id"),
                    }
                stored = state["artifacts"].get(str(prior.get("artifact_id") or ""))
                release = dict((stored or {}).get("release") or {})
                return {
                    "ok": True,
                    "artifact_id": prior.get("artifact_id"),
                    "run_id": (stored or {}).get("run_id"),
                    "release_status": (stored or {}).get("release_status"),
                    "release": release,
                    "idempotent_replay": True,
                }

    ready = mark_artifact_release_ready(
        workspace_id,
        artifact_id,
        actor=actor,
        request_id=request_id,
    )
    if not ready.get("ok"):
        return ready

    released_at = _now()
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        if scope and (prior := _active_idempotency_record(state["idempotency"], scope)):
            if prior.get("body_hash") != body_hash:
                return {
                    "ok": False, "error": "IDEMPOTENCY_CONFLICT",
                    "resource_id": prior.get("release_id") or prior.get("artifact_id"),
                }
            replayed = state["artifacts"].get(str(prior.get("artifact_id") or "")) or {}
            return {
                "ok": True,
                "artifact_id": prior.get("artifact_id"),
                "run_id": replayed.get("run_id"),
                "release_status": replayed.get("release_status"),
                "release": dict(replayed.get("release") or {}),
                "idempotent_replay": True,
            }
        stored = state["artifacts"].get(artifact_id)
        if not stored:
            return {"ok": False, "error": "ARTIFACT_NOT_FOUND"}
        if stored.get("release_status") == "released":
            existing_release = dict(stored.get("release") or {})
            existing_review_binding = {
                "review_id": existing_release.get("review_id"),
                "review_release_id": existing_release.get("review_release_id"),
                "review_release_hash": existing_release.get(
                    "review_release_hash"
                ),
                "review_release_basis_hash": existing_release.get(
                    "review_release_basis_hash"
                ),
                "review_event_chain_hash": existing_release.get(
                    "review_event_chain_hash"
                ),
                "review_target_sha256": existing_release.get(
                    "review_target_sha256"
                ),
            }
            current_review_binding = {
                "review_id": normalized_review_id,
                "review_release_id": review_binding.get("review_release_id"),
                "review_release_hash": review_binding.get("review_release_hash"),
                "review_release_basis_hash": review_binding.get(
                    "review_release_basis_hash"
                ),
                "review_event_chain_hash": review_binding.get(
                    "review_event_chain_hash"
                ),
                "review_target_sha256": (
                    review_binding.get("target") or {}
                ).get("target_sha256"),
            }
            if (
                existing_release.get("actor") != actor
                or str(existing_release.get("note") or "") != str(note or "")
                or existing_review_binding != current_review_binding
            ):
                return {
                    "ok": False, "error": "ARTIFACT_ALREADY_RELEASED",
                    "resource_id": existing_release.get("release_id") or artifact_id,
                }
            if scope:
                state["idempotency"][scope] = _idempotency_record(
                    scope, body_hash,
                    artifact_id=artifact_id,
                    release_id=existing_release.get("release_id"),
                )
                _save(workspace_id, state)
            return {
                "ok": True, "artifact_id": artifact_id,
                "run_id": stored.get("run_id"), "release_status": "released",
                "release": existing_release, "idempotent_replay": True,
            }
        if stored.get("release_status") != "release_ready":
            return {
                "ok": False, "error": "ARTIFACT_NOT_READY",
                "status": stored.get("release_status"),
            }
        release_id = f"release_{uuid.uuid4().hex}"
        release = {
            "release_id": release_id,
            "artifact_id": artifact_id,
            "run_id": stored.get("run_id"),
            "spec_hash": stored.get("spec_hash"),
            "fact_revision": stored.get("fact_revision"),
            "spec_snapshot_hash": stored.get("spec_snapshot_hash"),
            "evidence_binding_version": stored.get("evidence_binding_version"),
            "evidence_binding_hash": stored.get("evidence_binding_hash"),
            "review_id": normalized_review_id,
            "review_release_id": review_binding.get("review_release_id"),
            "review_release_hash": review_binding.get("review_release_hash"),
            "review_release_basis_hash": review_binding.get("review_release_basis_hash"),
            "review_event_chain_hash": review_binding.get("review_event_chain_hash"),
            "review_target_sha256": (review_binding.get("target") or {}).get("target_sha256"),
            "actor": actor,
            "note": str(note or ""),
            "request_id": request_id,
            "released_at": released_at,
        }
        release["release_hash"] = _hash(release)
        stored.update({
            "release_status": "released",
            "released_at": released_at,
            "released_by": actor,
            "release": release,
            "updated_at": released_at,
        })
        stored.setdefault("release_history", []).append(copy.deepcopy(release))
        stored.setdefault("state_history", []).append(_history_event(
            "released", actor=actor, request_id=request_id,
            run_id=stored.get("run_id"), artifact_id=artifact_id,
            release_id=release_id,
        ))
        run = state["runs"].get(str(stored.get("run_id") or ""))
        if run:
            run["lifecycle_status"] = "released"
            run["released_at"] = released_at
            run["released_by"] = actor
            run.setdefault("state_history", []).append(_history_event(
                "released", actor=actor, request_id=request_id,
                artifact_id=artifact_id, release_id=release_id,
            ))
        if scope:
            state["idempotency"][scope] = _idempotency_record(
                scope, body_hash, artifact_id=artifact_id, release_id=release_id,
            )
        _save(workspace_id, state)
    return {
        "ok": True,
        "artifact_id": artifact_id,
        "run_id": artifact.get("run_id"),
        "release_status": "released",
        "release": release,
    }


def _asset_signoff_basis(artifact: Mapping[str, Any]) -> dict[str, Any]:
    release = artifact.get("release") or {}
    return {
        "artifact_id": artifact.get("artifact_id"),
        "run_id": artifact.get("run_id"),
        "spec_hash": artifact.get("spec_hash"),
        "fact_revision": artifact.get("fact_revision"),
        "spec_snapshot_hash": artifact.get("spec_snapshot_hash"),
        "evidence_binding_hash": artifact.get("evidence_binding_hash"),
        "report_data_hash": artifact.get("report_data_hash"),
        "artifact_files_hash": _hash(artifact.get("files") or []),
        "release_id": release.get("release_id") if isinstance(release, Mapping) else "",
        "internal_release_hash": (
            _hash(_asset_release_integrity_material(release))
            if isinstance(release, Mapping) else ""
        ),
    }


def _require_current_released_asset_artifact(
    workspace_id: str,
    artifact_id: str,
) -> dict[str, Any]:
    artifact = get_artifact(workspace_id, artifact_id)
    if not artifact:
        raise AcquisitionSignoffError("ARTIFACT_NOT_FOUND", "正式工件不存在")
    if (
        artifact.get("status") != "succeeded"
        or artifact.get("ok") is not True
        or artifact.get("integrity_status") != "passed"
        or artifact.get("release_status") != "released"
    ):
        raise AcquisitionSignoffError(
            "SIGNOFF_REQUIRES_CURRENT_INTERNAL_RELEASE",
            "仅已完成内部发布且当前完整的资产收购正式工件可登记或核验签章证据",
            details={
                "status": artifact.get("status"),
                "release_status": artifact.get("release_status"),
                "integrity_status": artifact.get("integrity_status"),
                "failures": artifact.get("failures") or [],
            },
        )
    run = get_run(
        workspace_id,
        str(artifact.get("run_id") or ""),
    )
    if run.get("review_status") != "approved":
        raise AcquisitionSignoffError(
            "RUN_NOT_APPROVABLE", "签章证据只能绑定当前已批准财务 run",
        )
    return artifact


def _update_asset_release_history_signoff(
    artifact: dict[str, Any],
    release_id: str,
    signoff: dict[str, Any],
) -> None:
    for item in artifact.get("release_history") or []:
        if str(item.get("release_id") or "") == release_id:
            item["professional_signoff"] = copy.deepcopy(signoff)
            return


def register_asset_external_professional_signoff(
    workspace_id: str,
    artifact_id: str,
    *,
    actor: Mapping[str, Any],
    evidence_hash: str,
    evidence_reference: str,
    signer: str,
    signer_credential: str,
    signed_at: str,
    note: str = "",
    operation_id: str = "",
) -> dict[str, Any]:
    """Record external evidence metadata; this function never signs a report."""

    validated_actor = _validated_signoff_actor(actor)
    normalized_hash = _normalize_signoff_evidence_hash(evidence_hash)
    normalized_reference = _required_signoff_text(
        evidence_reference, field="evidence_reference", maximum=1000,
    )
    normalized_signer = _required_signoff_text(
        signer, field="signer", maximum=300,
    )
    normalized_signer_credential = _required_signoff_text(
        signer_credential,
        field="signer_credential",
        maximum=500,
    )
    normalized_signed_at = _normalize_signoff_time(signed_at, field="signed_at")
    normalized_note = str(note or "").strip()[:2000]
    normalized_operation = str(operation_id or "").strip()
    if normalized_operation and not _SAFE_SIGNOFF_OPERATION_ID.fullmatch(
        normalized_operation
    ):
        raise AcquisitionSignoffError(
            "INVALID_OPERATION_ID", "签章证据登记 operation_id 不合法",
        )
    current = _require_current_released_asset_artifact(
        workspace_id,
        artifact_id,
    )
    registration_body = {
        "artifact_id": artifact_id,
        "evidence_hash": normalized_hash,
        "evidence_reference": normalized_reference,
        "signer": normalized_signer,
        "signer_credential": normalized_signer_credential,
        "signed_at": normalized_signed_at,
        "note": normalized_note,
    }
    registration_hash = _hash(registration_body)
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        stored = state["artifacts"].get(artifact_id)
        if not isinstance(stored, dict):
            raise AcquisitionSignoffError("ARTIFACT_NOT_FOUND", "正式工件不存在")
        if _asset_signoff_basis(stored) != _asset_signoff_basis(current):
            raise AcquisitionSignoffError(
                "ARTIFACT_MISMATCH", "正式工件或发布绑定在登记前发生变化",
            )
        release = copy.deepcopy(stored.get("release") or {})
        existing = release.get("professional_signoff") or {}
        if isinstance(existing, dict) and existing:
            if (
                existing.get("registration_hash") == registration_hash
                and str((existing.get("registered_by") or {}).get("actor_id") or "")
                == str(validated_actor.get("actor_id") or "")
            ):
                return {**copy.deepcopy(stored), "idempotent_replay": True}
            raise AcquisitionSignoffError(
                "PROFESSIONAL_SIGNOFF_ALREADY_REGISTERED",
                "该资产收购发布记录已经登记其他专业签章证据",
                details={"signoff_id": existing.get("signoff_id")},
            )
        release_failures = _asset_signoff_integrity_failures(stored)
        if release_failures:
            raise AcquisitionSignoffError(
                "ARTIFACT_MISMATCH",
                "内部发布记录完整性校验失败",
                details={"failures": release_failures},
            )
        now = _now()
        basis = _asset_signoff_basis(stored)
        signoff = {
            "schema_version": "asset_external_professional_signoff.v1",
            "signoff_id": f"signoff_{uuid.uuid4().hex}",
            **basis,
            **registration_body,
            "registration_hash": registration_hash,
            "registration_operation_id": normalized_operation,
            "registered_by": validated_actor,
            "registered_at": now,
            "verified": False,
            "manual_external_evidence": True,
            "evidence_hash_source": "human_supplied_external_evidence_digest",
            "system_signature_performed": False,
        }
        signoff["registration_integrity_hash"] = _hash(
            _asset_signoff_registration_integrity_material(signoff)
        )
        release["professional_signoff"] = copy.deepcopy(signoff)
        stored["release"] = release
        stored.setdefault("professional_signoff_history", []).append({
            "event": "registered",
            "signoff": copy.deepcopy(signoff),
            "created_at": now,
        })
        _update_asset_release_history_signoff(
            stored,
            str(release.get("release_id") or ""),
            signoff,
        )
        stored.setdefault("state_history", []).append(_history_event(
            "external_professional_signoff_registered",
            actor=str(validated_actor.get("actor_id") or ""),
            artifact_id=artifact_id,
            signoff_id=signoff["signoff_id"],
            manual_external_evidence=True,
            system_signature_performed=False,
        ))
        stored["updated_at"] = now
        state["artifacts"][artifact_id] = stored
        _save(workspace_id, state)
    result = get_artifact(workspace_id, artifact_id)
    if result.get("integrity_status") != "passed":
        raise AcquisitionSignoffError(
            "ARTIFACT_MISMATCH",
            "签章证据登记后的完整性复验失败",
            details={"failures": result.get("failures") or []},
        )
    return result


def verify_asset_external_professional_signoff(
    workspace_id: str,
    artifact_id: str,
    *,
    actor: Mapping[str, Any],
    signoff_id: str,
    evidence_hash: str,
    verification_reference: str,
    note: str = "",
    operation_id: str = "",
) -> dict[str, Any]:
    """Record a separately attributed human check of external evidence."""

    validated_actor = _validated_signoff_actor(actor)
    normalized_signoff_id = str(signoff_id or "").strip()
    if not _SAFE_PROFESSIONAL_SIGNOFF_ID.fullmatch(normalized_signoff_id):
        raise AcquisitionSignoffError(
            "INVALID_SIGNOFF_EVIDENCE", "signoff_id 不合法",
        )
    normalized_hash = _normalize_signoff_evidence_hash(evidence_hash)
    normalized_reference = _required_signoff_text(
        verification_reference,
        field="verification_reference",
        maximum=1000,
    )
    normalized_note = str(note or "").strip()[:2000]
    normalized_operation = str(operation_id or "").strip()
    if normalized_operation and not _SAFE_SIGNOFF_OPERATION_ID.fullmatch(
        normalized_operation
    ):
        raise AcquisitionSignoffError(
            "INVALID_OPERATION_ID", "签章证据核验 operation_id 不合法",
        )
    current = _require_current_released_asset_artifact(
        workspace_id,
        artifact_id,
    )
    with _state_guard(workspace_id):
        state = _load(workspace_id)
        stored = state["artifacts"].get(artifact_id)
        if not isinstance(stored, dict):
            raise AcquisitionSignoffError("ARTIFACT_NOT_FOUND", "正式工件不存在")
        if _asset_signoff_basis(stored) != _asset_signoff_basis(current):
            raise AcquisitionSignoffError(
                "ARTIFACT_MISMATCH", "正式工件或发布绑定在核验前发生变化",
            )
        release = copy.deepcopy(stored.get("release") or {})
        signoff = copy.deepcopy(release.get("professional_signoff") or {})
        if not signoff:
            raise AcquisitionSignoffError(
                "PROFESSIONAL_SIGNOFF_NOT_REGISTERED",
                "尚未登记外部专业签章证据",
            )
        if str(signoff.get("signoff_id") or "") != normalized_signoff_id:
            raise AcquisitionSignoffError(
                "PROFESSIONAL_SIGNOFF_MISMATCH", "signoff_id 与当前登记不一致",
            )
        if str(signoff.get("evidence_hash") or "") != normalized_hash:
            raise AcquisitionSignoffError(
                "PROFESSIONAL_SIGNOFF_MISMATCH", "证据哈希与当前登记不一致",
            )
        verifier_id = str(validated_actor.get("actor_id") or "")
        registrar_id = str((signoff.get("registered_by") or {}).get("actor_id") or "")
        release_actor_id = str(release.get("actor") or "")
        if verifier_id in {registrar_id, release_actor_id}:
            raise AcquisitionSignoffError(
                "PROFESSIONAL_SIGNOFF_DUTY_SEPARATION_REQUIRED",
                "签章证据核验人不得同时是登记人或内部发布执行人",
                details={
                    "registrar_actor_id": registrar_id,
                    "release_actor_id": release_actor_id,
                },
            )
        verification_body = {
            "signoff_id": normalized_signoff_id,
            "evidence_hash": normalized_hash,
            "verification_reference": normalized_reference,
            "note": normalized_note,
            "verified_by": verifier_id,
        }
        verification_hash = _hash(verification_body)
        if signoff.get("verified") is True:
            if (
                signoff.get("verification_hash") == verification_hash
                and str((signoff.get("verified_by") or {}).get("actor_id") or "")
                == verifier_id
            ):
                return {**copy.deepcopy(stored), "idempotent_replay": True}
            raise AcquisitionSignoffError(
                "PROFESSIONAL_SIGNOFF_ALREADY_VERIFIED",
                "该签章证据已经由其他核验记录确认",
            )
        integrity_failures = _asset_signoff_integrity_failures(
            stored,
        )
        if integrity_failures:
            raise AcquisitionSignoffError(
                "ARTIFACT_MISMATCH",
                "签章证据登记完整性校验失败",
                details={"failures": integrity_failures},
            )
        now = _now()
        signoff.update({
            "verified": True,
            "verified_by": validated_actor,
            "verified_at": now,
            "verification_reference": normalized_reference,
            "verification_note": normalized_note,
            "verification_hash": verification_hash,
            "verification_operation_id": normalized_operation,
            "verification_kind": "independent_human_evidence_check",
            "evidence_hash_verification": "independent_human_digest_match",
            "system_signature_performed": False,
        })
        signoff["verification_integrity_hash"] = _hash(
            _asset_signoff_verification_integrity_material(signoff)
        )
        release["professional_signoff"] = copy.deepcopy(signoff)
        stored["release"] = release
        stored.setdefault("professional_signoff_history", []).append({
            "event": "verified",
            "signoff": copy.deepcopy(signoff),
            "created_at": now,
        })
        _update_asset_release_history_signoff(
            stored,
            str(release.get("release_id") or ""),
            signoff,
        )
        stored.setdefault("state_history", []).append(_history_event(
            "external_professional_signoff_verified",
            actor=verifier_id,
            artifact_id=artifact_id,
            signoff_id=normalized_signoff_id,
            verification_kind="independent_human_evidence_check",
            system_signature_performed=False,
        ))
        stored["updated_at"] = now
        state["artifacts"][artifact_id] = stored
        _save(workspace_id, state)
    result = get_artifact(workspace_id, artifact_id)
    if result.get("integrity_status") != "passed":
        raise AcquisitionSignoffError(
            "ARTIFACT_MISMATCH",
            "签章证据核验后的完整性复验失败",
            details={"failures": result.get("failures") or []},
        )
    return result


def _xlsx_summary_values(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("xl/worksheets/sheet1.xml"))
    values: dict[str, str] = {}
    for row in root.iter():
        if not row.tag.endswith("}row"):
            continue
        cells: list[str] = []
        for cell in row:
            if not cell.tag.endswith("}c"):
                continue
            inline = "".join(node.text or "" for node in cell.iter() if node.tag.endswith("}t"))
            raw = next((node.text or "" for node in cell if node.tag.endswith("}v")), "")
            cells.append(inline if inline else raw)
        if len(cells) >= 2:
            values[cells[0]] = cells[1]
    return values


def _same_number(actual: str, expected: Any) -> bool:
    try:
        return math.isclose(float(actual), float(expected), rel_tol=1e-12, abs_tol=1e-9)
    except (TypeError, ValueError):
        return False


def _check_artifact_consistency(
    run: dict[str, Any], markdown: str, docx_path: Path, xlsx_path: Path,
    *, report_data_path: Path | None = None,
) -> dict[str, Any]:
    from docx import Document  # type: ignore

    document = Document(io.BytesIO(docx_path.read_bytes()))
    docx_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    result = run.get("result") or {}
    indicators = result.get("indicators") or {}
    asset_type = str(result.get("asset_type") or "hotel_lease")
    is_solar = asset_type == "solar_power"
    max_price_analysis = run.get("max_acquisition_price_analysis") or {}
    max_price_result = max_price_analysis.get("result") or {}
    checks: list[dict[str, Any]] = []

    binding_tokens = [
        str(run.get("run_id") or ""), str(run.get("spec_hash") or ""),
        str(run.get("model_version") or ""), str(run.get("evidence_binding_hash") or ""),
    ]
    numeric_tokens = [
        f"{float(result.get('purchase_price_wan') or 0):,.2f}",
        f"{float(result.get('total_acquisition_cost_wan') or 0):,.2f}",
    ]
    numeric_fields = [
        (indicators.get("project_irr_pct"), _pct),
        (indicators.get("equity_irr_pct"), _pct),
        (indicators.get("npv_wan"), _num),
        (indicators.get("minimum_dscr"), _num),
    ]
    if not is_solar:
        numeric_fields.append((indicators.get("minimum_tenant_rent_coverage"), _num))
    for value, formatter in numeric_fields:
        if value is not None:
            numeric_tokens.append(formatter(value))
    for token in binding_tokens + numeric_tokens:
        checks.append({
            "artifact": "markdown", "field": token, "passed": bool(token and token in markdown),
        })
        checks.append({
            "artifact": "docx", "field": token, "passed": bool(token and token in docx_text),
        })

    summary = _xlsx_summary_values(xlsx_path)
    string_expectations = {
        "资产收购财务模型": run.get("run_id"),
        "模型版本": run.get("model_version"),
        "Spec哈希": run.get("spec_hash"),
        "证据绑定哈希": run.get("evidence_binding_hash"),
        "证据绑定版本": run.get("evidence_binding_version"),
        "批准人": run.get("approved_by"),
        "批准时间": run.get("approved_at"),
        "最高价求解哈希": max_price_analysis.get("analysis_hash"),
        "最高价决策状态": max_price_analysis.get("decision_status") or "pending_business_thresholds",
    }
    number_expectations = {
        "收购价格(万元)": result.get("purchase_price_wan"),
        "总收购成本(万元)": result.get("total_acquisition_cost_wan"),
        "项目IRR(%)": indicators.get("project_irr_pct"),
        "资本金IRR(%)": indicators.get("equity_irr_pct"),
        "NPV(万元)": indicators.get("npv_wan"),
        "最低DSCR": indicators.get("minimum_dscr"),
        "最低ICR": indicators.get("minimum_icr"),
        "最高可接受收购价(万元)": max_price_result.get("max_acquisition_price_wan"),
        "最高价目标IRR": (max_price_analysis.get("parameters") or {}).get("target_irr"),
        "最高价最低DSCR": (max_price_analysis.get("parameters") or {}).get("min_dscr"),
    }
    if is_solar:
        solar = result.get("solar_operation") or {}
        number_expectations.update({
            "装机容量(MW)": solar.get("installed_capacity_mw"),
            "基准发电量(MWh)": solar.get("base_generation_mwh"),
            "上网电价(元/kWh)": solar.get("tariff_yuan_per_kwh"),
            "限电率": solar.get("curtailment_rate"),
            "年衰减率": solar.get("degradation_rate"),
        })
    else:
        number_expectations.update({
            "最低租金覆盖率": indicators.get("minimum_tenant_rent_coverage"),
            "租约覆盖年限": indicators.get("lease_coverage_years"),
            "合同收入占比": indicators.get("contract_income_ratio"),
            "未锁定收入占比": indicators.get("unlocked_income_ratio"),
        })
    for label, expected in string_expectations.items():
        checks.append({
            "artifact": "xlsx", "field": label, "expected": expected,
            "actual": summary.get(label), "passed": summary.get(label) == str(expected or ""),
        })
    for label, expected in number_expectations.items():
        passed = summary.get(label, "") == "" if expected is None else _same_number(summary.get(label, ""), expected)
        checks.append({
            "artifact": "xlsx", "field": label, "expected": expected,
            "actual": summary.get(label), "passed": passed,
        })
    if report_data_path is not None:
        try:
            report_data = json.loads(report_data_path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError, TypeError):
            report_data = {}
        bindings = report_data.get("bindings") or {}
        for field in (
            "run_id", "spec_hash", "input_hash", "model_version",
            "spec_snapshot_hash", "evidence_binding_hash", "evidence_binding_version",
        ):
            checks.append({
                "artifact": "report_data", "field": field,
                "expected": run.get(field), "actual": bindings.get(field),
                "passed": bindings.get(field) == run.get(field),
            })
        expected_hash = report_data.get("report_data_hash")
        payload = {key: value for key, value in report_data.items() if key != "report_data_hash"}
        checks.append({
            "artifact": "report_data", "field": "report_data_hash",
            "expected": _hash(payload), "actual": expected_hash,
            "passed": expected_hash == _hash(payload),
        })
    return {"status": "passed" if all(row["passed"] for row in checks) else "failed", "checks": checks}


def _file_hash(path: Path) -> str:
    hasher = hashlib.sha256()
    with path.open("rb") as source:
        for chunk in iter(lambda: source.read(1024 * 1024), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


def _xlsx_col(index: int) -> str:
    out = ""
    while index:
        index, rem = divmod(index - 1, 26)
        out = chr(65 + rem) + out
    return out


def _xml_cell(row: int, column: int, value: Any) -> str:
    from xml.sax.saxutils import escape

    ref = f"{_xlsx_col(column)}{row}"
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return f'<c r="{ref}"><v>{value}</v></c>'
    return f'<c r="{ref}" t="inlineStr"><is><t>{escape(str(value or ""))}</t></is></c>'


def _sheet_xml(rows: list[list[Any]]) -> str:
    body = []
    for row_index, values in enumerate(rows, 1):
        cells = "".join(_xml_cell(row_index, col, value) for col, value in enumerate(values, 1))
        body.append(f'<row r="{row_index}">{cells}</row>')
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        f'<sheetData>{"".join(body)}</sheetData></worksheet>'
    )


def _write_minimal_xlsx(
    path: Path, run: dict[str, Any], *, report_data: dict[str, Any] | None = None,
) -> None:
    result = run.get("result") or {}
    indicators = result.get("indicators") or {}
    report_data = report_data or {}
    asset_type = str(result.get("asset_type") or report_data.get("asset_type") or "hotel_lease")
    is_solar = asset_type == "solar_power"
    max_price_analysis = report_data.get("maximum_acceptable_price") or {}
    max_price_result = max_price_analysis.get("result") or {}
    summary = [
        ["资产收购财务模型", run.get("run_id")],
        ["模型版本", run.get("model_version")],
        ["Spec哈希", run.get("spec_hash")],
        ["证据绑定哈希", run.get("evidence_binding_hash")],
        ["证据绑定版本", run.get("evidence_binding_version")],
        ["收购价格(万元)", result.get("purchase_price_wan")],
        ["总收购成本(万元)", result.get("total_acquisition_cost_wan")],
        ["项目IRR(%)", indicators.get("project_irr_pct")],
        ["资本金IRR(%)", indicators.get("equity_irr_pct")],
        ["NPV(万元)", indicators.get("npv_wan")],
        ["最低DSCR", indicators.get("minimum_dscr")],
        ["最低ICR", indicators.get("minimum_icr")],
        ["维修资本开支覆盖", indicators.get("maintenance_capex_coverage")],
        ["退出价值NPV占比", indicators.get("exit_value_npv_ratio")],
        ["最高可接受收购价(万元)", max_price_result.get("max_acquisition_price_wan")],
        ["最高价目标IRR", (max_price_analysis.get("parameters") or {}).get("target_irr")],
        ["最高价最低DSCR", (max_price_analysis.get("parameters") or {}).get("min_dscr")],
        ["最高价求解哈希", max_price_analysis.get("analysis_hash")],
        ["最高价决策状态", max_price_analysis.get("decision_status")],
        ["批准人", run.get("approved_by")],
        ["批准时间", run.get("approved_at")],
    ]
    if is_solar:
        solar = result.get("solar_operation") or {}
        summary[13:13] = [
            ["装机容量(MW)", solar.get("installed_capacity_mw")],
            ["基准发电量(MWh)", solar.get("base_generation_mwh")],
            ["上网电价(元/kWh)", solar.get("tariff_yuan_per_kwh")],
            ["限电率", solar.get("curtailment_rate")],
            ["年衰减率", solar.get("degradation_rate")],
        ]
    else:
        summary[13:13] = [
            ["最低租金覆盖率", indicators.get("minimum_tenant_rent_coverage")],
            ["租约覆盖年限", indicators.get("lease_coverage_years")],
            ["合同收入占比", indicators.get("contract_income_ratio")],
            ["未锁定收入占比", indicators.get("unlocked_income_ratio")],
        ]
    project = result.get("project_cashflows_wan") or []
    equity = result.get("equity_cashflows_wan") or []
    cashflows = [["年度", "项目现金流(万元)", "资本金现金流(万元)"]]
    for index in range(max(len(project), len(equity))):
        cashflows.append([index, project[index] if index < len(project) else "", equity[index] if index < len(equity) else ""])
    def json_text(value: Any) -> str:
        if isinstance(value, (dict, list)):
            return json.dumps(value, ensure_ascii=False, sort_keys=True, default=str)
        return str(value or "")

    parties = [["主体ID", "名称", "角色", "状态", "证据"]] + [
        [row.get("entity_id"), row.get("name"), json_text(row.get("roles")), row.get("status"), json_text(row.get("evidence_ids"))]
        for row in report_data.get("party_relationships") or []
    ]
    assets = [["资产ID", "类型", "是否纳入", "面积(㎡)", "状态", "冲突", "裁决", "证据"]] + [
        [row.get("scope_id"), row.get("type"), row.get("included"), row.get("area_sqm"), row.get("status"),
         json_text(row.get("conflicts")), row.get("resolution"), json_text(row.get("evidence_ids"))]
        for row in report_data.get("asset_boundary") or []
    ]
    if is_solar:
        operations = [[
            "年度", "理论发电量(MWh)", "上网电量(MWh)", "上网电价(元/kWh)",
            "售电收入(万元)", "运维费(万元)", "维护性资本开支(万元)",
            "所得税(万元)", "债务服务(万元)", "项目现金流(万元)", "资本金现金流(万元)",
        ]] + [
            [
                row.get("year"), row.get("gross_generation_mwh"), row.get("sold_generation_mwh"),
                row.get("tariff_yuan_per_kwh"), row.get("revenue_wan"), row.get("operating_cost_wan"),
                row.get("maintenance_capex_wan"), row.get("income_tax_wan"), row.get("debt_service_wan"),
                row.get("project_cf_wan"), row.get("equity_cf_wan"),
            ]
            for row in report_data.get("solar_operating_ledger") or []
        ]
    else:
        leases = [["单元ID", "位置", "面积(㎡)", "出租人", "承租人", "起始日", "终止日", "基础租金(万元)", "证据"]] + [
            [row.get("unit_id"), row.get("asset_location"), row.get("area_sqm"), row.get("lessor_id"), row.get("lessee_id"),
             row.get("start_date"), row.get("end_date"), row.get("base_rent_wan"), json_text(row.get("evidence_ids"))]
            for row in report_data.get("lease_ledger") or []
        ]
    history = [["主体", "开始日", "结束日", "报表类型", "来源格式", "人工复核", "勾稽", "异常", "来源定位"]] + [
        [row.get("entity_id"), row.get("period_start"), row.get("period_end"), row.get("statement_type"), row.get("source_format"),
         row.get("manual_review_status"), json_text(row.get("reconciliation")), json_text(row.get("anomalies")), json_text(row.get("source_locators"))]
        for row in report_data.get("historical_financial_comparison") or []
    ]
    scenarios = ([
        ["矩阵ID", "场景ID", "变更", "收购价(万元)", "上网电价(元/kWh)", "年发电量(MWh)", "利用小时", "年运维费(万元)", "融资比例", "指标"]
    ] if is_solar else [
        ["矩阵ID", "场景ID", "变更", "收购价(万元)", "市场租金", "入住率", "融资比例", "指标"]
    ])
    for matrix in report_data.get("scenario_matrices") or []:
        for row in matrix.get("rows") or []:
            common = [
                matrix.get("matrix_id"), row.get("scenario_id"), json_text(row.get("changes")),
                row.get("purchase_price_wan"),
            ]
            scenarios.append(common + ([
                row.get("tariff_yuan_per_kwh"), row.get("annual_generation_mwh"),
                row.get("utilization_hours"), row.get("annual_opex_wan"),
                row.get("financing_ratio"), json_text(row.get("indicators")),
            ] if is_solar else [
                json_text(row.get("market_rent")), json_text(row.get("occupancy")),
                row.get("financing_ratio"), json_text(row.get("indicators")),
            ]))
    risks = [["类型", "编码/内容", "状态", "裁决/说明", "证据"]]
    risks.extend([
        ["red_flag", row.get("code"), row.get("status"), row.get("resolution"), json_text(row.get("evidence_ids"))]
        for row in report_data.get("red_flags") or []
    ])
    risks.extend([["closing_condition", item, "", "", ""] for item in report_data.get("closing_conditions") or []])
    risks.extend([["veto_item", item, "", "", ""] for item in report_data.get("veto_items") or []])
    evidence = [["字段", "证据ID", "来源", "复核状态"]] + [
        [row.get("field"), json_text(row.get("evidence_ids")), json_text(row.get("source")), row.get("review_status")]
        for row in report_data.get("source_processing_ledger") or []
    ]
    sheets: list[tuple[str, list[list[Any]]]] = [("收购摘要", summary), ("现金流", cashflows)]
    if report_data:
        sheets.extend([("主体关系", parties), ("资产边界", assets)])
        sheets.append(("光伏运营", operations) if is_solar else ("租约台账", leases))
        sheets.extend([
            ("历史财务", history), ("情景矩阵", scenarios),
            ("风险与条件", risks), ("证据台账", evidence),
        ])
    content_types = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
        + ''.join(
            f'<Override PartName="/xl/worksheets/sheet{index}.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
            for index in range(1, len(sheets) + 1)
        )
        + '</Types>'
    )
    from xml.sax.saxutils import escape

    workbook = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<sheets>'
        + ''.join(
            f'<sheet name="{escape(name)}" sheetId="{index}" r:id="rId{index}"/>'
            for index, (name, _rows) in enumerate(sheets, 1)
        )
        + '</sheets></workbook>'
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        + ''.join(
            f'<Relationship Id="rId{index}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet{index}.xml"/>'
            for index in range(1, len(sheets) + 1)
        )
        + '</Relationships>'
    )
    root_rels = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>'
        '</Relationships>'
    )
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", root_rels)
        archive.writestr("xl/workbook.xml", workbook)
        archive.writestr("xl/_rels/workbook.xml.rels", rels)
        for index, (_name, rows) in enumerate(sheets, 1):
            archive.writestr(f"xl/worksheets/sheet{index}.xml", _sheet_xml(rows))


def _num(value: Any) -> str:
    try:
        return f"{float(value):,.2f}"
    except (TypeError, ValueError):
        return "—"


def _pct(value: Any) -> str:
    try:
        return f"{float(value):.2f}%"
    except (TypeError, ValueError):
        return "—"


def _pct_ratio(value: Any) -> str:
    try:
        return f"{float(value) * 100:.2f}%"
    except (TypeError, ValueError):
        return "—"
