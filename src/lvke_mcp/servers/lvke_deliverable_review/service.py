"""Unified deliverable-review application service.

All decisions are projected from immutable preparation objects and append-only
review events. Existing finance/report validators are evidence inputs, never a
substitute for the unified review verdict.
"""

from __future__ import annotations

import base64
import hashlib
import io
import json
import os
import re
import tempfile
import threading
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable
from urllib.parse import quote, unquote

from lvke_mcp.runtime.storage import (
    JSONArtifactStore, canonical_json, paginate_resource_entries, require_safe_id,
    sha256_json, utc_now,
)
from lvke_mcp.runtime.workspace import workspace_root
# MCP 本地工具服务不做身份认证。
from lvke_mcp.servers.lvke_deliverable_review import financial_checks, report_checks, rules
from lvke_mcp.servers.lvke_deliverable_review.contracts import (
    DEPLOYMENT_MODES, FINDING_STATUSES, SEVERITIES, SEVERITY_ORDER,
    finding_blocks, normalize_project_context, normalize_target,
    require_write_context, verdict_for,
)
from lvke_mcp.servers.lvke_deliverable_review.store import STORE

REPO_ROOT = Path(__file__).resolve().parents[2]
PREPARATION_STORE = JSONArtifactStore(
    "deliverable-review", "preparations", "rvprep", "preparations"
)
EXPORT_STORE = JSONArtifactStore(
    "deliverable-review", "exports", "rvexp", "exports"
)
RELEASE_STORE = JSONArtifactStore(
    "deliverable-review", "releases", "rvrel", "releases"
)
STANDARD_APPLICABILITY_STORE = JSONArtifactStore(
    "deliverable-review", "standard_applicabilities", "stdapp", "standard-applicabilities"
)
STANDARD_EVIDENCE_STORE = JSONArtifactStore(
    "deliverable-review", "standard_evidence", "stdev", "standard-evidence"
)
_REPORT_ARTIFACT_DOMAINS = {"generic_feasibility", "asset_acquisition"}
_ASYNC_THREADS: dict[tuple[str, str], threading.Thread] = {}
_ASYNC_LOCK = threading.Lock()


def _ok(**data: Any) -> dict[str, Any]:
    status = str(data.pop("status", "ok"))
    business_success = status in {"ok", "accepted"}
    return {
        "success": business_success, "transport_success": True,
        "system_success": True, "business_success": business_success,
        "completed": status == "ok", "outcome": status, "status": status,
        **data, "resource_uris": list(data.get("resource_uris") or []),
        "warnings": list(data.get("warnings") or []),
        "blockers": list(data.get("blockers") or []),
        "next_actions": list(data.get("next_actions") or []),
    }


def _blocked(code: str, message: str, **data: Any) -> dict[str, Any]:
    blockers = list(data.pop("blockers", []) or [code])
    return _ok(status="blocked", code=code, message=message, blockers=blockers, **data)


def _write(operation: str, args: dict[str, Any], callback: Callable[[str], dict[str, Any]]) -> dict[str, Any]:
    try:
        workspace_id, _, key = require_write_context(args)
        scoped_operation = operation

        def execute_once() -> dict[str, Any]:
            cached = STORE.idempotent(workspace_id, scoped_operation, key, args)
            if cached is not None:
                if operation != "review_release":
                    return cached
                # A terminal release response is idempotent, but it must not
                # conceal later target, event-chain, record, or file damage.
                verified = callback(workspace_id)
                return cached if verified.get("status") == "ok" else verified
            response = callback(workspace_id)
            STORE.remember(workspace_id, scoped_operation, key, args, response)
            return response

        with STORE.mutation_guard(workspace_id, scoped_operation, key):
            review_id = str(args.get("review_id") or "").strip()
            if review_id and operation != "review_release":
                with STORE.mutation_guard(
                    workspace_id,
                    "review_terminal_mutation",
                    review_id,
                ):
                    return execute_once()
            return execute_once()
    except ValueError as exc:
        code = str(exc)
        return _blocked(code, _message(code))


def _message(code: str) -> str:
    messages = {
        "workspace_id_required": "缺少 workspace_id",
        "idempotency_key_required": "写操作必须提供有效 idempotency_key",
        "idempotency_key_conflict": "同一幂等键已用于不同请求",
        "target_required": "缺少审查目标", "target_type_invalid": "目标类型不受支持",
        "target_id_required": "目标必须包含 target_id", "preparation_not_found": "审查准备对象不存在",
        "review_not_found": "审查运行不存在", "finding_not_found": "finding 不存在",
        "retest_operation_conflict": "复测操作事件与已固化的操作意图冲突",
        "retest_preparation_unavailable": "复测绑定的审查准备对象不可用",
        "retest_child_review_unavailable": "复测子审查不可用",
        "project_type_invalid": "project_type 不受支持",
        "transaction_structure_invalid": "transaction_structure 不受支持",
        "transaction_structure_project_type_mismatch": "交易结构与项目类型不匹配",
        "asset_type_invalid": "asset_type 不受支持",
        "evidence_track_invalid": "evidence_track 不受支持",
        "standard_catalog_invalid": "标准适用性目录不可用",
        "standard_applicability_not_found": "标准适用性对象不存在",
        "standard_requirement_not_found": "标准需求不存在或不适用于当前项目",
        "standard_evidence_resource_invalid": "标准证据必须是当前工作区内受支持的不可变 Resource",
        "standard_evidence_hash_mismatch": "标准证据内容 hash 与不可变 Resource 不一致",
        "standard_evidence_track_mismatch": "标准证据轨与适用性对象不一致",
    }
    return messages.get(code, code.replace("_", " "))


def _safe_file(workspace_id: str, raw: str) -> Path | None:
    try:
        candidate = Path(raw).expanduser().resolve()
        root = workspace_root(require_safe_id(workspace_id, "workspace_id")).resolve()
        candidate.relative_to(root)
    except (OSError, ValueError):
        return None
    return candidate if candidate.is_file() else None


def _legacy_blockers(values: Any) -> list[str]:
    blockers: list[str] = []
    for value in values or []:
        if isinstance(value, dict):
            text = str(value.get("code") or value.get("rule") or value.get("message") or "")
        else:
            text = str(value or "")
        if text:
            blockers.append(text[:500])
    return sorted(set(blockers))


def _legacy_gate_result(
    passed: bool | None,
    source: str,
    *,
    blockers: Any = (),
    details: dict[str, Any] | None = None,
) -> dict[str, Any]:
    return {
        "verdict": "pass" if passed is True else ("fail" if passed is False else "unknown"),
        "passed": passed,
        "source": source,
        "blockers": _legacy_blockers(blockers),
        "details": deepcopy(details or {}),
    }


def _legacy_gate_snapshot(
    workspace_id: str,
    resolved: dict[str, Any],
) -> dict[str, Any]:
    """Freeze the legacy engineering/publish decision for shadow comparison.

    The snapshot is derived from the resolved target or by invoking the exact
    read-only legacy validator. Callers cannot submit their own legacy verdict.
    Unknown is preserved when an old surface never had a corresponding gate.
    """

    target_type = str(resolved.get("target_type") or "")
    target_id = str(resolved.get("target_id") or "")
    snapshot = resolved.get("snapshot") if isinstance(resolved.get("snapshot"), dict) else {}
    validation = _legacy_gate_result(None, "legacy_gate_unavailable")
    publish = _legacy_gate_result(None, "legacy_publish_gate_unavailable")

    try:
        if target_type == "finance_run":
            from lvke_mcp.servers.lvke_finance_tables import service as tables_service

            result = tables_service.validate(workspace_id, target_id)
            assessment = result.get("validation") if isinstance(result.get("validation"), dict) else {}
            valid = assessment.get("valid")
            formal = result.get("formal_delivery_ready")
            validation = _legacy_gate_result(
                valid if isinstance(valid, bool) else None,
                "tables_validate",
                blockers=result.get("blockers") or assessment.get("blockers"),
                details={"status": result.get("status"), "run_id": result.get("run_id")},
            )
            publish = _legacy_gate_result(
                formal if isinstance(formal, bool) else None,
                "tables_validate.formal_delivery_ready",
                blockers=result.get("blockers") or assessment.get("gate_blockers"),
            )
        elif target_type == "finance_tables_package":
            payload = snapshot.get("payload") if isinstance(snapshot.get("payload"), dict) else {}
            assessment = payload.get("validation") if isinstance(payload.get("validation"), dict) else {}
            valid = assessment.get("valid")
            formal = payload.get("formal_delivery_ready")
            validation = _legacy_gate_result(
                valid if isinstance(valid, bool) else None,
                "finance_tables_package.validation",
                blockers=assessment.get("blockers"),
                details={"package_status": snapshot.get("status")},
            )
            publish = _legacy_gate_result(
                formal if isinstance(formal, bool) else None,
                "finance_tables_package.formal_delivery_ready",
                blockers=assessment.get("gate_blockers"),
            )
        elif target_type in {"finance_xlsx", "finance_xlsx_source"}:
            validation = _legacy_gate_result(None, "external_xlsx_has_no_legacy_validator")
            publish = _legacy_gate_result(None, "external_xlsx_has_no_legacy_publish_gate")
        elif target_type == "acquisition_run":
            available = snapshot.get("available") is True
            succeeded = str(snapshot.get("status") or "") == "succeeded"
            consistent = snapshot.get("consistency_ok") is True
            open_blockers = [
                row for row in snapshot.get("issues") or []
                if isinstance(row, dict)
                and row.get("blocking") is True
                and str(row.get("status") or "open") == "open"
            ]
            validation = _legacy_gate_result(
                available and succeeded and consistent,
                "acquisition_run.consistency",
                blockers=open_blockers,
                details={
                    "status": snapshot.get("status"),
                    "consistency_ok": snapshot.get("consistency_ok"),
                    "formal_spec_valid": snapshot.get("formal_spec_valid"),
                },
            )
            publish = _legacy_gate_result(
                available
                and succeeded
                and consistent
                and str(snapshot.get("review_status") or "") == "approved"
                and not open_blockers,
                "acquisition_run.approval_gate",
                blockers=open_blockers,
                details={"review_status": snapshot.get("review_status")},
            )
        elif target_type == "acquisition_tables_package":
            payload = snapshot.get("payload") if isinstance(snapshot.get("payload"), dict) else {}
            integrity = payload.get("integrity") if isinstance(payload.get("integrity"), dict) else {}
            integrity_passed = str(integrity.get("status") or "") == "passed"
            validation = _legacy_gate_result(
                integrity_passed,
                "acquisition_tables_package.integrity",
                blockers=integrity.get("blockers"),
            )
            run_id = str((resolved.get("bindings") or {}).get("finance_run_id") or "")
            from lvke_mcp.servers.lvke_asset_acquisition import backend as acquisition_service

            run = (
                acquisition_service.get_run(
                    workspace_id,
                    run_id,
                )
                if run_id
                else {}
            )
            run_approved = bool(run) and str(run.get("review_status") or "") == "approved"
            publish = _legacy_gate_result(
                integrity_passed and run_approved,
                "acquisition_tables_package.integrity_and_run_approval",
                blockers=integrity.get("blockers"),
                details={"run_id": run_id, "run_review_status": run.get("review_status")},
            )
        elif target_type == "report_revision":
            from lvke_mcp.servers.lvke_report_generation import service as report_service

            result = report_service.validate(workspace_id, target_id)
            valid = result.get("valid")
            validation = _legacy_gate_result(
                valid if isinstance(valid, bool) else None,
                "report_validate",
                blockers=result.get("blockers"),
                details={
                    "status": result.get("status"),
                    "report_revision_id": result.get("report_revision_id"),
                },
            )
            publish = _legacy_gate_result(
                None,
                "report_revision_had_no_standalone_legacy_release_gate",
            )
        elif target_type == "report_artifact":
            validation = _legacy_gate_result(
                True,
                "artifact_current_and_integrity_gate",
                details={
                    "artifact_family": snapshot.get("artifact_family"),
                    "artifact_id": snapshot.get("artifact_id"),
                },
            )
            publish = _legacy_gate_result(
                None,
                "artifact_release_readiness_not_embedded_in_snapshot",
            )
        elif target_type == "combined_deliverable":
            component_snapshots = [
                _legacy_gate_snapshot(
                    workspace_id,
                    component,
                )
                for component in snapshot.get("components") or []
                if isinstance(component, dict)
            ]

            def combined(kind: str) -> dict[str, Any]:
                verdicts = [str((row.get(kind) or {}).get("verdict") or "unknown") for row in component_snapshots]
                passed: bool | None
                if any(value == "fail" for value in verdicts):
                    passed = False
                elif verdicts and all(value == "pass" for value in verdicts):
                    passed = True
                else:
                    passed = None
                return _legacy_gate_result(
                    passed,
                    f"combined_component_{kind}",
                    details={"component_verdicts": verdicts},
                )

            validation = combined("validation")
            publish = combined("publish")
        else:
            validation = _legacy_gate_result(None, "target_type_has_no_legacy_validator")
            publish = _legacy_gate_result(None, "target_type_has_no_legacy_publish_gate")
    except Exception:  # noqa: BLE001 - comparison must remain honest and non-blocking
        validation = _legacy_gate_result(
            None,
            "legacy_gate_lookup_failed",
            blockers=["legacy_gate_lookup_failed"],
        )
        publish = _legacy_gate_result(
            None,
            "legacy_publish_gate_lookup_failed",
            blockers=["legacy_gate_lookup_failed"],
        )

    body = {
        "schema_version": "deliverable_review_legacy_gate_snapshot.v1",
        "target": {
            key: resolved.get(key)
            for key in ("target_type", "target_id", "target_sha256")
        },
        "captured_at": utc_now(),
        "validation": validation,
        "publish": publish,
    }
    return {**body, "content_hash": sha256_json(body)}


def _immutable_artifact_files(raw_files: Any) -> list[dict[str, Any]]:
    """Keep only immutable manifest fields used to bind a reviewed artifact."""

    files: list[dict[str, Any]] = []
    for raw in raw_files or []:
        if not isinstance(raw, dict):
            continue
        name = str(raw.get("name") or raw.get("filename") or "")
        files.append({
            "name": name,
            "role": str(raw.get("role") or ""),
            "media_type": str(raw.get("media_type") or ""),
            "size_bytes": raw.get("size_bytes"),
            "sha256": str(raw.get("sha256") or ""),
        })
    return sorted(files, key=lambda row: (row["name"], row["role"], row["sha256"]))


def _generic_artifact_snapshot(record: dict[str, Any]) -> dict[str, Any]:
    return {
        "artifact_family": "generic",
        "artifact_id": str(record.get("artifact_id") or ""),
        "kind": str(record.get("kind") or ""),
        "template_version": str(record.get("template_version") or ""),
        "basis_fingerprint": str(record.get("basis_fingerprint") or ""),
        "document_revision_id": str(record.get("document_revision_id") or ""),
        "finance_run_id": str(record.get("finance_run_id") or record.get("run_id") or ""),
        "spec_hash": str(record.get("spec_hash") or ""),
        "manifest_hash": str(record.get("manifest_hash") or ""),
        "index_hash": str(record.get("index_hash") or ""),
        "files": _immutable_artifact_files(record.get("files")),
    }


def _acquisition_artifact_snapshot(record: dict[str, Any]) -> dict[str, Any]:
    return {
        "artifact_family": "asset_acquisition",
        "artifact_id": str(record.get("artifact_id") or ""),
        "type": str(record.get("type") or "asset_acquisition"),
        "run_id": str(record.get("run_id") or ""),
        "spec_hash": str(record.get("spec_hash") or ""),
        "fact_revision": str(record.get("fact_revision") or ""),
        "spec_snapshot_hash": str(record.get("spec_snapshot_hash") or ""),
        "evidence_binding_version": str(record.get("evidence_binding_version") or ""),
        "evidence_binding_hash": str(record.get("evidence_binding_hash") or ""),
        "template_version": str(record.get("template_version") or ""),
        "report_data_hash": str(record.get("report_data_hash") or ""),
        "numeric_consistency": str(record.get("numeric_consistency") or ""),
        "files": _immutable_artifact_files(record.get("files")),
    }


def _string_ids(*values: Any) -> list[str]:
    return sorted({
        str(item).strip()
        for value in values
        for item in (value if isinstance(value, (list, tuple, set)) else [])
        if str(item).strip()
    })


def _artifact_upstream_bindings(record: dict[str, Any]) -> dict[str, Any]:
    containers = [
        record,
        record.get("bindings") if isinstance(record.get("bindings"), dict) else {},
        record.get("upstream") if isinstance(record.get("upstream"), dict) else {},
        record.get("basis") if isinstance(record.get("basis"), dict) else {},
    ]
    evidence_ids = _string_ids(
        *(container.get("evidence_pack_ids") for container in containers)
    )
    research_ids = _string_ids(
        *(container.get("research_package_ids") for container in containers),
        *(container.get("research_pack_ids") for container in containers),
    )
    bindings: dict[str, Any] = {
        "evidence_pack_ids": evidence_ids,
        "research_package_ids": research_ids,
    }
    for key in ("finance_run_id", "finance_tables_package_id", "report_revision_id"):
        value = next(
            (
                str(container.get(key) or "").strip()
                for container in containers
                if str(container.get(key) or "").strip()
            ),
            "",
        )
        if value:
            bindings[key] = value
    return bindings


def _linked_generic_report_revision(
    workspace_id: str,
    artifact: dict[str, Any],
) -> tuple[dict[str, Any], dict[str, Any]]:
    native_revision_id = str(artifact.get("document_revision_id") or "")
    if not native_revision_id:
        return {}, {}
    try:
        from lvke_mcp.servers.lvke_report_generation.service import REVISION_STORE

        records = REVISION_STORE.list(workspace_id)
    except (OSError, ValueError):
        return {}, {}
    artifact_created_at = _parse_timestamp(artifact.get("created_at"))
    artifact_run_id = str(
        artifact.get("finance_run_id") or artifact.get("run_id") or ""
    )
    candidates: list[dict[str, Any]] = []
    for record in records:
        payload = record.get("payload") or {}
        upstream = payload.get("upstream") or {}
        if str(payload.get("native_revision_id") or "") != native_revision_id:
            continue
        upstream_run_id = str(upstream.get("run_id") or "")
        if artifact_run_id and upstream_run_id and upstream_run_id != artifact_run_id:
            continue
        created_at = _parse_timestamp(record.get("created_at"))
        if artifact_created_at and created_at and created_at > artifact_created_at:
            continue
        candidates.append(record)
    if not candidates:
        return {}, {}
    candidates.sort(
        key=lambda row: (
            str(row.get("created_at") or ""),
            str(row.get("object_id") or ""),
        ),
        reverse=True,
    )
    revision = candidates[0]
    payload = revision.get("payload") or {}
    upstream = payload.get("upstream") or {}
    bindings = _artifact_upstream_bindings({
        "report_revision_id": revision.get("object_id"),
        "upstream": upstream,
    })
    snapshot = {
        "report_revision_id": str(revision.get("object_id") or ""),
        "native_revision_id": native_revision_id,
        "content_hash": str(revision.get("content_hash") or ""),
        "basis_hash": str(revision.get("basis_hash") or ""),
    }
    return snapshot, bindings


def _resolve_report_artifact(
    workspace_id: str,
    artifact_id: str,
    *,
    artifact_domain: str,
) -> tuple[dict[str, Any] | None, dict[str, Any], list[str]]:
    """Resolve one explicitly selected artifact domain to an immutable target."""

    if artifact_domain == "generic_feasibility":
        try:
            from lvke_mcp.domains.reports import artifacts as deliverable_artifacts

            generic = deliverable_artifacts.get_artifact(
                workspace_id, artifact_id,
            )
        except Exception:  # noqa: BLE001 - target is unavailable or invalid
            generic = {}
        if not generic:
            return None, {}, ["report_artifact_not_found"]
        if (
            generic.get("status") not in {"succeeded", "released"}
            or generic.get("current") is not True
            or generic.get("integrity_status") != "passed"
        ):
            return None, {}, ["report_artifact_not_current"]
        snapshot = _generic_artifact_snapshot(generic)
        revision_snapshot, revision_bindings = _linked_generic_report_revision(
            workspace_id,
            generic,
        )
        bindings = _artifact_upstream_bindings(generic)
        for key in ("evidence_pack_ids", "research_package_ids"):
            bindings[key] = _string_ids(bindings.get(key), revision_bindings.get(key))
        for key in ("finance_tables_package_id", "report_revision_id"):
            if revision_bindings.get(key):
                bindings[key] = revision_bindings[key]
        bindings["finance_run_id"] = str(
            generic.get("finance_run_id") or generic.get("run_id") or ""
        )
        snapshot["upstream_bindings"] = deepcopy(bindings)
        if revision_snapshot:
            snapshot["report_revision"] = revision_snapshot
        return snapshot, bindings, []
    if artifact_domain == "asset_acquisition":
        try:
            from lvke_mcp.servers.lvke_asset_acquisition import backend as acquisition_service

            acquisition = acquisition_service.get_artifact(
                workspace_id,
                artifact_id,
            )
        except Exception:  # noqa: BLE001 - acquisition storage may be unavailable
            acquisition = {}
        if not acquisition:
            return None, {}, ["report_artifact_not_found"]
        if (
            acquisition.get("status") != "succeeded"
            or acquisition.get("ok") is not True
            or acquisition.get("integrity_status") != "passed"
        ):
            return None, {}, ["report_artifact_not_current"]
        snapshot = _acquisition_artifact_snapshot(acquisition)
        bindings = _artifact_upstream_bindings(acquisition)
        bindings["finance_run_id"] = str(acquisition.get("run_id") or "")
        snapshot["upstream_bindings"] = deepcopy(bindings)
        return snapshot, bindings, []
    return None, {}, ["report_artifact_domain_invalid"]


def _combined_bindings_manifest(
    components: list[dict[str, Any]],
) -> dict[str, Any]:
    """Preserve every component binding while retaining scalar compatibility.

    A combined delivery may legitimately contain more than one run, table
    package, evidence pack or report artifact.  A plain ``dict.update`` loses
    earlier component bindings and makes the preparation's upstream snapshot
    incomplete.  The component manifest is authoritative; the top-level
    aggregate remains for existing single-run consumers.
    """

    aggregate: dict[str, Any] = {}
    component_bindings: list[dict[str, Any]] = []
    conflicts: dict[str, list[Any]] = {}
    for component in components:
        raw_bindings = deepcopy(component.get("bindings") or {})
        component_bindings.append({
            "target_type": str(component.get("target_type") or ""),
            "target_id": str(component.get("target_id") or ""),
            "target_sha256": str(component.get("target_sha256") or ""),
            "bindings": raw_bindings,
        })
        for key, raw_value in raw_bindings.items():
            if raw_value in (None, "", [], {}):
                continue
            if isinstance(raw_value, list):
                current = aggregate.setdefault(key, [])
                if not isinstance(current, list):
                    conflicts.setdefault(key, [deepcopy(current)])
                    values = conflicts[key]
                else:
                    values = current
                for value in raw_value:
                    if value not in values:
                        values.append(deepcopy(value))
                continue
            if key not in aggregate:
                aggregate[key] = deepcopy(raw_value)
                continue
            if aggregate[key] != raw_value:
                values = conflicts.setdefault(key, [deepcopy(aggregate[key])])
                if raw_value not in values:
                    values.append(deepcopy(raw_value))
    aggregate["component_bindings"] = component_bindings
    if conflicts:
        aggregate["component_binding_conflicts"] = conflicts
    return aggregate


def _resolve_target(
    workspace_id: str,
    target: dict[str, Any],
) -> tuple[dict[str, Any] | None, list[str]]:
    target = deepcopy(target)
    target_type = target["target_type"]
    target_id = target["target_id"]
    blockers: list[str] = []
    payload: Any = None
    bindings: dict[str, Any] = {}
    resource_uri = ""
    if target_type == "finance_run":
        from lvke_mcp.domains.finance.run_service import get_workspace_finance_run
        payload = get_workspace_finance_run(
            workspace_id,
            run_id=target_id,
            view="full",
        )
        if not payload.get("available") or str(payload.get("run_id") or "") != target_id:
            blockers.append("finance_run_not_found")
        bindings["finance_run_id"] = target_id
    elif target_type == "finance_tables_package":
        from lvke_mcp.servers.lvke_finance_tables.service import PACKAGE_STORE
        record = PACKAGE_STORE.get(
            workspace_id,
            target_id,
        )
        payload = record
        if record is None:
            blockers.append("finance_tables_package_not_found")
        else:
            bindings["finance_run_id"] = str((record.get("payload") or {}).get("run_id") or "")
            resource_uri = str(record.get("resource_uri") or "")
    elif target_type == "finance_xlsx":
        path = _safe_file(workspace_id, str(target.get("file_path") or target_id))
        if path is None or path.suffix.lower() not in {".xlsx", ".xlsm"}:
            blockers.append("finance_xlsx_not_found_or_outside_workspace")
        else:
            digest = "sha256:" + hashlib.sha256(path.read_bytes()).hexdigest()
            payload = {"path": str(path), "size": path.stat().st_size, "sha256": digest}
    elif target_type == "finance_xlsx_source":
        from lvke_mcp.servers.lvke_source_files.backend import resolve_source_workbook_for_review
        source_file_id = str(target.get("source_file_id") or "").strip()
        if not source_file_id:
            blockers.append("finance_xlsx_source_file_id_required")
        else:
            resolution = resolve_source_workbook_for_review(
                workspace_id,
                source_file_id,
            )
            if not resolution.get("ok"):
                blockers.append(str(resolution.get("code") or "source_workbook_not_found"))
            else:
                payload = {
                    "path": str(resolution.get("path") or ""),
                    "size": int(resolution.get("size") or 0),
                    "sha256": str(resolution.get("sha256") or ""),
                    "source_file_id": str(resolution.get("source_file_id") or source_file_id),
                    "original_filename": str(resolution.get("original_filename") or ""),
                }
                bindings["source_file_id"] = str(resolution.get("source_file_id") or source_file_id)
    elif target_type == "acquisition_run":
        from lvke_mcp.servers.lvke_asset_acquisition.backend import get_run
        payload = get_run(
            workspace_id,
            target_id,
        )
        if not payload.get("available") or str(payload.get("run_id") or "") != target_id:
            blockers.append("acquisition_run_not_found")
        bindings["finance_run_id"] = target_id
    elif target_type == "acquisition_tables_package":
        from lvke_mcp.servers.lvke_asset_acquisition.tables import get_package_record
        record = get_package_record(workspace_id, target_id)
        payload = record
        if record is None:
            blockers.append("acquisition_tables_package_not_found")
        else:
            bindings["finance_run_id"] = str((record.get("payload") or {}).get("run_id") or "")
            resource_uri = str(record.get("resource_uri") or "")
    elif target_type == "report_revision":
        from lvke_mcp.servers.lvke_report_generation.service import REVISION_STORE
        record = REVISION_STORE.get(
            workspace_id,
            target_id,
        )
        payload = record
        if record is None:
            blockers.append("report_revision_not_found")
        else:
            record_payload = record.get("payload") or {}
            native_revision_id = str(record_payload.get("native_revision_id") or "")
            document = record_payload.get("document_snapshot")
            if not isinstance(document, dict):
                try:
                    from lvke_mcp.domains.reports.doc_service import read_document

                    document = read_document(
                        workspace_id,
                        revision_id=native_revision_id,
                    )
                except Exception:  # noqa: BLE001
                    blockers.append("report_native_revision_not_found")
                    document = {}
            payload = {"revision_record": record, "document": document}
            upstream = record_payload.get("upstream") or {}
            bindings = {
                "finance_run_id": str(upstream.get("run_id") or ""),
                "finance_tables_package_id": str(upstream.get("finance_tables_package_id") or ""),
                "evidence_pack_ids": list(upstream.get("evidence_pack_ids") or []),
                "research_package_ids": list(upstream.get("research_package_ids") or []),
            }
            resource_uri = str(record.get("resource_uri") or "")
    elif target_type == "report_artifact":
        artifact_domain = str(target.get("artifact_domain") or "").strip()
        if artifact_domain not in _REPORT_ARTIFACT_DOMAINS:
            blockers.append(
                "report_artifact_domain_required"
                if not artifact_domain else "report_artifact_domain_invalid"
            )
        else:
            target["artifact_domain"] = artifact_domain
            payload, bindings, artifact_blockers = _resolve_report_artifact(
                workspace_id,
                target_id,
                artifact_domain=artifact_domain,
            )
            blockers.extend(artifact_blockers)
    elif target_type == "combined_deliverable":
        component_targets = target.get("components") or []
        if not isinstance(component_targets, list) or len(component_targets) < 2:
            blockers.append("combined_components_required")
            payload = {"components": []}
        else:
            components = []
            for raw_component in component_targets:
                try:
                    component = normalize_target(raw_component)
                except ValueError as exc:
                    blockers.append(f"combined_component_invalid:{exc}")
                    continue
                if component["target_type"] == "combined_deliverable":
                    blockers.append("combined_component_recursive")
                    continue
                resolved, component_blockers = _resolve_target(
                    workspace_id,
                    component,
                )
                blockers.extend(f"{component['target_id']}:{item}" for item in component_blockers)
                if resolved:
                    components.append(resolved)
            payload = {"components": components}
            bindings = _combined_bindings_manifest(components)
    if payload is None:
        payload = {}
    target_hash = (
        str(payload.get("sha256") or "")
        if target_type in {"finance_xlsx", "finance_xlsx_source"} and isinstance(payload, dict)
        else sha256_json(payload)
    )
    return ({
        "target_type": target_type, "target_id": target_id, "target_sha256": target_hash,
        "snapshot": payload, "bindings": bindings, "resource_uri": resource_uri,
        "target_spec": target,
    } if not blockers else None), blockers


def _acquisition_run_snapshot(run: dict[str, Any]) -> dict[str, Any]:
    """Project acquisition-run business state without release bookkeeping."""

    snapshot = deepcopy(run)
    release_states = {"release_ready", "released"}
    history = [
        deepcopy(event)
        for event in snapshot.get("state_history") or []
        if not isinstance(event, dict)
        or str(event.get("status") or "") not in release_states
    ]
    if "state_history" in snapshot:
        snapshot["state_history"] = history
    if str(snapshot.get("lifecycle_status") or "") in release_states:
        previous_status = next(
            (
                str(event.get("status") or "")
                for event in reversed(history)
                if isinstance(event, dict) and str(event.get("status") or "")
            ),
            "",
        )
        if previous_status:
            snapshot["lifecycle_status"] = previous_status
        else:
            snapshot.pop("lifecycle_status", None)
    for field in ("release_status", "release_ready_at", "released_at", "release", "release_history"):
        snapshot.pop(field, None)
    snapshot = {
        key: value for key, value in snapshot.items()
        if not key.endswith("_by")
    }
    return snapshot


def _binding_snapshot(
    workspace_id: str,
    bindings: dict[str, Any],
) -> dict[str, Any]:
    """Resolve immutable upstream IDs to content hashes for invalidation checks."""

    snapshot: dict[str, Any] = {}
    finance_run_id = str(bindings.get("finance_run_id") or "")
    if finance_run_id:
        run: dict[str, Any] = {}
        acquisition_run = False
        try:
            from lvke_mcp.domains.finance.run_service import get_workspace_finance_run

            run = get_workspace_finance_run(
                workspace_id,
                run_id=finance_run_id,
                view="full",
            )
        except Exception:  # noqa: BLE001 - acquisition IDs use a separate service
            pass
        if not run.get("available"):
            try:
                from lvke_mcp.servers.lvke_asset_acquisition.backend import get_run

                run = get_run(
                    workspace_id,
                    finance_run_id,
                )
                acquisition_run = True
            except Exception:  # noqa: BLE001
                run = {}
        run_snapshot = _acquisition_run_snapshot(run) if acquisition_run else run
        snapshot["finance_run"] = {
            "id": finance_run_id,
            "content_hash": (
                sha256_json(run_snapshot)
                if run and run.get("available") else None
            ),
        }

    table_id = str(bindings.get("finance_tables_package_id") or "")
    if table_id:
        record = None
        try:
            from lvke_mcp.servers.lvke_finance_tables.service import PACKAGE_STORE

            record = PACKAGE_STORE.get(
                workspace_id,
                table_id,
            )
        except (ValueError, OSError):
            record = None
        if record is None:
            try:
                from lvke_mcp.servers.lvke_asset_acquisition.tables import get_package_record

                record = get_package_record(workspace_id, table_id)
            except (ValueError, OSError):
                record = None
        snapshot["finance_tables_package"] = {
            "id": table_id,
            "content_hash": (record or {}).get("content_hash"),
            "basis_hash": (record or {}).get("basis_hash"),
        }

    report_revision_id = str(bindings.get("report_revision_id") or "")
    if report_revision_id:
        try:
            from lvke_mcp.servers.lvke_report_generation.service import REVISION_STORE

            record = REVISION_STORE.get(
                workspace_id,
                report_revision_id,
            )
        except (OSError, ValueError):
            record = None
        snapshot["report_revision"] = {
            "id": report_revision_id,
            "content_hash": (record or {}).get("content_hash"),
            "basis_hash": (record or {}).get("basis_hash"),
        }

    try:
        from lvke_mcp.servers.lvke_data_analysis.service import EVIDENCE_STORE
    except Exception:  # noqa: BLE001
        EVIDENCE_STORE = None  # type: ignore[assignment]
    evidence_rows = []
    for object_id in bindings.get("evidence_pack_ids") or []:
        record = (
            EVIDENCE_STORE.get(
                workspace_id,
                str(object_id),
            )
            if EVIDENCE_STORE
            else None
        )
        evidence_rows.append({
            "id": str(object_id), "content_hash": (record or {}).get("content_hash"),
            "basis_hash": (record or {}).get("basis_hash"),
        })
    if evidence_rows:
        snapshot["evidence_packs"] = evidence_rows

    try:
        from lvke_mcp.servers.lvke_deep_research.package_service import PACKAGE_STORE as RESEARCH_STORE
    except Exception:  # noqa: BLE001
        RESEARCH_STORE = None  # type: ignore[assignment]
    research_rows = []
    for object_id in bindings.get("research_package_ids") or []:
        record = (
            RESEARCH_STORE.get(
                workspace_id,
                str(object_id),
            )
            if RESEARCH_STORE
            else None
        )
        research_rows.append({
            "id": str(object_id), "content_hash": (record or {}).get("content_hash"),
            "basis_hash": (record or {}).get("basis_hash"), "status": (record or {}).get("status"),
        })
    if research_rows:
        snapshot["research_packages"] = research_rows
    component_snapshots: list[dict[str, Any]] = []
    for component in bindings.get("component_bindings") or []:
        if not isinstance(component, dict):
            continue
        component_snapshots.append({
            "target_type": str(component.get("target_type") or ""),
            "target_id": str(component.get("target_id") or ""),
            "target_sha256": str(component.get("target_sha256") or ""),
            "upstream": _binding_snapshot(
                workspace_id,
                component.get("bindings")
                if isinstance(component.get("bindings"), dict)
                else {},
            ),
        })
    if component_snapshots:
        snapshot["components"] = component_snapshots
    return snapshot


def _mandatory_findings(standards: dict[str, Any]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for package in standards.get("packages") or []:
        package_id = str(package.get("package_id") or "")
        required = list(package.get("review_findings_required") or [])
        for index, text in enumerate(required, start=1):
            if package_id == "PKG-STD-021":
                rule_id = "HM-HOTEL-004"
                role = "legal"
                category = "land_use_compliance"
            elif package_id == "PKG-STD-022":
                rule_id = "HM-MINE-003"
                role = "legal"
                category = "mineral_rights_and_land"
            else:
                rule_id = f"CORE.REQUIRED.{package_id}.{index}"
                role = "business"
                category = "required_manual_review"
            rows.append({
                "rule_id": rule_id,
                "instance_key": f"{package_id}:{index}",
                "package_id": package_id,
                "message": str(text),
                "category": category,
                "review_area": role,
            })
    return rows


def prepare(args: dict[str, Any]) -> dict[str, Any]:
    def execute(workspace_id: str) -> dict[str, Any]:
        target = normalize_target(args.get("target"))
        resolved, blockers = _resolve_target(
            workspace_id,
            target,
        )
        if blockers or resolved is None:
            return _blocked(blockers[0], "审查目标无法完整解析", blockers=blockers)
        component_types = [
            str(item.get("target_type") or "")
            for item in (resolved.get("snapshot") or {}).get("components") or []
            if isinstance(item, dict)
        ]
        project_context = normalize_project_context(
            args.get("project_context"),
            target_type=target["target_type"],
        )
        pack = rules.compose(
            target["target_type"], args.get("rule_pack_ids") or [],
            args.get("industry_overlays") or [], component_types=component_types,
            project_context=project_context,
        )
        standards = rules.standards_snapshot(REPO_ROOT, pack["standard_package_ids"])
        upstream_snapshot = _binding_snapshot(
            workspace_id,
            resolved["bindings"],
        )
        legacy_gate_snapshot = _legacy_gate_snapshot(
            workspace_id,
            resolved,
        )
        mandatory_findings = _mandatory_findings(standards)
        basis = _preparation_basis({
            "target": {key: resolved[key] for key in ("target_type", "target_id", "target_sha256")},
            "bindings": resolved["bindings"], "upstream_snapshot": upstream_snapshot,
            "rule_pack": pack, "standards": standards,
            "legacy_gate_snapshot": legacy_gate_snapshot,
            "project_context": project_context,
            "engine_version": rules.ENGINE_VERSION, "recalculation_environment_version": rules.RECALC_ENV_VERSION,
        })
        preparation_payload = {
            **basis,
            "target_spec": resolved["target_spec"],
            "target_snapshot": resolved["snapshot"],
            "mandatory_findings": mandatory_findings,
        }
        record = PREPARATION_STORE.put(
            workspace_id, preparation_payload,
            producer="lvke-deliverable-review.review_prepare", status="ok",
            source_ids=[target["target_id"], *[str(value) for value in resolved["bindings"].values() if isinstance(value, str)]],
            basis=basis, schema_version="deliverable_review_preparation.v1",
        )
        verified_record, integrity_reasons = _verified_preparation_record(
            workspace_id,
            str(record.get("object_id") or ""),

            expected_basis_hash=sha256_json(basis),
            expected_content_hash=sha256_json(preparation_payload),
        )
        if verified_record is None:
            return _blocked(
                "preparation_integrity_failed",
                "审查准备对象写入后完整性校验失败",
                integrity_reasons=integrity_reasons,
            )
        record = verified_record
        warnings = [f"标准包未完成：{item}" for item in standards["incomplete"]]
        return _ok(
            review_preparation_id=record["object_id"], 
            target=basis["target"], bindings=basis["bindings"],
            rule_pack=pack, standards=standards, review_scope=pack["applicable_rules"],
            project_context=project_context,
            selected_rule_packs=pack.get("selected_rule_packs") or [],
            excluded_rule_packs=pack.get("excluded_rule_packs") or [],
            excluded_rules=pack.get("excluded_rules") or [],
            legacy_gate_snapshot=legacy_gate_snapshot,
            mandatory_findings=mandatory_findings,
            resource_uris=[record["resource_uri"]], warnings=warnings, blockers=[],
            next_actions=["调用 review_start 创建不可变审查运行"],
        )
    return _write("review_prepare", args, execute)


def _review_uri(workspace_id: str, review_id: str) -> str:
    return f"lvke://deliverable-review/workspaces/{workspace_id}/reviews/{review_id}"


def _metrics_uri(workspace_id: str) -> str:
    return f"lvke://deliverable-review/workspaces/{workspace_id}/metrics/current"


def _finding_uri(workspace_id: str, review_id: str, finding_id: str) -> str:
    return f"{_review_uri(workspace_id, review_id)}/findings/{finding_id}"


_PREPARATION_BASIS_FIELDS = (
    "target",
    "bindings",
    "upstream_snapshot",
    "rule_pack",
    "standards",
    "project_context",
    "legacy_gate_snapshot",
    "engine_version",
    "recalculation_environment_version",
)


def _preparation_basis(payload: dict[str, Any]) -> dict[str, Any]:
    return {
        key: deepcopy(payload.get(key))
        for key in _PREPARATION_BASIS_FIELDS
    }


def _verified_preparation_record(
    workspace_id: str,
    preparation_id: str,
    *,
    expected_basis_hash: str = "",
    expected_content_hash: str = "",
) -> tuple[dict[str, Any] | None, list[str]]:
    """Read and verify a content-addressed review preparation record."""

    try:
        workspace_id = require_safe_id(workspace_id, "workspace_id")
        preparation_id = require_safe_id(
            preparation_id,
            "review_preparation_id",
        )
        record = PREPARATION_STORE.get(workspace_id, preparation_id)
    except (OSError, ValueError):
        return None, ["preparation_record_unavailable"]
    if not isinstance(record, dict):
        return None, ["preparation_record_unavailable"]
    payload = record.get("payload")
    if not isinstance(payload, dict):
        return None, ["preparation_payload_invalid"]
    content_hash = sha256_json(payload)
    basis_hash = sha256_json(_preparation_basis(payload))
    expected_object_id = (
        f"{PREPARATION_STORE.id_prefix}_"
        f"{content_hash.removeprefix('sha256:')[:24]}"
    )
    reasons: list[str] = []
    if expected_basis_hash and basis_hash != expected_basis_hash:
        reasons.append("preparation_basis_binding_mismatch")
    if expected_content_hash and content_hash != expected_content_hash:
        reasons.append("preparation_content_binding_mismatch")
    if record.get("content_hash") != content_hash:
        reasons.append("preparation_content_hash_mismatch")
    if record.get("basis_hash") != basis_hash:
        reasons.append("preparation_basis_hash_mismatch")
    if record.get("object_id") != expected_object_id:
        reasons.append("preparation_object_id_mismatch")
    if preparation_id != expected_object_id:
        reasons.append("preparation_reference_id_mismatch")
    if record.get("workspace_id") != workspace_id:
        reasons.append("preparation_workspace_mismatch")
    if record.get("resource_uri") != PREPARATION_STORE.uri(
        workspace_id,
        expected_object_id,
    ):
        reasons.append("preparation_resource_uri_mismatch")
    if record.get("producer") != "lvke-deliverable-review.review_prepare":
        reasons.append("preparation_producer_mismatch")
    if record.get("schema_version") != "deliverable_review_preparation.v1":
        reasons.append("preparation_schema_mismatch")
    if record.get("status") != "ok":
        reasons.append("preparation_status_invalid")
    return (
        deepcopy(record) if not reasons else None,
        sorted(set(reasons)),
    )


def _standard_basis(preparation_payload: dict[str, Any]) -> list[dict[str, Any]]:
    basis: list[dict[str, Any]] = []
    for row in ((preparation_payload.get("standards") or {}).get("packages") or []):
        artifacts = row.get("artifacts") or []
        if not artifacts:
            basis.append({
                "standard_package_id": row.get("package_id"), "title": row.get("title"),
                "content_hash": row.get("source_manifest_sha256"), "gate_status": row.get("gate_status"),
            })
            continue
        for artifact in artifacts:
            basis.append({
                "standard_package_id": row.get("package_id"), "title": row.get("title"),
                "standard_artifact_id": artifact.get("artifact_id"),
                "publisher": artifact.get("publisher"), "document_number": artifact.get("document_number"),
                "publication_date": artifact.get("publication_date"),
                "source_url": artifact.get("source_url") or artifact.get("official_page_url"),
                "content_hash": artifact.get("sha256"), "gate_status": row.get("gate_status"),
            })
    return basis


def _required_finding_rows(
    preparation_payload: dict[str, Any], standard_basis: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    output: list[dict[str, Any]] = []
    pack = preparation_payload.get("rule_pack") or {}
    for row in preparation_payload.get("mandatory_findings") or []:
        location = {
            "target_id": (preparation_payload.get("target") or {}).get("target_id"),
            "standard_package_id": row.get("package_id"),
            "required_finding_instance": row.get("instance_key"),
        }
        item = rules.finding(
            str(row.get("rule_id") or "CORE.REQUIRED.MANUAL"), "P0", str(row.get("message") or "强制人工审查项"),
            category=str(row.get("category") or "required_manual_review"),
            expected="取得可回读原件并完成质量核验", actual="尚未提供或尚未完成核验",
            target_location=location, standard_basis=[
                item for item in standard_basis
                if item.get("standard_package_id") == row.get("package_id")
            ],
            review_area=str(row.get("review_area") or "legal"),
            remediation="补充原件、精确定位和内容哈希；生成新目标版本后复测并完成质量核验",
        )
        item["rule_pack_id"] = pack.get("rule_pack_id")
        item["rule_pack_version"] = pack.get("version")
        item["waiver_allowed"] = False
        item["required_source_document"] = True
        output.append(item)
    return output


def _professional_rule_finding(
    preparation_payload: dict[str, Any],
    source_rule: dict[str, Any],
    standard_basis: list[dict[str, Any]],
) -> dict[str, Any]:
    target = preparation_payload.get("target") or {}
    standard = source_rule.get("standard") or {}
    package_id = str(standard.get("package_id") or "")
    artifact_id = str(standard.get("artifact_id") or "")
    matching_basis = [
        deepcopy(row)
        for row in standard_basis
        if str(row.get("standard_package_id") or "") == package_id
        and (
            not artifact_id
            or str(row.get("standard_artifact_id") or "") == artifact_id
        )
    ]
    if not matching_basis:
        matching_basis = [{}]
    precise_basis = [
        {
            **row,
            "standard_package_id": package_id,
            "standard_artifact_id": artifact_id,
            "content_hash": standard.get("sha256") or row.get("content_hash"),
            "locator": standard.get("locator"),
            "quote": standard.get("quote"),
        }
        for row in matching_basis
    ]
    target_kinds = set(source_rule.get("target_kinds") or [])
    location: dict[str, Any] = {
        "target_type": target.get("target_type"),
        "target_id": target.get("target_id"),
        "professional_rule_id": source_rule.get("rule_id"),
    }
    if target.get("target_type") == "combined_deliverable":
        location["components"] = [
            {
                "target_type": component.get("target_type"),
                "target_id": component.get("target_id"),
            }
            for component in (
                (preparation_payload.get("target_snapshot") or {}).get("components")
                or []
            )
            if component.get("target_type") in target_kinds
        ]
    item = rules.finding(
        str(source_rule.get("rule_id") or "PROFESSIONAL.REVIEW"),
        _severity(
            source_rule.get("severity"),
            blocking=bool(source_rule.get("blocking")),
        ),
        f"待专业核验：{source_rule.get('title') or source_rule.get('rule_id')}",
        category="professional_review_pending",
        blocking=bool(source_rule.get("blocking")),
        expected=str(source_rule.get("requirement") or "完成证据化专业核验"),
        actual="尚未提交逐规则专业核验结论",
        target_location=location,
        standard_basis=precise_basis,
        review_area=str(source_rule.get("review_area") or ""),
        remediation="逐规则提交带内容哈希和精确定位的核验证据",
    )
    item["manual_review_required"] = True
    item["waiver_allowed"] = False
    item["professional_rule"] = {
        key: deepcopy(source_rule.get(key))
        for key in (
            "rule_id", "title", "requirement", "check_kind", "review_area",
            "severity", "blocking", "target_kinds", "on_unavailable",
        )
    }
    return item


def _severity(value: Any, *, blocking: bool = False) -> str:
    text = str(value or "").lower()
    if text in {"p0", "critical", "fatal"}:
        return "P0"
    if text in {"p1", "high", "error", "major"} or blocking:
        return "P1"
    if text in {"p3", "low", "info", "minor"}:
        return "P3"
    return "P2"


def _run_from_preparation(
    workspace_id: str,
    preparation_payload: dict[str, Any],
) -> dict[str, Any]:
    target = preparation_payload.get("target") or {}
    target_type = str(target.get("target_type") or "")
    snapshot = preparation_payload.get("target_snapshot") or {}
    if target_type == "finance_run":
        return snapshot if isinstance(snapshot, dict) else {}
    if target_type == "acquisition_run":
        run = snapshot if isinstance(snapshot, dict) else {}
        if run and not isinstance(run.get("spec"), dict):
            try:
                from lvke_mcp.servers.lvke_asset_acquisition.backend import get_spec

                spec_row = get_spec(
                    workspace_id,
                    str(run.get("spec_id") or ""),
                )
                run = {**run, "spec": deepcopy(spec_row.get("spec") or {})}
            except Exception:  # noqa: BLE001 - missing spec remains an explicit incomplete input
                pass
        return run
    run_id = str((preparation_payload.get("bindings") or {}).get("finance_run_id") or "")
    if not run_id:
        return {}
    if target_type == "acquisition_tables_package":
        from lvke_mcp.servers.lvke_asset_acquisition.backend import get_run, get_spec

        run = get_run(workspace_id, run_id)
        spec_row = (
            get_spec(
                workspace_id,
                str(run.get("spec_id") or ""),
            )
            if run
            else {}
        )
        return {**run, "spec": deepcopy(spec_row.get("spec") or {})} if run else {}
    from lvke_mcp.domains.finance.run_service import get_workspace_finance_run

    run = get_workspace_finance_run(
        workspace_id,
        run_id=run_id,
        view="full",
    )
    if run.get("available") and str(run.get("run_id") or "") == run_id:
        return run
    try:
        from lvke_mcp.servers.lvke_asset_acquisition.backend import get_run, get_spec

        acquisition_run = get_run(
            workspace_id,
            run_id,
        )
        if acquisition_run.get("available"):
            spec_row = get_spec(
                workspace_id,
                str(acquisition_run.get("spec_id") or ""),
            )
            return {**acquisition_run, "spec": deepcopy(spec_row.get("spec") or {})}
    except Exception:  # noqa: BLE001 - caller records unavailable bound run
        pass
    return run


def _existing_issue_findings(run: dict[str, Any], standard_basis: list[dict[str, Any]]) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    if run and run.get("consistency_ok") is False:
        findings.append(rules.finding(
            "FIN.EXISTING.CONSISTENCY", "P0", "既有财务一致性门禁未通过",
            category="financial_consistency", expected=True, actual=False,
            review_area="finance", remediation="修正财务模型并生成新 run 后复测",
            standard_basis=standard_basis,
        ))
    legacy_issues = [
        *(run.get("issues") or []),
        *((run.get("audit") or {}).get("issues") or []),
    ]
    seen_issue_ids: set[str] = set()
    for issue in legacy_issues:
        if not isinstance(issue, dict):
            continue
        is_failed = issue.get("ok") is False or (
            str(issue.get("status") or "open") == "open"
            and bool(issue.get("blocking"))
        )
        if not is_failed:
            continue
        source_id = str(issue.get("issue_id") or "")
        issue_identity = source_id or sha256_json(issue)
        if issue_identity in seen_issue_ids:
            continue
        seen_issue_ids.add(issue_identity)
        rule_id = str(issue.get("rule") or "FIN.EXISTING.ISSUE")
        detail = issue.get("detail")
        message = (
            str(detail.get("message") or detail.get("detail") or rule_id)
            if isinstance(detail, dict) else str(detail or rule_id)
        )
        findings.append(rules.finding(
            f"FIN.LEGACY.{rule_id}", _severity(issue.get("severity"), blocking=bool(issue.get("blocking"))),
            message, category="legacy_finance_issue", blocking=bool(issue.get("blocking")),
            actual=detail, target_location={"run_id": run.get("run_id"), "source_issue_id": source_id},
            evidence=[{"source_issue_id": source_id, "audit_history": issue.get("history") or []}],
            standard_basis=standard_basis, review_area="finance",
            remediation="沿用原问题证据关闭流程整改，并以新 run 复测", source_issue_id=source_id,
        ))
    return findings


def _project_metadata_findings(
    preparation_payload: dict[str, Any], run: dict[str, Any], standard_basis: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    target = preparation_payload.get("target") or {}
    target_type = str(target.get("target_type") or "")
    snapshot = preparation_payload.get("target_snapshot") or {}
    if target_type in {"finance_run", "finance_tables_package"}:
        source = run
    else:
        revision_payload = ((snapshot.get("revision_record") or {}).get("payload") or {}) if isinstance(snapshot, dict) else {}
        source = {
            **dict(revision_payload.get("project_metadata") or {}),
            **dict((preparation_payload.get("target_spec") or {}).get("project_metadata") or {}),
        }
        if run:
            source = {**run, **source}
    aliases = {
        "project_type": ("project_type", "invest_type"),
        "industry": ("industry",),
        "valuation_date": ("valuation_date", "base_date", "as_of_date"),
        "currency": ("currency", "currency_code"),
        "amount_unit": ("amount_unit", "unit"),
        "tax_basis": ("tax_basis", "tax_inclusive_basis", "price_tax_basis"),
        "forecast_period": ("forecast_period", "calc_years", "forecast_years"),
    }
    missing = [
        field for field, candidates in aliases.items()
        if not any(source.get(candidate) not in (None, "", []) for candidate in candidates)
    ]
    if not missing:
        return []
    return [rules.finding(
        "PROJECT.METADATA.COMPLETE", "P1", "项目审查元数据不完整",
        category="project_metadata", expected=sorted(aliases), actual={"missing": missing},
        target_location={"target_id": target.get("target_id"), "fields": missing},
        standard_basis=standard_basis, review_area="business",
        remediation="补齐项目类型、行业、估值基准日、币种、金额单位、税口径和预测期间后生成新版本",
    )]


def _acquisition_input_findings(
    run: dict[str, Any], target_id: str, standard_basis: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    spec = run.get("spec") or {}
    transaction = dict(
        run.get("transaction") or run.get("acquisition") or spec.get("transaction") or {}
    )
    assets = run.get("assets") or run.get("asset_scope") or transaction.get("assets")
    required = {
        "valuation_date": transaction.get("valuation_date") or run.get("valuation_date"),
        "purchase_price": transaction.get("purchase_price") or transaction.get("consideration"),
        "transaction_tax": transaction.get("transaction_tax") or transaction.get("taxes"),
        "asset_scope": assets,
    }
    missing = [key for key, value in required.items() if value in (None, "", [], {})]
    if not missing:
        return []
    return [rules.finding(
        "ACQ.TRANSACTION.INPUTS", "P1", "资产收购交易输入或资产边界不完整",
        category="acquisition_inputs", expected=sorted(required), actual={"missing": missing},
        target_location={"target_id": target_id, "fields": missing}, standard_basis=standard_basis,
        review_area="business", remediation="补齐交易价格、税费、估值基准日和资产边界原始依据",
    )]


def _hotel_acquisition_run_findings(
    run: dict[str, Any], target_id: str, standard_basis: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], set[str]]:
    spec = run.get("spec") or {}
    transaction = spec.get("transaction") or run.get("transaction") or {}
    findings: list[dict[str, Any]] = []
    mode = str(transaction.get("operating_mode") or run.get("operating_mode") or "")
    if mode not in {"owner_lessor", "mixed_owner_operator"}:
        findings.append(rules.finding(
            "HOTEL.OPERATING_MODEL", "P0", "酒店收购经营模式未确认或不受模型支持",
            category="operating_assumption",
            expected=["owner_lessor", "mixed_owner_operator"], actual=mode or None,
            target_location={"target_id": target_id, "field": "transaction.operating_mode"},
            standard_basis=standard_basis, review_area="business",
            remediation="确认纯出租或混合自营模式，并按确认模式完整重算收入、成本、税费和现金流",
        ))
    assets = transaction.get("asset_scope") or run.get("asset_scope") or []
    licenses = transaction.get("licenses") or []
    parties = spec.get("project_parties") or []
    license_holders = [
        row for row in parties
        if isinstance(row, dict) and "license_holder" in (row.get("roles") or [])
    ]
    asset_evidence_ok = bool(assets) and all(
        isinstance(row, dict) and list(row.get("evidence_ids") or [])
        for row in assets
    )
    license_evidence_ok = bool(licenses or license_holders) and all(
        isinstance(row, dict) and list(row.get("evidence_ids") or [])
        for row in [*licenses, *license_holders]
    )
    if not asset_evidence_ok or not license_evidence_ok:
        findings.append(rules.finding(
            "HOTEL.RIGHTS.LICENSES", "P0", "酒店资产权属或经营许可缺少逐项原件证据绑定",
            category="rights_and_licenses",
            expected={"asset_scope_with_evidence": True, "licenses_with_evidence": True},
            actual={
                "asset_scope_count": len(assets), "asset_evidence_complete": asset_evidence_ok,
                "license_count": len(licenses) + len(license_holders),
                "license_evidence_complete": license_evidence_ok,
            },
            target_location={"target_id": target_id, "fields": ["transaction.asset_scope", "transaction.licenses", "project_parties"]},
            standard_basis=standard_basis, review_area="legal",
            remediation="逐项绑定权证、许可、主体授权及用途原件，并核对证载主体、地址、范围和有效期",
        ))
    return findings, {"HOTEL.RIGHTS.LICENSES", "HOTEL.OPERATING_MODEL"}


def _number(value: Any) -> float | None:
    try:
        return float(value) if value not in (None, "") else None
    except (TypeError, ValueError):
        return None


def _finance_recalculation_findings(run: dict[str, Any], standard_basis: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[str], dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    incomplete: list[str] = []
    coverage: dict[str, Any] = {"finance_recalculations": []}
    if not run or not run.get("available"):
        return findings, ["bound_finance_run_unavailable"], coverage
    investment = run.get("investment") or {}
    funding = run.get("funding") or {}
    total = _number(investment.get("total"))
    components = [_number(investment.get(key)) for key in ("construction", "interest", "working_capital")]
    if total is not None and all(value is not None for value in components):
        calculated = sum(float(value) for value in components if value is not None)
        difference = round(total - calculated, 6)
        tolerance = max(0.01, abs(total) * 1e-8)
        coverage["finance_recalculations"].append("total_investment")
        if abs(difference) > tolerance:
            item = rules.finding(
                "FIN.INVESTMENT.BALANCE", "P0", "总投资与建设投资、建设期利息和流动资金不平",
                category="financial_recalculation", expected=calculated, actual=total,
                difference=difference, tolerance=tolerance, target_location={"run_id": run.get("run_id"), "field": "investment.total"},
                standard_basis=standard_basis, review_area="finance", remediation="修正投资构成并重新生成财务 run",
            )
            item["calculation_trace"] = ["total = construction + interest + working_capital", f"{total} - {calculated} = {difference}"]
            findings.append(item)
    else:
        incomplete.append("investment_balance_inputs_missing")
    sources = [_number(funding.get(key)) for key in ("capital", "loan", "subsidy")]
    if total is not None and all(value is not None for value in sources):
        calculated = sum(float(value) for value in sources if value is not None)
        difference = round(total - calculated, 6)
        tolerance = max(0.01, abs(total) * 1e-8)
        coverage["finance_recalculations"].append("funding_balance")
        if abs(difference) > tolerance:
            item = rules.finding(
                "FIN.FUNDING.BALANCE", "P0", "总资金来源与投资需求不平",
                category="financial_recalculation", expected=total, actual=calculated,
                difference=round(calculated - total, 6), tolerance=tolerance,
                target_location={"run_id": run.get("run_id"), "field": "funding"},
                standard_basis=standard_basis, review_area="finance", remediation="修正资金筹措构成并重新生成财务 run",
            )
            item["calculation_trace"] = ["funding = capital + loan + subsidy", f"{calculated} - {total} = {calculated-total}"]
            findings.append(item)
    else:
        incomplete.append("funding_balance_inputs_missing")
    cashflows = ((run.get("operating") or {}).get("cashflows") or [])
    if cashflows and all(_number(value) is not None for value in cashflows):
        numeric = [float(value) for value in cashflows]
        signs = [1 if value > 0 else -1 if value < 0 else 0 for value in numeric]
        sign_changes = sum(1 for left, right in zip(signs, signs[1:]) if left and right and left != right)
        if sign_changes > 1:
            findings.append(rules.finding(
                "FIN.IRR.MULTIPLE_SIGN_CHANGES", "P1", "现金流存在多次符号变化，IRR 可能多解",
                category="financial_recalculation", actual=sign_changes, expected=1,
                target_location={"run_id": run.get("run_id"), "field": "operating.cashflows"},
                standard_basis=standard_basis, review_area="finance", remediation="披露多重 IRR 风险并以 NPV 等指标交叉判断",
            ))
        try:
            from lvke_mcp.servers.finance_calc.calculations import irr, npv
            independent_irr = irr(numeric) * 100.0
            reported_irr = _number((run.get("indicators") or {}).get("project_irr_pct"))
            coverage["finance_recalculations"].append("project_irr")
            if reported_irr is None:
                incomplete.append("reported_project_irr_missing")
            elif abs(reported_irr - independent_irr) > 0.01:
                findings.append(rules.finding(
                    "FIN.IRR.INDEPENDENT_RECALC", "P0", "项目 IRR 与独立复算不一致",
                    category="financial_recalculation", expected=round(independent_irr, 6), actual=reported_irr,
                    difference=round(reported_irr-independent_irr, 6), tolerance=0.01,
                    target_location={"run_id": run.get("run_id"), "field": "indicators.project_irr_pct"},
                    standard_basis=standard_basis, review_area="finance", remediation="核对现金流期间、符号和 IRR 口径",
                ))
            residual = npv(numeric, independent_irr / 100.0)
            if abs(residual) > max(0.01, sum(abs(value) for value in numeric) * 1e-8):
                findings.append(rules.finding(
                    "FIN.IRR.NPV.RESIDUAL", "P0", "IRR 对应 NPV 残差超出容差",
                    category="financial_recalculation", expected=0.0, actual=residual, tolerance=0.01,
                    target_location={"run_id": run.get("run_id"), "field": "operating.cashflows"},
                    standard_basis=standard_basis, review_area="finance",
                ))
            rate = _number(run.get("benchmark_rate"))
            reported_npv = _number((run.get("indicators") or {}).get("npv_wan"))
            if rate is not None and reported_npv is not None:
                independent_npv = npv(numeric, rate)
                coverage["finance_recalculations"].append("project_npv")
                if abs(reported_npv-independent_npv) > max(0.01, abs(independent_npv)*1e-8):
                    findings.append(rules.finding(
                        "FIN.NPV.INDEPENDENT_RECALC", "P0", "项目 NPV 与独立复算不一致",
                        category="financial_recalculation", expected=round(independent_npv, 6), actual=reported_npv,
                        difference=round(reported_npv-independent_npv, 6), tolerance=0.01,
                        target_location={"run_id": run.get("run_id"), "field": "indicators.npv_wan"},
                        standard_basis=standard_basis, review_area="finance", remediation="核对折现率与现金流时点口径",
                    ))
            elif rate is None:
                incomplete.append("npv_discount_rate_missing")
        except ValueError:
            incomplete.append("independent_irr_unavailable")
    else:
        incomplete.append("project_cashflows_missing")
    return findings, sorted(set(incomplete)), coverage


def _flatten_numbers(value: Any, *, path: str = "", output: list[dict[str, Any]] | None = None) -> list[dict[str, Any]]:
    output = output if output is not None else []
    if isinstance(value, dict):
        for key, item in value.items():
            if key in {"target_snapshot", "spec_json", "result_snapshot"} and path:
                continue
            _flatten_numbers(item, path=f"{path}.{key}".strip("."), output=output)
    elif isinstance(value, list):
        for index, item in enumerate(value):
            _flatten_numbers(item, path=f"{path}[{index}]", output=output)
    elif isinstance(value, (int, float)) and not isinstance(value, bool):
        output.append({"path": path, "value": float(value)})
    return output


_CLAIM_PATTERN = re.compile(
    r"(?P<number>-?\d{1,3}(?:,\d{3})*(?:\.\d+)?|-?\d+(?:\.\d+)?)\s*"
    r"(?P<unit>亿元|万元|元|%|％|平方米|㎡|间|吨|万吨|年|个月)"
)
_FINANCE_WORDS = {
    "投资", "收入", "成本", "利润", "现金流", "贷款", "借款", "偿债",
    "IRR", "NPV", "收益率", "回收期", "资本金", "税", "折旧", "利息",
}


def _claim_value(number: str, unit: str) -> tuple[float, str]:
    value = float(number.replace(",", ""))
    if unit == "亿元":
        return value * 10000.0, "万元"
    if unit == "元":
        return value / 10000.0, "万元"
    if unit in {"%", "％"}:
        return value, "%"
    return value, unit


def _document_from_snapshot(snapshot: dict[str, Any]) -> dict[str, Any]:
    document = snapshot.get("document") if isinstance(snapshot, dict) else {}
    if isinstance(document, dict) and isinstance(document.get("content"), str):
        return document
    return {}


def _report_content(workspace_id: str, preparation_payload: dict[str, Any]) -> str:
    snapshot = preparation_payload.get("target_snapshot") or {}
    document = _document_from_snapshot(snapshot)
    content = str(document.get("content") or "")
    if content:
        return content
    if str((preparation_payload.get("target") or {}).get("target_type") or "") != "report_artifact":
        return ""
    return _report_artifact_text(
        workspace_id,
        snapshot,
    )


def _report_evidence_packs(workspace_id: str, preparation_payload: dict[str, Any]) -> list[dict[str, Any]]:
    try:
        from lvke_mcp.servers.lvke_data_analysis.service import EVIDENCE_STORE
    except Exception:  # noqa: BLE001
        return []
    records: list[dict[str, Any]] = []
    for evidence_id in (preparation_payload.get("bindings") or {}).get("evidence_pack_ids") or []:
        try:
            record = EVIDENCE_STORE.get(
                workspace_id,
                str(evidence_id),
            )
        except (OSError, ValueError):
            record = None
        if record is not None:
            records.append(record)
    return records


def _expected_report_sections(preparation_payload: dict[str, Any]) -> list[str]:
    snapshot = preparation_payload.get("target_snapshot") or {}
    revision = (snapshot.get("revision_record") or {}).get("payload") or {}
    upstream = revision.get("upstream") or {}
    return [str(item) for item in (upstream.get("outline") or []) if str(item).strip()]


def _report_findings(
    workspace_id: str, preparation_payload: dict[str, Any], run: dict[str, Any],
    standard_basis: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[str], dict[str, Any]]:
    target = preparation_payload.get("target") or {}
    target_type = str(target.get("target_type") or "")
    content = _report_content(workspace_id, preparation_payload)
    findings: list[dict[str, Any]] = []
    incomplete: list[str] = []
    coverage: dict[str, Any] = {"claim_count": 0, "financial_claim_count": 0, "matched_financial_claims": 0}
    if target_type == "report_revision":
        try:
            from lvke_mcp.servers.lvke_report_generation.service import validate as validate_report
            validation = validate_report(workspace_id, str(target.get("target_id") or ""))
        except Exception:  # noqa: BLE001
            validation = {"valid": False, "blockers": ["report_validator_failed"]}
        for blocker in validation.get("blockers") or []:
            code = str(blocker.get("code") if isinstance(blocker, dict) else blocker)
            findings.append(rules.finding(
                f"REPORT.LEGACY.{code}", "P1", f"既有研报校验阻断：{code}", category="report_validation",
                actual=blocker, target_location={"report_revision_id": target.get("target_id")},
                standard_basis=standard_basis, review_area="report", remediation="按既有报告校验提示修订后生成新 revision",
            ))
    if not content:
        return findings, ["report_content_unreadable"], coverage
    paragraphs = [line.strip() for line in content.splitlines() if line.strip()]
    seen: dict[str, int] = {}
    for index, paragraph in enumerate(paragraphs, start=1):
        normalized = re.sub(r"\s+", " ", re.sub(r"^#+\s*", "", paragraph)).strip()
        if len(normalized) >= 20:
            fingerprint = sha256_json(normalized)
            if fingerprint in seen:
                findings.append(rules.finding(
                    "REPORT.DUPLICATE.PARAGRAPH", "P2", "正文存在重复段落", category="report_integrity",
                    actual=normalized[:300], target_location={"paragraph": index, "duplicate_of": seen[fingerprint], "text_anchor": normalized[:80]},
                    standard_basis=standard_basis, review_area="report", remediation="删除旧修订残留或重复正文",
                ))
            else:
                seen[fingerprint] = index
        if re.search(r"(?:TODO|TBD|待补充|待确认|XX+|【[^】]*(?:待|占位)[^】]*】|\{\{[^}]+}})", paragraph, re.I):
            findings.append(rules.finding(
                "REPORT.PLACEHOLDER", "P1", "正文仍含空白占位或待确认内容", category="report_integrity", actual=paragraph[:300],
                target_location={"paragraph": index, "text_anchor": paragraph[:80]}, standard_basis=standard_basis,
                review_area="report", remediation="补齐正式内容及依据后生成新修订",
            ))
    pack = preparation_payload.get("rule_pack") or {}
    overlays = {
        str(item.get("rule_pack_id") or "")
        for item in (pack.get("components") or [])
        if isinstance(item, dict)
    }
    deterministic, missing, deterministic_metrics, executed = report_checks.review_report(
        content=content,
        target_id=str(target.get("target_id") or ""),
        run=run,
        evidence_packs=_report_evidence_packs(workspace_id, preparation_payload),
        expected_sections=_expected_report_sections(preparation_payload),
        overlays=overlays,
        standard_basis=standard_basis,
        review_as_of=str(preparation_payload.get("review_as_of") or ""),
        evidence_track=str((preparation_payload.get("project_context") or {}).get("evidence_track") or "real"),
    )
    findings.extend(deterministic)
    incomplete.extend(missing)
    coverage = {
        **coverage,
        **deterministic_metrics,
        "executed_rules": sorted(executed),
    }
    return findings, sorted(set(incomplete)), coverage


def _report_artifact_text(
    workspace_id: str,
    artifact: dict[str, Any],
) -> str:
    try:
        artifact_id = str(artifact.get("artifact_id") or "")
        artifact_family = str(artifact.get("artifact_family") or "generic")
        chunks: list[str] = []
        for item in artifact.get("files") or []:
            name = str(item.get("name") or item.get("filename") or "")
            if Path(name).suffix.lower() not in {".md", ".txt", ".docx"}:
                continue
            if artifact_family == "asset_acquisition":
                from lvke_mcp.servers.lvke_asset_acquisition import backend as acquisition_service

                resolved = acquisition_service.read_artifact_candidate_download(
                    workspace_id,
                    artifact_id,
                    name,
                )
                if resolved.get("ok") is not True:
                    continue
                content = resolved.get("content") or b""
            else:
                from lvke_mcp.domains.reports import artifacts as deliverable_artifacts

                resolved = deliverable_artifacts.read_artifact_candidate_download(
                    workspace_id, artifact_id, name, 
                )
                content = resolved.get("content") or b""
            if name.lower().endswith(".docx"):
                from docx import Document
                document = Document(io.BytesIO(content))
                chunks.extend(paragraph.text for paragraph in document.paragraphs)
                for table in document.tables:
                    for row in table.rows:
                        chunks.append(" | ".join(cell.text.strip() for cell in row.cells))
            else:
                chunks.append(content.decode("utf-8", errors="replace") if isinstance(content, bytes) else str(content))
        return "\n".join(chunks)
    except Exception:  # noqa: BLE001
        return ""


_EXTERNAL_GAP_REASON_MARKERS = (
    "formal_evidence",
    "evidence_pack",
    "source_",
    "standard_package_incomplete",
    "quality_review",
    "quality_note",
)
_LOCAL_IMPLEMENTATION_REASON_MARKERS = (
    "rule_not_executed",
    "review_engine_failed",
    "integrity",
    "lineage",
    "content_hash",
    "schema",
    "parser_unavailable",
    "target_reresolution_failed",
)
_EXTERNAL_GAP_CATEGORIES = {
    "acquisition_inputs",
    "citation_quality",
    "contract_consistency",
    "evidence",
    "evidence_conflict",
    "land_use_compliance",
    "market_evidence",
    "mineral_permits",
    "operating_assumption",
    "professional_review_pending",
    "project_metadata",
    "rights_and_area",
    "rights_and_licenses",
}
_LOCAL_IMPLEMENTATION_CATEGORIES = {
    "benchmark_applicability",
    "borrowing_cost",
    "combined_conclusion_consistency",
    "combined_numeric_consistency",
    "debt_service",
    "depreciation",
    "finance_consistency",
    "financial_consistency",
    "financial_recalculation",
    "financial_sustainability",
    "funding",
    "legacy_finance_issue",
    "report_finance_binding",
    "report_integrity",
    "report_internal_consistency",
    "report_validation",
    "sensitivity",
    "tax",
    "working_capital",
}


def _summarize_track_coverage(
    metrics: dict[str, Any],
    incomplete_reasons: list[str],
    findings: list[dict[str, Any]],
    *,
    evidence_track: str,
) -> dict[str, Any]:
    """Expose dual-track review counts without changing the gate verdict."""

    report_metrics = [
        row for row in metrics.get("report") or []
        if isinstance(row, dict)
    ]
    formal_count = sum(int(row.get("formal_evidence_claim_count") or 0) for row in report_metrics)
    technical_count = sum(int(row.get("technical_fixture_claim_count") or 0) for row in report_metrics)
    external_reason_count = sum(
        1 for reason in incomplete_reasons
        if (
            evidence_track == "real"
            and any(marker in str(reason) for marker in _EXTERNAL_GAP_REASON_MARKERS)
        )
    )
    local_reason_count = sum(
        1 for reason in incomplete_reasons
        if any(marker in str(reason) for marker in _LOCAL_IMPLEMENTATION_REASON_MARKERS)
    )
    external_finding_count = sum(
        1 for finding in findings
        if str(finding.get("category") or "") in _EXTERNAL_GAP_CATEGORIES
    )
    local_finding_count = sum(
        1 for finding in findings
        if str(finding.get("category") or "") in _LOCAL_IMPLEMENTATION_CATEGORIES
    )
    return {
        "evidence_track": evidence_track,
        "formal_evidence_claim_count": formal_count,
        "technical_fixture_claim_count": technical_count,
        "external_data_gap_count": external_reason_count + external_finding_count,
        "local_implementation_issue_count": local_reason_count + local_finding_count,
    }


def _component_preparation(parent: dict[str, Any], component: dict[str, Any]) -> dict[str, Any]:
    child = deepcopy(parent)
    child["target"] = {
        key: component.get(key) for key in ("target_type", "target_id", "target_sha256")
    }
    child["target_spec"] = component.get("target_spec") or child["target"]
    child["target_snapshot"] = component.get("snapshot") or {}
    child["bindings"] = component.get("bindings") or {}
    return child


def _execute_rules(
    workspace_id: str,
    preparation_payload: dict[str, Any],
    mode: str,
) -> dict[str, Any]:
    target = preparation_payload.get("target") or {}
    target_type = str(target.get("target_type") or "")
    pack = preparation_payload.get("rule_pack") or {}
    standard_basis = _standard_basis(preparation_payload)
    findings: list[dict[str, Any]] = _required_finding_rows(preparation_payload, standard_basis)
    incomplete = [f"standard_package_incomplete:{item}" for item in ((preparation_payload.get("standards") or {}).get("incomplete") or [])]
    metrics: dict[str, Any] = {}
    executed_rules: set[str] = {"CORE.TARGET.RESOLVED", "CORE.UPSTREAM.COMPLETE", "CORE.STANDARDS.LOCKED"}
    executed_rules.update(str(row.get("rule_id") or "") for row in findings)
    manual_routed_rules: list[str] = []
    applicable_rules = set(pack.get("applicable_rules") or [])
    reviewed_finance_runs: set[str] = set()
    professional_target_types = {target_type}
    if target_type == "combined_deliverable":
        professional_target_types.update(
            str(component.get("target_type") or "")
            for component in (
                (preparation_payload.get("target_snapshot") or {}).get("components")
                or []
            )
        )
    for source_rule in pack.get("rule_sources") or []:
        if source_rule.get("check_kind") != "professional":
            continue
        if not professional_target_types.intersection(
            set(source_rule.get("target_kinds") or [])
        ):
            continue
        rule_id = str(source_rule.get("rule_id") or "")
        executed_rules.add(rule_id)
        manual_routed_rules.append(rule_id)
        findings.append(_professional_rule_finding(
            preparation_payload,
            source_rule,
            standard_basis,
        ))

    def review_one(child: dict[str, Any]) -> None:
        nonlocal findings, incomplete
        child_type = str((child.get("target") or {}).get("target_type") or "")
        run = _run_from_preparation(
            workspace_id,
            child,
        )
        run_key = str(run.get("run_id") or run.get("id") or "") if run else ""
        review_finance_run = bool(run) and (
            not run_key or run_key not in reviewed_finance_runs
        )
        finance_targets = {
            "finance_run", "finance_tables_package", "acquisition_run",
            "acquisition_tables_package",
        }
        if child_type in finance_targets and review_finance_run:
            findings.extend(_existing_issue_findings(run, standard_basis))
            if run.get("available"):
                executed_rules.add("FIN.EXISTING.CHECKS")
            is_acquisition = child_type.startswith("acquisition_")
            if not is_acquisition:
                recalculated, missing, finance_metrics = _finance_recalculation_findings(
                    run, standard_basis,
                )
                findings.extend(recalculated)
                incomplete.extend(missing)
                metrics.setdefault("finance", []).append(finance_metrics)
                recalculation_rules = set(finance_metrics.get("finance_recalculations") or [])
                if "total_investment" in recalculation_rules:
                    executed_rules.add("FIN.INVESTMENT.BALANCE")
                if "funding_balance" in recalculation_rules:
                    executed_rules.add("FIN.FUNDING.BALANCE")
        if review_finance_run and (
            child_type in finance_targets
            or child_type == "report_revision"
        ):
            deterministic_rows, deterministic_missing, deterministic_executed, deterministic_metrics = (
                financial_checks.review_finance_run(
                    run,
                    target_id=str((child.get("target") or {}).get("target_id") or ""),
                    target_type=child_type,
                    applicable_rules=applicable_rules,
                    source_rule_rows=pack.get("rule_sources") or [],
                    standard_basis=standard_basis,
                )
            )
            findings.extend(deterministic_rows)
            incomplete.extend(deterministic_missing)
            executed_rules.update(deterministic_executed)
            metrics.setdefault("deterministic_finance", []).append(deterministic_metrics)
        if review_finance_run and run_key:
            reviewed_finance_runs.add(run_key)
        if "PROJECT.METADATA.COMPLETE" in applicable_rules:
            findings.extend(_project_metadata_findings(child, run, standard_basis))
            executed_rules.add("PROJECT.METADATA.COMPLETE")
        if child_type in {"acquisition_run", "acquisition_tables_package"} and "ACQ.TRANSACTION.INPUTS" in applicable_rules:
            findings.extend(_acquisition_input_findings(
                run, str((child.get("target") or {}).get("target_id") or ""), standard_basis,
            ))
            executed_rules.add("ACQ.TRANSACTION.INPUTS")
        if child_type in {"acquisition_run", "acquisition_tables_package"} and {
            "HOTEL.RIGHTS.LICENSES", "HOTEL.OPERATING_MODEL",
        }.intersection(applicable_rules):
            hotel_rows, hotel_executed = _hotel_acquisition_run_findings(
                run,
                str((child.get("target") or {}).get("target_id") or ""),
                standard_basis,
            )
            findings.extend(hotel_rows)
            executed_rules.update(hotel_executed)
        if child_type in {"finance_xlsx", "finance_xlsx_source"}:
            path = Path(str((child.get("target_snapshot") or {}).get("path") or ""))
            scanned, missing, xlsx_metrics = rules.scan_xlsx(path, deep=mode == "deep")
            for item in scanned:
                item["standard_basis"] = item.get("standard_basis") or standard_basis
                item["coverage_rule_id"] = "FIN.XLSX.INTEGRITY"
            findings.extend(scanned)
            incomplete.extend(missing)
            executed_rules.add("FIN.XLSX.INTEGRITY")
            if mode == "deep":
                recalculated, recalc_missing, recalc_metrics = rules.recalculate_xlsx(path)
                for item in recalculated:
                    item["standard_basis"] = item.get("standard_basis") or standard_basis
                    item["coverage_rule_id"] = "FIN.XLSX.RECALC"
                findings.extend(recalculated)
                incomplete.extend(recalc_missing)
                xlsx_metrics["recalculation"] = recalc_metrics
                if recalc_metrics.get("available") and not recalc_missing:
                    executed_rules.add("FIN.XLSX.RECALC")
            else:
                incomplete.append("deep_recalculation_not_executed")
            metrics.setdefault("xlsx", []).append(xlsx_metrics)
        if child_type in {"report_revision", "report_artifact"}:
            report_rows, missing, report_metrics = _report_findings(
                workspace_id, child, run, standard_basis,
            )
            findings.extend(report_rows)
            incomplete.extend(missing)
            metrics.setdefault("report", []).append(report_metrics)
            executed_rules.update(report_metrics.get("executed_rules") or [])
            executed_rules.update({"REPORT.PLACEHOLDER", "REPORT.DUPLICATE.PARAGRAPH"})
            if child_type == "report_revision":
                executed_rules.add("REPORT.EXISTING.VALIDATION")

    if target_type == "combined_deliverable":
        components = (preparation_payload.get("target_snapshot") or {}).get("components") or []
        finance_present = False
        report_present = False
        combined_reports: list[dict[str, Any]] = []
        combined_runs: list[dict[str, Any]] = []
        for component in components:
            component_type = str(component.get("target_type") or "")
            finance_present = finance_present or component_type in {
                "finance_run", "finance_tables_package", "finance_xlsx", "finance_xlsx_source",
                "acquisition_run", "acquisition_tables_package",
            }
            report_present = report_present or component_type in {"report_revision", "report_artifact"}
            child = _component_preparation(preparation_payload, component)
            if component_type in {"report_revision", "report_artifact"}:
                content = _report_content(workspace_id, child)
                if content:
                    combined_reports.append({
                        "target_id": (child.get("target") or {}).get("target_id"),
                        "content": content,
                    })
            elif component_type in {
                "finance_run", "finance_tables_package", "acquisition_run",
                "acquisition_tables_package",
            }:
                run = _run_from_preparation(
                    workspace_id,
                    child,
                )
                run_id = str(run.get("run_id") or run.get("id") or "") if run else ""
                if run and not any(
                    str(item.get("run_id") or item.get("id") or "") == run_id
                    for item in combined_runs
                ):
                    combined_runs.append(run)
            review_one(child)
        if not finance_present:
            incomplete.append("combined_finance_component_missing")
        if not report_present:
            incomplete.append("combined_report_component_missing")
        combined_rows, combined_missing, combined_metrics, combined_executed = (
            report_checks.review_combined(
                report_contents=combined_reports,
                finance_runs=combined_runs,
                target_id=str(target.get("target_id") or ""),
                standard_basis=standard_basis,
            )
        )
        findings.extend(combined_rows)
        incomplete.extend(combined_missing)
        metrics["combined"] = combined_metrics
        executed_rules.update(combined_executed)
        executed_rules.update({"COMBINED.BINDINGS.COMPLETE", "COMBINED.UPSTREAM.VERDICTS"})
    else:
        review_one(preparation_payload)

    unique: dict[str, dict[str, Any]] = {}
    for item in findings:
        item["rule_pack_id"] = pack.get("rule_pack_id")
        item["rule_pack_version"] = pack.get("version")
        unique[str(item["finding_id"])] = item
    findings = sorted(unique.values(), key=lambda row: (SEVERITY_ORDER.get(str(row.get("severity")), 9), str(row.get("finding_id"))))
    incomplete = sorted(set(incomplete))
    applicable = list(pack.get("applicable_rules") or [])
    if target_type == "combined_deliverable":
        component_types = {
            str(component.get("target_type") or "")
            for component in ((preparation_payload.get("target_snapshot") or {}).get("components") or [])
        }
        if not component_types.intersection({"finance_xlsx", "finance_xlsx_source"}):
            applicable = [rule_id for rule_id in applicable if not rule_id.startswith("FIN.XLSX.")]
        if not component_types.intersection({
            "finance_run", "finance_tables_package", "acquisition_run", "acquisition_tables_package",
        }):
            applicable = [
                rule_id for rule_id in applicable
                if rule_id not in {"FIN.EXISTING.CHECKS", "FIN.INVESTMENT.BALANCE", "FIN.FUNDING.BALANCE"}
            ]
        if "report_revision" not in component_types:
            applicable = [rule_id for rule_id in applicable if rule_id != "REPORT.EXISTING.VALIDATION"]
        source_targets = {
            str(row.get("rule_id") or ""): set(row.get("target_kinds") or [])
            for row in (pack.get("rule_sources") or [])
        }
        applicable = [
            rule_id for rule_id in applicable
            if rule_id not in source_targets
            or bool(source_targets[rule_id].intersection(component_types | {"combined_deliverable"}))
        ]
    coverage = {
        "applicable_rule_count": len(applicable), "executed_rule_count": len(set(applicable) & executed_rules),
        "coverage_ratio": round(len(set(applicable) & executed_rules) / len(applicable), 6) if applicable else 1.0,
        "applicable_rules": applicable, "executed_rules": sorted(executed_rules),
        "not_executed_rules": sorted(set(applicable) - executed_rules), "metrics": metrics,
        "manual_routed_rules": sorted(set(manual_routed_rules)),
        "deterministic_rule_count": sum(
            1 for row in (pack.get("rule_sources") or []) if row.get("check_kind") == "deterministic"
        ),
        "professional_rule_count": len(set(manual_routed_rules)),
        "rule_source_evidence": {
            str(row.get("rule_id") or ""): deepcopy(row.get("standard") or {})
            for row in (pack.get("rule_sources") or [])
        },
    }
    if coverage["not_executed_rules"]:
        incomplete.extend(f"rule_not_executed:{rule_id}" for rule_id in coverage["not_executed_rules"])
        incomplete = sorted(set(incomplete))
    coverage.update(_summarize_track_coverage(
        metrics,
        incomplete,
        findings,
        evidence_track=str(
            (preparation_payload.get("project_context") or {}).get("evidence_track")
            or "real"
        ),
    ))
    return {
        "findings": findings, "incomplete_reasons": incomplete, "coverage": coverage,
        "overall_verdict": verdict_for(findings, incomplete),
    }


def _preparation_execution_integrity_reasons(
    events: list[dict[str, Any]],
    preparation_payload: dict[str, Any],
) -> list[str]:
    if not events or events[0].get("event_type") != "review_created":
        return ["review_created_event_missing"]
    created = events[0].get("payload") or {}
    if not isinstance(created, dict):
        return ["review_created_payload_invalid"]
    reasons: list[str] = []
    expected_basis_hash = str(created.get("preparation_basis_hash") or "")
    expected_content_hash = str(
        created.get("preparation_content_hash") or ""
    )
    if not expected_basis_hash:
        reasons.append("preparation_basis_binding_missing")
    elif sha256_json(
        _preparation_basis(preparation_payload)
    ) != expected_basis_hash:
        reasons.append("preparation_basis_binding_mismatch")
    if (
        expected_content_hash
        and sha256_json(preparation_payload) != expected_content_hash
    ):
        reasons.append("preparation_content_binding_mismatch")
    for field in (
        "target",
        "target_spec",
        "bindings",
        "upstream_snapshot",
        "rule_pack",
        "standards",
        "legacy_gate_snapshot",
        "engine_version",
        "recalculation_environment_version",
    ):
        if deepcopy(created.get(field)) != deepcopy(
            preparation_payload.get(field)
        ):
            reasons.append(f"preparation_review_binding_mismatch:{field}")
    return sorted(set(reasons))


def _run_review(
    workspace_id: str,
    review_id: str,
    preparation_payload: dict[str, Any] | None,
    mode: str,
    preparation_integrity_reasons: list[str] | None = None,
) -> None:
    with STORE.mutation_guard(workspace_id, "review_engine_execute", review_id):
        events = STORE.events(workspace_id, review_id)
        if any(
            event.get("event_type") in {"review_completed", "review_failed"}
            for event in events
        ):
            return
        chain_ok, chain_reasons = STORE.verify_event_chain(
            workspace_id,
            review_id,
        )
        integrity_reasons = list(preparation_integrity_reasons or [])
        if not chain_ok:
            integrity_reasons.extend(chain_reasons or ["review_event_chain_invalid"])
        if isinstance(preparation_payload, dict):
            integrity_reasons.extend(
                _preparation_execution_integrity_reasons(
                    events,
                    preparation_payload,
                )
            )
        running_events = [
            event for event in events
            if event.get("event_type") == "review_running"
        ]
        resumed = bool(running_events)
        first_running_payload = (
            (running_events[0].get("payload") or {}) if running_events else {}
        )
        review_as_of = str(
            first_running_payload.get("review_as_of")
            or first_running_payload.get("started_at")
            or utc_now()
        )
        STORE.append(
            workspace_id,
            review_id,
            "review_running",
            {
                "started_at": utc_now(),
                "review_as_of": review_as_of,
                "resumed": resumed,
            },
        )
        if preparation_payload is None or integrity_reasons:
            incomplete_reason = (
                "review_preparation_integrity_failed"
                if integrity_reasons
                else "review_preparation_unavailable"
            )
            STORE.append(
                workspace_id,
                review_id,
                "review_failed",
                {
                    "completed_at": utc_now(),
                    "incomplete_reason": incomplete_reason,
                    "integrity_reasons": sorted(set(integrity_reasons)),
                },
            )
            return
        try:
            execution_payload = {
                **preparation_payload,
                "review_as_of": review_as_of,
            }
            result = _execute_rules(
                workspace_id,
                execution_payload,
                mode,
            )
        except Exception:  # noqa: BLE001 - fail closed and avoid leaking exception text
            STORE.append(
                workspace_id, review_id, "review_failed",
                {"completed_at": utc_now(), "incomplete_reason": "review_engine_failed"},
            )
        else:
            # Keep terminal persistence outside the engine exception handler.
            # If the event is durably written and the caller then crashes, a
            # retry observes the terminal event instead of appending a false
            # review_failed event after review_completed.
            STORE.append(
                workspace_id, review_id, "review_completed",
                {**result, "completed_at": utc_now()},
            )


def _run_async_review(
    workspace_id: str,
    review_id: str,
    preparation_payload: dict[str, Any] | None,
    mode: str,
    preparation_integrity_reasons: list[str] | None = None,
) -> None:
    key = (workspace_id, review_id)
    try:
        _run_review(
            workspace_id,
            review_id,
            preparation_payload,
            mode,
            preparation_integrity_reasons,
        )
    finally:
        with _ASYNC_LOCK:
            if _ASYNC_THREADS.get(key) is threading.current_thread():
                _ASYNC_THREADS.pop(key, None)


def _schedule_async_review(
    workspace_id: str,
    review_id: str,
    preparation_payload: dict[str, Any] | None,
    mode: str,
    preparation_integrity_reasons: list[str] | None = None,
) -> bool:
    key = (workspace_id, review_id)
    with _ASYNC_LOCK:
        existing = _ASYNC_THREADS.get(key)
        if existing is not None and existing.is_alive():
            return False
        thread = threading.Thread(
            target=_run_async_review,
            args=(
                workspace_id,
                review_id,
                preparation_payload,
                mode,
                preparation_integrity_reasons,
            ),
            name=f"deliverable-review-{review_id[-8:]}",
            daemon=True,
        )
        _ASYNC_THREADS[key] = thread
        thread.start()
    return True


def _resume_async_review_if_needed(workspace_id: str, state: dict[str, Any]) -> bool:
    if (
        state.get("execution") != "async"
        or state.get("reviewed")
        or state.get("invalidated")
        or state.get("released")
    ):
        return False
    preparation_id = str(state.get("review_preparation_id") or "")
    preparation, integrity_reasons = _verified_preparation_record(
        workspace_id,
        preparation_id,
        expected_basis_hash=str(
            state.get("preparation_basis_hash") or ""
        ),
        expected_content_hash=str(
            state.get("preparation_content_hash") or ""
        ),
    )
    payload = (preparation or {}).get("payload")
    return _schedule_async_review(
        workspace_id,
        str(state.get("review_id") or ""),
        payload if isinstance(payload, dict) else None,
        str(state.get("mode") or "deep"),
        integrity_reasons,
    )


def start(args: dict[str, Any]) -> dict[str, Any]:
    def execute(workspace_id: str) -> dict[str, Any]:
        preparation_id = str(args.get("review_preparation_id") or "")
        mode = str(args.get("mode") or "quick")
        execution = str(args.get("execution") or ("async" if mode == "deep" else "sync"))
        deployment_mode = str(args.get("deployment_mode") or "enforced")
        if mode not in {"quick", "deep"}:
            return _blocked("review_mode_invalid", "mode 必须为 quick 或 deep")
        if execution not in {"sync", "async"} or (mode == "quick" and execution == "async"):
            return _blocked("review_execution_invalid", "快速审查必须同步；深度审查支持同步或异步")
        if deployment_mode not in DEPLOYMENT_MODES:
            return _blocked(
                "review_deployment_mode_invalid",
                "deployment_mode 必须为 enforced 或 shadow",
            )
        start_identity = {
            "operation": "review_start",
            "idempotency_key": str(args.get("idempotency_key") or ""),
        }
        start_key_hash = sha256_json(start_identity)
        start_request_hash = sha256_json(args)
        review_id = (
            "review_"
            + start_key_hash.removeprefix("sha256:")[:32]
        )
        existing_events = STORE.events(workspace_id, review_id)
        if existing_events:
            first = existing_events[0]
            existing_created = first.get("payload") or {}
            if (
                first.get("event_type") != "review_created"
                or existing_created.get("start_key_hash") != start_key_hash
                or existing_created.get("start_request_hash") != start_request_hash
                or existing_created.get("review_preparation_id") != preparation_id
            ):
                # Raising keeps _write from caching the conflicting request
                # over the recoverable operation recorded in the event log.
                raise ValueError("idempotency_key_conflict")
        engine_terminal = any(
            event.get("event_type") in {"review_completed", "review_failed"}
            for event in existing_events
        )
        preparation: dict[str, Any] | None = None
        payload: dict[str, Any] | None = None
        preparation_integrity_reasons: list[str] = []
        if not engine_terminal:
            preparation, preparation_integrity_reasons = (
                _verified_preparation_record(
                    workspace_id,
                    preparation_id,

                    expected_basis_hash=str(
                        (existing_created if existing_events else {}).get(
                            "preparation_basis_hash"
                        )
                        or ""
                    ),
                    expected_content_hash=str(
                        (existing_created if existing_events else {}).get(
                            "preparation_content_hash"
                        )
                        or ""
                    ),
                )
            )
            candidate_payload = (preparation or {}).get("payload")
            payload = (
                candidate_payload
                if isinstance(candidate_payload, dict) else None
            )
        if not existing_events:
            if preparation is None or payload is None:
                if preparation_integrity_reasons and not set(
                    preparation_integrity_reasons
                ).intersection({
                    "preparation_record_unavailable",
                    "preparation_owner_mismatch",
                }):
                    return _blocked(
                        "preparation_integrity_failed",
                        "审查准备对象完整性校验失败",
                        integrity_reasons=preparation_integrity_reasons,
                    )
                return _blocked("preparation_not_found", _message("preparation_not_found"))
            created = {
                "review_preparation_id": preparation_id,
                "preparation_basis_hash": preparation.get("basis_hash"),
                "preparation_content_hash": preparation.get("content_hash"),
                "start_key_hash": start_key_hash,
                "start_request_hash": start_request_hash,
                "target": payload.get("target"),
                "target_spec": payload.get("target_spec"),
                "bindings": payload.get("bindings"),
                "upstream_snapshot": payload.get("upstream_snapshot"),
                "rule_pack": payload.get("rule_pack"),
                "project_context": payload.get("project_context"),
                "standards": payload.get("standards"),
                "legacy_gate_snapshot": payload.get("legacy_gate_snapshot"),
                "mode": mode,
                "execution": execution,
                "deployment_mode": deployment_mode,
                "engine_version": payload.get("engine_version"),
                "recalculation_environment_version": payload.get("recalculation_environment_version"),
                "created_at": utc_now(),
            }
            STORE.append(workspace_id, review_id, "review_created", created)
        if execution == "sync" and not engine_terminal:
            _run_review(
                workspace_id,
                review_id,
                payload,
                mode,
                preparation_integrity_reasons,
            )
        elif execution == "async" and not engine_terminal:
            _schedule_async_review(
                workspace_id,
                review_id,
                payload,
                mode,
                preparation_integrity_reasons,
            )
        current = _project(workspace_id, review_id, check_freshness=False)
        response_status = (
            "accepted" if execution == "async"
            else _review_envelope_status(current)
        )
        return _ok(
            status=response_status, review_id=review_id,
            review_status=current.get("review_status"), overall_verdict=current.get("overall_verdict"),
            deployment_mode=current.get("deployment_mode"),
            shadow_comparison=current.get("shadow_comparison") or {},
            release_ready=bool(current.get("release_ready")),
            resource_uris=[_review_uri(workspace_id, review_id)],
            blockers=current.get("blockers") or [], warnings=[],
            next_actions=["调用 review_get 查询深度审查进度"] if execution == "async" else ["处理 findings 或调用 review_release 固化质量检查结果"],
        )
    return _write("review_start", args, execute)


def _parse_timestamp(value: Any) -> datetime | None:
    text = str(value or "").strip()
    if not text:
        return None
    try:
        parsed = datetime.fromisoformat(text.replace("Z", "+00:00"))
    except ValueError:
        return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=timezone.utc)
    return parsed.astimezone(timezone.utc)


def _finding_match_key(row: dict[str, Any]) -> str:
    volatile_keys = {
        "run_id", "target_id", "report_revision_id", "document",
        "workbook", "file_path", "formula",
    }

    def stable_location(value: Any) -> Any:
        if isinstance(value, dict):
            return {
                key: stable_location(item)
                for key, item in value.items()
                if key not in volatile_keys
            }
        if isinstance(value, list):
            return [stable_location(item) for item in value]
        return value

    location = stable_location(row.get("target_location") or {})
    return sha256_json({
        "rule_id": row.get("rule_id"), "category": row.get("category"),
        "location": location, "source_issue_id": row.get("source_issue_id"),
    })


def _finding_coverage_rule_id(row: dict[str, Any]) -> str:
    explicit = str(row.get("coverage_rule_id") or "")
    if explicit:
        return explicit
    rule_id = str(row.get("rule_id") or "")
    if rule_id in {
        "FIN.XLSX.EMPTY_FORMULA_CACHE",
        "FIN.XLSX.RECALCULATED.ERROR",
    }:
        return "FIN.XLSX.RECALC"
    if rule_id.startswith("FIN.XLSX."):
        return "FIN.XLSX.INTEGRITY"
    return rule_id


def _classify_retest_operations(
    events: list[dict[str, Any]],
    review_id: str,
) -> dict[str, Any]:
    operation_rows: dict[str, list[dict[str, Any]]] = {}
    operation_event_types = {
        "retest_started",
        "retest_prepared",
        "retest_child_started",
        "finding_retested",
        "retest_linked",
        "retest_completed",
        "retest_failed",
    }
    for event in events:
        if event.get("event_type") not in operation_event_types:
            continue
        operation_id = str((event.get("payload") or {}).get("operation_id") or "")
        if operation_id:
            operation_rows.setdefault(operation_id, []).append(event)

    completed: set[str] = set()
    pending: set[str] = set()
    failed: dict[str, dict[str, Any]] = {}
    invalid: dict[str, str] = {}
    for operation_id, rows in operation_rows.items():
        intents = [row for row in rows if row.get("event_type") == "retest_started"]
        failures = [row for row in rows if row.get("event_type") == "retest_failed"]
        completions = [row for row in rows if row.get("event_type") == "retest_completed"]
        links = [row for row in rows if row.get("event_type") == "retest_linked"]
        findings = [row for row in rows if row.get("event_type") == "finding_retested"]
        if failures:
            if len(failures) != 1:
                invalid[operation_id] = "duplicate_failure_event"
            else:
                failed[operation_id] = deepcopy(failures[0].get("payload") or {})
            continue
        if intents:
            if len(intents) != 1:
                invalid[operation_id] = "duplicate_intent_event"
                continue
            if not completions:
                pending.add(operation_id)
                continue
            intent = intents[0].get("payload") or {}
            parent_completions = [
                row for row in completions
                if (row.get("payload") or {}).get("side") == "parent"
            ]
            expected_finding_ids = {
                str(item) for item in intent.get("expected_finding_ids") or []
            }
            actual_finding_ids = [
                str((row.get("payload") or {}).get("finding_id") or "")
                for row in findings
            ]
            valid_link = links[0].get("payload") or {} if len(links) == 1 else {}
            completion_payload = (
                parent_completions[0].get("payload") or {}
                if len(parent_completions) == 1 else {}
            )
            valid = bool(
                len(completions) == 1
                and len(parent_completions) == 1
                and len(links) == 1
                and str(intent.get("parent_review_id") or "") == review_id
                and str(valid_link.get("parent_review_id") or "") == review_id
                and len(actual_finding_ids) == len(expected_finding_ids)
                and set(actual_finding_ids) == expected_finding_ids
                and {
                    str(item)
                    for item in completion_payload.get("expected_finding_ids") or []
                } == expected_finding_ids
                and completion_payload.get("link_hash") == sha256_json(valid_link)
                and completion_payload.get("completed") is True
            )
            if valid:
                completed.add(operation_id)
            else:
                invalid[operation_id] = "incomplete_parent_completion"
            continue
        if not completions:
            pending.add(operation_id)
            continue
        child_completions = [
            row for row in completions
            if (row.get("payload") or {}).get("side") == "child"
        ]
        valid_link = links[0].get("payload") or {} if len(links) == 1 else {}
        completion_payload = (
            child_completions[0].get("payload") or {}
            if len(child_completions) == 1 else {}
        )
        valid = bool(
            len(completions) == 1
            and len(child_completions) == 1
            and len(links) == 1
            and str(valid_link.get("child_review_id") or "") == review_id
            and completion_payload.get("link_hash") == sha256_json(valid_link)
            and completion_payload.get("completed") is True
        )
        if valid:
            completed.add(operation_id)
        else:
            invalid[operation_id] = "incomplete_child_completion"
    return {
        "completed": completed,
        "pending": pending,
        "failed": failed,
        "invalid": invalid,
    }


def _gate_difference(legacy_verdict: str, unified_verdict: str) -> str:
    if legacy_verdict not in {"pass", "fail"} or unified_verdict not in {"pass", "fail"}:
        return "unavailable"
    if legacy_verdict == unified_verdict:
        return f"both_{legacy_verdict}"
    if legacy_verdict == "pass":
        return "legacy_pass_unified_block"
    return "legacy_block_unified_pass"


def _shadow_comparison(state: dict[str, Any], enforced_release_ready: bool) -> dict[str, Any]:
    legacy = state.get("legacy_gate_snapshot") or {}
    automated = str(state.get("automated_gate_verdict") or "unknown")
    formal = "pass" if enforced_release_ready else (
        "fail" if automated in {"pass", "fail"} else "unknown"
    )
    legacy_validation = str((legacy.get("validation") or {}).get("verdict") or "unknown")
    legacy_publish = str((legacy.get("publish") or {}).get("verdict") or "unknown")
    return {
        "schema_version": "deliverable_review_shadow_comparison.v1",
        "legacy_snapshot_hash": legacy.get("content_hash"),
        "legacy_validation_verdict": legacy_validation,
        "legacy_publish_verdict": legacy_publish,
        "unified_automated_verdict": automated,
        "unified_formal_gate_verdict": formal,
        "validation_difference": _gate_difference(legacy_validation, automated),
        "publish_difference": _gate_difference(legacy_publish, formal),
        "release_effect": "record_only_release_forbidden",
    }


def _project_events(workspace_id: str, review_id: str) -> dict[str, Any]:
    events = STORE.events(workspace_id, review_id)
    if not events:
        raise ValueError("review_not_found")
    chain_ok, chain_reasons = STORE.verify_event_chain(workspace_id, review_id)
    retest_operations = _classify_retest_operations(events, review_id)
    completed_retest_operations = retest_operations["completed"]
    retest_terminal_events = [
        event
        for event in events
        if (
            event.get("event_type") == "retest_failed"
            and str((event.get("payload") or {}).get("operation_id") or "")
            in retest_operations["failed"]
        ) or (
            event.get("event_type") == "retest_completed"
            and (event.get("payload") or {}).get("side") == "parent"
            and str((event.get("payload") or {}).get("operation_id") or "")
            in completed_retest_operations
        )
    ]
    latest_retest_terminal = max(
        retest_terminal_events,
        key=lambda row: int(row.get("sequence") or 0),
        default=None,
    )
    active_failed_retest_operations: dict[str, dict[str, Any]] = {}
    if latest_retest_terminal is not None and latest_retest_terminal.get("event_type") == "retest_failed":
        latest_operation_id = str(
            (latest_retest_terminal.get("payload") or {}).get("operation_id") or ""
        )
        active_failed_retest_operations[latest_operation_id] = deepcopy(
            retest_operations["failed"][latest_operation_id]
        )
    state: dict[str, Any] = {
        "schema_version": "deliverable_review.v1", "workspace_id": workspace_id,
        "review_id": review_id, "review_status": "created", "overall_verdict": "incomplete",
        "target": {}, "bindings": {}, "rule_pack": {}, "standards": {}, "project_context": {}, "findings": [],
        "incomplete_reasons": [], "coverage": {}, "quality_notes": [], "exports": [], "retests": [],
        "released": False, "release": None, "invalidated": False,
        "deployment_mode": "enforced", "legacy_gate_snapshot": {},
        "automated_gate_verdict": "unknown", "shadow_comparison": {},
        "pending_retest_operation_ids": sorted(retest_operations["pending"]),
        "failed_retest_operations": deepcopy(retest_operations["failed"]),
        "active_failed_retest_operations": active_failed_retest_operations,
        "invalid_retest_operations": deepcopy(retest_operations["invalid"]),
    }
    findings: dict[str, dict[str, Any]] = {}
    quality_notes: list[dict[str, Any]] = []
    disposition_seen = False
    engine_completed = False
    engine_failed = False
    last_review_basis_change_sequence = 0
    release_sequence = 0
    for event in events:
        event_type = str(event.get("event_type") or "")
        payload = deepcopy(event.get("payload") or {})
        audit = {
            "sequence": event.get("sequence"), "event_type": event_type,

            "event_hash": event.get("event_hash"),
        }
        if event_type == "review_created":
            state.update({key: deepcopy(payload.get(key)) for key in (
                "review_preparation_id", "preparation_basis_hash",
                "preparation_content_hash", "target", "target_spec",
                "bindings", "upstream_snapshot", "rule_pack", "standards", "mode", "execution",
                "project_context",
                "deployment_mode", "legacy_gate_snapshot",
                "engine_version", "recalculation_environment_version", "created_at",
            )})
            state["deployment_mode"] = str(payload.get("deployment_mode") or "enforced")
            state["legacy_gate_snapshot"] = deepcopy(payload.get("legacy_gate_snapshot") or {})
        elif event_type == "review_running":
            state["review_status"] = "running"
            state["started_at"] = payload.get("started_at")
        elif event_type == "review_completed":
            last_review_basis_change_sequence = max(last_review_basis_change_sequence, int(event.get("sequence") or 0))
            engine_completed = True
            state["review_status"] = "findings_ready"
            state["completed_at"] = payload.get("completed_at")
            state["incomplete_reasons"] = list(payload.get("incomplete_reasons") or [])
            state["coverage"] = deepcopy(payload.get("coverage") or {})
            completed_findings = list(payload.get("findings") or [])
            completed_verdict = str(payload.get("overall_verdict") or "incomplete")
            state["automated_gate_verdict"] = (
                "pass"
                if completed_verdict in {"pass", "conditional_pass"}
                and not state["incomplete_reasons"]
                and not any(finding_blocks(row) for row in completed_findings)
                else "fail"
            )
            for raw in completed_findings:
                row = deepcopy(raw)
                row["history"] = [{**audit, "status": row.get("status", "open")}]
                findings[str(row.get("finding_id") or "")] = row
        elif event_type == "review_failed":
            last_review_basis_change_sequence = max(last_review_basis_change_sequence, int(event.get("sequence") or 0))
            engine_failed = True
            engine_completed = True
            state["review_status"] = "findings_ready"
            state["completed_at"] = payload.get("completed_at")
            reason = str(payload.get("incomplete_reason") or "review_engine_failed")
            state["incomplete_reasons"] = sorted(set([*state["incomplete_reasons"], reason]))
            state["automated_gate_verdict"] = "fail"
        elif event_type == "finding_disposition_recorded":
            last_review_basis_change_sequence = max(last_review_basis_change_sequence, int(event.get("sequence") or 0))
            finding_id = str(payload.get("finding_id") or "")
            row = findings.get(finding_id)
            if row is not None:
                disposition_seen = True
                row["status"] = payload.get("new_status")
                for key in (
                    "disposition", "note", "closure_basis", "before_value", "after_value",
                    "remediation_evidence", "false_positive_reason", "waiver_scope",
                    "waiver_expires_at", "waiver_invalidation_conditions",
                ):
                    if key in payload:
                        row[key] = deepcopy(payload.get(key))
                row.setdefault("history", []).append({**audit, **payload})
        elif event_type == "finding_retested":
            operation_id = str(payload.get("operation_id") or "")
            if operation_id and operation_id not in completed_retest_operations:
                continue
            last_review_basis_change_sequence = max(last_review_basis_change_sequence, int(event.get("sequence") or 0))
            finding_id = str(payload.get("finding_id") or "")
            row = findings.get(finding_id)
            if row is not None:
                disposition_seen = True
                row["status"] = str(payload.get("new_status") or row.get("status") or "open")
                row["retest_result"] = deepcopy(payload)
                row.setdefault("history", []).append({**audit, **payload})
        elif event_type == "quality_note_recorded":
            quality_notes.append({**payload, **audit})
        elif event_type == "retest_linked":
            operation_id = str(payload.get("operation_id") or "")
            if operation_id and operation_id not in completed_retest_operations:
                continue
            last_review_basis_change_sequence = max(last_review_basis_change_sequence, int(event.get("sequence") or 0))
            state["retests"].append({**payload, **audit})
        elif event_type == "retest_completed":
            operation_id = str(payload.get("operation_id") or "")
            if operation_id in completed_retest_operations:
                last_review_basis_change_sequence = max(
                    last_review_basis_change_sequence,
                    int(event.get("sequence") or 0),
                )
        elif event_type == "retest_failed":
            last_review_basis_change_sequence = max(
                last_review_basis_change_sequence,
                int(event.get("sequence") or 0),
            )
        elif event_type == "review_exported":
            state["exports"].append({**payload, **audit})
        elif event_type == "review_released":
            if not state["released"]:
                state["released"] = True
                state["release"] = {**payload, **audit}
                release_sequence = int(event.get("sequence") or 0)
        elif event_type == "review_invalidated":
            state["invalidated"] = True
            state["invalidation"] = {**payload, **audit}

    post_release_events = [
        {
            "sequence": event.get("sequence"),
            "event_type": event.get("event_type"),
            "event_hash": event.get("event_hash"),
        }
        for event in events
        if release_sequence and int(event.get("sequence") or 0) > release_sequence
    ]
    unexpected_post_release_events = [
        event
        for event in post_release_events
        if event.get("event_type") != "review_invalidated"
    ]
    state["post_release_events"] = post_release_events
    state["release_is_terminal"] = bool(
        state["released"] and not unexpected_post_release_events
    )
    if unexpected_post_release_events:
        state["invalidated"] = True
        state["incomplete_reasons"] = sorted(set([
            *state["incomplete_reasons"],
            *[
                f"post_release_mutation:{event.get('event_type')}"
                for event in unexpected_post_release_events
            ],
        ]))
        state["post_release_mutation"] = unexpected_post_release_events

    state["quality_notes"] = quality_notes
    ordered_findings = sorted(
        findings.values(), key=lambda row: (
            SEVERITY_ORDER.get(str(row.get("severity") or ""), 9), str(row.get("finding_id") or ""),
        ),
    )
    state["findings"] = ordered_findings
    if not chain_ok:
        state["incomplete_reasons"] = sorted(set([
            *state["incomplete_reasons"], *[f"event_chain:{reason}" for reason in chain_reasons],
        ]))
        state["invalidated"] = True
    state["incomplete_reasons"] = sorted(set([
        *state["incomplete_reasons"],
        *[
            f"retest_operation_failed:{operation_id}:{payload.get('code') or 'unknown'}"
            for operation_id, payload in active_failed_retest_operations.items()
        ],
        *[
            f"retest_operation_invalid:{operation_id}:{reason}"
            for operation_id, reason in retest_operations["invalid"].items()
        ],
    ]))
    state["retest_in_progress"] = bool(state["pending_retest_operation_ids"])

    verdict = verdict_for(ordered_findings, state["incomplete_reasons"])
    active_blockers = [row for row in ordered_findings if finding_blocks(row)]
    state["active_blocking_finding_ids"] = [str(row.get("finding_id") or "") for row in active_blockers]
    state["pending_quality_rule_ids"] = sorted(
        str(row.get("rule_id") or "")
        for row in ordered_findings
        if row.get("manual_review_required") is True
        and row.get("status") != "resolved"
    )
    if state["invalidated"]:
        verdict = "incomplete"
        state["review_status"] = "invalidated"
    elif state["pending_retest_operation_ids"]:
        state["review_status"] = "retest_required"
    elif state["active_failed_retest_operations"] or state["invalid_retest_operations"]:
        state["review_status"] = "retest_required"
    elif engine_completed:
        if active_blockers:
            state["review_status"] = "remediation_in_progress" if disposition_seen else "findings_ready"
        elif state["incomplete_reasons"]:
            state["review_status"] = "findings_ready"
        else:
            state["review_status"] = "approved"
    elif engine_failed:
        state["review_status"] = "findings_ready"
    state["overall_verdict"] = verdict
    # Findings remain visible as quality results. They are never authorization
    # or identity gates for creating an immutable MCP review package.
    enforced_release_ready = bool(
        not state["invalidated"] and engine_completed
        and not state["pending_retest_operation_ids"]
    )
    state["would_release_under_enforced_mode"] = enforced_release_ready
    state["release_ready"] = bool(enforced_release_ready)
    if state.get("deployment_mode") == "shadow":
        state["shadow_comparison"] = _shadow_comparison(
            state,
            enforced_release_ready,
        )
        # Technical verification is diagnostic information, not an authorization gate.
        deterministic_blocking_ids = [
            str(row.get("finding_id") or "")
            for row in active_blockers
            if row.get("manual_review_required") is not True
        ]
        state["technical_verification_verdict"] = (
            "technical_pass" if not deterministic_blocking_ids and engine_completed
            else "technical_fail"
        )
        state["technical_verification_blockers"] = deterministic_blocking_ids
    state["generated"] = True
    state["validated"] = bool(engine_completed and not engine_failed)
    state["reviewed"] = bool(engine_completed)
    state["approved"] = state["review_status"] == "approved" or bool(state["released"])
    state["formally_deliverable"] = bool(state["released"] and not state["invalidated"])
    state["event_chain_hash"] = STORE.event_chain_hash(workspace_id, review_id)
    state["event_chain_valid"] = chain_ok
    state["event_count"] = len(events)
    state["finding_counts"] = {
        severity: sum(1 for row in ordered_findings if row.get("severity") == severity)
        for severity in ("P0", "P1", "P2", "P3")
    }
    state["active_finding_counts"] = {
        severity: sum(
            1 for row in ordered_findings
            if row.get("severity") == severity
            and row.get("status") not in {"resolved", "rejected", "superseded", "waived"}
        ) for severity in ("P0", "P1", "P2", "P3")
    }
    blockers: list[str] = []
    if state["invalidated"]:
        blockers.append("review_invalidated")
    blockers.extend(str(item) for item in state["incomplete_reasons"])
    blockers.extend(f"blocking_finding:{item}" for item in state["active_blocking_finding_ids"])
    blockers.extend(
        f"retest_in_progress:{operation_id}"
        for operation_id in state["pending_retest_operation_ids"]
    )
    if state.get("deployment_mode") == "shadow":
        blockers.append("shadow_mode_release_forbidden")
    if str((state.get("project_context") or {}).get("evidence_track") or "real") != "real":
        blockers.append("non_real_evidence_track_release_forbidden")
    state["blockers"] = sorted(set(blockers))
    state["gate_status"] = {
        "review": state["review_status"], "verdict": state["overall_verdict"],
        "deployment_mode": state.get("deployment_mode"),
        "would_release_under_enforced_mode": state["would_release_under_enforced_mode"],
        "release_ready": state["release_ready"],
        "formally_deliverable": state["formally_deliverable"],
    }
    return state


def _freshness_reasons(workspace_id: str, state: dict[str, Any]) -> list[str]:
    reasons: list[str] = []
    target_spec = state.get("target_spec") or state.get("target") or {}
    try:
        normalized = normalize_target(target_spec)
        resolved, blockers = _resolve_target(
            workspace_id,
            normalized,
        )
    except (ValueError, OSError):
        resolved, blockers = None, ["target_reresolution_failed"]
    if blockers or resolved is None:
        reasons.extend(str(item) for item in blockers or ["target_unavailable"])
    else:
        if str(resolved.get("target_sha256") or "") != str((state.get("target") or {}).get("target_sha256") or ""):
            reasons.append("target_content_changed")
        current_upstream = _binding_snapshot(
            workspace_id,
            resolved.get("bindings") or {},
        )
        recorded_upstream = state.get("upstream_snapshot") or {}
        if recorded_upstream and sha256_json(current_upstream) != sha256_json(recorded_upstream):
            reasons.append("upstream_binding_changed")
    components = [
        str(row.get("rule_pack_id") or "")
        for row in ((state.get("rule_pack") or {}).get("components") or [])
        if str(row.get("rule_pack_id") or "")
    ]
    component_types = [
        str(row.get("target_type") or "")
        for row in ((target_spec.get("components") or []) if isinstance(target_spec, dict) else [])
        if isinstance(row, dict) and str(row.get("target_type") or "")
    ]
    try:
        current_pack = rules.compose(
            str((state.get("target") or {}).get("target_type") or ""),
            components,
            component_types=component_types,
            project_context=state.get("project_context") or {},
        )
        if current_pack.get("content_hash") != (state.get("rule_pack") or {}).get("content_hash"):
            reasons.append("rule_pack_changed")
        current_standards = rules.standards_snapshot(REPO_ROOT, current_pack.get("standard_package_ids") or [])
        if current_standards.get("content_hash") != (state.get("standards") or {}).get("content_hash"):
            reasons.append("standards_changed")
    except ValueError:
        reasons.append("rule_pack_unavailable")
    return sorted(set(reasons))


def _project(workspace_id: str, review_id: str, *, check_freshness: bool = True) -> dict[str, Any]:
    state = _project_events(workspace_id, require_safe_id(review_id, "review_id"))
    if check_freshness and not state.get("invalidated"):
        reasons = _freshness_reasons(workspace_id, state)
        if reasons:
            STORE.append(
                workspace_id, review_id, "review_invalidated",
                {"reasons": reasons, "invalidated_at": utc_now()}, "system:freshness-check",
            )
            state = _project_events(workspace_id, review_id)
    return state


def get_review(args: dict[str, Any] | str, review_id: str = "") -> dict[str, Any]:
    if isinstance(args, str):
        workspace_id = args
    else:
        workspace_id = str(args.get("workspace_id") or "")
        review_id = str(args.get("review_id") or "")
    try:
        workspace_id = require_safe_id(workspace_id, "workspace_id")
        state = _project(workspace_id, review_id)
    except ValueError as exc:
        code = "review_not_found" if str(exc) in {"review_not_found", "invalid review_id"} else str(exc)
        return _blocked(code, _message(code))
    _resume_async_review_if_needed(workspace_id, state)
    return _ok(
        status=_review_envelope_status(state),
        review=state, review_id=state["review_id"], review_status=state["review_status"],
        overall_verdict=state["overall_verdict"], release_ready=state["release_ready"],
        deployment_mode=state.get("deployment_mode"),
        shadow_comparison=state.get("shadow_comparison") or {},
        would_release_under_enforced_mode=bool(state.get("would_release_under_enforced_mode")),
        formally_deliverable=state["formally_deliverable"], finding_counts=state["finding_counts"],
        active_finding_counts=state["active_finding_counts"], coverage=state.get("coverage") or {},
        quality_notes=state.get("quality_notes") or [],
        blockers=state.get("blockers") or [], resource_uris=[_review_uri(workspace_id, review_id)],
        next_actions=_next_actions(state),
    )


def _review_envelope_status(state: dict[str, Any]) -> str:
    """Map the projected review verdict to the public business envelope."""

    if state.get("incomplete_reasons") or state.get("overall_verdict") == "incomplete":
        return "incomplete"
    if (
        state.get("invalidated")
        or state.get("overall_verdict") == "fail"
        or state.get("active_blocking_finding_ids")
        or state.get("active_blocking_finding_ids")
    ):
        return "blocked"
    return "ok"


def _next_actions(state: dict[str, Any]) -> list[str]:
    if state.get("invalidated"):
        return ["目标或审查依据已变化；调用 review_retest 创建新版本审查"]
    if state.get("pending_retest_operation_ids"):
        return ["使用原 idempotency_key 重试 review_retest 以恢复未完成的复测"]
    if state.get("active_failed_retest_operations") or state.get("invalid_retest_operations"):
        return ["排除复测失败原因后，使用新 idempotency_key 重新发起复测"]
    if state.get("incomplete_reasons"):
        return ["补齐无法完成的核查条件后，以新目标版本调用 review_retest"]
    if state.get("active_blocking_finding_ids"):
        return ["处置阻断 findings，完成整改后调用 review_retest"]
    if state.get("deployment_mode") == "shadow":
        return ["影子审查仅记录新旧门禁差异；继续采集指标，不得正式释放"]
    if state.get("release_ready") and not state.get("released"):
        return ["调用 review_release 固化正式审查包"]
    return []


def _metric_rate(numerator: int, denominator: int) -> float | None:
    return round(numerator / denominator, 6) if denominator else None


def _metric_percentile(values: list[float], fraction: float) -> float | None:
    if not values:
        return None
    ordered = sorted(values)
    index = min(len(ordered) - 1, max(0, int((len(ordered) - 1) * fraction)))
    return round(ordered[index], 6)


def _workspace_metrics_payload(
    workspace_id: str,
    *,
    deployment_mode: str = "",
    started_after: str = "",
    started_before: str = "",
) -> dict[str, Any]:
    if deployment_mode and deployment_mode not in DEPLOYMENT_MODES:
        raise ValueError("review_deployment_mode_invalid")
    after = _parse_timestamp(started_after)
    before = _parse_timestamp(started_before)
    if started_after and after is None:
        raise ValueError("metrics_started_after_invalid")
    if started_before and before is None:
        raise ValueError("metrics_started_before_invalid")
    if after and before and after > before:
        raise ValueError("metrics_time_range_invalid")

    states: list[dict[str, Any]] = []
    raw_events: dict[str, list[dict[str, Any]]] = {}
    for review_id in STORE.review_ids(workspace_id):
        try:
            state = _project_events(workspace_id, review_id)
        except ValueError:
            continue
        created = _parse_timestamp(state.get("created_at"))
        if deployment_mode and state.get("deployment_mode") != deployment_mode:
            continue
        if after and (created is None or created < after):
            continue
        if before and (created is None or created > before):
            continue
        states.append(state)
        raw_events[review_id] = STORE.events(workspace_id, review_id)

    completed = [state for state in states if state.get("reviewed")]
    shadow = [state for state in states if state.get("deployment_mode") == "shadow"]
    completed_shadow = [state for state in shadow if state.get("reviewed")]
    applicable_rule_count = 0
    executed_rule_count = 0
    high_risk_rule_count = 0
    omitted_high_risk_rule_count = 0
    uncheckable_review_count = 0
    uncheckable_reason_count = 0
    finding_count = 0
    false_positive_appeal_count = 0
    retested_finding_count = 0
    passed_retested_finding_count = 0
    durations: list[float] = []
    known_comparison_count = 0
    disagreement_count = 0

    for state in completed:
        coverage = state.get("coverage") or {}
        applicable = set(coverage.get("applicable_rules") or [])
        executed = set(coverage.get("executed_rules") or [])
        applicable_rule_count += len(applicable)
        executed_rule_count += len(applicable & executed)
        source_severity = {
            str(row.get("rule_id") or ""): str(row.get("severity") or "").upper()
            for row in ((state.get("rule_pack") or {}).get("rule_sources") or [])
            if isinstance(row, dict)
        }
        high_risk = {
            rule_id for rule_id in applicable
            if source_severity.get(str(rule_id)) in {"P0", "P1"}
        }
        high_risk_rule_count += len(high_risk)
        omitted_high_risk_rule_count += len(high_risk - executed)
        reasons = list(state.get("incomplete_reasons") or [])
        if reasons:
            uncheckable_review_count += 1
            uncheckable_reason_count += len(reasons)

        findings = list(state.get("findings") or [])
        finding_count += len(findings)
        for finding in findings:
            appealed = str(finding.get("status") or "") == "false_positive_appeal"
            if not appealed:
                appealed = any(
                    str(history.get("new_status") or "") == "false_positive_appeal"
                    or str(history.get("disposition") or "") in {
                        "reject", "rejected", "false_positive", "false_positive_appeal",
                    }
                    for history in finding.get("history") or []
                    if isinstance(history, dict)
                )
            false_positive_appeal_count += int(appealed)

        events = raw_events.get(str(state.get("review_id") or ""), [])
        running_times = [
            _parse_timestamp((event.get("payload") or {}).get("started_at") or event.get("created_at"))
            for event in events if event.get("event_type") == "review_running"
        ]
        completed_times = [
            _parse_timestamp((event.get("payload") or {}).get("completed_at") or event.get("created_at"))
            for event in events
            if event.get("event_type") in {"review_completed", "review_failed"}
        ]
        starts = [value for value in running_times if value is not None]
        finishes = [value for value in completed_times if value is not None]
        if starts and finishes:
            seconds = (min(finishes) - min(starts)).total_seconds()
            if seconds >= 0:
                durations.append(seconds)
        for event in events:
            if event.get("event_type") != "finding_retested":
                continue
            payload = event.get("payload") or {}
            retested_finding_count += 1
            passed_retested_finding_count += int(payload.get("retest_passed") is True)

    for state in completed_shadow:
        difference = str((state.get("shadow_comparison") or {}).get("validation_difference") or "")
        if difference == "unavailable" or not difference:
            continue
        known_comparison_count += 1
        disagreement_count += int(difference in {
            "legacy_pass_unified_block", "legacy_block_unified_pass",
        })

    shadow_starts = [
        value for value in (_parse_timestamp(state.get("created_at")) for state in shadow)
        if value is not None
    ]
    generated_at = utc_now()
    generated_at_dt = _parse_timestamp(generated_at) or datetime.now(timezone.utc)
    first_shadow_at = min(shadow_starts) if shadow_starts else None
    shadow_elapsed_days = (
        round(max(0.0, (generated_at_dt - first_shadow_at).total_seconds()) / 86400, 6)
        if first_shadow_at else 0.0
    )
    shadow_days = sorted({value.date().isoformat() for value in shadow_starts})
    duration_requirement_met = bool(
        first_shadow_at and shadow_elapsed_days >= 14 and completed_shadow
    )
    high_risk_rate = _metric_rate(
        omitted_high_risk_rule_count,
        high_risk_rule_count,
    )
    warnings: list[str] = []
    if high_risk_rate is None:
        warnings.append("筛选范围内没有声明严重度的适用 P0/P1 规则，遗漏率不可计算")
    if shadow and not duration_requirement_met:
        warnings.append("影子观察期尚未达到连续 14 天，不得据此启用强制门禁")

    indicators = {
        "rule_coverage": {
            "rate": _metric_rate(executed_rule_count, applicable_rule_count),
            "executed_rule_count": executed_rule_count,
            "applicable_rule_count": applicable_rule_count,
            "definition": "已执行适用规则数 / 适用规则数（按审查运行加权）",
        },
        "uncheckable": {
            "rate": _metric_rate(uncheckable_review_count, len(completed)),
            "review_count": uncheckable_review_count,
            "completed_review_count": len(completed),
            "reason_count": uncheckable_reason_count,
            "definition": "存在 incomplete_reasons 的已完成审查数 / 已完成审查数",
        },
        "p0_p1_omission": {
            "rate": high_risk_rate,
            "omitted_rule_count": omitted_high_risk_rule_count,
            "applicable_declared_rule_count": high_risk_rule_count,
            "definition": "未执行的适用 P0/P1 声明规则数 / 适用 P0/P1 声明规则数；不冒充人工金标假阴性率",
        },
        "false_positive_appeal": {
            "rate": _metric_rate(false_positive_appeal_count, finding_count),
            "appealed_finding_count": false_positive_appeal_count,
            "finding_count": finding_count,
            "definition": "曾进入误报申诉流程的 finding 数 / finding 总数",
        },
        "remediation_retest_pass": {
            "rate": _metric_rate(passed_retested_finding_count, retested_finding_count),
            "passed_finding_count": passed_retested_finding_count,
            "retested_finding_count": retested_finding_count,
            "definition": "复测未再复现且规则确已执行的 finding 数 / 已复测 finding 数",
        },
        "review_duration_seconds": {
            "sample_count": len(durations),
            "mean": round(sum(durations) / len(durations), 6) if durations else None,
            "p50": _metric_percentile(durations, 0.50),
            "p95": _metric_percentile(durations, 0.95),
            "max": round(max(durations), 6) if durations else None,
        },
        "shadow_gate_disagreement": {
            "rate": _metric_rate(disagreement_count, known_comparison_count),
            "disagreement_count": disagreement_count,
            "comparable_review_count": known_comparison_count,
            "definition": "旧工程校验与统一自动审查结论不一致的影子审查数 / 可比较影子审查数",
        },
    }
    return {
        "schema_version": "deliverable_review_metrics.v1",
        "workspace_id": workspace_id,
        "generated_at": generated_at,
        "filters": {
            "deployment_mode": deployment_mode or None,
            "started_after": started_after or None,
            "started_before": started_before or None,
        },
        "review_count": len(states),
        "completed_review_count": len(completed),
        "shadow_review_count": len(shadow),
        "completed_shadow_review_count": len(completed_shadow),
        "indicators": indicators,
        "shadow_period": {
            "first_review_at": first_shadow_at.isoformat() if first_shadow_at else None,
            "observed_at": generated_at,
            "elapsed_days": shadow_elapsed_days,
            "distinct_review_days": shadow_days,
            "minimum_required_days": 14,
            "duration_requirement_met": duration_requirement_met,
            "auto_enforcement_allowed": False,
            "governance_decision_eligible": bool(
                duration_requirement_met and completed_shadow and known_comparison_count
            ),
            "recommendation": (
                "eligible_for_governance_decision"
                if duration_requirement_met and completed_shadow and known_comparison_count
                else "continue_shadow"
            ),
        },
        "warnings": warnings,
    }


def workspace_metrics(args: dict[str, Any] | str) -> dict[str, Any]:
    if isinstance(args, str):
        args = {"workspace_id": args}
    workspace_id = str(args.get("workspace_id") or "")
    try:
        workspace_id = require_safe_id(workspace_id, "workspace_id")
        metrics = _workspace_metrics_payload(
            workspace_id,
            deployment_mode=str(args.get("deployment_mode") or ""),
            started_after=str(args.get("started_after") or ""),
            started_before=str(args.get("started_before") or ""),
        )
    except ValueError as exc:
        return _blocked(str(exc), _message(str(exc)))
    return _ok(
        metrics=metrics,
        resource_uris=[_metrics_uri(workspace_id)],
        warnings=metrics.get("warnings") or [],
        blockers=[],
        next_actions=[],
    )


def list_findings(args: dict[str, Any]) -> dict[str, Any]:
    workspace_id = str(args.get("workspace_id") or "")
    review_id = str(args.get("review_id") or "")
    try:
        workspace_id = require_safe_id(workspace_id, "workspace_id")
        state = _project(workspace_id, review_id)
        severity = str(args.get("severity") or "")
        status = str(args.get("status") or "")
        category = str(args.get("category") or "")
        role = str(args.get("review_area") or "")
        location_query = str(args.get("location") or "").lower()
        if severity and severity not in SEVERITIES:
            raise ValueError("severity_invalid")
        if status and status not in FINDING_STATUSES:
            raise ValueError("finding_status_invalid")
        rows = []
        for finding in state.get("findings") or []:
            if severity and finding.get("severity") != severity:
                continue
            if status and finding.get("status") != status:
                continue
            if category and finding.get("category") != category:
                continue
            if role and finding.get("review_area") != role:
                continue
            if location_query and location_query not in canonical_json(finding.get("target_location") or {}).lower():
                continue
            rows.append({**finding, "uri": _finding_uri(workspace_id, review_id, str(finding["finding_id"]))})
        page = paginate_resource_entries(
            rows, cursor=str(args.get("cursor") or ""), limit=int(args.get("limit") or 50),
        )
    except ValueError as exc:
        code = str(exc)
        if code == "invalid review_id":
            code = "review_not_found"
        return _blocked(code, _message(code))
    findings = [{key: value for key, value in row.items() if key != "uri"} for row in page["resources"]]
    return _ok(
        review_id=review_id, findings=findings, total_matching=len(rows),
        next_cursor=page["next_cursor"], has_more=page["has_more"], snapshot_hash=page["snapshot_hash"],
        resource_uris=[row["uri"] for row in page["resources"]], blockers=[], next_actions=[],
    )


def get_finding(args: dict[str, Any]) -> dict[str, Any]:
    workspace_id = str(args.get("workspace_id") or "")
    review_id = str(args.get("review_id") or "")
    finding_id = str(args.get("finding_id") or "")
    try:
        workspace_id = require_safe_id(workspace_id, "workspace_id")
        state = _project(workspace_id, review_id)
    except ValueError:
        return _blocked("review_not_found", _message("review_not_found"))
    row = next((item for item in state.get("findings") or [] if item.get("finding_id") == finding_id), None)
    if row is None:
        return _blocked("finding_not_found", _message("finding_not_found"))
    return _ok(
        review_id=review_id, finding=deepcopy(row),
        resource_uris=[_finding_uri(workspace_id, review_id, finding_id)], blockers=[], next_actions=[],
    )


def _require_open_review(
    workspace_id: str,
    review_id: str,
) -> tuple[dict[str, Any] | None, dict[str, Any] | None]:
    try:
        state = _project(workspace_id, review_id)
    except ValueError:
        return None, _blocked("review_not_found", _message("review_not_found"))
    if state.get("invalidated"):
        return None, _blocked("review_invalidated", "目标或审查依据已变化，旧审查不可继续处置", review_id=review_id)
    if state.get("released"):
        return None, _blocked("review_already_released", "正式审查包已固化，不可修改其审查历史", review_id=review_id)
    if not state.get("reviewed"):
        return None, _blocked("review_not_ready", "审查引擎尚未形成 findings", review_id=review_id)
    if state.get("review_status") in {"approved", "rejected", "waived"}:
        return None, _blocked("review_terminal", "审查已进入终态；后续变化必须创建新版本复测", review_id=review_id)
    return state, None


def _evidence_is_precise(rows: Any) -> bool:
    if not isinstance(rows, list) or not rows:
        return False
    for row in rows:
        if not isinstance(row, dict):
            return False
        source_id = str(row.get("file_id") or row.get("source_id") or row.get("url") or "")
        locator = str(
            row.get("locator") or row.get("page") or row.get("paragraph")
            or row.get("cell") or row.get("range") or ""
        )
        content_hash = str(row.get("content_hash") or row.get("sha256") or "")
        if not source_id or not locator or not content_hash:
            return False
    return True


def _target_version_scope(value: dict[str, Any]) -> dict[str, Any]:
    target = value.get("target") if isinstance(value.get("target"), dict) else value
    target_spec = (
        value.get("target_spec")
        if isinstance(value.get("target_spec"), dict)
        else target
    )
    target_type = str(target.get("target_type") or "")
    scope: dict[str, Any] = {"target_type": target_type}
    if target_type == "report_artifact":
        artifact_domain = str(target_spec.get("artifact_domain") or "")
        scope["artifact_domain"] = artifact_domain
    elif target_type == "finance_xlsx":
        # External workbooks have no repository-managed supersession chain;
        # target_id is therefore their caller-owned stable logical identity.
        scope["logical_target_id"] = str(target_spec.get("target_id") or "")
    elif target_type == "finance_xlsx_source":
        # A retest uploads a corrected workbook, so source_file_id necessarily
        # changes; lineage keys on the caller-owned stable logical identity,
        # exactly as finance_xlsx keys on target_id rather than the file path.
        scope["logical_target_id"] = str(target_spec.get("target_id") or "")
    elif target_type == "combined_deliverable":
        component_scopes = [
            _target_version_scope(component)
            for component in (target_spec.get("components") or [])
            if isinstance(component, dict)
        ]
        scope["components"] = sorted(
            component_scopes,
            key=sha256_json,
        )
    return scope


def _retest_target_scope_matches(
    parent: dict[str, Any],
    resolved: dict[str, Any],
) -> bool:
    return _target_version_scope(parent) == _target_version_scope(resolved)


def _successful_retest_closes_finding(
    state: dict[str, Any],
    finding_id: str,
    retest_review_id: str,
) -> bool:
    review_id = str(state.get("review_id") or "")
    return any(
        row.get("completed") is True
        and str(row.get("parent_review_id") or "") == review_id
        and str(row.get("child_review_id") or "") == retest_review_id
        and finding_id in {
            str(item) for item in row.get("closed_finding_ids") or []
        }
        and finding_id not in {
            str(item) for item in row.get("remaining_finding_ids") or []
        }
        for row in state.get("retests") or []
    )


def disposition_finding(args: dict[str, Any]) -> dict[str, Any]:
    def execute(workspace_id: str) -> dict[str, Any]:
        review_id = str(args.get("review_id") or "")
        finding_id = str(args.get("finding_id") or "")
        state, blocked = _require_open_review(
            workspace_id, review_id,
        )
        if blocked is not None or state is None:
            return blocked or _blocked("review_not_found", _message("review_not_found"))
        finding = next((row for row in state["findings"] if row.get("finding_id") == finding_id), None)
        if finding is None:
            return _blocked("finding_not_found", _message("finding_not_found"))
        disposition = str(args.get("disposition") or "").strip()
        aliases = {
            "confirm": "confirmed", "confirmed": "confirmed",
            "remediate": "remediation_in_progress", "remediation_in_progress": "remediation_in_progress",
            "reject": "false_positive_appeal", "rejected": "false_positive_appeal",
            "false_positive": "false_positive_appeal", "false_positive_appeal": "false_positive_appeal",
            "appeal_waiver": "waiver_requested", "compliance_waiver": "waiver_requested",
            "waiver_requested": "waiver_requested", "resolve": "resolved", "resolved": "resolved",
        }
        new_status = aliases.get(disposition)
        if new_status is None:
            return _blocked("disposition_invalid", "disposition 必须为确认、驳回申诉、整改中、误报申诉、合规豁免申请或整改关闭")
        note = str(args.get("note") or "").strip()
        if not note:
            return _blocked("disposition_note_required", "finding 处置必须提供依据说明")
        evidence = args.get("remediation_evidence") or args.get("evidence") or []
        # MCP 本地工具:不做角色授权校验。
        if new_status == "false_positive_appeal":
            reason = str(args.get("false_positive_reason") or "").strip()
            if not reason or not _evidence_is_precise(evidence):
                return _blocked("false_positive_evidence_required", "误报申诉必须提供理由及带哈希和精确定位的证据")
        if new_status == "waiver_requested":
            if finding.get("severity") == "P0" or finding.get("waiver_allowed") is False:
                return _blocked("p0_waiver_forbidden", "P0 finding 不可豁免")
            if finding.get("severity") != "P1":
                return _blocked("waiver_not_applicable", "合规豁免仅用于规则允许的 P1 finding")
            expiry = _parse_timestamp(args.get("waiver_expires_at"))
            if expiry is None or expiry <= datetime.now(timezone.utc):
                return _blocked("waiver_expiry_required", "P1 豁免必须设置未来有效期")
            if not str(args.get("waiver_scope") or "").strip() or not list(args.get("waiver_invalidation_conditions") or []):
                return _blocked("waiver_scope_required", "P1 豁免必须限定范围并声明失效条件")
            if not _evidence_is_precise(evidence):
                return _blocked("waiver_evidence_required", "P1 豁免申请必须绑定精确证据")
        if new_status == "resolved":
            if not str(args.get("closure_basis") or "").strip():
                return _blocked("closure_basis_required", "关闭 finding 必须提供关闭依据")
            if "before_value" not in args or "after_value" not in args:
                return _blocked("before_after_values_required", "关闭 finding 必须记录整改前值与整改后值")
            if not _evidence_is_precise(evidence):
                return _blocked("closure_evidence_required", "关闭 finding 必须绑定带哈希和精确定位的批准证据")
            if finding_blocks(finding):
                retest_review_id = str(args.get("retest_review_id") or "")
                if not _successful_retest_closes_finding(
                    state, finding_id, retest_review_id,
                ):
                    return _blocked("successful_retest_required", "阻断 finding 只能由更新目标版本的成功复测关闭")
        payload = {
            "finding_id": finding_id, "disposition": disposition, "new_status": new_status,
            "note": note, "remediation_evidence": deepcopy(evidence),
        }
        for key in (
            "closure_basis", "before_value", "after_value", "false_positive_reason",
            "waiver_scope", "waiver_expires_at", "waiver_invalidation_conditions", "retest_review_id",
        ):
            if key in args:
                payload[key] = deepcopy(args.get(key))
        STORE.append(workspace_id, review_id, "finding_disposition_recorded", payload)
        current = _project(workspace_id, review_id, check_freshness=False)
        updated = next(row for row in current["findings"] if row.get("finding_id") == finding_id)
        return _ok(
            review_id=review_id, finding_id=finding_id, finding_status=updated.get("status"),
            review_status=current["review_status"], overall_verdict=current["overall_verdict"],
            release_ready=current["release_ready"], blockers=current["blockers"],
            resource_uris=[_finding_uri(workspace_id, review_id, finding_id)],
            next_actions=_next_actions(current),
        )
    return _write("review_disposition_finding", args, execute)


def attest(args: dict[str, Any]) -> dict[str, Any]:
    def execute(workspace_id: str) -> dict[str, Any]:
        review_id = str(args.get("review_id") or "")
        state, blocked = _require_open_review(
            workspace_id, review_id,
        )
        if blocked is not None or state is None:
            return blocked or _blocked("review_not_found", _message("review_not_found"))
        verdict = str(args.get("verdict") or "")
        if verdict not in {"approve", "reject"}:
            return _blocked("quality_verdict_invalid", "质量结论必须为 approve 或 reject")
        note = str(args.get("note") or "").strip()
        if not note:
            return _blocked("quality_note_required", "质量结论必须提供依据说明")
        quality_note_id = "qnote_" + sha256_json({
            "review_id": review_id,
            "verdict": verdict,
            "note": note,
            "basis": state.get("event_chain_hash"),
        }).removeprefix("sha256:")[:24]
        STORE.append(workspace_id, review_id, "quality_note_recorded", {
            "quality_note_id": quality_note_id,
            "verdict": verdict,
            "note": note,
            "basis_event_chain_hash": state.get("event_chain_hash"),
            "recorded_at": utc_now(),
        })
        current = _project(workspace_id, review_id, check_freshness=False)
        return _ok(
            review_id=review_id, quality_note_id=quality_note_id, verdict=verdict,
            review_status=current["review_status"],
            overall_verdict=current["overall_verdict"], release_ready=current["release_ready"],
            blockers=current["blockers"], resource_uris=[_review_uri(workspace_id, review_id)],
            next_actions=_next_actions(current),
        )
    return _write("review_attest", args, execute)


def _retest_operation_identity(
    args: dict[str, Any],
) -> tuple[str, str, str]:
    key_identity = {
        "operation": "review_retest",
        "idempotency_key": str(args.get("idempotency_key") or ""),
    }
    key_hash = sha256_json(key_identity)
    operation_id = "retestop_" + key_hash.removeprefix("sha256:")[:32]
    return operation_id, key_hash, sha256_json(args)


def _find_retest_intent(
    workspace_id: str,
    operation_id: str,
) -> tuple[str, dict[str, Any]] | None:
    matches: list[tuple[str, dict[str, Any]]] = []
    for review_id in STORE.review_ids(workspace_id):
        for event in STORE.events(workspace_id, review_id):
            payload = event.get("payload") or {}
            if (
                event.get("event_type") == "retest_started"
                and payload.get("operation_id") == operation_id
            ):
                matches.append((review_id, event))
    if len(matches) > 1:
        raise ValueError("retest_operation_conflict")
    return matches[0] if matches else None


def _append_retest_event_once(
    workspace_id: str,
    review_id: str,
    event_type: str,
    payload: dict[str, Any],

    *,
    identity_fields: tuple[str, ...] = (),
) -> dict[str, Any]:
    operation_id = str(payload.get("operation_id") or "")
    matches = [
        event
        for event in STORE.events(workspace_id, review_id)
        if event.get("event_type") == event_type
        and str((event.get("payload") or {}).get("operation_id") or "") == operation_id
        and all(
            (event.get("payload") or {}).get(field) == payload.get(field)
            for field in identity_fields
        )
    ]
    if matches:
        if len(matches) != 1 or sha256_json(matches[0].get("payload") or {}) != sha256_json(payload):
            raise ValueError("retest_operation_conflict")
        return matches[0]
    return STORE.append(workspace_id, review_id, event_type, payload)


def _retest_failure(
    workspace_id: str,
    parent_review_id: str,
    intent: dict[str, Any],

    code: str,
    message: str,
) -> dict[str, Any]:
    payload = {
        "operation_id": intent["operation_id"],
        "code": code,
        "message": message,
        "failed_at_operation_started_at": intent["operation_started_at"],
    }
    _append_retest_event_once(
        workspace_id,
        parent_review_id,
        "retest_failed",
        payload,

    )
    return _blocked(code, message, review_id=parent_review_id)


def retest(args: dict[str, Any]) -> dict[str, Any]:
    def execute(workspace_id: str) -> dict[str, Any]:
        requested_parent_review_id = str(args.get("review_id") or "")
        operation_id, key_hash, request_hash = _retest_operation_identity(
            args,
        )
        located_intent = _find_retest_intent(workspace_id, operation_id)
        if located_intent is not None:
            intent_parent_review_id, intent_event = located_intent
            intent = deepcopy(intent_event.get("payload") or {})
            if (
                intent_parent_review_id != requested_parent_review_id
                or intent.get("key_hash") != key_hash
                or intent.get("request_hash") != request_hash
            ):
                # Do not let a conflicting retry become the cached owner of
                # the original operation's idempotency key.
                raise ValueError("idempotency_key_conflict")
            parent_review_id = intent_parent_review_id
        else:
            intent = {}
            parent_review_id = requested_parent_review_id

        try:
            parent = _project(workspace_id, parent_review_id)
        except ValueError:
            return _blocked("review_not_found", _message("review_not_found"))
        if intent:
            failed = next(
                (
                    event.get("payload") or {}
                    for event in STORE.events(workspace_id, parent_review_id)
                    if event.get("event_type") == "retest_failed"
                    and (event.get("payload") or {}).get("operation_id") == operation_id
                ),
                None,
            )
            if failed is not None:
                return _blocked(
                    str(failed.get("code") or "retest_operation_failed"),
                    str(failed.get("message") or "复测操作未完成"),
                    review_id=parent_review_id,
                )
        else:
            if not parent.get("reviewed"):
                return _blocked("review_not_ready", "原审查尚未完成，不能复测")
            if parent.get("released"):
                return _blocked(
                    "review_already_released",
                    "正式审查包已固化；复测必须从未发布审查创建新版本链，不能修改原正式包",
                    review_id=parent_review_id,
                )
            if parent.get("review_status") in {"approved", "rejected", "waived"}:
                return _blocked(
                    "review_terminal",
                    "审查已进入终态；请重新准备目标并创建独立的新版本审查",
                    review_id=parent_review_id,
                )
            if parent.get("pending_retest_operation_ids"):
                return _blocked(
                    "retest_already_in_progress",
                    "已有复测操作未完成；必须使用原 idempotency_key 恢复，不能并行创建新复测链",
                    review_id=parent_review_id,
                    pending_retest_operation_ids=parent["pending_retest_operation_ids"],
                )
            evidence = args.get("remediation_evidence") or []
            if not _evidence_is_precise(evidence):
                return _blocked("retest_evidence_required", "复测必须提供带内容哈希和精确定位的整改证据")
            target = args.get("target")
            if not isinstance(target, dict):
                return _blocked("retest_target_required", "复测必须显式提供新目标版本")
            try:
                normalized = normalize_target(target)
                resolved, target_blockers = _resolve_target(
                    workspace_id,
                    normalized,
                )
            except ValueError as exc:
                return _blocked(str(exc), _message(str(exc)))
            if target_blockers or resolved is None:
                return _blocked(target_blockers[0], "复测目标无法完整解析", blockers=target_blockers)
            if not _retest_target_scope_matches(parent, resolved):
                return _blocked(
                    "retest_target_scope_mismatch",
                    "复测目标必须保持原目标类型、逻辑身份、报告业务域及租户范围",
                )
            if resolved.get("target_sha256") == (parent.get("target") or {}).get("target_sha256"):
                return _blocked("retest_target_not_newer", "复测对象内容必须新于原目标，不能复用相同哈希")
            old_components = [
                str(row.get("rule_pack_id") or "")
                for row in ((parent.get("rule_pack") or {}).get("components") or [])
                if str(row.get("rule_pack_id") or "")
            ]
            requested_packs = list(args.get("rule_pack_ids") or old_components)
            mode = str(args.get("mode") or parent.get("mode") or "quick")
            if mode not in {"quick", "deep"}:
                return _blocked("review_mode_invalid", "mode 必须为 quick 或 deep")
            prepare_args = {
                "workspace_id": workspace_id,

                "idempotency_key": f"{operation_id}:prepare",
                "target": deepcopy(target),
                "rule_pack_ids": requested_packs,
                "industry_overlays": list(args.get("industry_overlays") or []),
            }
            parent_basis = {
                "target": deepcopy(parent.get("target") or {}),
                "target_spec": deepcopy(parent.get("target_spec") or {}),
                "rule_pack": deepcopy(parent.get("rule_pack") or {}),
                "mode": str(parent.get("mode") or "quick"),
                "deployment_mode": str(parent.get("deployment_mode") or "enforced"),
                "findings": deepcopy(parent.get("findings") or []),
                "event_chain_hash": str(parent.get("event_chain_hash") or ""),
            }
            intent = {
                "operation_id": operation_id,
                "key_hash": key_hash,
                "request_hash": request_hash,
                "parent_review_id": parent_review_id,
                "parent_basis": parent_basis,
                "resolved_target": {
                    key: deepcopy(resolved.get(key))
                    for key in ("target_type", "target_id", "target_sha256", "target_spec")
                },
                "prepare_args": prepare_args,
                "mode": mode,
                "expected_finding_ids": sorted(
                    str(row.get("finding_id") or "")
                    for row in parent_basis["findings"]
                    if str(row.get("finding_id") or "")
                ),
                "remediation_evidence": deepcopy(evidence),
                "operation_started_at": utc_now(),
            }
            _append_retest_event_once(
                workspace_id,
                parent_review_id,
                "retest_started",
                intent,

            )

        operation_events = [
            event
            for event in STORE.events(workspace_id, parent_review_id)
            if (event.get("payload") or {}).get("operation_id") == operation_id
        ]
        prepared_event = next(
            (event for event in operation_events if event.get("event_type") == "retest_prepared"),
            None,
        )
        if prepared_event is None:
            prepared = prepare(deepcopy(intent["prepare_args"]))
            if prepared.get("status") == "blocked":
                return _retest_failure(
                    workspace_id,
                    parent_review_id,
                    intent,

                    str(prepared.get("code") or "retest_preparation_failed"),
                    str(prepared.get("message") or "复测审查准备失败"),
                )
            intended_target_sha256 = str((intent.get("resolved_target") or {}).get("target_sha256") or "")
            if str((prepared.get("target") or {}).get("target_sha256") or "") != intended_target_sha256:
                return _retest_failure(
                    workspace_id,
                    parent_review_id,
                    intent,

                    "retest_target_changed_during_operation",
                    "复测目标在操作恢复期间发生变化，已按失败关闭",
                )
            prepared_payload = {
                "operation_id": operation_id,
                "review_preparation_id": str(prepared.get("review_preparation_id") or ""),
                "target_sha256": intended_target_sha256,
            }
            prepared_event = _append_retest_event_once(
                workspace_id,
                parent_review_id,
                "retest_prepared",
                prepared_payload,

            )
        prepared_payload = prepared_event.get("payload") or {}
        preparation_id = str(prepared_payload.get("review_preparation_id") or "")

        operation_events = [
            event
            for event in STORE.events(workspace_id, parent_review_id)
            if (event.get("payload") or {}).get("operation_id") == operation_id
        ]
        child_started_event = next(
            (event for event in operation_events if event.get("event_type") == "retest_child_started"),
            None,
        )
        if child_started_event is None:
            parent_basis = intent.get("parent_basis") or {}
            started = start({
                "workspace_id": workspace_id,

                "idempotency_key": f"{operation_id}:start",
                "review_preparation_id": preparation_id,
                "mode": str(intent.get("mode") or "quick"),
                "execution": "sync",
                "deployment_mode": str(
                    (intent.get("parent_basis") or {}).get("deployment_mode") or "enforced"
                ),
            })
            if started.get("status") == "blocked":
                return _retest_failure(
                    workspace_id,
                    parent_review_id,
                    intent,

                    str(started.get("code") or "retest_child_review_failed"),
                    str(started.get("message") or "复测子审查创建失败"),
                )
            child_review_id = str(started.get("review_id") or "")
            try:
                child = _project(workspace_id, child_review_id, check_freshness=False)
            except ValueError:
                return _retest_failure(
                    workspace_id,
                    parent_review_id,
                    intent,

                    "retest_child_review_unavailable",
                    _message("retest_child_review_unavailable"),
                )
            intended_target_sha256 = str((intent.get("resolved_target") or {}).get("target_sha256") or "")
            if str((child.get("target") or {}).get("target_sha256") or "") != intended_target_sha256:
                return _retest_failure(
                    workspace_id,
                    parent_review_id,
                    intent,

                    "retest_child_target_mismatch",
                    "复测子审查未绑定操作意图中的目标哈希",
                )
            child_started_event = _append_retest_event_once(
                workspace_id,
                parent_review_id,
                "retest_child_started",
                {
                    "operation_id": operation_id,
                    "review_preparation_id": preparation_id,
                    "child_review_id": child_review_id,
                    "child_target_sha256": intended_target_sha256,
                },

            )
        child_review_id = str((child_started_event.get("payload") or {}).get("child_review_id") or "")
        try:
            child = _project(workspace_id, child_review_id, check_freshness=False)
        except ValueError:
            return _retest_failure(
                workspace_id,
                parent_review_id,
                intent,

                "retest_child_review_unavailable",
                _message("retest_child_review_unavailable"),
            )
        parent_basis = intent.get("parent_basis") or {}
        old_findings = list(parent_basis.get("findings") or [])
        new_by_match = {_finding_match_key(row): row for row in child.get("findings") or []}
        new_by_rule: dict[tuple[str, str], list[dict[str, Any]]] = {}
        for finding in child.get("findings") or []:
            new_by_rule.setdefault(
                (
                    str(finding.get("rule_id") or ""),
                    str(finding.get("category") or ""),
                ),
                [],
            ).append(finding)
        executed_rules = set((child.get("coverage") or {}).get("executed_rules") or [])
        closed: list[str] = []
        remaining: list[str] = []
        for old_finding in old_findings:
            old_id = str(old_finding.get("finding_id") or "")
            new_finding = new_by_match.get(_finding_match_key(old_finding))
            if new_finding is None:
                conservative_matches = new_by_rule.get(
                    (
                        str(old_finding.get("rule_id") or ""),
                        str(old_finding.get("category") or ""),
                    ),
                    [],
                )
                new_finding = conservative_matches[0] if conservative_matches else None
            coverage_rule_id = _finding_coverage_rule_id(old_finding)
            can_conclude_absence = coverage_rule_id in executed_rules and not any(
                str(reason) == f"rule_not_executed:{coverage_rule_id}"
                for reason in child.get("incomplete_reasons") or []
            )
            retest_passed = new_finding is None and can_conclude_absence
            if retest_passed:
                closed.append(old_id)
                after_value = "finding_not_reproduced"
            else:
                remaining.append(old_id)
                after_value = deepcopy((new_finding or {}).get("actual"))
            _append_retest_event_once(
                workspace_id,
                parent_review_id,
                "finding_retested",
                {
                    "operation_id": operation_id,
                    "finding_id": old_id,
                    "new_status": "remediation_in_progress",
                    "retest_passed": retest_passed,
                    "parent_review_id": parent_review_id,
                    "retest_review_id": child_review_id,
                    "before_value": deepcopy(old_finding.get("actual")),
                    "after_value": after_value,
                    "remediation_evidence": deepcopy(intent.get("remediation_evidence") or []),
                    "same_rule_pack": (
                        (child.get("rule_pack") or {}).get("content_hash")
                        == (parent_basis.get("rule_pack") or {}).get("content_hash")
                    ),
                    "retested_at": intent.get("operation_started_at"),
                },

                identity_fields=("finding_id",),
            )
        closed = sorted(set(closed))
        remaining = sorted(set(remaining))
        link = {
            "operation_id": operation_id,
            "parent_review_id": parent_review_id,
            "child_review_id": child_review_id,
            "parent_target_sha256": (parent_basis.get("target") or {}).get("target_sha256"),
            "child_target_sha256": (child.get("target") or {}).get("target_sha256"),
            "parent_rule_pack_hash": (parent_basis.get("rule_pack") or {}).get("content_hash"),
            "child_rule_pack_hash": (child.get("rule_pack") or {}).get("content_hash"),
            "completed": True,
            "closed_finding_ids": closed,
            "remaining_finding_ids": remaining,
            "remediation_evidence": deepcopy(intent.get("remediation_evidence") or []),
            "retested_at": intent.get("operation_started_at"),
        }
        _append_retest_event_once(
            workspace_id,
            parent_review_id,
            "retest_linked",
            link,

        )
        _append_retest_event_once(
            workspace_id,
            child_review_id,
            "retest_linked",
            link,

        )
        completion = {
            "operation_id": operation_id,
            "parent_review_id": parent_review_id,
            "child_review_id": child_review_id,
            "expected_finding_ids": list(intent.get("expected_finding_ids") or []),
            "link_hash": sha256_json(link),
            "completed": True,
        }
        _append_retest_event_once(
            workspace_id,
            child_review_id,
            "retest_completed",
            {**completion, "side": "child"},

        )
        _append_retest_event_once(
            workspace_id,
            parent_review_id,
            "retest_completed",
            {**completion, "side": "parent"},

        )
        child = _project(workspace_id, child_review_id, check_freshness=False)
        return _ok(
            parent_review_id=parent_review_id,
            retest_review_id=child_review_id,
            review_id=child_review_id,
            review_status=child["review_status"],
            overall_verdict=child["overall_verdict"],
            release_ready=child["release_ready"],
            closed_finding_ids=closed,
            remaining_finding_ids=remaining,
            resource_uris=[
                _review_uri(workspace_id, parent_review_id),
                _review_uri(workspace_id, child_review_id),
            ],
            blockers=child["blockers"],
            next_actions=_next_actions(child),
        )

    return _write("review_retest", args, execute)


def _export_root(workspace_id: str, export_id: str) -> Path:
    return (
        workspace_root(require_safe_id(workspace_id, "workspace_id"))
        / "mcp_objects" / "deliverable-review" / "export-files"
        / require_safe_id(export_id, "export_id")
    )


def _write_once_bytes(path: Path, content: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists():
        if path.read_bytes() != content:
            raise ValueError("immutable_export_conflict")
        return
    descriptor, temporary_name = tempfile.mkstemp(prefix=f".{path.name}.", suffix=".tmp", dir=str(path.parent))
    try:
        with os.fdopen(descriptor, "wb") as stream:
            stream.write(content)
            stream.flush()
            os.fsync(stream.fileno())
        os.replace(temporary_name, path)
    finally:
        Path(temporary_name).unlink(missing_ok=True)


def _review_markdown(state: dict[str, Any]) -> str:
    lines = [
        f"# 交付物审查报告 {state['review_id']}", "",
        f"- 目标：`{(state.get('target') or {}).get('target_type')}` / `{(state.get('target') or {}).get('target_id')}`",
        f"- 目标哈希：`{(state.get('target') or {}).get('target_sha256')}`",
        f"- 规则包：`{(state.get('rule_pack') or {}).get('rule_pack_id')}` / `{(state.get('rule_pack') or {}).get('version')}`",
        f"- 总体结论：`{state.get('overall_verdict')}`",
        f"- 审查状态：`{state.get('review_status')}`",
        f"- 结果可固化：`{str(bool(state.get('release_ready'))).lower()}`", "",
        "## Findings", "",
    ]
    if not state.get("findings"):
        lines.append("无 findings。")
    for row in state.get("findings") or []:
        lines.extend([
            f"### {row.get('severity')} {row.get('rule_id')} ({row.get('status')})", "",
            str(row.get("message") or ""), "",
            f"- 定位：`{canonical_json(row.get('target_location') or {})}`",
            f"- 期望：`{canonical_json(row.get('expected'))}`",
            f"- 实际：`{canonical_json(row.get('actual'))}`",
            f"- 差额/容差：`{canonical_json(row.get('difference'))}` / `{canonical_json(row.get('tolerance'))}`",
            f"- 检查分类：`{row.get('review_area') or '-'}`",
            f"- 整改建议：{row.get('remediation') or '-'}", "",
        ])
    lines.extend(["## 质量说明", ""])
    for row in state.get("quality_notes") or []:
        lines.append(
            f"- `{row.get('quality_note_id')}` / `{row.get('verdict')}` / `{row.get('note')}`"
        )
    lines.extend(["", "## 不可核查项", ""])
    if state.get("incomplete_reasons"):
        lines.extend(f"- `{item}`" for item in state["incomplete_reasons"])
    else:
        lines.append("无。")
    return "\n".join(lines) + "\n"


def _review_docx(state: dict[str, Any]) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading(f"交付物审查报告 {state['review_id']}", level=1)
    document.add_paragraph(f"总体结论：{state.get('overall_verdict')}  审查状态：{state.get('review_status')}")
    document.add_paragraph(f"目标：{canonical_json(state.get('target') or {})}")
    document.add_heading("Findings", level=2)
    for row in state.get("findings") or []:
        document.add_heading(f"{row.get('severity')} {row.get('rule_id')} ({row.get('status')})", level=3)
        document.add_paragraph(str(row.get("message") or ""))
        document.add_paragraph(f"定位：{canonical_json(row.get('target_location') or {})}")
        document.add_paragraph(f"期望：{canonical_json(row.get('expected'))}")
        document.add_paragraph(f"实际：{canonical_json(row.get('actual'))}")
        document.add_paragraph(f"整改建议：{row.get('remediation') or '-'}")
    document.add_heading("质量说明", level=2)
    for row in state.get("quality_notes") or []:
        document.add_paragraph(
            f"{row.get('quality_note_id')} / {row.get('verdict')} / {row.get('note')}",
            style="List Bullet",
        )
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _findings_xlsx(state: dict[str, Any]) -> bytes:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "findings"
    headers = [
        "finding_id", "rule_id", "rule_pack_version", "category", "severity", "blocking",
        "status", "confidence", "message", "expected", "actual", "difference", "tolerance",
        "target_location", "evidence", "standard_basis", "review_area", "remediation",
    ]
    sheet.append(headers)
    for row in state.get("findings") or []:
        sheet.append([
            row.get(key) if isinstance(row.get(key), (str, int, float, bool, type(None)))
            else canonical_json(row.get(key))
            for key in headers
        ])
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    for column in sheet.columns:
        letter = column[0].column_letter
        sheet.column_dimensions[letter].width = min(60, max(12, max(len(str(cell.value or "")) for cell in column) + 2))
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _export_file_uri(workspace_id: str, export_id: str, filename: str) -> str:
    return f"lvke://deliverable-review/workspaces/{workspace_id}/exports/{export_id}/files/{quote(filename)}"


def _export_resource_uri(workspace_id: str, export_id: str) -> str:
    return f"lvke://deliverable-review/workspaces/{workspace_id}/exports/{export_id}"


def _export_review_locked(
    workspace_id: str,
    review_id: str,
    requested_formats: Any,
) -> dict[str, Any]:
    """Export one review while the caller holds its terminal-mutation lock."""

    try:
        state = _project(workspace_id, review_id)
    except ValueError:
        return _blocked("review_not_found", _message("review_not_found"))
    if not state.get("reviewed"):
        return _blocked("review_not_ready", "审查尚未形成可导出的不可变结论")
    if state.get("released"):
        return _blocked(
            "review_already_released",
            "正式审查包已固化；请读取 release 绑定的不可变导出文件",
            review_id=review_id,
            resource_uris=list((state.get("release") or {}).get("resource_uris") or []),
        )
    for previous in state.get("exports") or []:
        integrity_reasons = _export_integrity_reasons(
            workspace_id,
            {
                "review_id": review_id,
                "export_id": previous.get("export_id"),
                "export_record_id": previous.get("export_record_id"),
                "export_record_hash": previous.get("export_record_hash"),
                "export_basis_hash": previous.get("export_basis_hash"),
                "export_files": previous.get("files"),
            },
        )
        if integrity_reasons:
            return _blocked(
                "review_export_integrity_failed",
                "既有审查导出记录或文件完整性校验失败",
                review_id=review_id,
                integrity_reasons=integrity_reasons,
            )
    requested = list(
        requested_formats or ["json", "markdown", "docx", "xlsx"]
    )
    allowed = {"json", "markdown", "docx", "xlsx"}
    if not requested or any(item not in allowed for item in requested):
        return _blocked("export_format_invalid", "导出格式仅支持 json、markdown、docx、xlsx")
    export_basis = {
        "review_id": review_id,
        "event_chain_hash": state.get("event_chain_hash"),
        "formats": sorted(set(requested)),
        "review_status": state.get("review_status"),
        "overall_verdict": state.get("overall_verdict"),
    }
    export_id = (
        "rvexp_" + sha256_json(export_basis).removeprefix("sha256:")[:24]
    )
    output = _export_root(workspace_id, export_id)
    payloads: dict[str, tuple[str, bytes, str]] = {
        "json": (
            "review.json",
            json.dumps(
                state, ensure_ascii=False, indent=2, default=str,
            ).encode("utf-8"),
            "application/json",
        ),
        "markdown": (
            "review.md",
            _review_markdown(state).encode("utf-8"),
            "text/markdown; charset=utf-8",
        ),
        "docx": (
            "review.docx",
            _review_docx(state),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        "xlsx": (
            "findings.xlsx",
            _findings_xlsx(state),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ),
    }
    files = []
    for format_name in sorted(set(requested)):
        filename, content, media_type = payloads[format_name]
        _write_once_bytes(output / filename, content)
        files.append({
            "format": format_name,
            "filename": filename,
            "media_type": media_type,
            "bytes": len(content),
            "sha256": "sha256:" + hashlib.sha256(content).hexdigest(),
            "uri": _export_file_uri(workspace_id, export_id, filename),
        })
    record = EXPORT_STORE.put(
        workspace_id,
        {
            **export_basis,
            "export_id": export_id,
            "files": files,
        },
        producer="lvke-deliverable-review.review_export",
        source_ids=[review_id],
        basis=export_basis,
        schema_version="deliverable_review_export.v1",
    )
    export_envelope = {
        "review_id": review_id,
        "export_id": export_id,
        "export_record_id": record.get("object_id"),
        "export_record_hash": record.get("content_hash"),
        "export_basis_hash": record.get("basis_hash"),
        "export_files": files,
    }
    integrity_reasons = _export_integrity_reasons(
        workspace_id,
        export_envelope,
    )
    if integrity_reasons:
        return _blocked(
            "review_export_integrity_failed",
            "新生成的审查导出记录或文件完整性校验失败",
            review_id=review_id,
            integrity_reasons=integrity_reasons,
        )
    STORE.append(
        workspace_id,
        review_id,
        "review_exported",
        {
            "export_id": export_id,
            "export_record_id": record["object_id"],
            "export_record_hash": record["content_hash"],
            "export_basis_hash": record["basis_hash"],
            "files": files,
            "exported_at": utc_now(),
        },
    )
    return _ok(
        review_id=review_id,
        export_id=export_id,
        export_record_id=record["object_id"],
        export_record_hash=record["content_hash"],
        export_basis_hash=record["basis_hash"],
        files=files,
        resource_uris=[_export_resource_uri(workspace_id, export_id), *[row["uri"] for row in files]],
        blockers=[],
        next_actions=_next_actions(
            _project(workspace_id, review_id, check_freshness=False)
        ),
    )


def export_review(args: dict[str, Any]) -> dict[str, Any]:
    def execute(workspace_id: str) -> dict[str, Any]:
        return _export_review_locked(
            workspace_id,
            str(args.get("review_id") or ""),
            args.get("formats"),
        )

    return _write("review_export", args, execute)


def release(args: dict[str, Any]) -> dict[str, Any]:
    def execute(workspace_id: str) -> dict[str, Any]:
        review_id = str(args.get("review_id") or "")
        with STORE.mutation_guard(
            workspace_id,
            "review_terminal_mutation",
            review_id,
        ):
            try:
                state = _project(workspace_id, review_id)
            except ValueError:
                return _blocked("review_not_found", _message("review_not_found"))
            if state.get("deployment_mode") == "shadow":
                return _blocked(
                    "review_shadow_mode_release_forbidden",
                    "影子期审查仅记录新旧门禁差异，不得固化为正式审查包",
                    review_id=review_id,
                    shadow_comparison=state.get("shadow_comparison") or {},
                )
            if str((state.get("project_context") or {}).get("evidence_track") or "real") != "real":
                return _blocked(
                    "review_non_real_evidence_release_forbidden",
                    "technical_fixture 或 controlled_assumption 轨不得固化正式审查包",
                    review_id=review_id,
                    evidence_track=(state.get("project_context") or {}).get("evidence_track"),
                )
            if state.get("released"):
                target = state.get("target") or {}
                target_spec = state.get("target_spec") or {}
                verified = require_released_review_for_target(
                    workspace_id,
                    review_id,
                    str(target.get("target_type") or ""),
                    str(target.get("target_id") or ""),
                    artifact_domain=str(
                        target_spec.get("artifact_domain") or ""
                    ),

                    target_spec=target_spec,
                )
                if verified.get("status") != "ok":
                    return verified
                release_row = state.get("release") or {}
                return _ok(
                    review_id=review_id,
                    release_id=release_row.get("release_id"),
                    release_record_id=release_row.get("release_record_id"),
                    release_record_hash=release_row.get("release_record_hash"),
                    released_at=release_row.get("released_at"),
                    release_note=release_row.get("release_note"),
                    formally_deliverable=True,
                    resource_uris=list(
                        release_row.get("resource_uris")
                        or [_review_uri(workspace_id, review_id)]
                    ),
                    warnings=["正式审查包此前已固化并重新通过完整性校验"],
                    blockers=[],
                    next_actions=[],
                )
            export_result = _export_review_locked(
                workspace_id,
                review_id,
                ["json", "markdown", "docx", "xlsx"],
            )
            if export_result.get("status") == "blocked":
                return export_result
            state = _project(workspace_id, review_id)
            if state.get("invalidated"):
                return _blocked(
                    "review_invalidated",
                    "目标或审查依据在导出期间发生变化，正式审查包未固化",
                    review_id=review_id,
                    invalidation=state.get("invalidation") or {},
                )
            export_envelope = {
                "review_id": review_id,
                "export_id": export_result.get("export_id"),
                "export_record_id": export_result.get("export_record_id"),
                "export_record_hash": export_result.get("export_record_hash"),
                "export_basis_hash": export_result.get("export_basis_hash"),
                "export_files": export_result.get("files") or [],
            }
            export_integrity_reasons = _export_integrity_reasons(
                workspace_id,
                export_envelope,
            )
            if export_integrity_reasons:
                return _blocked(
                    "review_export_integrity_failed",
                    "正式审查包引用的导出记录或文件完整性校验失败",
                    review_id=review_id,
                    integrity_reasons=export_integrity_reasons,
                )
            release_note = str(args.get("note") or "").strip()
            release_basis = {
                "review_id": review_id,
                "target": state.get("target"),
                "rule_pack_hash": (state.get("rule_pack") or {}).get(
                    "content_hash"
                ),
                "standards_hash": (state.get("standards") or {}).get(
                    "content_hash"
                ),
                "event_chain_hash": state.get("event_chain_hash"),
                "export_id": export_result.get("export_id"),
                "export_record_id": export_result.get("export_record_id"),
                "export_record_hash": export_result.get("export_record_hash"),
                "export_basis_hash": export_result.get("export_basis_hash"),
            }
            if release_note:
                release_basis["release_note"] = release_note
            release_id = (
                "release_"
                + sha256_json(release_basis).removeprefix("sha256:")[:24]
            )
            release_payload = {
                **release_basis,
                "release_id": release_id,
                "released_at": utc_now(),
                "immutable_review_snapshot": state,
                "export_files": export_result.get("files") or [],
                "notice": (
                    "固化审查包记录的是工具检查结果，不代表人工签字或安全审批。"
                ),
            }
            record = RELEASE_STORE.put(
                workspace_id,
                release_payload,
                producer="lvke-deliverable-review.review_release",
                source_ids=[
                    review_id,
                    str(export_result.get("export_id") or ""),
                ],
                basis=release_basis,
                schema_version="deliverable_review_release.v1",
            )
            expected_release_record_id = (
                f"{RELEASE_STORE.id_prefix}_"
                f"{sha256_json(release_payload).removeprefix('sha256:')[:24]}"
            )
            if (
                record.get("payload") != release_payload
                or record.get("content_hash") != sha256_json(release_payload)
                or record.get("basis_hash") != sha256_json(release_basis)
                or record.get("object_id") != expected_release_record_id
                or record.get("workspace_id") != workspace_id
            ):
                return _blocked(
                    "review_release_integrity_failed",
                    "正式审查包记录写入后完整性校验失败",
                    review_id=review_id,
                )
            resource_uris = [
                record["resource_uri"],
                *list(export_result.get("resource_uris") or []),
            ]
            STORE.append(
                workspace_id,
                review_id,
                "review_released",
                {
                    "release_id": release_id,
                    "release_record_id": record["object_id"],
                    "release_record_hash": record["content_hash"],
                    "released_at": (record.get("payload") or {}).get(
                        "released_at"
                    ),
                    **(
                        {"release_note": release_note}
                        if release_note else {}
                    ),
                    "resource_uris": resource_uris,
                },
            )
            current = _project(
                workspace_id,
                review_id,
            )
            if current.get("formally_deliverable") is not True:
                code = (
                    "review_invalidated"
                    if current.get("invalidated")
                    else "review_release_integrity_failed"
                )
                return _blocked(
                    code,
                    "正式审查包固化后完整性复验失败，不得宣称发布成功",
                    review_id=review_id,
                    release_id=release_id,
                    release_record_id=record["object_id"],
                    formally_deliverable=False,
                    resource_uris=resource_uris,
                    invalidation=current.get("invalidation") or {},
                )
            return _ok(
                review_id=review_id,
                release_id=release_id,
                release_record_id=record["object_id"],
                release_record_hash=record["content_hash"],
                released_at=(record.get("payload") or {}).get("released_at"),
                release_note=release_note or None,
                formally_deliverable=current.get("formally_deliverable"),
                resource_uris=resource_uris,
                warnings=["发布记录仅表示 MCP 检查结果已固化"],
                blockers=[],
                next_actions=[],
            )
    return _write("review_release", args, execute)


def _release_export_integrity_reasons(
    workspace_id: str,
    release_payload: dict[str, Any],
) -> list[str]:
    # Records produced before export files were embedded remain readable; all
    # current release records include the field and are verified strictly.
    if "export_files" not in release_payload:
        return []
    export_id = str(release_payload.get("export_id") or "")
    rows = release_payload.get("export_files")
    if not export_id or not isinstance(rows, list) or not rows:
        return ["release_export_manifest_invalid"]
    try:
        root = _export_root(workspace_id, export_id)
    except ValueError:
        return ["release_export_id_invalid"]
    reasons: list[str] = []
    seen: set[str] = set()
    for row in rows:
        if not isinstance(row, dict):
            reasons.append("release_export_entry_invalid")
            continue
        filename = str(row.get("filename") or "")
        if (
            not filename
            or filename in seen
            or Path(filename).name != filename
            or filename in {".", ".."}
        ):
            reasons.append("release_export_filename_invalid")
            continue
        seen.add(filename)
        path = root / filename
        if not path.is_file():
            reasons.append(f"release_export_missing:{filename}")
            continue
        try:
            content = path.read_bytes()
            expected_bytes = int(row.get("bytes") or -1)
        except (OSError, TypeError, ValueError):
            reasons.append(f"release_export_unreadable:{filename}")
            continue
        if expected_bytes != len(content):
            reasons.append(f"release_export_size_mismatch:{filename}")
        digest = "sha256:" + hashlib.sha256(content).hexdigest()
        if str(row.get("sha256") or "") != digest:
            reasons.append(f"release_export_hash_mismatch:{filename}")
    return sorted(set(reasons))


def _export_record_integrity_reasons(
    workspace_id: str,
    export_envelope: dict[str, Any],
) -> list[str]:
    export_record_id = str(export_envelope.get("export_record_id") or "")
    if not export_record_id:
        return ["release_export_record_id_missing"]
    try:
        record = EXPORT_STORE.get(workspace_id, export_record_id)
    except ValueError:
        record = None
    if record is None:
        return ["release_export_record_missing"]
    payload = record.get("payload")
    if not isinstance(payload, dict):
        return ["release_export_record_payload_invalid"]
    export_basis = {
        key: deepcopy(payload.get(key))
        for key in (
            "review_id",
            "event_chain_hash",
            "formats",
            "review_status",
            "overall_verdict",
        )
    }
    content_hash = sha256_json(payload)
    expected_record_id = (
        f"{EXPORT_STORE.id_prefix}_"
        f"{content_hash.removeprefix('sha256:')[:24]}"
    )
    expected_uri = EXPORT_STORE.uri(workspace_id, expected_record_id)
    declared_record_hash = str(
        export_envelope.get("export_record_hash") or ""
    )
    declared_basis_hash = str(
        export_envelope.get("export_basis_hash") or ""
    )
    reasons: list[str] = []
    if record.get("content_hash") != content_hash:
        reasons.append("release_export_record_content_hash_mismatch")
    if record.get("basis_hash") != sha256_json(export_basis):
        reasons.append("release_export_record_basis_hash_mismatch")
    if record.get("object_id") != expected_record_id:
        reasons.append("release_export_record_object_id_mismatch")
    if export_record_id != expected_record_id:
        reasons.append("release_export_record_reference_mismatch")
    if record.get("workspace_id") != workspace_id:
        reasons.append("release_export_record_workspace_mismatch")
    if record.get("resource_uri") != expected_uri:
        reasons.append("release_export_record_uri_mismatch")
    if record.get("producer") != "lvke-deliverable-review.review_export":
        reasons.append("release_export_record_producer_mismatch")
    if record.get("schema_version") != "deliverable_review_export.v1":
        reasons.append("release_export_record_schema_mismatch")
    if declared_record_hash and declared_record_hash != content_hash:
        reasons.append("release_export_record_declared_hash_mismatch")
    if declared_basis_hash and declared_basis_hash != sha256_json(export_basis):
        reasons.append("release_export_record_declared_basis_mismatch")
    if str(payload.get("review_id") or "") != str(
        export_envelope.get("review_id") or ""
    ):
        reasons.append("release_export_record_review_mismatch")
    if str(payload.get("export_id") or "") != str(
        export_envelope.get("export_id") or ""
    ):
        reasons.append("release_export_record_export_id_mismatch")
    if payload.get("files") != export_envelope.get("export_files"):
        reasons.append("release_export_record_files_mismatch")
    if str(export_envelope.get("review_id") or "") not in {
        str(item) for item in record.get("source_ids") or []
    }:
        reasons.append("release_export_record_source_mismatch")
    return sorted(set(reasons))


def _export_integrity_reasons(
    workspace_id: str,
    export_envelope: dict[str, Any],
) -> list[str]:
    return sorted(set([
        *_export_record_integrity_reasons(workspace_id, export_envelope),
        *_release_export_integrity_reasons(workspace_id, export_envelope),
    ]))


def require_released_review_for_target(
    workspace_id: str,
    review_id: str,
    target_type: str,
    target_id: str,
    *,
    artifact_domain: str = "",
    target_spec: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Validate one explicit review against the current exact target.

    Formal release callers must never search for an arbitrary eligible review.
    This helper re-resolves the current target, runs the freshness projection,
    compares the immutable target hash and verifies the stored review-release
    envelope before returning binding material for the legacy release record.
    """

    try:
        workspace_id = require_safe_id(workspace_id, "workspace_id")
        review_id = require_safe_id(review_id, "review_id")
        artifact_scope = {"artifact_domain": artifact_domain}
        normalized = normalize_target(
            deepcopy(target_spec)
            if isinstance(target_spec, dict)
            else {
                "target_type": target_type,
                "target_id": target_id,
                **(
                    artifact_scope
                    if target_type == "report_artifact" else {}
                ),
            }
        )
        if (
            normalized.get("target_type") != target_type
            or normalized.get("target_id") != target_id
        ):
            return _blocked(
                "review_target_scope_mismatch",
                "完整目标规格与待发布目标类型或逻辑 ID 不一致",
                review_id=review_id,
            )
        resolved, target_blockers = _resolve_target(
            workspace_id,
            normalized,
        )
        if target_blockers or resolved is None:
            return _blocked(
                "review_target_unavailable",
                "待发布目标不存在、已失效或完整性校验未通过",
                review_id=review_id,
                blockers=target_blockers or ["review_target_unavailable"],
            )
        state = _project(workspace_id, review_id)
    except ValueError as exc:
        code = "review_not_found" if str(exc) in {"review_not_found", "invalid review_id"} else str(exc)
        return _blocked(code, _message(code), review_id=str(review_id or ""))

    reviewed_target = state.get("target") or {}
    if target_type == "report_artifact":
        reviewed_spec = state.get("target_spec") or {}
        expected_spec = resolved.get("target_spec") or {}
        expected_scope = {
            "artifact_domain": expected_spec.get("artifact_domain"),
        }
        actual_scope = {
            "artifact_domain": reviewed_spec.get("artifact_domain"),
        }
        if actual_scope != expected_scope:
            return _blocked(
                "review_target_scope_mismatch",
                "review_id 未绑定当前待发布工件的业务域范围",
                review_id=review_id,
                expected_scope=expected_scope,
                actual_scope=actual_scope,
            )
    expected_target = {
        key: resolved.get(key)
        for key in ("target_type", "target_id", "target_sha256")
    }
    actual_target = {
        key: reviewed_target.get(key)
        for key in ("target_type", "target_id", "target_sha256")
    }
    if actual_target != expected_target:
        return _blocked(
            "review_target_mismatch",
            "review_id 未绑定当前待发布工件的精确内容哈希",
            review_id=review_id,
            expected_target=expected_target,
            actual_target=actual_target,
        )
    if state.get("invalidated"):
        return _blocked(
            "review_invalidated",
            "审查目标或依据已变化，原审查不可用于正式发布",
            review_id=review_id,
            invalidation=state.get("invalidation") or {},
        )
    if state.get("event_chain_valid") is not True:
        return _blocked(
            "review_event_chain_invalid",
            "审查事件链完整性校验失败",
            review_id=review_id,
        )
    if state.get("released") is not True:
        return _blocked(
            "review_not_released",
            "统一审查尚未固化正式审查包",
            review_id=review_id,
        )
    if state.get("formally_deliverable") is not True:
        return _blocked(
            "review_not_formally_deliverable",
            "统一审查未通过正式交付门禁",
            review_id=review_id,
            overall_verdict=state.get("overall_verdict"),
            review_status=state.get("review_status"),
        )

    release_projection = state.get("release") or {}
    release_record_id = str(release_projection.get("release_record_id") or "")
    try:
        release_record = (
            RELEASE_STORE.get(workspace_id, release_record_id)
            if release_record_id else None
        )
    except ValueError:
        release_record = None
    if release_record is None:
        return _blocked(
            "review_release_record_missing",
            "正式审查包记录不存在",
            review_id=review_id,
        )
    release_payload = release_record.get("payload") or {}
    release_hash = str(release_record.get("content_hash") or "")
    release_basis_keys = [
        "review_id",
        "target",
        "rule_pack_hash",
        "standards_hash",
        "event_chain_hash",
        "export_id",
    ]
    release_basis_keys.extend(
        key
        for key in (
            "export_record_id",
            "export_record_hash",
            "export_basis_hash",
            "release_note",
        )
        if key in release_payload
    )
    release_basis = {
        key: deepcopy(release_payload.get(key))
        for key in release_basis_keys
    }
    expected_basis_hash = sha256_json(release_basis)
    expected_release_id = (
        "release_" + expected_basis_hash.removeprefix("sha256:")[:24]
    )
    expected_record_id = (
        f"{RELEASE_STORE.id_prefix}_"
        f"{release_hash.removeprefix('sha256:')[:24]}"
    )
    if (
        release_hash != sha256_json(release_payload)
        or release_record_id != expected_record_id
        or str(release_record.get("object_id") or "") != expected_record_id
        or str(release_record.get("workspace_id") or "") != workspace_id
        or str(release_record.get("basis_hash") or "") != expected_basis_hash
        or (
            str(release_projection.get("release_record_hash") or "")
            and str(release_projection.get("release_record_hash") or "")
            != release_hash
        )
        or str(release_payload.get("review_id") or "") != review_id
        or (release_payload.get("target") or {}) != reviewed_target
        or str(release_payload.get("release_id") or "")
        != str(release_projection.get("release_id") or "")
        or str(release_payload.get("release_id") or "") != expected_release_id
        or (
            "release_note" in release_projection
            and release_projection.get("release_note")
            != release_payload.get("release_note")
        )
    ):
        return _blocked(
            "review_release_integrity_failed",
            "正式审查包哈希或目标绑定校验失败",
            review_id=review_id,
        )
    export_integrity_reasons = (
        _export_integrity_reasons(
            workspace_id,
            {
                "review_id": review_id,
                "export_id": release_payload.get("export_id"),
                "export_record_id": release_payload.get("export_record_id"),
                "export_record_hash": release_payload.get("export_record_hash"),
                "export_basis_hash": release_payload.get("export_basis_hash"),
                "export_files": release_payload.get("export_files"),
            },
        )
        if "export_record_id" in release_payload
        else _release_export_integrity_reasons(
            workspace_id,
            release_payload,
        )
    )
    if export_integrity_reasons:
        return _blocked(
            "review_release_integrity_failed",
            "正式审查包绑定的导出文件不存在或哈希校验失败",
            review_id=review_id,
            integrity_reasons=export_integrity_reasons,
        )
    return _ok(
        review_id=review_id,

        target=expected_target,
        review_release_id=release_projection.get("release_id"),
        review_release_record_id=release_record_id,
        review_release_hash=release_hash,
        review_release_basis_hash=release_record.get("basis_hash"),
        review_event_chain_hash=state.get("event_chain_hash"),
        released_at=release_projection.get("released_at"),
        formally_deliverable=True,
        resource_uris=list(release_projection.get("resource_uris") or []),
        blockers=[],
        next_actions=[],
    )


def latest_review_for_target(
    workspace_id: str,
    target_type: str,
    target_id: str,
    *,
    artifact_domain: str = "",
) -> dict[str, Any] | None:
    """Return the latest review state after applying freshness projection."""

    try:
        workspace_id = require_safe_id(workspace_id, "workspace_id")
        normalized = normalize_target({
            "target_type": target_type,
            "target_id": target_id,
        })
    except ValueError:
        return None
    if normalized["target_type"] == "report_artifact":
        if artifact_domain not in _REPORT_ARTIFACT_DOMAINS:
            return None
    elif artifact_domain:
        return None

    candidates: list[dict[str, Any]] = []
    for review_id in STORE.review_ids(workspace_id):
        try:
            state = _project(workspace_id, review_id)
        except (OSError, ValueError):
            continue
        target = state.get("target") or {}
        if (
            target.get("target_type") != normalized["target_type"]
            or str(target.get("target_id") or "") != normalized["target_id"]
        ):
            continue
        if normalized["target_type"] == "report_artifact":
            target_spec = state.get("target_spec") or {}
            if target_spec.get("artifact_domain") != artifact_domain:
                continue
        candidates.append(state)
    if not candidates:
        return None
    candidates.sort(
        key=lambda row: (
            _parse_timestamp(row.get("created_at"))
            or datetime.min.replace(tzinfo=timezone.utc),
            str(row.get("review_id") or ""),
        ),
        reverse=True,
    )
    return deepcopy(candidates[0])


def released_review_for_target(
    workspace_id: str,
    target_type: str,
    target_id: str,
    *,
    target_sha256: str = "",
    artifact_domain: str = "",
) -> dict[str, Any] | None:
    """Return a fresh formally-deliverable review for an exact target binding."""

    if target_type == "report_artifact" and artifact_domain not in _REPORT_ARTIFACT_DOMAINS:
        return None
    candidates: list[dict[str, Any]] = []
    for review_id in STORE.review_ids(require_safe_id(workspace_id, "workspace_id")):
        try:
            state = _project(workspace_id, review_id)
        except ValueError:
            continue
        target = state.get("target") or {}
        if target.get("target_type") != target_type or str(target.get("target_id") or "") != str(target_id):
            continue
        if target_type == "report_artifact":
            target_spec = state.get("target_spec") or {}
            if target_spec.get("artifact_domain") != artifact_domain:
                continue
        if target_sha256 and target.get("target_sha256") != target_sha256:
            continue
        if state.get("formally_deliverable"):
            candidates.append(state)
    if not candidates:
        return None
    candidates.sort(key=lambda row: str(((row.get("release") or {}).get("released_at") or "")), reverse=True)
    return candidates[0]


def _standard_catalog() -> dict[str, Any]:
    path = REPO_ROOT / "config" / "review_standard_requirements.json"
    try:
        document = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        raise ValueError("standard_catalog_invalid") from None
    requirements = document.get("requirements") if isinstance(document, dict) else None
    if not isinstance(requirements, list) or not requirements:
        raise ValueError("standard_catalog_invalid")
    seen: set[str] = set()
    normalized: list[dict[str, Any]] = []
    for raw in requirements:
        if not isinstance(raw, dict):
            raise ValueError("standard_catalog_invalid")
        requirement_id = str(raw.get("requirement_id") or "").strip()
        if not requirement_id or requirement_id in seen:
            raise ValueError("standard_catalog_invalid")
        seen.add(requirement_id)
        normalized.append(deepcopy(raw))
    body = {
        "schema_version": str(document.get("schema_version") or ""),
        "catalog_version": str(document.get("catalog_version") or ""),
        "requirements": normalized,
    }
    return {**body, "content_hash": sha256_json(body)}


def _standard_requirement_applicability(
    requirement: dict[str, Any],
    project_context: dict[str, Any],
    facilities: list[dict[str, Any]],
) -> tuple[bool, str, list[str]]:
    project_types = {str(item) for item in requirement.get("applicable_project_types") or []}
    if project_types and str(project_context.get("project_type") or "") not in project_types:
        return False, "project_type_not_applicable", []
    asset_types = {str(item) for item in requirement.get("applicable_asset_types") or []}
    if asset_types and str(project_context.get("asset_type") or "") not in asset_types:
        return False, "asset_type_not_applicable", []
    facility_types = {str(item) for item in requirement.get("applicable_facility_types") or []}
    if not facility_types:
        return True, "project_context_match", []
    matched = [
        str(item.get("facility_id") or item.get("name") or item.get("facility_type") or "")
        for item in facilities
        if str(item.get("facility_type") or "") in facility_types
    ]
    if matched:
        return True, "facility_inventory_match", matched
    if not facilities:
        # Missing equipment inventory must widen the pending scope rather than
        # silently exclude a potentially mandatory large-facility standard.
        return True, "facility_inventory_pending", []
    return False, "facility_type_not_present", []


def resolve_standards(args: dict[str, Any]) -> dict[str, Any]:
    def execute(workspace_id: str) -> dict[str, Any]:
        raw_context = args.get("project_context") if isinstance(args.get("project_context"), dict) else {}
        target_type = str(raw_context.get("target_type") or "report_revision")
        project_context = normalize_project_context(raw_context, target_type=target_type)
        facilities = [
            {
                "facility_id": str(item.get("facility_id") or "").strip(),
                "name": str(item.get("name") or "").strip(),
                "facility_type": str(item.get("facility_type") or "").strip(),
                "model": str(item.get("model") or "").strip(),
                "quantity": int(item.get("quantity") or 1),
            }
            for item in (args.get("facilities") or [])
            if isinstance(item, dict)
        ]
        catalog = _standard_catalog()
        applicable: list[dict[str, Any]] = []
        excluded: list[dict[str, Any]] = []
        for requirement in catalog["requirements"]:
            selected, reason, matched_facility_ids = _standard_requirement_applicability(
                requirement, project_context, facilities,
            )
            row = {
                **deepcopy(requirement),
                "applicability_reason": reason,
                "matched_facility_ids": matched_facility_ids,
                "evidence_status": "pending_evidence" if selected else "not_applicable",
            }
            (applicable if selected else excluded).append(row)
        payload = {
            "project_context": project_context,
            "facilities": facilities,
            "applicable_requirements": applicable,
            "excluded_requirements": excluded,
            "catalog_version": catalog["catalog_version"],
            "catalog_content_hash": catalog["content_hash"],

        }
        record = STANDARD_APPLICABILITY_STORE.put(
            workspace_id,
            payload,
            producer="lvke-deliverable-review.review_resolve_standards",
            status="ok",
            source_ids=[row["requirement_id"] for row in applicable],
            basis={
                "project_context": project_context,
                "facilities": facilities,
                "catalog_content_hash": catalog["content_hash"],
            },
            schema_version="standard_applicability.v1",
        )
        pending_inventory = any(
            row.get("applicability_reason") == "facility_inventory_pending"
            for row in applicable
        )
        return _ok(
            standard_applicability_id=record["object_id"],
            project_context=project_context,
            applicable_requirements=applicable,
            excluded_requirements=excluded,
            applicable_requirement_count=len(applicable),
            excluded_requirement_count=len(excluded),
            standards_content_hash=record["content_hash"],
            catalog_content_hash=catalog["content_hash"],
            compliance_conclusion="not_determined",
            resource_uris=[record["resource_uri"]],
            warnings=["设备设施清单缺失，涉及设备类型的标准仅能判为待确认"] if pending_inventory else [],
            blockers=[],
            next_actions=["调用 review_list_requirements 查看证明材料需求并绑定不可变证据"],
        )
    scoped_args = dict(args)
    if not str(scoped_args.get("idempotency_key") or "").strip():
        basis = {
            "workspace_id": scoped_args.get("workspace_id"),
            "project_context": scoped_args.get("project_context"),
            "facilities": scoped_args.get("facilities"),
        }
        scoped_args["idempotency_key"] = (
            "standards-" + sha256_json(basis).removeprefix("sha256:")[:40]
        )
    return _write("review_resolve_standards", scoped_args, execute)


def _standard_applicability_record(
    workspace_id: str,
    applicability_id: str,
) -> dict[str, Any] | None:
    try:
        record = STANDARD_APPLICABILITY_STORE.get(
            workspace_id, applicability_id, 
        )
    except ValueError:
        return None
    if not record or record.get("content_hash") != sha256_json(record.get("payload") or {}):
        return None
    return record


def _standard_evidence_rows(
    workspace_id: str,
    applicability_id: str,
) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for record in STANDARD_EVIDENCE_STORE.list(workspace_id):
        payload = record.get("payload") if isinstance(record.get("payload"), dict) else {}
        if str(payload.get("standard_applicability_id") or "") == applicability_id:
            rows.append({**deepcopy(payload), "standard_evidence_id": record.get("object_id")})
    return rows


def list_standard_requirements(args: dict[str, Any]) -> dict[str, Any]:
    workspace_id = str(args.get("workspace_id") or "")
    applicability_id = str(args.get("standard_applicability_id") or "")
    try:
        workspace_id = require_safe_id(workspace_id, "workspace_id")
        applicability_id = require_safe_id(applicability_id, "standard_applicability_id")
    except ValueError as exc:
        return _blocked(str(exc), _message(str(exc)))
    record = _standard_applicability_record(workspace_id, applicability_id)
    if record is None:
        return _blocked("standard_applicability_not_found", _message("standard_applicability_not_found"))
    payload = record.get("payload") or {}
    evidence = _standard_evidence_rows(workspace_id, applicability_id)
    by_requirement: dict[str, list[dict[str, Any]]] = {}
    for row in evidence:
        by_requirement.setdefault(str(row.get("requirement_id") or ""), []).append(row)
    requirements = [
        {**deepcopy(row), "evidence_attachments": by_requirement.get(str(row.get("requirement_id") or ""), [])}
        for row in payload.get("applicable_requirements") or []
    ]
    return _ok(
        standard_applicability_id=applicability_id,
        project_context=payload.get("project_context") or {},
        requirements=requirements,
        excluded_requirements=deepcopy(payload.get("excluded_requirements") or []),
        requirement_count=len(requirements),
        resource_uris=[record["resource_uri"], *[str(row.get("resource_uri") or "") for row in evidence if row.get("resource_uri")]],
        warnings=[], blockers=[], next_actions=["为待补证要求调用 review_attach_requirement_evidence"],
    )


def _resolve_standard_evidence_resource(
    uri: str,
    workspace_id: str,
) -> tuple[dict[str, Any], str] | None:
    if uri.startswith(f"lvke://data-acquisition/workspaces/{workspace_id}/"):
        from lvke_mcp.servers.lvke_data_acquisition import service as acquisition_service

        record = acquisition_service.resolve_resource(uri, workspace_id)
        if not isinstance(record, dict):
            return None
        return record, "real"
    if uri.startswith(f"lvke://data-analysis/workspaces/{workspace_id}/"):
        from lvke_mcp.servers.lvke_data_analysis import service as analysis_service

        record = analysis_service.resolve_resource(uri, workspace_id)
        if not isinstance(record, dict):
            return None
        payload = record.get("payload") if isinstance(record.get("payload"), dict) else {}
        evidence_track = str(payload.get("evidence_track") or "real")
        return record, evidence_track
    return None


def attach_requirement_evidence(args: dict[str, Any]) -> dict[str, Any]:
    def execute(workspace_id: str) -> dict[str, Any]:
        applicability_id = require_safe_id(
            str(args.get("standard_applicability_id") or ""),
            "standard_applicability_id",
        )
        requirement_id = str(args.get("requirement_id") or "").strip()
        record = _standard_applicability_record(workspace_id, applicability_id)
        if record is None:
            return _blocked("standard_applicability_not_found", _message("standard_applicability_not_found"))
        payload = record.get("payload") or {}
        requirement = next((
            row for row in payload.get("applicable_requirements") or []
            if str(row.get("requirement_id") or "") == requirement_id
        ), None)
        if requirement is None:
            return _blocked("standard_requirement_not_found", _message("standard_requirement_not_found"))
        resource_uri = str(args.get("resource_uri") or "").strip()
        resolved = _resolve_standard_evidence_resource(resource_uri, workspace_id)
        if resolved is None:
            return _blocked("standard_evidence_resource_invalid", _message("standard_evidence_resource_invalid"))
        source_record, source_track = resolved
        supplied_hash = str(args.get("content_hash") or "").lower()
        if supplied_hash and not supplied_hash.startswith("sha256:"):
            supplied_hash = f"sha256:{supplied_hash}"
        actual_hash = str(source_record.get("content_hash") or "").lower()
        if actual_hash != supplied_hash:
            return _blocked("standard_evidence_hash_mismatch", _message("standard_evidence_hash_mismatch"))
        requested_track = str(args.get("evidence_track") or "")
        applicability_track = str((payload.get("project_context") or {}).get("evidence_track") or "real")
        if requested_track != applicability_track or source_track != requested_track:
            return _blocked("standard_evidence_track_mismatch", _message("standard_evidence_track_mismatch"))
        evidence_payload = {
            "standard_applicability_id": applicability_id,
            "requirement_id": requirement_id,
            "resource_uri": resource_uri,
            "locator": str(args.get("locator") or "").strip(),
            "content_hash": actual_hash,
            "evidence_track": requested_track,

        }
        evidence_record = STANDARD_EVIDENCE_STORE.put(
            workspace_id,
            evidence_payload,
            producer="lvke-deliverable-review.review_attach_requirement_evidence",
            source_ids=[applicability_id, requirement_id, str(source_record.get("object_id") or "")],
            basis=evidence_payload,
            schema_version="standard_requirement_evidence.v1",
        )
        return _ok(
            standard_applicability_id=applicability_id,
            standard_evidence_id=evidence_record["object_id"],
            requirement_id=requirement_id,
            evidence_track=requested_track,
            evidence_status={
                "technical_fixture": "satisfied_technical_fixture",
                "real": "evidence_attached_pending_review",
                "controlled_assumption": "unable_to_determine",
            }.get(requested_track, "unable_to_determine"),
            formal_evidence_candidate=requested_track == "real",
            compliance_conclusion="not_determined",
            resource_uris=[evidence_record["resource_uri"], resource_uri],
            warnings=[], blockers=[], next_actions=["调用 review_validate_standards 汇总证据状态"],
        )
    return _write("review_attach_requirement_evidence", args, execute)


def validate_standards(args: dict[str, Any]) -> dict[str, Any]:
    workspace_id = str(args.get("workspace_id") or "")
    applicability_id = str(args.get("standard_applicability_id") or "")
    try:
        workspace_id = require_safe_id(workspace_id, "workspace_id")
        applicability_id = require_safe_id(applicability_id, "standard_applicability_id")
    except ValueError as exc:
        return _blocked(str(exc), _message(str(exc)))
    record = _standard_applicability_record(workspace_id, applicability_id)
    if record is None:
        return _blocked("standard_applicability_not_found", _message("standard_applicability_not_found"))
    payload = record.get("payload") or {}
    evidence_track = str((payload.get("project_context") or {}).get("evidence_track") or "real")
    attachments = _standard_evidence_rows(workspace_id, applicability_id)
    attached_ids = {str(row.get("requirement_id") or "") for row in attachments}
    requirements: list[dict[str, Any]] = []
    for row in payload.get("applicable_requirements") or []:
        requirement_id = str(row.get("requirement_id") or "")
        if requirement_id not in attached_ids:
            evidence_status = "pending_evidence"
        elif evidence_track == "technical_fixture":
            evidence_status = "satisfied_technical_fixture"
        elif evidence_track == "real":
            evidence_status = "evidence_attached_pending_review"
        else:
            evidence_status = "unable_to_determine"
        requirements.append({**deepcopy(row), "evidence_status": evidence_status})
    status_counts = {
        status: sum(1 for row in requirements if row.get("evidence_status") == status)
        for status in (
            "satisfied_technical_fixture",
            "evidence_attached_pending_review",
            "pending_evidence",
            "unable_to_determine",
        )
    }
    unresolved = status_counts["pending_evidence"] + status_counts["unable_to_determine"]
    technical_complete = bool(
        evidence_track == "technical_fixture" and requirements and not unresolved
    )
    result_status = "ok" if technical_complete else "partial"
    return _ok(
        status=result_status,
        standard_applicability_id=applicability_id,
        evidence_track=evidence_track,
        requirements=requirements,
        excluded_requirements=deepcopy(payload.get("excluded_requirements") or []),
        status_counts=status_counts,
        technical_validation_complete=technical_complete,
        formal_compliance_determined=False,
        compliance_conclusion="not_determined",
        formal_evidence_claim_count=(
            status_counts["evidence_attached_pending_review"]
            if evidence_track == "real" else 0
        ),
        technical_fixture_claim_count=(
            status_counts["satisfied_technical_fixture"]
            if evidence_track == "technical_fixture" else 0
        ),
        external_data_gap_count=(
            status_counts["pending_evidence"] + status_counts["unable_to_determine"]
        ),
        local_implementation_issue_count=0,
        resource_uris=[record["resource_uri"], *[
            str(row.get("resource_uri") or "") for row in attachments if row.get("resource_uri")
        ]],
        warnings=["标准证据状态不能替代专业合规审查或法定审批"],
        blockers=[] if technical_complete else ["standard_evidence_or_professional_review_incomplete"],
        next_actions=(
            ["技术夹具链已完成；不得据此形成正式合规或 release 结论"]
            if technical_complete else ["补充真实不可变证据并完成质量核验"]
        ),
    )


def _resource_entry(uri: str, resource_type: str, name: str, description: str) -> dict[str, Any]:
    return {
        "uri": uri, "resource_type": resource_type, "name": name,
        "description": description, "mime_type": "application/json",
    }


def list_resources(args: dict[str, Any] | str) -> dict[str, Any]:
    if isinstance(args, str):
        args = {"workspace_id": args}
    workspace_id = str(args.get("workspace_id") or "")
    try:
        workspace_id = require_safe_id(workspace_id, "workspace_id")
        entries: dict[str, dict[str, Any]] = {}
        for candidate in PREPARATION_STORE.list(workspace_id):
            record, _integrity_reasons = _verified_preparation_record(
                workspace_id,
                str(candidate.get("object_id") or ""),
            )
            if record is None:
                continue
            uri = str(record.get("resource_uri") or "")
            entries[uri] = _resource_entry(uri, "preparation", str(record.get("object_id") or ""), "不可变审查准备对象")
        for review_id in STORE.review_ids(workspace_id):
            try:
                state = _project_events(workspace_id, review_id)
            except ValueError:
                continue
            uri = _review_uri(workspace_id, review_id)
            entries[uri] = _resource_entry(uri, "review", review_id, "append-only 审查运行投影")
            for finding in state.get("findings") or []:
                finding_id = str(finding.get("finding_id") or "")
                finding_uri = _finding_uri(workspace_id, review_id, finding_id)
                entries[finding_uri] = _resource_entry(finding_uri, "finding", finding_id, "审查 finding 及完整历史")
        for record in EXPORT_STORE.list(workspace_id):
            payload = record.get("payload") or {}
            review_id = str(payload.get("review_id") or "")
            export_id = str(payload.get("export_id") or "")
            uri = _export_resource_uri(workspace_id, export_id)
            entries[uri] = _resource_entry(uri, "export", export_id, "不可变审查导出清单")
            for file_row in payload.get("files") or []:
                file_uri = str(file_row.get("uri") or "")
                entries[file_uri] = {
                    **_resource_entry(file_uri, "export_file", str(file_row.get("filename") or ""), "审查报告导出文件"),
                    "mime_type": file_row.get("media_type") or "application/octet-stream",
                }
        for record in RELEASE_STORE.list(workspace_id):
            payload = record.get("payload") or {}
            review_id = str(payload.get("review_id") or "")
            uri = str(record.get("resource_uri") or "")
            entries[uri] = _resource_entry(uri, "release", str(record.get("object_id") or ""), "正式不可变审查包")
        for record in STANDARD_APPLICABILITY_STORE.list(workspace_id):
            uri = str(record.get("resource_uri") or "")
            entries[uri] = _resource_entry(
                uri, "standard_applicability", str(record.get("object_id") or ""),
                "标准适用性解析及排除原因",
            )
        for record in STANDARD_EVIDENCE_STORE.list(workspace_id):
            uri = str(record.get("resource_uri") or "")
            entries[uri] = _resource_entry(
                uri, "standard_evidence", str(record.get("object_id") or ""),
                "标准需求绑定的不可变证据索引",
            )
        from lvke_mcp.servers.lvke_deliverable_review import rubrics

        for store, resource_type, description in (
            (rubrics.ASSESSMENT_STORE, "rubric_assessment", "不可变章节 rubric 评分"),
            (rubrics.COMPARISON_STORE, "rubric_comparison", "修订前后 rubric 对比"),
        ):
            for record in store.list(workspace_id):
                uri = str(record.get("resource_uri") or "")
                entries[uri] = _resource_entry(
                    uri,
                    resource_type,
                    str(record.get("object_id") or ""),
                    description,
                )
        for pack in rules.registry():
            pack_id = str(pack.get("rule_pack_id") or "")
            uri = f"lvke://deliverable-review/workspaces/{workspace_id}/rule-packs/{quote(pack_id)}"
            entries[uri] = _resource_entry(uri, "rule_pack", pack_id, "版本化规则包")
        standards_uri = f"lvke://deliverable-review/workspaces/{workspace_id}/standards/current"
        entries[standards_uri] = _resource_entry(standards_uri, "standards", "current", "当前标准来源锁定清单")
        metrics_uri = _metrics_uri(workspace_id)
        entries[metrics_uri] = _resource_entry(
            metrics_uri,
            "metrics",
            "current",
            "工作区统一审查上线指标与影子门禁差异",
        )
        resource_type = str(args.get("resource_type") or "")
        page = paginate_resource_entries(
            (row for row in entries.values() if not resource_type or row["resource_type"] == resource_type),
            cursor=str(args.get("cursor") or ""), limit=int(args.get("limit") or 50),
        )
    except ValueError as exc:
        return _blocked(str(exc), _message(str(exc)))
    return _ok(
        resources=page["resources"], next_cursor=page["next_cursor"], has_more=page["has_more"],
        snapshot_hash=page["snapshot_hash"], resource_uris=[row["uri"] for row in page["resources"]],
        blockers=[], next_actions=[],
    )


def resolve_resource(
    uri: str,
    workspace_id: str | None = None,
) -> tuple[str | bytes, str] | None:
    prefix = "lvke://deliverable-review/workspaces/"
    if not str(uri).startswith(prefix):
        return None
    parts = str(uri)[len(prefix):].split("/")
    if len(parts) < 3:
        return None
    uri_workspace = parts[0]
    if workspace_id is not None and uri_workspace != workspace_id:
        return None
    try:
        require_safe_id(uri_workspace, "workspace_id")
    except ValueError:
        return None
    segment = parts[1]
    if segment == "preparations" and len(parts) == 3:
        record, _integrity_reasons = _verified_preparation_record(
            uri_workspace,
            parts[2],
        )
        if not record:
            return None
        return (json.dumps(record, ensure_ascii=False, indent=2, default=str), "application/json") if record else None
    if segment == "reviews" and len(parts) in {3, 5}:
        try:
            state = _project(uri_workspace, parts[2])
        except ValueError:
            return None
        if len(parts) == 3:
            return json.dumps(state, ensure_ascii=False, indent=2, default=str), "application/json"
        if parts[3] != "findings":
            return None
        finding = next((row for row in state.get("findings") or [] if row.get("finding_id") == parts[4]), None)
        return (json.dumps(finding, ensure_ascii=False, indent=2, default=str), "application/json") if finding else None
    if segment == "exports" and len(parts) == 3:
        requested_export_id = parts[2]
        record = next((
            row for row in EXPORT_STORE.list(uri_workspace)
            if str((row.get("payload") or {}).get("export_id") or "")
            == requested_export_id
        ), None)
        if record is None:
            # Read-only compatibility for historical content-addressed record URIs.
            record = EXPORT_STORE.get(uri_workspace, requested_export_id)
        if not record:
            return None
        if record:
            projected = {
                **record,
                "canonical_resource_uri": _export_resource_uri(
                    uri_workspace,
                    str((record.get("payload") or {}).get("export_id") or requested_export_id),
                ),
            }
            return json.dumps(projected, ensure_ascii=False, indent=2, default=str), "application/json"
        return None
    if segment == "exports" and len(parts) == 5 and parts[3] == "files":
        export_id = parts[2]
        filename = unquote(parts[4])
        if Path(filename).name != filename:
            return None
        record = next((
            row for row in EXPORT_STORE.list(uri_workspace)
            if str((row.get("payload") or {}).get("export_id") or "") == export_id
        ), None)
        if not record:
            return None
        file_row = next((
            row for row in ((record or {}).get("payload") or {}).get("files") or []
            if row.get("filename") == filename
        ), None)
        path = _export_root(uri_workspace, export_id) / filename
        if file_row is None or not path.is_file():
            return None
        content = path.read_bytes()
        if "sha256:" + hashlib.sha256(content).hexdigest() != file_row.get("sha256"):
            return None
        return content, str(file_row.get("media_type") or "application/octet-stream")
    if segment == "releases" and len(parts) == 3:
        record = RELEASE_STORE.get(uri_workspace, parts[2])
        if not record:
            return None
        return (json.dumps(record, ensure_ascii=False, indent=2, default=str), "application/json") if record else None
    if segment == "standard-applicabilities" and len(parts) == 3:
        record = STANDARD_APPLICABILITY_STORE.resolve_uri(uri)
        if record is None or str(record.get("workspace_id") or "") != uri_workspace:
            return None
        return json.dumps(record, ensure_ascii=False, indent=2, default=str), "application/json"
    if segment == "standard-evidence" and len(parts) == 3:
        record = STANDARD_EVIDENCE_STORE.resolve_uri(uri)
        if record is None or str(record.get("workspace_id") or "") != uri_workspace:
            return None
        return json.dumps(record, ensure_ascii=False, indent=2, default=str), "application/json"
    if segment in {"rubric-assessments", "rubric-comparisons"} and len(parts) == 3:
        from lvke_mcp.servers.lvke_deliverable_review import rubrics

        resolved = rubrics.resolve_rubric_resource(uri)
        if resolved is None:
            return None
        record, _object_type = resolved
        if str(record.get("workspace_id") or "") != uri_workspace:
            return None
        return json.dumps(record, ensure_ascii=False, indent=2, default=str), "application/json"
    if segment == "rule-packs" and len(parts) == 3:
        pack_id = unquote(parts[2])
        record = next((row for row in rules.registry() if row.get("rule_pack_id") == pack_id), None)
        return (json.dumps(record, ensure_ascii=False, indent=2), "application/json") if record else None
    if segment == "standards" and parts[2] == "current" and len(parts) == 3:
        path = REPO_ROOT / "config" / "review_standards.lock.json"
        return (path.read_text(encoding="utf-8"), "application/json") if path.is_file() else None
    if segment == "metrics" and parts[2] == "current" and len(parts) == 3:
        metrics = _workspace_metrics_payload(uri_workspace)
        return json.dumps(metrics, ensure_ascii=False, indent=2, default=str), "application/json"
    return None


def read_resource(args: dict[str, Any]) -> dict[str, Any]:
    workspace_id = str(args.get("workspace_id") or "")
    uri = str(args.get("uri") or "")
    try:
        workspace_id = require_safe_id(workspace_id, "workspace_id")
    except ValueError as exc:
        return _blocked(str(exc), _message(str(exc)))
    resolved = resolve_resource(uri, workspace_id)
    if resolved is None:
        return _blocked("resource_not_found", "资源不存在或不属于当前工作区")
    content, mime_type = resolved
    encoded = isinstance(content, bytes)
    return _ok(
        uri=uri, mime_type=mime_type, content_encoding="base64" if encoded else "utf-8",
        content=base64.b64encode(content).decode("ascii") if encoded else content,
        resource_uris=[uri], blockers=[], next_actions=[],
    )
