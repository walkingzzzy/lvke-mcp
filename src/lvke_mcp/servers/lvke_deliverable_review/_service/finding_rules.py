"""各规则组的 finding 生成：必需行、专业规则、既有 issue、项目元信息、收购输入、财务复算与报告。"""

from __future__ import annotations

import io
import re
from copy import deepcopy
from pathlib import Path
from typing import Any
from urllib.parse import quote

from lvke_mcp.runtime.storage import sha256_json
from lvke_mcp.servers.lvke_deliverable_review import report_checks, rules

from .base import (
    _number,
    _severity,
)


def _required_finding_rows(
    preparation_payload: dict[str, Any], standard_basis: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    track = str((preparation_payload.get("project_context") or {}).get("evidence_track") or "real")
    if track == "sim_a_formal":
        return []
    output: list[dict[str, Any]] = []
    pack = preparation_payload.get("rule_pack") or {}
    for row in preparation_payload.get("mandatory_findings") or []:
        location = {
            "target_id": (preparation_payload.get("target") or {}).get("target_id"),
            "standard_package_id": row.get("package_id"),
            "required_finding_instance": row.get("instance_key"),
        }
        item = rules.finding(
            str(row.get("rule_id") or "CORE.REQUIRED.MANUAL"), "P0", str(row.get("message") or "强制人工审查项"),
            category=str(row.get("category") or "required_manual_review"),
            expected="取得可回读原件并完成质量核验", actual="尚未提供或尚未完成核验",
            target_location=location, standard_basis=[
                item for item in standard_basis
                if item.get("standard_package_id") == row.get("package_id")
            ],
            review_area=str(row.get("review_area") or "legal"),
            remediation="补充原件、精确定位和内容哈希；生成新目标版本后复测并完成质量核验",
        )
        item["rule_pack_id"] = pack.get("rule_pack_id")
        item["rule_pack_version"] = pack.get("version")
        item["waiver_allowed"] = False
        item["required_source_document"] = True
        output.append(item)
    return output


def _professional_rule_finding(
    preparation_payload: dict[str, Any],
    source_rule: dict[str, Any],
    standard_basis: list[dict[str, Any]],
) -> dict[str, Any]:
    target = preparation_payload.get("target") or {}
    standard = source_rule.get("standard") or {}
    package_id = str(standard.get("package_id") or "")
    artifact_id = str(standard.get("artifact_id") or "")
    matching_basis = [
        deepcopy(row)
        for row in standard_basis
        if str(row.get("standard_package_id") or "") == package_id
        and (
            not artifact_id
            or str(row.get("standard_artifact_id") or "") == artifact_id
        )
    ]
    if not matching_basis:
        matching_basis = [{}]
    precise_basis = [
        {
            **row,
            "standard_package_id": package_id,
            "standard_artifact_id": artifact_id,
            "content_hash": standard.get("sha256") or row.get("content_hash"),
            "locator": standard.get("locator"),
            "quote": standard.get("quote"),
        }
        for row in matching_basis
    ]
    target_kinds = set(source_rule.get("target_kinds") or [])
    location: dict[str, Any] = {
        "target_type": target.get("target_type"),
        "target_id": target.get("target_id"),
        "professional_rule_id": source_rule.get("rule_id"),
    }
    if target.get("target_type") == "combined_deliverable":
        location["components"] = [
            {
                "target_type": component.get("target_type"),
                "target_id": component.get("target_id"),
            }
            for component in (
                (preparation_payload.get("target_snapshot") or {}).get("components")
                or []
            )
            if component.get("target_type") in target_kinds
        ]
    item = rules.finding(
        str(source_rule.get("rule_id") or "PROFESSIONAL.REVIEW"),
        _severity(
            source_rule.get("severity"),
            blocking=bool(source_rule.get("blocking")),
        ),
        f"待专业核验：{source_rule.get('title') or source_rule.get('rule_id')}",
        category="professional_review_pending",
        blocking=bool(source_rule.get("blocking")),
        expected=str(source_rule.get("requirement") or "完成证据化专业核验"),
        actual="尚未提交逐规则专业核验结论",
        target_location=location,
        standard_basis=precise_basis,
        review_area=str(source_rule.get("review_area") or ""),
        remediation="逐规则提交带内容哈希和精确定位的核验证据",
    )
    item["manual_review_required"] = True
    track = str((preparation_payload.get("project_context") or {}).get("evidence_track") or "real")
    item["waiver_allowed"] = track == "sim_a_formal"
    item["professional_rule"] = {
        key: deepcopy(source_rule.get(key))
        for key in (
            "rule_id", "title", "requirement", "check_kind", "review_area",
            "severity", "blocking", "target_kinds", "on_unavailable",
        )
    }
    return item


def _existing_issue_findings(run: dict[str, Any], standard_basis: list[dict[str, Any]]) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    if run and run.get("consistency_ok") is False:
        findings.append(rules.finding(
            "FIN.EXISTING.CONSISTENCY", "P0", "既有财务一致性门禁未通过",
            category="financial_consistency", expected=True, actual=False,
            review_area="finance", remediation="修正财务模型并生成新 run 后复测",
            standard_basis=standard_basis,
        ))
    legacy_issues = [
        *(run.get("issues") or []),
        *((run.get("audit") or {}).get("issues") or []),
    ]
    seen_issue_ids: set[str] = set()
    for issue in legacy_issues:
        if not isinstance(issue, dict):
            continue
        is_failed = issue.get("ok") is False or (
            str(issue.get("status") or "open") == "open"
            and bool(issue.get("blocking"))
        )
        if not is_failed:
            continue
        source_id = str(issue.get("issue_id") or "")
        issue_identity = source_id or sha256_json(issue)
        if issue_identity in seen_issue_ids:
            continue
        seen_issue_ids.add(issue_identity)
        rule_id = str(issue.get("rule") or "FIN.EXISTING.ISSUE")
        detail = issue.get("detail")
        message = (
            str(detail.get("message") or detail.get("detail") or rule_id)
            if isinstance(detail, dict) else str(detail or rule_id)
        )
        findings.append(rules.finding(
            f"FIN.LEGACY.{rule_id}", _severity(issue.get("severity"), blocking=bool(issue.get("blocking"))),
            message, category="legacy_finance_issue", blocking=bool(issue.get("blocking")),
            actual=detail, target_location={"run_id": run.get("run_id"), "source_issue_id": source_id},
            evidence=[{"source_issue_id": source_id, "audit_history": issue.get("history") or []}],
            standard_basis=standard_basis, review_area="finance",
            remediation="沿用原问题证据关闭流程整改，并以新 run 复测", source_issue_id=source_id,
        ))
    return findings


def _project_metadata_findings(
    preparation_payload: dict[str, Any], run: dict[str, Any], standard_basis: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    target = preparation_payload.get("target") or {}
    target_type = str(target.get("target_type") or "")
    snapshot = preparation_payload.get("target_snapshot") or {}
    if target_type in {"finance_run", "finance_tables_package"}:
        nested = {}
        for candidate in (
            run.get("project_metadata"),
            (run.get("finance_inputs") or {}).get("project_metadata") if isinstance(run.get("finance_inputs"), dict) else None,
            (run.get("input_revision") or {}).get("project_metadata") if isinstance(run.get("input_revision"), dict) else None,
            run.get("project_context"),
        ):
            if isinstance(candidate, dict):
                nested.update(candidate)
        source = {**run, **nested}
        if nested.get("invest_type") and not source.get("project_type"):
            source["project_type"] = nested.get("invest_type")
    else:
        revision_payload = ((snapshot.get("revision_record") or {}).get("payload") or {}) if isinstance(snapshot, dict) else {}
        source = {
            **dict((revision_payload.get("upstream") or {}).get("project_metadata") or {}),
            **dict(revision_payload.get("project_metadata") or {}),
            **dict((preparation_payload.get("target_spec") or {}).get("project_metadata") or {}),
        }
        if run:
            source = {**run, **source}
    aliases = {
        "project_type": ("project_type", "invest_type"),
        "industry": ("industry",),
        "valuation_date": ("valuation_date", "base_date", "as_of_date"),
        "currency": ("currency", "currency_code"),
        "amount_unit": ("amount_unit", "unit"),
        "tax_basis": ("tax_basis", "tax_inclusive_basis", "price_tax_basis"),
        "forecast_period": ("forecast_period", "calc_years", "forecast_years", "calc_period_years"),
    }
    missing = [
        field for field, candidates in aliases.items()
        if not any(source.get(candidate) not in (None, "", []) for candidate in candidates)
    ]
    if not missing:
        return []
    return [rules.finding(
        "PROJECT.METADATA.COMPLETE", "P1", "项目审查元数据不完整",
        category="project_metadata", expected=sorted(aliases), actual={"missing": missing},
        target_location={"target_id": target.get("target_id"), "fields": missing},
        standard_basis=standard_basis, review_area="business",
        remediation="补齐项目类型、行业、估值基准日、币种、金额单位、税口径和预测期间后生成新版本",
    )]


def _acquisition_input_findings(
    run: dict[str, Any], target_id: str, standard_basis: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    spec = run.get("spec") or {}
    transaction = dict(
        run.get("transaction") or run.get("acquisition") or spec.get("transaction") or {}
    )
    assets = run.get("assets") or run.get("asset_scope") or transaction.get("assets")
    required = {
        "valuation_date": transaction.get("valuation_date") or run.get("valuation_date"),
        "purchase_price": transaction.get("purchase_price") or transaction.get("consideration"),
        "transaction_tax": transaction.get("transaction_tax") or transaction.get("taxes"),
        "asset_scope": assets,
    }
    missing = [key for key, value in required.items() if value in (None, "", [], {})]
    if not missing:
        return []
    return [rules.finding(
        "ACQ.TRANSACTION.INPUTS", "P1", "资产收购交易输入或资产边界不完整",
        category="acquisition_inputs", expected=sorted(required), actual={"missing": missing},
        target_location={"target_id": target_id, "fields": missing}, standard_basis=standard_basis,
        review_area="business", remediation="补齐交易价格、税费、估值基准日和资产边界原始依据",
    )]


def _hotel_acquisition_run_findings(
    run: dict[str, Any], target_id: str, standard_basis: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], set[str]]:
    spec = run.get("spec") or {}
    transaction = spec.get("transaction") or run.get("transaction") or {}
    findings: list[dict[str, Any]] = []
    mode = str(transaction.get("operating_mode") or run.get("operating_mode") or "")
    if mode not in {"owner_lessor", "mixed_owner_operator"}:
        findings.append(rules.finding(
            "HOTEL.OPERATING_MODEL", "P0", "酒店收购经营模式未确认或不受模型支持",
            category="operating_assumption",
            expected=["owner_lessor", "mixed_owner_operator"], actual=mode or None,
            target_location={"target_id": target_id, "field": "transaction.operating_mode"},
            standard_basis=standard_basis, review_area="business",
            remediation="确认纯出租或混合自营模式，并按确认模式完整重算收入、成本、税费和现金流",
        ))
    assets = transaction.get("asset_scope") or run.get("asset_scope") or []
    licenses = transaction.get("licenses") or []
    parties = spec.get("project_parties") or []
    license_holders = [
        row for row in parties
        if isinstance(row, dict) and "license_holder" in (row.get("roles") or [])
    ]
    asset_evidence_ok = bool(assets) and all(
        isinstance(row, dict) and list(row.get("evidence_ids") or [])
        for row in assets
    )
    license_evidence_ok = bool(licenses or license_holders) and all(
        isinstance(row, dict) and list(row.get("evidence_ids") or [])
        for row in [*licenses, *license_holders]
    )
    if not asset_evidence_ok or not license_evidence_ok:
        findings.append(rules.finding(
            "HOTEL.RIGHTS.LICENSES", "P0", "酒店资产权属或经营许可缺少逐项原件证据绑定",
            category="rights_and_licenses",
            expected={"asset_scope_with_evidence": True, "licenses_with_evidence": True},
            actual={
                "asset_scope_count": len(assets), "asset_evidence_complete": asset_evidence_ok,
                "license_count": len(licenses) + len(license_holders),
                "license_evidence_complete": license_evidence_ok,
            },
            target_location={"target_id": target_id, "fields": ["transaction.asset_scope", "transaction.licenses", "project_parties"]},
            standard_basis=standard_basis, review_area="legal",
            remediation="逐项绑定权证、许可、主体授权及用途原件，并核对证载主体、地址、范围和有效期",
        ))
    return findings, {"HOTEL.RIGHTS.LICENSES", "HOTEL.OPERATING_MODEL"}


def _finance_recalculation_findings(run: dict[str, Any], standard_basis: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[str], dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    incomplete: list[str] = []
    coverage: dict[str, Any] = {"finance_recalculations": []}
    if not run or not run.get("available"):
        return findings, ["bound_finance_run_unavailable"], coverage
    investment = run.get("investment") or {}
    funding = run.get("funding") or {}
    total = _number(investment.get("total"))
    components = [_number(investment.get(key)) for key in ("construction", "interest", "working_capital")]
    if total is not None and all(value is not None for value in components):
        calculated = sum(float(value) for value in components if value is not None)
        difference = round(total - calculated, 6)
        tolerance = max(0.01, abs(total) * 1e-8)
        coverage["finance_recalculations"].append("total_investment")
        if abs(difference) > tolerance:
            item = rules.finding(
                "FIN.INVESTMENT.BALANCE", "P0", "总投资与建设投资、建设期利息和流动资金不平",
                category="financial_recalculation", expected=calculated, actual=total,
                difference=difference, tolerance=tolerance, target_location={"run_id": run.get("run_id"), "field": "investment.total"},
                standard_basis=standard_basis, review_area="finance", remediation="修正投资构成并重新生成财务 run",
            )
            item["calculation_trace"] = ["total = construction + interest + working_capital", f"{total} - {calculated} = {difference}"]
            findings.append(item)
    else:
        incomplete.append("investment_balance_inputs_missing")
    sources = [_number(funding.get(key)) for key in ("capital", "loan", "subsidy")]
    if total is not None and all(value is not None for value in sources):
        calculated = sum(float(value) for value in sources if value is not None)
        difference = round(total - calculated, 6)
        tolerance = max(0.01, abs(total) * 1e-8)
        coverage["finance_recalculations"].append("funding_balance")
        if abs(difference) > tolerance:
            item = rules.finding(
                "FIN.FUNDING.BALANCE", "P0", "总资金来源与投资需求不平",
                category="financial_recalculation", expected=total, actual=calculated,
                difference=round(calculated - total, 6), tolerance=tolerance,
                target_location={"run_id": run.get("run_id"), "field": "funding"},
                standard_basis=standard_basis, review_area="finance", remediation="修正资金筹措构成并重新生成财务 run",
            )
            item["calculation_trace"] = ["funding = capital + loan + subsidy", f"{calculated} - {total} = {calculated-total}"]
            findings.append(item)
    else:
        incomplete.append("funding_balance_inputs_missing")
    cashflows = ((run.get("operating") or {}).get("cashflows") or [])
    if cashflows and all(_number(value) is not None for value in cashflows):
        numeric = [float(value) for value in cashflows]
        signs = [1 if value > 0 else -1 if value < 0 else 0 for value in numeric]
        sign_changes = sum(1 for left, right in zip(signs, signs[1:]) if left and right and left != right)
        if sign_changes > 1:
            findings.append(rules.finding(
                "FIN.IRR.MULTIPLE_SIGN_CHANGES", "P1", "现金流存在多次符号变化，IRR 可能多解",
                category="financial_recalculation", actual=sign_changes, expected=1,
                target_location={"run_id": run.get("run_id"), "field": "operating.cashflows"},
                standard_basis=standard_basis, review_area="finance", remediation="披露多重 IRR 风险并以 NPV 等指标交叉判断",
            ))
        try:
            from lvke_mcp.domains.finance.calculations import irr, npv
            independent_irr = irr(numeric) * 100.0
            reported_irr = _number((run.get("indicators") or {}).get("project_irr_pct"))
            coverage["finance_recalculations"].append("project_irr")
            if reported_irr is None:
                incomplete.append("reported_project_irr_missing")
            elif abs(reported_irr - independent_irr) > 0.01:
                findings.append(rules.finding(
                    "FIN.IRR.INDEPENDENT_RECALC", "P0", "项目 IRR 与独立复算不一致",
                    category="financial_recalculation", expected=round(independent_irr, 6), actual=reported_irr,
                    difference=round(reported_irr-independent_irr, 6), tolerance=0.01,
                    target_location={"run_id": run.get("run_id"), "field": "indicators.project_irr_pct"},
                    standard_basis=standard_basis, review_area="finance", remediation="核对现金流期间、符号和 IRR 口径",
                ))
            residual = npv(numeric, independent_irr / 100.0)
            if abs(residual) > max(0.01, sum(abs(value) for value in numeric) * 1e-8):
                findings.append(rules.finding(
                    "FIN.IRR.NPV.RESIDUAL", "P0", "IRR 对应 NPV 残差超出容差",
                    category="financial_recalculation", expected=0.0, actual=residual, tolerance=0.01,
                    target_location={"run_id": run.get("run_id"), "field": "operating.cashflows"},
                    standard_basis=standard_basis, review_area="finance",
                ))
            rate = _number(run.get("benchmark_rate"))
            reported_npv = _number((run.get("indicators") or {}).get("npv_wan"))
            if rate is not None and reported_npv is not None:
                independent_npv = npv(numeric, rate)
                coverage["finance_recalculations"].append("project_npv")
                if abs(reported_npv-independent_npv) > max(0.01, abs(independent_npv)*1e-8):
                    findings.append(rules.finding(
                        "FIN.NPV.INDEPENDENT_RECALC", "P0", "项目 NPV 与独立复算不一致",
                        category="financial_recalculation", expected=round(independent_npv, 6), actual=reported_npv,
                        difference=round(reported_npv-independent_npv, 6), tolerance=0.01,
                        target_location={"run_id": run.get("run_id"), "field": "indicators.npv_wan"},
                        standard_basis=standard_basis, review_area="finance", remediation="核对折现率与现金流时点口径",
                    ))
            elif rate is None:
                incomplete.append("npv_discount_rate_missing")
        except ValueError:
            incomplete.append("independent_irr_unavailable")
    else:
        incomplete.append("project_cashflows_missing")
    return findings, sorted(set(incomplete)), coverage


_CLAIM_PATTERN = re.compile(
    r"(?P<number>-?\d{1,3}(?:,\d{3})*(?:\.\d+)?|-?\d+(?:\.\d+)?)\s*"
    r"(?P<unit>亿元|万元|元|%|％|平方米|㎡|间|吨|万吨|年|个月)"
)


_FINANCE_WORDS = {
    "投资", "收入", "成本", "利润", "现金流", "贷款", "借款", "偿债",
    "IRR", "NPV", "收益率", "回收期", "资本金", "税", "折旧", "利息",
}


def _claim_value(number: str, unit: str) -> tuple[float, str]:
    value = float(number.replace(",", ""))
    if unit == "亿元":
        return value * 10000.0, "万元"
    if unit == "元":
        return value / 10000.0, "万元"
    if unit in {"%", "％"}:
        return value, "%"
    return value, unit


def _document_from_snapshot(snapshot: dict[str, Any]) -> dict[str, Any]:
    document = snapshot.get("document") if isinstance(snapshot, dict) else {}
    if isinstance(document, dict) and isinstance(document.get("content"), str):
        return document
    return {}


def _report_content(workspace_id: str, preparation_payload: dict[str, Any]) -> str:
    snapshot = preparation_payload.get("target_snapshot") or {}
    document = _document_from_snapshot(snapshot)
    content = str(document.get("content") or "")
    if content:
        return content
    if str((preparation_payload.get("target") or {}).get("target_type") or "") != "report_artifact":
        return ""
    return _report_artifact_text(
        workspace_id,
        snapshot,
    )


def _report_evidence_packs(workspace_id: str, preparation_payload: dict[str, Any]) -> list[dict[str, Any]]:
    try:
        from lvke_mcp.adapters.data_analysis_repository import EVIDENCE_STORE
    except Exception:  # noqa: BLE001
        return []
    records: list[dict[str, Any]] = []
    for evidence_id in (preparation_payload.get("bindings") or {}).get("evidence_pack_ids") or []:
        try:
            record = EVIDENCE_STORE.get(
                workspace_id,
                str(evidence_id),
            )
        except (OSError, ValueError):
            record = None
        if record is not None:
            records.append(record)
    return records


def _expected_report_sections(preparation_payload: dict[str, Any]) -> list[str]:
    snapshot = preparation_payload.get("target_snapshot") or {}
    revision = (snapshot.get("revision_record") or {}).get("payload") or {}
    upstream = revision.get("upstream") or {}
    return [str(item) for item in (upstream.get("outline") or []) if str(item).strip()]


def _report_findings(
    workspace_id: str, preparation_payload: dict[str, Any], run: dict[str, Any],
    standard_basis: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[str], dict[str, Any]]:
    target = preparation_payload.get("target") or {}
    target_type = str(target.get("target_type") or "")
    content = _report_content(workspace_id, preparation_payload)
    findings: list[dict[str, Any]] = []
    incomplete: list[str] = []
    coverage: dict[str, Any] = {"claim_count": 0, "financial_claim_count": 0, "matched_financial_claims": 0}
    if target_type == "report_revision":
        try:
            from lvke_mcp.domains.reports.validation import validate_report

            validation = validate_report(
                workspace_id, str(target.get("target_id") or "")
            )
        except Exception:  # noqa: BLE001
            validation = {"valid": False, "blockers": ["report_validator_failed"]}
        for blocker in validation.get("blockers") or []:
            code = str(blocker.get("code") if isinstance(blocker, dict) else blocker)
            findings.append(rules.finding(
                f"REPORT.LEGACY.{code}", "P1", f"既有研报校验阻断：{code}", category="report_validation",
                actual=blocker, target_location={"report_revision_id": target.get("target_id")},
                standard_basis=standard_basis, review_area="report", remediation="按既有报告校验提示修订后生成新 revision",
            ))
    if not content:
        return findings, ["report_content_unreadable"], coverage
    paragraphs = [line.strip() for line in content.splitlines() if line.strip()]
    seen: dict[str, int] = {}
    for index, paragraph in enumerate(paragraphs, start=1):
        normalized = re.sub(r"\s+", " ", re.sub(r"^#+\s*", "", paragraph)).strip()
        if len(normalized) >= 20:
            fingerprint = sha256_json(normalized)
            if fingerprint in seen:
                findings.append(rules.finding(
                    "REPORT.DUPLICATE.PARAGRAPH", "P2", "正文存在重复段落", category="report_integrity",
                    actual=normalized[:300], target_location={"paragraph": index, "duplicate_of": seen[fingerprint], "text_anchor": normalized[:80]},
                    standard_basis=standard_basis, review_area="report", remediation="删除旧修订残留或重复正文",
                ))
            else:
                seen[fingerprint] = index
        if re.search(r"(?:TODO|TBD|待补充|待确认|XX+|【[^】]*(?:待|占位)[^】]*】|\{\{[^}]+}})", paragraph, re.I):
            findings.append(rules.finding(
                "REPORT.PLACEHOLDER", "P1", "正文仍含空白占位或待确认内容", category="report_integrity", actual=paragraph[:300],
                target_location={"paragraph": index, "text_anchor": paragraph[:80]}, standard_basis=standard_basis,
                review_area="report", remediation="补齐正式内容及依据后生成新修订",
            ))
    pack = preparation_payload.get("rule_pack") or {}
    overlays = {
        str(item.get("rule_pack_id") or "")
        for item in (pack.get("components") or [])
        if isinstance(item, dict)
    }
    deterministic, missing, deterministic_metrics, executed = report_checks.review_report(
        content=content,
        target_id=str(target.get("target_id") or ""),
        run=run,
        evidence_packs=_report_evidence_packs(workspace_id, preparation_payload),
        expected_sections=_expected_report_sections(preparation_payload),
        overlays=overlays,
        standard_basis=standard_basis,
        review_as_of=str(preparation_payload.get("review_as_of") or ""),
        evidence_track=str((preparation_payload.get("project_context") or {}).get("evidence_track") or "real"),
    )
    findings.extend(deterministic)
    incomplete.extend(missing)
    coverage = {
        **coverage,
        **deterministic_metrics,
        "executed_rules": sorted(executed),
    }
    return findings, sorted(set(incomplete)), coverage


def _report_artifact_text(
    workspace_id: str,
    artifact: dict[str, Any],
) -> str:
    try:
        artifact_id = str(artifact.get("artifact_id") or "")
        artifact_family = str(artifact.get("artifact_family") or "generic")
        chunks: list[str] = []
        for item in artifact.get("files") or []:
            name = str(item.get("name") or item.get("filename") or "")
            if Path(name).suffix.lower() not in {".md", ".txt", ".docx"}:
                continue
            if artifact_family == "asset_acquisition":
                from lvke_mcp.domains.asset_acquisition import backend as acquisition_service

                resolved = acquisition_service.read_artifact_candidate_download(
                    workspace_id,
                    artifact_id,
                    name,
                )
                if resolved.get("ok") is not True:
                    continue
                content = resolved.get("content") or b""
            else:
                from lvke_mcp.domains.reports import artifacts as deliverable_artifacts

                resolved = deliverable_artifacts.read_artifact_download(
                    workspace_id, artifact_id, name,
                )
                content = resolved.get("content") or b""
            if name.lower().endswith(".docx"):
                from docx import Document
                document = Document(io.BytesIO(content))
                chunks.extend(paragraph.text for paragraph in document.paragraphs)
                for table in document.tables:
                    for row in table.rows:
                        chunks.append(" | ".join(cell.text.strip() for cell in row.cells))
            else:
                chunks.append(content.decode("utf-8", errors="replace") if isinstance(content, bytes) else str(content))
        return "\n".join(chunks)
    except Exception:  # noqa: BLE001
        return ""


_EXTERNAL_GAP_REASON_MARKERS = (
    "formal_evidence",
    "evidence_pack",
    "source_",
    "standard_package_incomplete",
    "quality_review",
    "quality_note",
)


_LOCAL_IMPLEMENTATION_REASON_MARKERS = (
    "rule_not_executed",
    "review_engine_failed",
    "integrity",
    "lineage",
    "content_hash",
    "schema",
    "parser_unavailable",
    "target_reresolution_failed",
)


_EXTERNAL_GAP_CATEGORIES = {
    "acquisition_inputs",
    "citation_quality",
    "contract_consistency",
    "evidence",
    "evidence_conflict",
    "land_use_compliance",
    "market_evidence",
    "mineral_permits",
    "operating_assumption",
    "professional_review_pending",
    "project_metadata",
    "rights_and_area",
    "rights_and_licenses",
}


_LOCAL_IMPLEMENTATION_CATEGORIES = {
    "benchmark_applicability",
    "borrowing_cost",
    "combined_conclusion_consistency",
    "combined_numeric_consistency",
    "debt_service",
    "depreciation",
    "finance_consistency",
    "financial_consistency",
    "financial_recalculation",
    "financial_sustainability",
    "funding",
    "legacy_finance_issue",
    "report_finance_binding",
    "report_integrity",
    "report_internal_consistency",
    "report_validation",
    "sensitivity",
    "tax",
    "working_capital",
}


def _summarize_track_coverage(
    metrics: dict[str, Any],
    incomplete_reasons: list[str],
    findings: list[dict[str, Any]],
    *,
    evidence_track: str,
) -> dict[str, Any]:
    """Expose dual-track review counts without changing the gate verdict."""

    report_metrics = [
        row for row in metrics.get("report") or []
        if isinstance(row, dict)
    ]
    formal_count = sum(int(row.get("formal_evidence_claim_count") or 0) for row in report_metrics)
    technical_count = sum(int(row.get("technical_fixture_claim_count") or 0) for row in report_metrics)
    reconstructed_count = sum(int(row.get("source_reconstructed_claim_count") or 0) for row in report_metrics)
    external_reason_count = sum(
        1 for reason in incomplete_reasons
        if (
            evidence_track == "real"
            and any(marker in str(reason) for marker in _EXTERNAL_GAP_REASON_MARKERS)
        )
    )
    local_reason_count = sum(
        1 for reason in incomplete_reasons
        if any(marker in str(reason) for marker in _LOCAL_IMPLEMENTATION_REASON_MARKERS)
    )
    external_finding_count = sum(
        1 for finding in findings
        if str(finding.get("category") or "") in _EXTERNAL_GAP_CATEGORIES
    )
    local_finding_count = sum(
        1 for finding in findings
        if str(finding.get("category") or "") in _LOCAL_IMPLEMENTATION_CATEGORIES
    )
    return {
        "evidence_track": evidence_track,
        "formal_evidence_claim_count": formal_count,
        "technical_fixture_claim_count": technical_count,
        "source_reconstructed_claim_count": reconstructed_count,
        "external_data_gap_count": external_reason_count + external_finding_count,
        "local_implementation_issue_count": local_reason_count + local_finding_count,
    }
