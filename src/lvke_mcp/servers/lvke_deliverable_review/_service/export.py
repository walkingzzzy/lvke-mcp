"""审查导出：markdown/docx/xlsx 与导出完整性判定。"""

from __future__ import annotations

import hashlib
import io
import json
import os
import tempfile
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from urllib.parse import quote

from lvke_mcp.runtime.storage import canonical_json, require_safe_id, sha256_json, utc_now
from lvke_mcp.runtime.soffice import resolve_soffice_binary, run_soffice_convert
from lvke_mcp.runtime.formal_promotion import (
    FormalLineageError,
    SIM_A_FORMAL,
    validate_object_formal_lineage,
)
from lvke_mcp.runtime.workspace import deliverable_dir
from lvke_mcp.servers.lvke_deliverable_review.contracts import normalize_target
from lvke_mcp.servers.lvke_deliverable_review.store import STORE

from .base import (
    EXPORT_STORE,
    _REPORT_ARTIFACT_DOMAINS,
    _blocked,
    _message,
    _next_actions,
    _ok,
    _parse_timestamp,
    _write,
)

from .events import (
    _project,
)


def _export_root(workspace_id: str, export_id: str) -> Path:
    """审查导出（json/md/docx/xlsx）落盘根，统一到仓库 ``lvke产出/``。"""
    return (
        deliverable_dir(
            require_safe_id(workspace_id, "workspace_id"),
            "review",
            "exports",
        )
        / require_safe_id(export_id, "export_id")
    )


def _write_once_bytes(path: Path, content: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists():
        if path.read_bytes() != content:
            raise ValueError("immutable_export_conflict")
        return
    descriptor, temporary_name = tempfile.mkstemp(prefix=f".{path.name}.", suffix=".tmp", dir=str(path.parent))
    try:
        with os.fdopen(descriptor, "wb") as stream:
            stream.write(content)
            stream.flush()
            os.fsync(stream.fileno())
        os.replace(temporary_name, path)
    finally:
        Path(temporary_name).unlink(missing_ok=True)


def _review_markdown(state: dict[str, Any]) -> str:
    lines = [
        f"# 交付物审查报告 {state['review_id']}", "",
        f"- 目标：`{(state.get('target') or {}).get('target_type')}` / `{(state.get('target') or {}).get('target_id')}`",
        f"- 目标哈希：`{(state.get('target') or {}).get('target_sha256')}`",
        f"- 规则包：`{(state.get('rule_pack') or {}).get('rule_pack_id')}` / `{(state.get('rule_pack') or {}).get('version')}`",
        f"- 总体结论：`{state.get('overall_verdict')}`",
        f"- 校验状态：`{state.get('validation_status')}`",
        f"- 校验完成：`{str(bool(state.get('validation_complete'))).lower()}`", "",
        "## Findings", "",
    ]
    if not state.get("findings"):
        lines.append("无 findings。")
    for row in state.get("findings") or []:
        lines.extend([
            f"### {row.get('severity')} {row.get('rule_id')} ({row.get('status')})", "",
            str(row.get("message") or ""), "",
            f"- 定位：`{canonical_json(row.get('target_location') or {})}`",
            f"- 期望：`{canonical_json(row.get('expected'))}`",
            f"- 实际：`{canonical_json(row.get('actual'))}`",
            f"- 差额/容差：`{canonical_json(row.get('difference'))}` / `{canonical_json(row.get('tolerance'))}`",
            f"- 检查分类：`{row.get('review_area') or '-'}`",
            f"- 整改建议：{row.get('remediation') or '-'}", "",
        ])
    lines.extend(["## 不可核查项", ""])
    if state.get("incomplete_reasons"):
        lines.extend(f"- `{item}`" for item in state["incomplete_reasons"])
    else:
        lines.append("无。")
    return "\n".join(lines) + "\n"


def _review_docx(state: dict[str, Any]) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading(f"交付物审查报告 {state['review_id']}", level=1)
    document.add_paragraph(f"总体结论：{state.get('overall_verdict')}  审查状态：{state.get('review_status')}")
    document.add_paragraph(f"目标：{canonical_json(state.get('target') or {})}")
    document.add_heading("Findings", level=2)
    for row in state.get("findings") or []:
        document.add_heading(f"{row.get('severity')} {row.get('rule_id')} ({row.get('status')})", level=3)
        document.add_paragraph(str(row.get("message") or ""))
        document.add_paragraph(f"定位：{canonical_json(row.get('target_location') or {})}")
        document.add_paragraph(f"期望：{canonical_json(row.get('expected'))}")
        document.add_paragraph(f"实际：{canonical_json(row.get('actual'))}")
        document.add_paragraph(f"整改建议：{row.get('remediation') or '-'}")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _findings_xlsx(state: dict[str, Any]) -> bytes:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "findings"
    headers = [
        "finding_id", "rule_id", "rule_pack_version", "category", "severity", "blocking",
        "status", "confidence", "message", "expected", "actual", "difference", "tolerance",
        "target_location", "evidence", "standard_basis", "review_area", "remediation",
    ]
    sheet.append(headers)
    for row in state.get("findings") or []:
        sheet.append([
            row.get(key) if isinstance(row.get(key), (str, int, float, bool, type(None)))
            else canonical_json(row.get(key))
            for key in headers
        ])
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    dimensions = workbook.create_sheet("dimension_results")
    dimension_headers = [
        "dimension", "status", "deterministic_status", "semantic_status",
        "role_confirmed", "compliance_status", "finding_count", "limitations",
        "assessment_id", "confirmation_id",
    ]
    dimensions.append(dimension_headers)
    for row in state.get("dimension_results") or []:
        dimensions.append([
            row.get(key) if isinstance(row.get(key), (str, int, float, bool, type(None)))
            else canonical_json(row.get(key))
            for key in dimension_headers
        ])
    dimensions.freeze_panes = "A2"

    standards = workbook.create_sheet("standards_snapshot")
    standards.append([
        "package_id", "title", "gate_status", "source_manifest_sha256",
        "artifact_id", "publisher", "document_number", "publication_date",
        "official_page_url", "source_url", "sha256",
    ])
    for package in (state.get("standards") or {}).get("packages") or []:
        artifacts = list(package.get("artifacts") or []) or [{}]
        for artifact in artifacts:
            standards.append([
                package.get("package_id"), package.get("title"), package.get("gate_status"),
                package.get("source_manifest_sha256"), artifact.get("artifact_id"),
                artifact.get("publisher"), artifact.get("document_number"),
                artifact.get("publication_date"), artifact.get("official_page_url"),
                artifact.get("source_url"), artifact.get("sha256"),
            ])

    audit = workbook.create_sheet("audit_manifest")
    audit.append(["field", "value"])
    for key in (
        "review_id", "schema_version", "review_package_id", "review_dossier_id",
        "review_profile", "review_mode", "event_chain_hash", "suite_overall_verdict",
        "formal_suite_review_complete", "incomplete_reasons", "suite_hard_gate_blockers",
    ):
        value = state.get(key)
        audit.append([key, value if isinstance(value, (str, int, float, bool, type(None))) else canonical_json(value)])
    for column in sheet.columns:
        letter = column[0].column_letter
        sheet.column_dimensions[letter].width = min(60, max(12, max(len(str(cell.value or "")) for cell in column) + 2))
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _annotated_review_docx(state: dict[str, Any]) -> bytes:
    """Produce a locator-oriented issue copy without mutating the source report."""

    from docx import Document

    document = Document()
    document.add_heading("问题定位版研报审查", level=1)
    document.add_paragraph(
        "本文件按原始 locator 汇总批注，不修改外部原件，也不代表法律或执业签署。"
    )
    for row in state.get("findings") or []:
        document.add_heading(
            f"{row.get('severity')} {row.get('rule_id')}", level=2,
        )
        document.add_paragraph(f"定位：{canonical_json(row.get('target_location') or {})}")
        document.add_paragraph(str(row.get("message") or ""))
        document.add_paragraph(f"整改建议：{row.get('remediation') or '-'}")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _docx_to_pdf(content: bytes) -> bytes:
    binary = resolve_soffice_binary("LVKE_REVIEW_SOFFICE", "SOFFICE", "LIBREOFFICE")
    if not binary:
        raise FileNotFoundError("review_pdf_worker_unavailable")
    with tempfile.TemporaryDirectory(prefix="lvke-review-pdf-") as directory:
        root = Path(directory)
        source = root / "review.docx"
        source.write_bytes(content)
        run_soffice_convert(
            source=source,
            convert_to="pdf",
            outdir=root,
            binary=binary,
            timeout=180,
        )
        output = root / "review.pdf"
        if not output.is_file():
            raise OSError("review_pdf_not_generated")
        return output.read_bytes()


def _export_file_uri(workspace_id: str, export_id: str, filename: str) -> str:
    return f"lvke://deliverable-review/workspaces/{workspace_id}/exports/{export_id}/files/{quote(filename)}"


def _export_resource_uri(workspace_id: str, export_id: str) -> str:
    return f"lvke://deliverable-review/workspaces/{workspace_id}/exports/{export_id}"


def _export_review_locked(
    workspace_id: str,
    review_id: str,
    requested_formats: Any,
) -> dict[str, Any]:
    """Export one validation snapshot while holding its mutation lock."""

    try:
        state = _project(workspace_id, review_id)
    except ValueError:
        return _blocked("review_not_found", _message("review_not_found"))
    quality_issues: list[str] = []
    evidence_metadata = (
        state.get("evidence_metadata")
        if isinstance(state.get("evidence_metadata"), dict)
        else {}
    )
    formal_lineage: dict[str, Any] = {}
    formal_review = (
        str((state.get("project_context") or {}).get("evidence_track") or "")
        == SIM_A_FORMAL
        or str(evidence_metadata.get("evidence_policy") or "") == SIM_A_FORMAL
    )
    if formal_review:
        try:
            formal_lineage = validate_object_formal_lineage(
                workspace_id,
                evidence_metadata,
            )
        except FormalLineageError as exc:
            return _blocked(
                exc.code,
                f"审查导出前正式 promotion 谱系无效：{exc.message}",
                review_id=review_id,
            )
        if any(
            evidence_metadata.get(key) != value
            for key, value in formal_lineage.items()
        ):
            return _blocked(
                "formal_lineage_metadata_mismatch",
                "审查导出前 promotion 元数据不是规范值",
                review_id=review_id,
            )
    if not state.get("validation_complete"):
        quality_issues.append("review_validation_incomplete")
    for previous in state.get("exports") or []:
        integrity_reasons = _export_integrity_reasons(
            workspace_id,
            {
                "review_id": review_id,
                "export_id": previous.get("export_id"),
                "export_record_id": previous.get("export_record_id"),
                "export_record_hash": previous.get("export_record_hash"),
                "export_basis_hash": previous.get("export_basis_hash"),
                "export_files": previous.get("files"),
            },
        )
        if integrity_reasons:
            return _blocked(
                "review_export_integrity_failed",
                "既有审查导出记录或文件完整性校验失败",
                review_id=review_id,
                integrity_reasons=integrity_reasons,
            )
    suite_review = str((state.get("target") or {}).get("target_type") or "") == "review_package"
    if suite_review:
        from .suite_review import (
            DIMENSION_CONFIRMATION_STORE,
            DOSSIER_STORE,
            REVIEW_PACKAGE_STORE,
            SUITE_ASSESSMENT_STORE,
            _verified_record,
            package_integrity_reasons,
        )

        package_id = str((state.get("target") or {}).get("target_id") or "")
        package_record = _verified_record(REVIEW_PACKAGE_STORE, workspace_id, package_id)
        if package_record is None:
            return _blocked("review_package_not_found", "导出前 ReviewPackage 不存在或完整性无效")
        package_reasons = package_integrity_reasons(workspace_id, package_record)
        if package_reasons:
            return _blocked(
                package_reasons[0],
                "导出前 ReviewPackage 完整性或正式谱系校验失败",
                integrity_reasons=package_reasons,
            )
        state = deepcopy(state)
        state["review_package"] = deepcopy(package_record)
        state["review_assessments"] = [
            deepcopy(row)
            for row in SUITE_ASSESSMENT_STORE.list(workspace_id)
            if str((row.get("payload") or {}).get("review_id") or "") == review_id
        ]
        state["review_dimension_confirmations"] = [
            deepcopy(row)
            for row in DIMENSION_CONFIRMATION_STORE.list(workspace_id)
            if str((row.get("payload") or {}).get("review_id") or "") == review_id
        ]
        state["review_dossiers"] = [
            deepcopy(row)
            for row in DOSSIER_STORE.list(workspace_id)
            if str((row.get("payload") or {}).get("review_id") or "") == review_id
        ]
    requested = list(
        requested_formats or (["json", "markdown", "docx", "xlsx", "annotated_docx"] if suite_review else ["json", "markdown", "docx", "xlsx"])
    )
    allowed = {"json", "markdown", "docx", "xlsx", "pdf", "annotated_docx"}
    if not requested or any(item not in allowed for item in requested):
        return _blocked("export_format_invalid", "导出格式仅支持 json、markdown、docx、xlsx、pdf、annotated_docx")
    formal_formats = {"docx", "xlsx", "pdf", "annotated_docx"}
    if any(item in formal_formats for item in requested) and not state.get("validation_complete"):
        return _blocked(
            "FORMAL_ARTIFACT_QUALIFICATION_REQUIRED",
            "DOCX/XLSX 正式导出要求审查验证完成；当前仅允许 JSON/Markdown 过程记录",
            review_id=review_id,
            quality_issues=["review_validation_incomplete"],
        )
    # A completed review can still be incomplete or explicitly blocked.  Do
    # not let a successful technical run, controlled-assumption input, or
    # unresolved finding masquerade as a formal review export.
    formal_issues: list[str] = []
    if suite_review:
        if str(state.get("suite_overall_verdict") or "") != "pass":
            formal_issues.append(f"suite_overall_verdict:{state.get('suite_overall_verdict') or 'incomplete'}")
        if not state.get("formal_suite_review_complete"):
            formal_issues.append("formal_suite_review_incomplete")
    else:
        formal_issues.extend(str(item) for item in state.get("quality_issues") or [])
        formal_issues.extend(str(item) for item in state.get("release_limitations") or [])
        for verdict_key in ("overall_verdict", "technical_verdict", "release_verdict"):
            verdict = str(state.get(verdict_key) or "").lower()
            if verdict and verdict not in {"pass", "passed", "ok", "complete", "completed"}:
                formal_issues.append(f"{verdict_key}:{verdict}")
    if formal_issues and any(item in formal_formats for item in requested):
        return _blocked(
            "FORMAL_ARTIFACT_QUALIFICATION_REQUIRED",
            "审查仍存在未关闭的质量或发布资格问题，禁止正式 DOCX/XLSX 导出",
            review_id=review_id,
            quality_issues=sorted(set(formal_issues)),
        )
    review_purpose = str((state.get("project_context") or {}).get("review_purpose") or "")
    if any(item in formal_formats for item in requested) and review_purpose != "project_delivery" and not suite_review:
        return _blocked(
            "FORMAL_ARTIFACT_QUALIFICATION_REQUIRED",
            "正式 DOCX/XLSX 审查导出要求 project_delivery 审查目的",
            review_id=review_id,
            quality_issues=["review_release_scope_not_formal"],
        )
    export_basis = {
        "review_id": review_id,
        "event_chain_hash": state.get("event_chain_hash"),
        "formats": sorted(set(requested)),
        "review_status": state.get("review_status"),
        "overall_verdict": state.get("overall_verdict"),
        **deepcopy(formal_lineage),
    }
    export_id = (
        "rvexp_" + sha256_json(export_basis).removeprefix("sha256:")[:24]
    )
    output = _export_root(workspace_id, export_id)
    review_docx = _review_docx(state)
    payloads: dict[str, tuple[str, bytes, str]] = {
        "json": (
            "review.json",
            json.dumps(
                state, ensure_ascii=False, indent=2, default=str,
            ).encode("utf-8"),
            "application/json",
        ),
        "markdown": (
            "review.md",
            _review_markdown(state).encode("utf-8"),
            "text/markdown; charset=utf-8",
        ),
        "docx": (
            "review.docx",
            review_docx,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        "xlsx": (
            "findings.xlsx",
            _findings_xlsx(state),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ),
        "annotated_docx": (
            "annotated-review.docx",
            _annotated_review_docx(state),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
    }
    if "pdf" in requested:
        try:
            payloads["pdf"] = (
                "review.pdf",
                _docx_to_pdf(review_docx),
                "application/pdf",
            )
        except (FileNotFoundError, OSError, TimeoutError):
            return _blocked(
                "review_pdf_export_unavailable",
                "PDF 审查报告要求可用的隔离 LibreOffice worker",
                review_id=review_id,
            )
    files = []
    for format_name in sorted(set(requested)):
        filename, content, media_type = payloads[format_name]
        _write_once_bytes(output / filename, content)
        files.append({
            "format": format_name,
            "filename": filename,
            "media_type": media_type,
            "bytes": len(content),
            "sha256": "sha256:" + hashlib.sha256(content).hexdigest(),
            "uri": _export_file_uri(workspace_id, export_id, filename),
        })
    record = EXPORT_STORE.put(
        workspace_id,
        {
            **export_basis,
            "export_id": export_id,
            "files": files,
        },
        producer="lvke-deliverable-review.review_export",
        source_ids=[review_id],
        basis=export_basis,
        schema_version="deliverable_review_export.v1",
    )
    export_envelope = {
        "review_id": review_id,
        "export_id": export_id,
        "export_record_id": record.get("object_id"),
        "export_record_hash": record.get("content_hash"),
        "export_basis_hash": record.get("basis_hash"),
        "export_files": files,
        **deepcopy(formal_lineage),
    }
    integrity_reasons = _export_integrity_reasons(
        workspace_id,
        export_envelope,
    )
    if integrity_reasons:
        return _blocked(
            "review_export_integrity_failed",
            "新生成的审查导出记录或文件完整性校验失败",
            review_id=review_id,
            integrity_reasons=integrity_reasons,
        )
    STORE.append(
        workspace_id,
        review_id,
        "review_exported",
        {
            "export_id": export_id,
            "export_record_id": record["object_id"],
            "export_record_hash": record["content_hash"],
            "export_basis_hash": record["basis_hash"],
            "files": files,
            **deepcopy(formal_lineage),
            "exported_at": utc_now(),
        },
    )
    return _ok(
        review_id=review_id,
        export_id=export_id,
        export_record_id=record["object_id"],
        export_record_hash=record["content_hash"],
        export_basis_hash=record["basis_hash"],
        files=files,
        resource_uris=[_export_resource_uri(workspace_id, export_id), *[row["uri"] for row in files]],
        quality_issues=quality_issues,
        warnings=[f"质量提示：{item}" for item in quality_issues],
        blockers=[],
        next_actions=_next_actions(
            _project(workspace_id, review_id, check_freshness=False)
        ),
    )


def export_review(args: dict[str, Any]) -> dict[str, Any]:
    def execute(workspace_id: str) -> dict[str, Any]:
        return _export_review_locked(
            workspace_id,
            str(args.get("review_id") or ""),
            args.get("formats"),
        )

    return _write("review_export", args, execute)


def _release_export_integrity_reasons(
    workspace_id: str,
    release_payload: dict[str, Any],
) -> list[str]:
    # Records produced before export files were embedded remain readable; all
    # current release records include the field and are verified strictly.
    if "export_files" not in release_payload:
        return []
    export_id = str(release_payload.get("export_id") or "")
    rows = release_payload.get("export_files")
    if not export_id or not isinstance(rows, list) or not rows:
        return ["release_export_manifest_invalid"]
    try:
        root = _export_root(workspace_id, export_id)
    except ValueError:
        return ["release_export_id_invalid"]
    reasons: list[str] = []
    seen: set[str] = set()
    for row in rows:
        if not isinstance(row, dict):
            reasons.append("release_export_entry_invalid")
            continue
        filename = str(row.get("filename") or "")
        if (
            not filename
            or filename in seen
            or Path(filename).name != filename
            or filename in {".", ".."}
        ):
            reasons.append("release_export_filename_invalid")
            continue
        seen.add(filename)
        path = root / filename
        if not path.is_file():
            reasons.append(f"release_export_missing:{filename}")
            continue
        try:
            content = path.read_bytes()
            expected_bytes = int(row.get("bytes") or -1)
        except (OSError, TypeError, ValueError):
            reasons.append(f"release_export_unreadable:{filename}")
            continue
        if expected_bytes != len(content):
            reasons.append(f"release_export_size_mismatch:{filename}")
        digest = "sha256:" + hashlib.sha256(content).hexdigest()
        if str(row.get("sha256") or "") != digest:
            reasons.append(f"release_export_hash_mismatch:{filename}")
    return sorted(set(reasons))


def _export_record_integrity_reasons(
    workspace_id: str,
    export_envelope: dict[str, Any],
) -> list[str]:
    export_record_id = str(export_envelope.get("export_record_id") or "")
    if not export_record_id:
        return ["release_export_record_id_missing"]
    try:
        record = EXPORT_STORE.get(workspace_id, export_record_id)
    except ValueError:
        record = None
    if record is None:
        return ["release_export_record_missing"]
    payload = record.get("payload")
    if not isinstance(payload, dict):
        return ["release_export_record_payload_invalid"]
    export_basis = {
        key: deepcopy(payload.get(key))
        for key in (
            "review_id",
            "event_chain_hash",
            "formats",
            "review_status",
            "overall_verdict",
            "evidence_policy",
            "evidence_origin",
            "project_fact_certified",
            "formal_promotion",
        )
        if key in payload
    }
    content_hash = sha256_json(payload)
    expected_record_id = (
        f"{EXPORT_STORE.id_prefix}_"
        f"{content_hash.removeprefix('sha256:')[:24]}"
    )
    expected_uri = EXPORT_STORE.uri(workspace_id, expected_record_id)
    declared_record_hash = str(
        export_envelope.get("export_record_hash") or ""
    )
    declared_basis_hash = str(
        export_envelope.get("export_basis_hash") or ""
    )
    reasons: list[str] = []
    if record.get("content_hash") != content_hash:
        reasons.append("release_export_record_content_hash_mismatch")
    if record.get("basis_hash") != sha256_json(export_basis):
        reasons.append("release_export_record_basis_hash_mismatch")
    if record.get("object_id") != expected_record_id:
        reasons.append("release_export_record_object_id_mismatch")
    if export_record_id != expected_record_id:
        reasons.append("release_export_record_reference_mismatch")
    if record.get("workspace_id") != workspace_id:
        reasons.append("release_export_record_workspace_mismatch")
    if record.get("resource_uri") != expected_uri:
        reasons.append("release_export_record_uri_mismatch")
    if record.get("producer") != "lvke-deliverable-review.review_export":
        reasons.append("release_export_record_producer_mismatch")
    if record.get("schema_version") != "deliverable_review_export.v1":
        reasons.append("release_export_record_schema_mismatch")
    if declared_record_hash and declared_record_hash != content_hash:
        reasons.append("release_export_record_declared_hash_mismatch")
    if declared_basis_hash and declared_basis_hash != sha256_json(export_basis):
        reasons.append("release_export_record_declared_basis_mismatch")
    if str(payload.get("review_id") or "") != str(
        export_envelope.get("review_id") or ""
    ):
        reasons.append("release_export_record_review_mismatch")
    if str(payload.get("export_id") or "") != str(
        export_envelope.get("export_id") or ""
    ):
        reasons.append("release_export_record_export_id_mismatch")
    if payload.get("files") != export_envelope.get("export_files"):
        reasons.append("release_export_record_files_mismatch")
    if str(export_envelope.get("review_id") or "") not in {
        str(item) for item in record.get("source_ids") or []
    }:
        reasons.append("release_export_record_source_mismatch")
    return sorted(set(reasons))


def _export_integrity_reasons(
    workspace_id: str,
    export_envelope: dict[str, Any],
) -> list[str]:
    return sorted(set([
        *_export_record_integrity_reasons(workspace_id, export_envelope),
        *_release_export_integrity_reasons(workspace_id, export_envelope),
    ]))


def latest_review_for_target(
    workspace_id: str,
    target_type: str,
    target_id: str,
    *,
    artifact_domain: str = "",
) -> dict[str, Any] | None:
    """Return the latest review state after applying freshness projection."""

    try:
        workspace_id = require_safe_id(workspace_id, "workspace_id")
        normalized = normalize_target({
            "target_type": target_type,
            "target_id": target_id,
        })
    except ValueError:
        return None
    if normalized["target_type"] == "report_artifact":
        if artifact_domain not in _REPORT_ARTIFACT_DOMAINS:
            return None
    elif artifact_domain:
        return None

    candidates: list[dict[str, Any]] = []
    for review_id in STORE.review_ids(workspace_id):
        try:
            state = _project(workspace_id, review_id)
        except (OSError, ValueError):
            continue
        target = state.get("target") or {}
        if (
            target.get("target_type") != normalized["target_type"]
            or str(target.get("target_id") or "") != normalized["target_id"]
        ):
            continue
        if normalized["target_type"] == "report_artifact":
            target_spec = state.get("target_spec") or {}
            if target_spec.get("artifact_domain") != artifact_domain:
                continue
        candidates.append(state)
    if not candidates:
        return None
    candidates.sort(
        key=lambda row: (
            _parse_timestamp(row.get("created_at"))
            or datetime.min.replace(tzinfo=timezone.utc),
            str(row.get("review_id") or ""),
        ),
        reverse=True,
    )
    return deepcopy(candidates[0])
