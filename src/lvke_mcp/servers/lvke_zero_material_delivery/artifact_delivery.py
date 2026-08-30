"""Deterministic preview report and immutable delivery manifests."""

from __future__ import annotations

import hashlib
import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lvke_mcp.adapters.zero_material_repository import (
    artifact_root as _artifact_root,
    resolve_report_file,
)


def _format_number(value: Any, suffix: str = "") -> str:
    if not isinstance(value, (int, float)) or isinstance(value, bool):
        return "未形成"
    return f"{value:,.2f}{suffix}"


def _first_value(mapping: dict[str, Any], *keys: str) -> Any:
    for key in keys:
        value = mapping.get(key)
        if value is not None:
            return value
    return None


def _finance_summary(
    workspace_id: str,
    finance_run_id: str,
) -> dict[str, Any]:
    from lvke_mcp.domains.finance.run_service import get_workspace_finance_run

    run = get_workspace_finance_run(
        workspace_id,
        run_id=finance_run_id,
        view="full",
    )
    indicators = dict(run.get("indicators") or {})
    investment = dict(run.get("investment") or {})
    total_investment_wan = _first_value(indicators, "total_investment_wan")
    if total_investment_wan is None:
        total_investment_wan = _first_value(investment, "total")
    return {
        "run_id": finance_run_id,
        "available": bool(run.get("available")),
        "assurance_level": str(run.get("assurance_level") or "estimate_preview"),
        "model_version": str(run.get("model_version") or ""),
        "template_version": str(run.get("template_version") or ""),
        "spec_hash": str(run.get("spec_hash") or ""),
        "input_hash": str(run.get("input_hash") or ""),
        "consistency_ok": bool(run.get("consistency_ok")),
        "total_investment_wan": total_investment_wan,
        "annual_revenue_wan": _first_value(
            indicators, "annual_revenue_wan", "revenue"
        ),
        "project_irr": _first_value(indicators, "project_irr", "project_irr_pct"),
        "project_npv": _first_value(indicators, "project_npv", "npv_wan"),
        "capital_irr": _first_value(indicators, "capital_irr", "capital_irr_pct"),
        "payback_years": _first_value(
            indicators,
            "payback_years",
            "dynamic_payback_years",
            "static_payback_years",
        ),
    }


def _resolve_report_profile(
    intent: dict[str, Any],
    domain: dict[str, Any],
) -> tuple[dict[str, Any], dict[str, Any], str]:
    """Reuse the profile frozen at intake, else re-resolve from the same selector.

    读取优先级刻意是「冻结快照 → 按 hash 校验的磁盘文件 → 按路由重解析」：

    1. 冻结快照（``selection.profile_snapshot``）随运行不可变留存，因此配置文件
       被升级、删除，或部署根目录变化后，旧运行仍能按**原**配置重放。这是"历史
       运行冻结、可重放"的落点。
    2. 快照缺失时（升级前建的运行）回落到磁盘，并强制 hash 相等——配置中途被改过
       就阻断，不用新配置续算同一个 run。
    3. 两者都没有时才按路由重解析（v1 老记录）。
    """

    from lvke_mcp.servers.lvke_zero_material_delivery._service.report_profiles import (
        ReportProfileError,
        load_profile_document,
        resolve_profile,
        verified_snapshot,
    )

    frozen = dict(intent.get("report_profile") or {})
    route = dict(domain.get("route") or {})
    if frozen.get("template_set_id") and frozen.get("profile_id"):
        # 快照采信统一走 verified_snapshot：它从内容复算 hash，而不是比对两个
        # 字面量（把章节改成 TAMPERED 同时保留原 hash，字面量仍然相等）。
        if isinstance(frozen.get("profile_snapshot"), dict):
            verified = verified_snapshot(frozen)
            if verified is None:
                return {}, {}, "report_profile_snapshot_hash_mismatch"
            return verified, frozen, ""
        try:
            document = load_profile_document(f"{frozen['profile_id']}.v1.json")
        except ReportProfileError as exc:
            return {}, {}, exc.code
        if str(document.get("content_hash") or "") != str(frozen.get("profile_content_hash") or ""):
            return {}, {}, "report_profile_hash_drifted"
        return document, frozen, ""
    try:
        resolved = resolve_profile(
            industry_code=str(
                route.get("industry_code")
                or dict(intent.get("industry") or {}).get("industry_code")
                or ""
            ),
            project_type=(
                "asset_acquisition"
                if str(route.get("finance_kind") or "") == "asset_acquisition"
                else "generic_feasibility"
            ),
            transaction_structure=str(route.get("transaction_structure") or "") or "new_build",
            asset_type=str(route.get("asset_type") or "general"),
            report_type=str(intent.get("report_type") or ""),
        )
    except ReportProfileError as exc:
        return {}, {}, exc.code
    return dict(resolved["profile"]), dict(resolved["selection"]), ""


def _docx_bytes(markdown: str, title: str) -> bytes:
    from lvke_mcp.domains.reports._doc_service.docx import append_markdown_pipe_table

    document = Document()
    heading = document.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("技术预估版，非正式发布").alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        text = lines[index].strip()
        if not text or text.startswith("# "):
            index += 1
            continue
        # 配置化章节树是两级的，必须先判更长的前缀：先判 "## " 会把 "### x"
        # 当成标题文本 "# x"，子节标题于是全部带一个多余的井号。
        if text.startswith("### "):
            document.add_heading(text[4:], level=2)
        elif text.startswith("## "):
            document.add_heading(text[3:], level=1)
        elif text.startswith("> "):
            paragraph = document.add_paragraph(text[2:])
            paragraph.style = document.styles["Intense Quote"]
        elif text.startswith("- "):
            document.add_paragraph(text[2:], style="List Bullet")
        elif text.startswith("|"):
            table_rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                row = lines[index].strip()
                if not row.replace("|", "").replace("-", "").replace(":", "").strip():
                    index += 1
                    continue
                table_rows.append([cell.strip() for cell in row.strip("|").split("|")])
                index += 1
            append_markdown_pipe_table(document, table_rows)
            continue
        else:
            document.add_paragraph(text)
        index += 1
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def build_delivery_artifacts(
    workspace_id: str,
    intent: dict[str, Any],
    assumption_package: dict[str, Any],
    source_run: dict[str, Any],
    domain: dict[str, Any],
    *,
    stores: dict[str, Any],
    service_version: str,
    delivery_run_id: str = "",
) -> dict[str, Any]:
    refs = {
        str(key): str(value)
        for key, value in dict(domain.get("object_refs") or {}).items()
        if value
    }
    finance_run_id = refs.get("finance_run_id", "")
    if not finance_run_id:
        return {
            "resource_uris": [],
            "object_refs": {},
            "blockers": ["technical_report_finance_run_required"],
        }
    profile, profile_selection, profile_error = _resolve_report_profile(intent, domain)
    if profile_error:
        return {
            "resource_uris": [],
            "object_refs": {},
            "blockers": [profile_error],
        }
    finance = _finance_summary(
        workspace_id,
        finance_run_id,
    )
    blockers = list(domain.get("blockers") or [])
    quality_issues = list(domain.get("quality_issues") or [])
    public_research = dict((domain.get("research") or {}).get("public_research") or {})
    skipped_fields = [
        dict(item)
        for item in source_run.get("skipped_fields") or []
        if isinstance(item, dict)
    ]
    from lvke_mcp.servers.lvke_zero_material_delivery._service.report_render import (
        build_slot_values,
        render_report_markdown,
    )

    slots = build_slot_values(
        intent=intent,
        assumption_package=assumption_package,
        finance=finance,
        blockers=blockers,
        quality_issues=quality_issues,
        public_research=public_research,
        skipped_fields=skipped_fields,
        report_profile=profile,
    )
    markdown, unresolved_slots = render_report_markdown(
        profile=profile,
        selection=profile_selection,
        slots=slots,
    )
    report_payload = {
        "object_type": "TechnicalReport",
        "title": f"{intent.get('project_name')}{str(profile.get('label') or '')}",
        "format_version": "zero-material-technical-report.v2",
        "assurance_level": "estimate_preview",
        "content_markdown": markdown,
        "report_profile": profile_selection,
        "unresolved_slots": unresolved_slots,
        "finance_summary": finance,
        "intent_id": intent.get("delivery_intent_id"),
        "assumption_package_id": assumption_package.get("assumption_package_id"),
        "finance_run_id": finance_run_id,
        "finance_tables_package_id": refs.get("finance_tables_package_id", ""),
        "research_task_id": refs.get("research_task_id", ""),
        "research_package_id": refs.get("research_package_id", ""),
        "evidence_pack_id": refs.get("evidence_pack_id", ""),
        "public_research": public_research,
        "validation_complete": False,
        "input_evidence_complete": False,
    }
    report_record = stores["report"].put(
        workspace_id,
        report_payload,
        producer="lvke-zero-material-delivery.technical_report",
        status="partial",
        source_ids=[
            *[value for value in refs.values() if value],
            *[str(item) for item in public_research.get("source_snapshot_ids") or [] if str(item)],
        ],
        basis={
            "intent_id": intent.get("delivery_intent_id"),
            "assumption_package_id": assumption_package.get("assumption_package_id"),
            "finance_run_id": finance_run_id,
            "evidence_pack_id": refs.get("evidence_pack_id", ""),
            "research_package_id": refs.get("research_package_id", ""),
            "source_snapshot_ids": list(public_research.get("source_snapshot_ids") or []),
        },
    )
    report_base_uri = report_record["resource_uri"]
    root = _artifact_root(workspace_id) / report_record["object_id"]
    root.mkdir(parents=True, exist_ok=True)
    markdown_bytes = markdown.encode("utf-8")
    docx_bytes = _docx_bytes(markdown, str(report_payload["title"]))
    files = {
        "report.md": markdown_bytes,
        "report.docx": docx_bytes,
    }
    file_manifest: list[dict[str, Any]] = []
    for name, content in files.items():
        path = root / name
        path.write_bytes(content)
        file_manifest.append(
            {
                "name": name,
                "content_hash": "sha256:" + hashlib.sha256(content).hexdigest(),
                "size": len(content),
                "resource_uri": f"{report_base_uri}/files/{name}",
            }
        )

    assumption_record = stores["assumption_register"].put(
        workspace_id,
        {
            "object_type": "AssumptionRegister",
            "assumption_package_id": assumption_package.get("assumption_package_id"),
            "revision": assumption_package.get("revision"),
            "fields": assumption_package.get("fields") or [],
            "validation_complete": False,
        },
        producer="lvke-zero-material-delivery.assumption_register",
        status="ok",
        source_ids=[str(assumption_package.get("assumption_package_id") or "")],
    )
    gap_record = stores["gap_register"].put(
        workspace_id,
        {
            "object_type": "GapRegister",
            "blockers": blockers,
            "missing_client_materials": ["合同", "测绘", "报价", "权属", "设计", "批复"],
            "research_status": str((domain.get("research") or {}).get("status") or ""),
            "validation_complete": False,
            "input_evidence_complete": False,
        },
        producer="lvke-zero-material-delivery.gap_register",
        status="partial",
        source_ids=[str(source_run.get("delivery_run_id") or "")],
    )
    evidence_record = stores["evidence_manifest"].put(
        workspace_id,
        {
            "object_type": "EvidenceManifest",
            "research_task_id": refs.get("research_task_id", ""),
            "evidence_pack_ids": [refs["evidence_pack_id"]] if refs.get("evidence_pack_id") else [],
            "research_package_ids": [refs["research_package_id"]] if refs.get("research_package_id") else [],
            "source_policy": "public_sources_only",
            "evidence_status": str(public_research.get("status") or "pending"),
            "source_snapshot_ids": list(public_research.get("source_snapshot_ids") or []),
            "discovery_set_id": str(public_research.get("discovery_set_id") or ""),
            "fallback_used": bool(public_research.get("fallback_used")),
            "controlled_assumptions_are_evidence": False,
            "formal_evidence_ready": False,
        },
        producer="lvke-zero-material-delivery.evidence_manifest",
        status="partial",
        source_ids=[
            value for value in [
                refs.get("research_task_id", ""),
                refs.get("evidence_pack_id", ""),
                refs.get("research_package_id", ""),
                *[str(item) for item in public_research.get("source_snapshot_ids") or []],
            ] if value
        ],
    )
    manifest_payload = {
        "object_type": "RunManifest",
        "schema_version": "zero-material-run-manifest.v2",
        "report_profile": profile_selection,
        "unresolved_slots": unresolved_slots,
        "skipped_fields": skipped_fields,
        "workspace_id": workspace_id,
        "delivery_run_id": delivery_run_id or source_run.get("delivery_run_id"),
        "intent_id": intent.get("delivery_intent_id"),
        "assumption_package_id": assumption_package.get("assumption_package_id"),
        "object_refs": {
            **refs,
            "technical_report_id": report_record["object_id"],
            "assumption_register_id": assumption_record["object_id"],
            "gap_register_id": gap_record["object_id"],
            "evidence_manifest_id": evidence_record["object_id"],
        },
        "finance_lineage": finance,
        "files": file_manifest,
        "artifact_uris": list(domain.get("resource_uris") or []),
        "service_version": service_version,
        "status": "estimate_preview",
        "blockers": blockers,
        "public_research": public_research,
        "validation_complete": False,
        "input_evidence_complete": False,
    }
    manifest_record = stores["manifest"].put(
        workspace_id,
        manifest_payload,
        producer="lvke-zero-material-delivery.run_manifest",
        status="partial",
        source_ids=[value for value in manifest_payload["object_refs"].values() if value],
        basis=manifest_payload,
    )
    resource_uris = [
        report_record["resource_uri"],
        *[item["resource_uri"] for item in file_manifest],
        assumption_record["resource_uri"],
        gap_record["resource_uri"],
        evidence_record["resource_uri"],
        manifest_record["resource_uri"],
    ]
    produced_components = {
        "report_markdown": any(item["name"] == "report.md" for item in file_manifest),
        "report_docx": any(item["name"] == "report.docx" for item in file_manifest),
        "finance_thirteen_tables": bool(refs.get("finance_tables_package_id")),
        "finance_xlsx": bool((domain.get("xlsx_export") or {}).get("xlsx_resource")),
        "finance_csv": bool((domain.get("csv_export") or {}).get("csv_resource_uris")),
        "evidence_register": bool(evidence_record["object_id"]),
        "assumption_register": bool(assumption_record["object_id"]),
    }
    required = [str(item) for item in profile.get("required_components") or []]
    component_status = {name: bool(produced_components.get(name)) for name in required}
    return {
        "resource_uris": resource_uris,
        "object_refs": {
            "technical_report_id": report_record["object_id"],
            "assumption_register_id": assumption_record["object_id"],
            "gap_register_id": gap_record["object_id"],
            "evidence_manifest_id": evidence_record["object_id"],
            "run_manifest_id": manifest_record["object_id"],
        },
        "manifest_uri": manifest_record["resource_uri"],
        "file_manifest": file_manifest,
        # 技术验收的确定性输入。在这里算而不是在 acceptance 里算：只有本函数
        # 真正知道哪些组件落盘成功、哪些配置槽位没解析到。
        "report_profile": profile_selection,
        "component_status": component_status,
        "unresolved_slots": unresolved_slots,
        # 技术验收要读 consistency_ok。它只在这份从不可变 FinanceRun 读出的摘要里，
        # run_model 的响应信封没有该字段。
        "finance_summary": finance,
        "blockers": [],
    }


__all__ = ["build_delivery_artifacts", "resolve_report_file"]
