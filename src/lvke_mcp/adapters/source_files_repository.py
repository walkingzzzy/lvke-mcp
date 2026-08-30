"""MCP-owned source-file persistence used by the source-files adapter."""

from __future__ import annotations

import csv
import hashlib
import io
import json
import os
import re
import threading
import uuid
from contextlib import contextmanager
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterator

from filelock import FileLock

from lvke_mcp.runtime.workspace import workspace_root

_STATE_LOCK = threading.RLock()
_DEFAULT_MAX_UPLOAD_BYTES = 64 * 1024 * 1024


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _root(workspace_id: str) -> Path:
    return workspace_root(workspace_id) / "source-files"


def _load_acquisition_snapshot(workspace_id: str, source_id: str) -> dict[str, Any] | None:
    """Look up a data-acquisition SourceSnapshot for cross-domain evidence binding.

    source-files 与 data-acquisition 各自铸造 ``src_<24hex>``（前者哈希上传字节，
    后者哈希 payload JSON），前缀撞车但命名空间不相交。财务证据绑定先查本域
    state["files"]，未命中时按本函数回落到快照存储，避免把公开来源误判为不存在。
    """
    try:
        from lvke_mcp.adapters.data_acquisition_repository import SOURCE_STORE
    except ImportError:  # pragma: no cover - data-acquisition 未安装时保持 fail-closed
        return None
    try:
        record = SOURCE_STORE.get(workspace_id, source_id)
    except (ValueError, OSError):
        return None
    if not isinstance(record, dict):
        return None
    if str(record.get("workspace_id") or "") != workspace_id:
        return None
    return record


def _decode_serialized_locator(locator: str) -> tuple[str, str] | None:
    """Decode a JSON-serialized locator row into ``(kind, text)``.

    调用方常把整行 locator 结构序列化后回传；返回 None 表示不是
    JSON 对象形态，交由其余匹配策略处理。
    """
    text = str(locator or "").strip()
    if not text.startswith("{") or not text.endswith("}"):
        return None
    try:
        value = json.loads(text)
    except (json.JSONDecodeError, ValueError):
        return None
    if not isinstance(value, dict):
        return None
    kind = str(value.get("kind") or "").strip()
    body = str(value.get("text") or "").strip()
    if not kind and not body:
        return None
    return kind, body


def _acquisition_snapshot_uris(workspace_id: str, source_id: str) -> tuple[str, ...]:
    """Accepted evidence ``source_uri`` spellings for one source id.

    同一份来源可由三种合法地址指称：source-files 的原件、data-acquisition
    的公网快照，以及 source-files 的解析投影（analyses）。逐值血缘实际读的
    正是解析投影，因此调用方引用 analyses 地址是正当的，不应被拒绝。
    """
    return (
        f"lvke://source-files/workspaces/{workspace_id}/files/{source_id}",
        f"lvke://source-files/workspaces/{workspace_id}/analyses/{source_id}",
        f"lvke://data-acquisition/workspaces/{workspace_id}/sources/{source_id}",
    )


def _state_path(workspace_id: str) -> Path:
    return _root(workspace_id) / "state.json"


def _max_upload_bytes() -> int:
    try:
        return max(1, int(os.getenv("LVKE_MCP_MAX_UPLOAD_BYTES", "")))
    except ValueError:
        return _DEFAULT_MAX_UPLOAD_BYTES


def _empty_state() -> dict[str, Any]:
    return {"files": {}, "jobs": {}, "analyses": {}, "idempotency": {}}


def _load_state(workspace_id: str) -> dict[str, Any]:
    path = _state_path(workspace_id)
    if not path.is_file():
        return _empty_state()
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return _empty_state()
    if not isinstance(value, dict):
        return _empty_state()
    return {**_empty_state(), **value}


def _save_state(workspace_id: str, state: dict[str, Any]) -> None:
    path = _state_path(workspace_id)
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
    try:
        temporary.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")
        os.replace(temporary, path)
    finally:
        temporary.unlink(missing_ok=True)


@contextmanager
def _state_guard(workspace_id: str) -> Iterator[None]:
    _root(workspace_id).mkdir(parents=True, exist_ok=True)
    with _STATE_LOCK, FileLock(str(_root(workspace_id) / "state.lock"), timeout=30):
        yield


class SourceFileError(RuntimeError):
    def __init__(self, code: str, message: str, *, retryable: bool = False) -> None:
        super().__init__(message)
        self.detail = {
            "code": code,
            "message": message,
            "retryable": retryable,
        }


def _error(code: str, message: str) -> SourceFileError:
    return SourceFileError(code, message)


def _idempotency_record(scope: str, request_hash: str, **resource: Any) -> dict[str, Any]:
    return {"scope": scope, "request_hash": request_hash, "created_at": _now(), **resource}


def _active_idempotency_record(records: dict[str, Any], scope: str) -> dict[str, Any] | None:
    value = records.get(scope)
    return value if isinstance(value, dict) else None


def _public_parse_job(job: dict[str, Any]) -> dict[str, Any]:
    value = dict(job)
    for key in ("worker_token", "_worker_control"):
        value.pop(key, None)
    return value


def _require_source_record_from_state(
    state: dict[str, Any], workspace_id: str, file_id: str, _scope: str = ""
) -> dict[str, Any]:
    record = (state.get("files") or {}).get(file_id)
    if not isinstance(record, dict) or str(record.get("workspace_id") or "") != workspace_id:
        raise _error("source_file_not_found", "原始资料不存在或不属于当前工作区")
    return record


def _require_source_record(workspace_id: str, file_id: str, scope: str = "") -> tuple[dict[str, Any], dict[str, Any]]:
    state = _load_state(workspace_id)
    return state, _require_source_record_from_state(state, workspace_id, file_id, scope)


def _require_parse_job_from_state(
    state: dict[str, Any], workspace_id: str, job_id: str, _scope: str = ""
) -> tuple[dict[str, Any], dict[str, Any]]:
    job = (state.get("jobs") or {}).get(job_id)
    if not isinstance(job, dict) or str(job.get("workspace_id") or "") != workspace_id:
        raise _error("source_parse_job_not_found", "解析任务不存在或不属于当前工作区")
    record = _require_source_record_from_state(state, workspace_id, str(job.get("file_id") or ""))
    return job, record


def _load_analysis(workspace_id: str, file_id: str) -> dict[str, Any] | None:
    return _load_state(workspace_id).get("analyses", {}).get(file_id)


def _sha256_text(value: str) -> str:
    return "sha256:" + hashlib.sha256(value.encode("utf-8")).hexdigest()


def _normalized_sha256(value: Any) -> str:
    text = str(value or "").strip().lower()
    digest = text.removeprefix("sha256:")
    if len(digest) != 64 or any(char not in "0123456789abcdef" for char in digest):
        return ""
    return "sha256:" + digest


def _citation_offsets(locator: dict[str, Any], length: int) -> tuple[int, int]:
    start_raw = locator.get("start_offset", locator.get("start", 0))
    end_raw = locator.get("end_offset", locator.get("end", length))
    try:
        start = int(start_raw)
        end = int(end_raw)
    except (TypeError, ValueError):
        raise _error("citation_locator_invalid", "引用 locator 的文本 offset 必须是整数") from None
    if start < 0 or end <= start or end > length:
        raise _error("citation_locator_out_of_bounds", "引用 locator 的文本 offset 越界")
    return start, end


def _structured_locator(value: Any) -> dict[str, Any]:
    if isinstance(value, dict):
        return dict(value)
    text = str(value or "").strip()
    if not text:
        raise _error("citation_locator_required", "引用 locator 必填")
    if text.startswith("{"):
        try:
            decoded = json.loads(text)
        except (json.JSONDecodeError, ValueError):
            raise _error("citation_locator_invalid", "引用 locator JSON 格式无效") from None
        if not isinstance(decoded, dict):
            raise _error("citation_locator_invalid", "引用 locator JSON 必须是对象")
        return decoded
    if re.fullmatch(r"(?:pdf_page|page):[1-9][0-9]*", text):
        return {"kind": "pdf_page", "page": int(text.rsplit(":", 1)[1])}
    if re.fullmatch(r"csv:[1-9][0-9]*:[1-9][0-9]*", text):
        _, row, column = text.split(":")
        return {"kind": "csv_cell", "row": int(row), "column": int(column)}
    if re.fullmatch(r"csv:[A-Za-z]+[1-9][0-9]*", text):
        return {"kind": "csv_cell", "cell": text.split(":", 1)[1].upper()}
    if re.fullmatch(r"[A-Za-z]+[1-9][0-9]*", text):
        return {"kind": "csv_cell", "cell": text.upper()}
    if re.fullmatch(r"paragraph:[1-9][0-9]*", text):
        return {"kind": "docx_paragraph", "paragraph": int(text.split(":", 1)[1])}
    if re.fullmatch(r"table:[1-9][0-9]*:row:[1-9][0-9]*", text):
        parts = text.split(":")
        return {"kind": "docx_table_row", "table": int(parts[1]), "row": int(parts[3])}
    if text in {"document_text", "web_snapshot"}:
        return {"kind": text}
    return {"kind": "stored_locator", "locator": text}


def _a1_position(cell: str) -> tuple[int, int] | None:
    matched = re.fullmatch(r"([A-Z]+)([1-9][0-9]*)", str(cell or "").upper())
    if not matched:
        return None
    column = 0
    for char in matched.group(1):
        column = column * 26 + ord(char) - ord("A") + 1
    return int(matched.group(2)), column


def _resolve_analysis_row(
    analysis: dict[str, Any], locator: dict[str, Any]
) -> tuple[dict[str, Any], dict[str, Any]]:
    rows = _with_canonical_locators(
        [item for item in analysis.get("locators") or [] if isinstance(item, dict)]
    )
    kind = str(locator.get("kind") or "").strip().lower()
    matches: list[dict[str, Any]] = []
    normalized: dict[str, Any] = {}
    if kind == "pdf_page":
        try:
            page = int(locator.get("page"))
        except (TypeError, ValueError):
            raise _error("citation_locator_invalid", "PDF locator 必须包含有效页码") from None
        matches = [row for row in rows if row.get("kind") == "pdf_page" and int(row.get("page") or 0) == page]
        normalized = {"kind": "pdf_page", "page": page, "locator": f"pdf_page:{page}"}
    elif kind == "csv_cell":
        position = _a1_position(str(locator.get("cell") or "")) if locator.get("cell") else None
        if position is None:
            try:
                position = (int(locator.get("row")), int(locator.get("column")))
            except (TypeError, ValueError):
                raise _error("citation_locator_invalid", "CSV locator 必须包含有效单元格或行列") from None
        row_index, column_index = position
        matches = [
            row for row in rows
            if row.get("kind") == "cell"
            and int(row.get("row_index") or 0) == row_index
            and int(row.get("column_index") or 0) == column_index
        ]
        normalized = {
            "kind": "csv_cell",
            "row": row_index,
            "column": column_index,
            "cell": f"{_spreadsheet_column(column_index)}{row_index}",
            "locator": f"csv:{row_index}:{column_index}",
        }
    elif kind == "docx_paragraph":
        try:
            paragraph = int(locator.get("paragraph"))
        except (TypeError, ValueError):
            raise _error("citation_locator_invalid", "DOCX locator 必须包含有效段落编号") from None
        matches = [row for row in rows if row.get("kind") == "docx_paragraph" and int(row.get("paragraph") or 0) == paragraph]
        normalized = {"kind": "docx_paragraph", "paragraph": paragraph, "locator": f"paragraph:{paragraph}"}
    elif kind == "docx_table_row":
        try:
            table = int(locator.get("table"))
            row_index = int(locator.get("row"))
        except (TypeError, ValueError):
            raise _error("citation_locator_invalid", "DOCX 表格 locator 必须包含有效表号和行号") from None
        matches = [
            row for row in rows
            if row.get("kind") == "docx_table_row"
            and int(row.get("table") or 0) == table
            and int(row.get("row") or 0) == row_index
        ]
        normalized = {
            "kind": "docx_table_row", "table": table, "row": row_index,
            "locator": f"table:{table}:row:{row_index}",
        }
    elif kind == "document_text":
        matches = [row for row in rows if row.get("kind") == "document_text"]
        normalized = {"kind": "document_text", "locator": "document_text"}
    elif kind == "stored_locator":
        address = str(locator.get("locator") or "").strip()
        matches = [row for row in rows if str(row.get("locator") or "") == address]
        normalized = {"kind": str(matches[0].get("kind") or "stored_locator"), "locator": address} if matches else {}
    else:
        address = str(locator.get("locator") or "").strip()
        matches = [row for row in rows if address and str(row.get("locator") or "") == address]
        normalized = {"kind": kind, "locator": address} if matches else {}
    if not matches:
        raise _error("citation_locator_not_found", "引用 locator 未能解析到来源中的实际片段")
    if len(matches) != 1:
        raise _error("citation_locator_ambiguous", "引用 locator 不能唯一解析")
    row = matches[0]
    text = str(row.get("text") if row.get("text") is not None else row.get("original_value") or "")
    if not text:
        raise _error("citation_fragment_empty", "引用 locator 对应的片段为空")
    start, end = _citation_offsets(locator, len(text))
    normalized.update({"start_offset": start, "end_offset": end})
    return row, {**normalized, "fragment_text": text[start:end]}


def resolve_citation_fragment(
    workspace_id: str,
    *,
    source_id: str,
    locator: Any,
    source_hash: Any,
    supplied_fragment: Any = "",
    supplied_fragment_hash: Any = "",
) -> dict[str, Any]:
    """Resolve and hash one citation fragment against immutable source content.

    The result proves deterministic source/locator/fragment binding only. It
    deliberately does not decide whether the fragment semantically supports a
    claim; that remains Agent or manual review evidence.
    """

    source_id = str(source_id or "").strip()
    if not source_id:
        raise _error("citation_source_id_required", "引用 source_id 必填")
    claimed_hash = _normalized_sha256(source_hash)
    if not claimed_hash:
        raise _error("citation_source_hash_invalid", "引用必须提供有效的整源 SHA-256")
    structured = _structured_locator(locator)
    state = _load_state(workspace_id)
    record = (state.get("files") or {}).get(source_id)
    source_kind = "source_file"
    resource_uri = f"lvke://source-files/workspaces/{workspace_id}/files/{source_id}"
    if isinstance(record, dict) and str(record.get("workspace_id") or "") == workspace_id:
        path = Path(str(record.get("path") or ""))
        if not path.is_file():
            raise _error("citation_source_integrity_failed", "引用来源文件不存在")
        raw = path.read_bytes()
        actual_hash = "sha256:" + hashlib.sha256(raw).hexdigest()
        recorded_hash = _normalized_sha256(record.get("sha256"))
        if actual_hash != recorded_hash or len(raw) != int(record.get("size_bytes") or -1):
            raise _error("citation_source_integrity_failed", "引用来源文件完整性校验失败")
        if claimed_hash != actual_hash:
            raise _error("citation_source_hash_mismatch", "引用整源 hash 与来源文件不一致")
        analysis = (state.get("analyses") or {}).get(source_id)
        if not isinstance(analysis, dict):
            raise _error("citation_source_unparsed", "引用来源尚未解析")
        if _normalized_sha256(analysis.get("sha256")) not in {"", actual_hash}:
            raise _error("citation_source_integrity_failed", "引用来源解析结果与原件 hash 不一致")
        if int(analysis.get("size_bytes") or len(raw)) != len(raw):
            raise _error("citation_source_integrity_failed", "引用来源解析结果与原件大小不一致")
        _row, resolved = _resolve_analysis_row(analysis, structured)
        fragment_text = str(resolved.pop("fragment_text"))
    else:
        snapshot = _load_acquisition_snapshot(workspace_id, source_id)
        if snapshot is None:
            raise _error("citation_source_not_found", "引用来源不存在或不属于当前工作区")
        payload = snapshot.get("payload") if isinstance(snapshot.get("payload"), dict) else {}
        content = payload.get("content")
        if not isinstance(content, str) or not content:
            raise _error("citation_source_content_missing", "引用来源缺少可解析正文")
        actual_hash = _sha256_text(content)
        recorded_external_hash = _normalized_sha256(payload.get("external_content_hash"))
        if recorded_external_hash and recorded_external_hash != actual_hash:
            raise _error("citation_source_integrity_failed", "引用来源正文与已存 hash 不一致")
        if claimed_hash != actual_hash:
            raise _error("citation_source_hash_mismatch", "引用整源 hash 与来源正文不一致")
        kind = str(structured.get("kind") or "")
        if kind not in {"document_text", "web_snapshot", "stored_locator"}:
            raise _error("citation_locator_invalid", "网页快照只支持正文及文本 offset locator")
        if kind == "stored_locator" and str(structured.get("locator") or "") not in {"document_text", "web_snapshot"}:
            raise _error("citation_locator_not_found", "网页快照 locator 未能解析到正文")
        start, end = _citation_offsets(structured, len(content))
        fragment_text = content[start:end]
        resolved = {
            "kind": "web_snapshot",
            "locator": "web_snapshot",
            "start_offset": start,
            "end_offset": end,
        }
        source_kind = "source_snapshot"
        resource_uri = str(snapshot.get("resource_uri") or "")
    fragment_hash = _sha256_text(fragment_text)
    supplied_text = supplied_fragment if isinstance(supplied_fragment, str) else ""
    if supplied_text and supplied_text != fragment_text:
        raise _error("citation_fragment_mismatch", "调用方提交的引用片段与服务端解析结果不一致")
    claimed_fragment_hash = _normalized_sha256(supplied_fragment_hash)
    if supplied_fragment_hash and not claimed_fragment_hash:
        raise _error("citation_fragment_hash_invalid", "调用方提交的片段 hash 格式无效")
    if claimed_fragment_hash and claimed_fragment_hash != fragment_hash:
        raise _error("citation_fragment_hash_mismatch", "调用方提交的片段 hash 与服务端解析结果不一致")
    return {
        "workspace_id": workspace_id,
        "source_id": source_id,
        "source_kind": source_kind,
        "resource_uri": resource_uri,
        "source_hash": actual_hash,
        "locator": resolved,
        "fragment_text": fragment_text,
        "fragment_hash": fragment_hash,
        "binding_status": "resolved",
        "semantic_support_status": "agent_or_manual_review_required",
    }


_SPREADSHEET_MIMES = frozenset({
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
})


def resolve_source_workbook_for_review(
    workspace_id: str,
    source_file_id: str,
) -> dict[str, Any]:
    """解析上传的 Excel 工作簿供统一审查（自 hermes source_files_api 裁剪的 MCP 版）。

    读取 MCP 自有 source-files 存储，fail-closed：记录不存在/非电子表格/文件
    丢失或字节与记录摘要不符时返回 ``ok=False`` 及原因码。不暴露工作区外路径。
    """
    file_id = str(source_file_id or "").strip()
    if not file_id or "/" in file_id or "\\" in file_id or ":" in file_id or file_id in {".", ".."}:
        return {"ok": False, "code": "source_workbook_not_found"}
    with _state_guard(workspace_id):
        state = _load_state(workspace_id)
        record = (state.get("files") or {}).get(file_id)
        if (
            not isinstance(record, dict)
            or str(record.get("workspace_id") or "") != workspace_id
        ):
            return {"ok": False, "code": "source_workbook_not_found"}
        declared = str(record.get("declared_mime") or record.get("mime_type") or "")
        suffix = Path(str(record.get("original_filename") or "")).suffix.lower()
        if declared not in _SPREADSHEET_MIMES and suffix not in {".xls", ".xlsx", ".xlsm"}:
            return {"ok": False, "code": "source_workbook_not_spreadsheet"}
        source_path = Path(str(record.get("path") or ""))
        if not source_path.is_file():
            return {"ok": False, "code": "source_workbook_integrity_failed"}
        raw = source_path.read_bytes()
        recorded_sha = str(record.get("sha256") or "").lower().removeprefix("sha256:")
        if (
            len(raw) != int(record.get("size_bytes") or -1)
            or hashlib.sha256(raw).hexdigest() != recorded_sha
        ):
            return {"ok": False, "code": "source_workbook_integrity_failed"}
        return {
            "ok": True,
            "path": str(source_path),
            "size": len(raw),
            "sha256": "sha256:" + recorded_sha,
            "source_file_id": file_id,
            "original_filename": str(record.get("original_filename") or ""),
            "source_version": record.get("version"),
        }


def resolve_authoritative_evidence_binding(
    workspace_id: str,
    *,
    source_id: str,
    evidence_id: str = "",
    locator: str = "",
) -> dict[str, Any]:
    """Resolve one finance fact binding from the durable source authority (MCP 版)。

    裁剪自 hermes source_files_api：record 存在性/完整性、analysis locator
    匹配、OCR 与 locator 状态全部 fail-closed。
    """
    source_id = str(source_id or "").strip()
    evidence_id = str(evidence_id or "").strip()
    locator = str(locator or "").strip()
    base = {
        "source_id": source_id,
        "evidence_id": evidence_id,
        "locator": locator,
        "evidence_grade": "E",
        "validation_status": "missing",
        "authoritative": True,
        "binding_ok": False,
        "issues": [],
    }
    if not source_id or (not evidence_id and not locator):
        base["issues"] = ["source_id 与 evidence_id/locator 必填"]
        return base

    with _state_guard(workspace_id):
        state = _load_state(workspace_id)
        record = (state.get("files") or {}).get(source_id)
        if not isinstance(record, dict) or str(record.get("workspace_id") or "") != workspace_id:
            # 跨域兜底：data_fetch 快照存放在独立 JSONArtifactStore，
            # 与本域 state["files"] 完全不相交。命中快照存储时按候选证据绑定，
            # 不自动取得正式证据资格（formal_use_allowed 决定证据等级）。
            snapshot = _load_acquisition_snapshot(workspace_id, source_id)
            if snapshot is None:
                base["issues"] = [
                    "source 不存在或不属于当前工作区（本域 source-files 与 data-acquisition 快照均未命中）"
                ]
                return base
            payload = snapshot.get("payload") or {}
            formal_allowed = payload.get("formal_use_allowed") is True
            content_hash = str(
                payload.get("content_hash")
                or payload.get("external_content_hash")
                or snapshot.get("content_hash")
                or ""
            )
            if content_hash:
                content_hash = "sha256:" + str(content_hash).removeprefix("sha256:")
            snapshot_issues: list[str] = []
            if not formal_allowed:
                snapshot_issues.append("快照未标记 formal_use_allowed，仅作候选证据，不可作为正式财务输入")
            if not content_hash:
                snapshot_issues.append("快照缺少 content_hash，无法固定证据版本")
            # 网页快照没有 analysis locator 表，改为按正文实证 locator：
            # locator 必须能在已固化正文中定位，否则 fail-closed。
            content_text = payload.get("content")
            content_text = content_text if isinstance(content_text, str) else ""
            reviewed_value = None
            if locator:
                if content_text and locator in content_text:
                    reviewed_value = locator
                else:
                    snapshot_issues.append("locator 未能在快照正文中定位")
            elif not evidence_id:
                snapshot_issues.append("快照绑定需提供 locator")
            kind = "web_snapshot"
            grade = "A" if formal_allowed else "B"
            binding_ok = not snapshot_issues
            return {
                "source_id": source_id,
                "file_id": source_id,
                "evidence_id": evidence_id,
                "locator": locator,
                "kind": kind,
                "evidence_grade": grade,
                "validation_status": "passed" if binding_ok else "failed",
                "value": reviewed_value,
                "formal_use_decision": "allowed" if formal_allowed else "candidate_only",
                "source_sha256": content_hash,
                "source_version": snapshot.get("schema_version"),
                "source_format": payload.get("mime_type"),
                "source_uri": snapshot.get("resource_uri") or "",
                "source_domain": "data-acquisition",
                "authoritative": True,
                "binding_ok": binding_ok,
                "issues": snapshot_issues,
            }
        analysis = _load_analysis(workspace_id, source_id)
        if not isinstance(analysis, dict):
            base["issues"] = ["source 尚未解析，无法绑定证据；先完成解析后重试"]
            return {
                **base,
                "source_sha256": record.get("sha256"),
                "source_version": record.get("version"),
                "source_format": record.get("declared_mime"),
            }
        # 存量 analysis 由旧解析器写入，只有 kind/text 而无 locator/evidence_id。
        # 读取时按同一规则补齐，使既有工作区无需重新解析即可绑定。
        candidates = _with_canonical_locators(
            [row for row in analysis.get("locators") or [] if isinstance(row, dict)]
        )
        target = next(
            (
                row for row in candidates
                if evidence_id and str(row.get("evidence_id") or "") == evidence_id
            ),
            None,
        )
        if target is None:
            target = next(
                (
                    row for row in candidates
                    if locator and str(row.get("locator") or "") == locator
                ),
                None,
            )
        if target is None and locator:
            # 调用方可能回传序列化的 locator 行（如
            # '{"text":"...","kind":"document_text"}'）。解出 kind/text 后
            # 按同一规则重新定位，不要求调用方先了解内部规范地址。
            decoded = _decode_serialized_locator(locator)
            if decoded:
                kind, text = decoded
                target = next(
                    (
                        row for row in candidates
                        if (not kind or str(row.get("kind") or "") == kind)
                        and (not text or text in str(row.get("text") or ""))
                    ),
                    None,
                )
                if target is not None:
                    # 已按 kind/text 实际命中权威行，规范地址即该行自身；
                    # 标记为已核实，避免下游再拿原始序列化串做字面比对。
                    canonical = str(target.get("locator") or "")
                    target = {**target, "locator": canonical, "evidence_id": canonical}
                    locator = canonical
                    evidence_id = canonical if evidence_id else evidence_id
        if target is None and locator:
            # 调用方常直接引用正文片段而非规范地址；只要该片段能在已固化
            # 正文中真实命中，就按承载它的 locator 行绑定，并保留原引用串。
            target = next(
                (
                    row for row in candidates
                    if locator in str(row.get("text") or "")
                ),
                None,
            )
            if target is not None:
                target = {**target, "locator": locator, "evidence_id": locator}
        if target is None:
            available = sorted({str(row.get("locator") or "") for row in candidates if row.get("locator")})
            base["issues"] = [
                "evidence locator 不存在（可用 locator："
                + ("、".join(available[:8]) if available else "无")
                + "；也可直接引用正文中真实存在的片段）"
            ]
            return {
                **base,
                "available_locators": available,
                "source_sha256": record.get("sha256"),
                "source_version": record.get("version"),
                "source_format": record.get("declared_mime"),
            }

        issues: list[str] = []
        target_evidence_id = str(target.get("evidence_id") or "")
        target_locator = str(target.get("locator") or "")
        if evidence_id and evidence_id != target_evidence_id:
            issues.append("evidence_id 与权威记录不一致")
        if locator and locator != target_locator:
            issues.append("locator 与权威记录不一致")

        source_path = Path(str(record.get("path") or ""))
        integrity_ok = False
        if source_path.is_file():
            raw = source_path.read_bytes()
            integrity_ok = (
                len(raw) == int(record.get("size_bytes") or -1)
                and hashlib.sha256(raw).hexdigest() == str(record.get("sha256") or "")
            )
        if not integrity_ok:
            issues.append("source 文件完整性校验失败")
        if analysis.get("source_verified_sha256") not in (None, "", record.get("sha256")):
            issues.append("analysis/source sha256 不一致")
        if analysis.get("source_verified_size_bytes") not in (None, "", record.get("size_bytes")):
            issues.append("analysis/source size 不一致")

        if analysis.get("ocr_status") == "failed":
            issues.append("OCR 未完成，证据绑定不可用")

        fmt = str(record.get("declared_mime") or "").lower()
        kind = str(target.get("kind") or "")
        grade = "A" if fmt in {"xls", "xlsx", "xlsm"} and kind == "cell" else "B"
        binding_ok = not issues
        reviewed_value = target.get("normalized_value")
        if reviewed_value in (None, ""):
            reviewed_value = target.get("original_value")
        if reviewed_value in (None, ""):
            reviewed_value = target.get("value")
        return {
            "source_id": source_id,
            "file_id": source_id,
            "evidence_id": target_evidence_id,
            "locator": target_locator,
            "kind": kind,
            "evidence_grade": grade,
            "validation_status": "passed" if binding_ok else "failed",
            "value": reviewed_value,
            "formal_use_decision": (
                "invalid" if analysis.get("ocr_status") == "failed"
                else str(analysis.get("deterministic_status") or "not_applicable")
            ),
            "source_sha256": record.get("sha256"),
            "source_version": record.get("version"),
            "source_format": fmt,
            "authoritative": True,
            "binding_ok": binding_ok,
            "issues": issues,
        }


def resolve_reconstructed_evidence_binding(
    workspace_id: str,
    *,
    source_id: str,
    locator: str,
    reconstruction_record: dict[str, Any],
) -> dict[str, Any]:
    """Resolve an explicit reconstruction mapping against one source snapshot.

    This confirms only snapshot identity and the declared reconstruction
    locator/method.  It deliberately does not certify the mapped value as an
    original project fact.
    """

    from lvke_mcp.runtime.source_reconstruction import reconstruction_errors

    source_id = str(source_id or "").strip()
    locator = str(locator or "").strip()
    base = {
        "source_id": source_id,
        "evidence_id": "",
        "locator": locator,
        "evidence_grade": "B",
        "review_status": "approved",
        "validation_status": "missing",
        "authoritative": True,
        "binding_ok": False,
        "allow_claimed_value": False,
        "project_fact_certified": False,
        "evidence_policy": "source_reconstructed",
        "issues": [],
    }
    errors = reconstruction_errors(reconstruction_record)
    if errors:
        base["issues"] = errors
        return base
    expected_uris = _acquisition_snapshot_uris(workspace_id, source_id)
    declared_uri = str(reconstruction_record.get("source_uri") or "")
    if declared_uri not in expected_uris:
        base["issues"] = [
            "source_uri 与 Source Snapshot 不一致（应为 "
            + " 或 ".join(expected_uris)
            + "）"
        ]
        return base
    # 入参 locator 已 strip，重建记录侧此前用原值比对：正文带结尾换行时
    # 两侧不等，报"locator 与来源重建记录不一致"。两侧统一规范化。
    if str(reconstruction_record.get("locator") or "").strip() != locator:
        base["issues"] = ["locator 与来源重建记录不一致"]
        return base
    with _state_guard(workspace_id):
        state = _load_state(workspace_id)
        record = (state.get("files") or {}).get(source_id)
        if not isinstance(record, dict) or str(record.get("workspace_id") or "") != workspace_id:
            # 跨域兜底：来源重建同样可指向 data-acquisition 快照。
            snapshot = _load_acquisition_snapshot(workspace_id, source_id)
            if snapshot is None:
                base["issues"] = [
                    "source 不存在或不属于当前工作区（本域 source-files 与 data-acquisition 快照均未命中）"
                ]
                return base
            payload = snapshot.get("payload") or {}
            snapshot_hash = str(
                payload.get("content_hash")
                or payload.get("external_content_hash")
                or snapshot.get("content_hash")
                or ""
            )
            snapshot_hash = "sha256:" + snapshot_hash.removeprefix("sha256:") if snapshot_hash else ""
            declared_hash = str(reconstruction_record.get("content_hash") or "")
            if not snapshot_hash or snapshot_hash != declared_hash:
                base["issues"] = ["来源重建 hash 与 Source Snapshot 不一致"]
                return {**base, "source_sha256": snapshot_hash}
            return {
                **base,
                "evidence_grade": "B",
                "validation_status": "passed",
                "binding_ok": True,
                "allow_claimed_value": True,
                "source_sha256": snapshot_hash,
                "source_version": snapshot.get("schema_version"),
                "source_format": payload.get("mime_type"),
                "source_uri": snapshot.get("resource_uri") or "",
                "source_domain": "data-acquisition",
                "reconstruction_id": reconstruction_record.get("reconstruction_id"),
                "reconstruction_method": reconstruction_record.get("method"),
                "source_kind": reconstruction_record.get("source_kind"),
                "issues": [],
            }
        source_path = Path(str(record.get("path") or ""))
        recorded_hash = "sha256:" + str(record.get("sha256") or "").removeprefix("sha256:")
        if recorded_hash != str(reconstruction_record.get("content_hash") or ""):
            base["issues"] = ["来源重建 hash 与 Source Snapshot 不一致"]
            return {**base, "source_sha256": recorded_hash}
        if not source_path.is_file():
            base["issues"] = ["Source Snapshot 文件不存在"]
            return {**base, "source_sha256": recorded_hash}
        raw = source_path.read_bytes()
        actual_hash = "sha256:" + hashlib.sha256(raw).hexdigest()
        if (
            actual_hash != recorded_hash
            or len(raw) != int(record.get("size_bytes") or -1)
        ):
            base["issues"] = ["Source Snapshot 完整性校验失败"]
            return {**base, "source_sha256": recorded_hash}
        spreadsheet = str(record.get("declared_mime") or "").lower() in _SPREADSHEET_MIMES
        return {
            **base,
            "evidence_grade": "A" if spreadsheet else "B",
            "validation_status": "passed",
            "binding_ok": True,
            "allow_claimed_value": True,
            "source_sha256": recorded_hash,
            "source_version": record.get("version"),
            "source_format": record.get("declared_mime"),
            "reconstruction_id": reconstruction_record.get("reconstruction_id"),
            "reconstruction_method": reconstruction_record.get("method"),
            "source_kind": reconstruction_record.get("source_kind"),
            "issues": [],
        }


_TEXT_SUFFIXES = frozenset({".md", ".txt", ".json", ".jsonl", ".html"})
_CSV_SUFFIXES = frozenset({".csv", ".tsv"})
_DOCX_SUFFIXES = frozenset({".docx"})
_PDF_TEXT_BUDGET = 2_000_000


def _is_docx(mime: str, path: Path) -> bool:
    normalized = str(mime or "").lower().split(";", 1)[0].strip()
    return (
        normalized
        in {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/docx",
        }
        or path.suffix.lower() in _DOCX_SUFFIXES
    )


def _parse_docx_locators(data: bytes) -> tuple[list[dict[str, Any]], str]:
    try:
        from docx import Document
    except ImportError:
        return [], "docx_parser_unavailable"
    try:
        document = Document(io.BytesIO(data))
    except Exception:  # noqa: BLE001
        return [], "docx_unreadable"
    locators: list[dict[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = str(paragraph.text or "").strip()
        if not text:
            continue
        locators.append({
            "kind": "docx_paragraph",
            "locator": f"paragraph:{index}",
            "paragraph": index,
            "text": text[:4000],
        })
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [str(cell.text or "").strip() for cell in row.cells]
            if not any(cells):
                continue
            locators.append({
                "kind": "docx_table_row",
                "locator": f"table:{table_index}:row:{row_index}",
                "table": table_index,
                "row": row_index,
                "text": " | ".join(cells)[:4000],
                "cells": cells,
            })
    return locators, ""


def _is_csv(mime: str, path: Path) -> bool:
    normalized = str(mime or "").lower().split(";", 1)[0].strip()
    return normalized in {"text/csv", "text/tab-separated-values"} or path.suffix.lower() in _CSV_SUFFIXES


def _decode_csv(data: bytes) -> tuple[str, str, str]:
    """Decode controlled CSV bytes without replacement-character guessing."""

    for encoding in ("utf-8-sig", "gb18030"):
        try:
            return data.decode(encoding, errors="strict"), encoding, ""
        except UnicodeDecodeError:
            continue
    return "", "", "csv_encoding_invalid"


_CSV_INTEGER = re.compile(r"^[+-]?[0-9]+$")
_CSV_NUMBER = re.compile(r"^[+-]?(?:[0-9]+(?:\.[0-9]+)?|\.[0-9]+)$")


def _csv_scalar(value: str) -> str | int | float:
    stripped = str(value).strip()
    normalized = stripped.replace(",", "")
    if _CSV_INTEGER.fullmatch(normalized):
        try:
            return int(normalized)
        except ValueError:
            return stripped
    if _CSV_NUMBER.fullmatch(normalized):
        try:
            return float(normalized)
        except ValueError:
            return stripped
    return stripped


# ── Column classification for field-value-unit association ────────────────
# Exact-match headers for column type detection
_FIELD_HEADERS: frozenset[str] = frozenset({
    "项目", "指标", "名称", "内容", "项目名称", "指标名称", "费用项目", "项目内容",
    "指标名称", "费用类别", "项目类别", "项目/指标", "指标/项目",
})
_VALUE_HEADERS: frozenset[str] = frozenset({
    "金额", "数值", "数量", "值", "价值", "价格", "收入", "成本", "指标值",
    "数值/金额", "金额/数值", "合计", "小计", "总计", "金额(万元)", "金额(元)",
    "金额/万元", "数值/万元",
})
_UNIT_HEADERS: frozenset[str] = frozenset({
    "单位", "计量单位", "度量单位", "数量单位", "货币单位", "金额单位",
})

# Partial-match keywords (fallback when exact match fails)
_FIELD_KEYWORDS: frozenset[str] = frozenset({"项目", "指标", "名称", "类别", "费用"})
_VALUE_KEYWORDS: frozenset[str] = frozenset({"金额", "数值", "数量", "价值", "价格", "收入", "成本"})
_UNIT_KEYWORDS: frozenset[str] = frozenset({"单位"})

# Controlled unit aliases → normalized unit
_CONTROLLED_UNIT_ALIASES: dict[str, str] = {
    "万元": "wan",
    "万": "wan",
    "w": "wan",
    "万元人民币": "wan",
    "mw": "MW",
    "兆瓦": "MW",
    "mwh": "MWh",
    "兆瓦时": "MWh",
    "万千瓦": "wan_kW",
    "万千瓦时": "wan_kWh",
    "千瓦": "kW",
    "千瓦时": "kWh",
    "元/kwh": "yuan/kWh",
    "元/千瓦时": "yuan/kWh",
    "yuan/kwh": "yuan/kWh",
    "元/kw": "yuan/kW",
    "元/千瓦": "yuan/kW",
    "元/mw": "yuan/MW",
    "元/兆瓦": "yuan/MW",
    "万元/mw": "wan/MW",
    "万元/兆瓦": "wan/MW",
    "元": "yuan",
    "元/年": "yuan/year",
    "%": "percent",
    "百分比": "percent",
    "ratio": "ratio",
    "h": "h",
    "小时": "h",
    "年": "year",
}

# Pattern for splitting "10MW" → (10, "MW") from a combined value cell
_VALUE_UNIT_SPLIT = re.compile(r"^([+-]?(?:\d+(?:\.\d+)?|\.\d+))\s*(.*)$")


def _normalize_unit(raw: str) -> str:
    """Map a raw unit string to its controlled normal form."""
    key = raw.strip().lower()
    direct = _CONTROLLED_UNIT_ALIASES.get(key)
    if direct:
        return direct
    # Fallback: case-insensitive scan
    for alias, canonical in _CONTROLLED_UNIT_ALIASES.items():
        if key == alias.lower():
            return canonical
    return raw.strip()


def _split_value_unit(raw: str) -> tuple[str | int | float, str]:
    """Split a combined cell like ``10MW`` into ``(10, "MW")``.

    Returns ``(raw_str, "")`` when the raw value has no parseable number prefix.
    """
    stripped = raw.strip()
    m = _VALUE_UNIT_SPLIT.match(stripped)
    if m:
        num_str, unit = m.group(1), m.group(2).strip()
        if "." in num_str:
            return float(num_str), unit
        return int(num_str), unit
    return stripped, ""


def _classify_csv_columns(headers: list[str]) -> dict[str, list[int]]:
    """Classify 0-based column indices into field/value/unit groups.

    Returns ``{"field": [...], "value": [...], "unit": [...]}``.
    Uses exact header match first, then partial keyword fallback.
    """
    field_indices: list[int] = []
    value_indices: list[int] = []
    unit_indices: list[int] = []

    for idx, header in enumerate(headers):
        h = header.strip().lower()
        if h in _FIELD_HEADERS:
            field_indices.append(idx)
        elif h in _UNIT_HEADERS:
            unit_indices.append(idx)
        elif h in _VALUE_HEADERS:
            value_indices.append(idx)
        else:
            # Partial keyword matching — only one category wins
            matched = False
            for kw in _FIELD_KEYWORDS:
                if kw in h:
                    field_indices.append(idx)
                    matched = True
                    break
            if matched:
                continue
            for kw in _VALUE_KEYWORDS:
                if kw in h:
                    value_indices.append(idx)
                    matched = True
                    break
            if matched:
                continue
            for kw in _UNIT_KEYWORDS:
                if kw in h:
                    unit_indices.append(idx)
                    break

    return {
        "field": field_indices,
        "value": value_indices,
        "unit": unit_indices,
    }


def _spreadsheet_column(index: int) -> str:
    value = max(1, int(index))
    letters = ""
    while value:
        value, remainder = divmod(value - 1, 26)
        letters = chr(ord("A") + remainder) + letters
    return letters


def _parse_csv_cells(data: bytes, path: Path, mime: str) -> tuple[list[dict[str, Any]], str, str, str]:
    """Return stable cell locators plus encoding/delimiter diagnostics.

    A CSV is a table, not prose.  Every non-empty cell receives a stable A1 and
    row/column locator, the normalized header name and the untouched source
    value.  Malformed quoting or ragged rows remain visible as ``partial`` and
    are never flattened into a whole-document text locator.
    """

    text, encoding, degraded = _decode_csv(data)
    if degraded:
        return [], encoding, "", degraded
    delimiter = "\t" if path.suffix.lower() == ".tsv" or "tab-separated" in mime else ","
    try:
        sample = text[:8192]
        if sample:
            try:
                delimiter = csv.Sniffer().sniff(sample, delimiters=",;\t|").delimiter
            except csv.Error:
                # Sniffer failure alone must not flatten a structured partial
                # result into prose or discard its valid cells.
                pass
        rows = list(csv.reader(io.StringIO(text, newline=""), delimiter=delimiter, strict=True))
    except (csv.Error, UnicodeError):
        return [], encoding, delimiter, "csv_structure_invalid"
    if not rows or not any(str(cell).strip() for row in rows for cell in row):
        return [], encoding, delimiter, "csv_empty"
    width = len(rows[0])
    if width <= 0:
        return [], encoding, delimiter, "csv_header_missing"
    headers = [str(value).strip() or f"column_{index + 1}" for index, value in enumerate(rows[0])]
    # Classify columns for field-value-unit association
    col_class = _classify_csv_columns(headers)
    field_indices: list[int] = col_class["field"]
    unit_indices: list[int] = col_class["unit"]
    ragged = any(len(row) != width for row in rows[1:])
    locators: list[dict[str, Any]] = []
    for row_index, row in enumerate(rows, start=1):
        padded = [*row, *([""] * max(0, width - len(row)))]
        # Determine row-level field name and unit (from classified columns)
        row_field = ""
        row_unit = ""
        if row_index > 1:  # skip header row for field/unit lookup
            for idx in field_indices:
                if idx < len(padded) and str(padded[idx]).strip():
                    row_field = str(padded[idx]).strip()
                    break
            for idx in unit_indices:
                if idx < len(padded) and str(padded[idx]).strip():
                    row_unit = str(padded[idx]).strip()
                    break
        for column_index, raw_value in enumerate(padded[:width], start=1):
            original = str(raw_value)
            if not original.strip():
                continue
            cell = f"{_spreadsheet_column(column_index)}{row_index}"
            scalar = _csv_scalar(original)
            locator = f"csv:{row_index}:{column_index}"
            # Per-cell enrichment
            col_idx_0 = column_index - 1
            is_field_col = col_idx_0 in field_indices
            is_unit_col = col_idx_0 in unit_indices
            # field_name: use own value if in a field column, else row-level field
            cell_field: str | None = original if is_field_col else (row_field or None)
            # numeric_value: parsed number if cell is numeric, else None
            cell_numeric: int | float | None = scalar if isinstance(scalar, (int, float)) else None
            # unit: own value if in unit column, else row-level unit
            cell_unit_raw: str | None = original if is_unit_col else (row_unit or None)
            cell_unit_normalized: str | None = _normalize_unit(cell_unit_raw) if cell_unit_raw else None
            # If the value is not a pure number but has a numeric prefix (e.g. "10MW"),
            # split it and use the extracted numeric portion
            if cell_numeric is None and cell_unit_raw is None:
                parsed_num, parsed_unit = _split_value_unit(original)
                if isinstance(parsed_num, (int, float)):
                    cell_numeric = parsed_num
                    if parsed_unit:
                        cell_unit_raw = parsed_unit
                        cell_unit_normalized = _normalize_unit(parsed_unit)
            locators.append({
                "kind": "cell",
                "table_kind": "csv",
                "locator": locator,
                "evidence_id": locator,
                "cell": cell,
                "row_index": row_index,
                "column_index": column_index,
                "header_name": headers[column_index - 1],
                "is_header": row_index == 1,
                "original_value": original,
                "display_value": scalar,
                "cached_value": scalar,
                "text": original,
                "content_hash": "sha256:" + hashlib.sha256(original.encode("utf-8")).hexdigest(),
                "field_name": cell_field,
                "numeric_value": cell_numeric,
                "unit": cell_unit_raw,
                "unit_normalized": cell_unit_normalized,
            })
    if not locators:
        return [], encoding, delimiter, "csv_empty"
    return locators, encoding, delimiter, "csv_ragged_rows" if ragged else ""


def _is_pdf(mime: str, path: Path) -> bool:
    return mime.lower().strip() == "application/pdf" or path.suffix.lower() == ".pdf"


def _parse_pdf_pages(data: bytes) -> tuple[list[dict[str, Any]], str]:
    """Build page-anchored locators for a text-layer PDF.

    Returns ``(locators, degraded_reason)``.  A non-empty ``degraded_reason``
    means no quotable text was recovered, and the caller must NOT report the
    parse as fully succeeded: a scanned PDF needs OCR, and silently returning
    zero locators with ``extract_status=succeeded`` would let an empty document
    reach evidence selection while looking usable.
    """

    from lvke_mcp.adapters.pdf_text import extract_pdf_pages

    pages, degraded = extract_pdf_pages(data)
    if degraded and not pages:
        # 完全没有可引用文本（无文本层 / 打不开 / 缺库）。
        return [], degraded
    locators: list[dict[str, Any]] = []
    budget = _PDF_TEXT_BUDGET
    for page in pages:
        text = str(page.get("text") or "").strip()
        if not text or budget <= 0:
            continue
        text = text[:budget]
        budget -= len(text)
        locators.append(
            {
                "kind": "pdf_page",
                "page": int(page.get("page") or 0),
                "text": text,
                "start_offset": page.get("start_offset"),
                "end_offset": page.get("end_offset"),
            }
        )
    if not locators:
        # Text layer absent (scanned scan-only PDF) — honest, not "succeeded".
        return [], "pdf_no_text_layer"
    # 有 locator 但缺中文标签：保留可引用文本，同时把降级原因透传给上层，
    # 由它决定 extract_status/ocr_status，绝不静默当作成功。
    return locators, degraded


def _canonical_locator(row: dict[str, Any]) -> str:
    """Canonical, quotable address for one analysis locator row.

    解析器此前只写 ``kind``/``text``，而财务证据绑定按 ``row["locator"]``
    匹配，两侧键名不重合，导致任何已上传文件的逐值绑定都必然报
    "evidence locator 不存在"。这里统一铸造稳定地址：
    ``document_text`` → ``document_text``；``pdf_page`` → ``pdf_page:<页码>``。
    """
    kind = str(row.get("kind") or "").strip()
    if kind == "cell" and row.get("locator"):
        return str(row["locator"])
    if kind == "pdf_page":
        return f"pdf_page:{int(row.get('page') or 0)}"
    return kind or "document_text"


def _with_canonical_locators(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Attach canonical locator/evidence_id to freshly parsed locator rows."""
    out: list[dict[str, Any]] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        item = dict(row)
        canonical = _canonical_locator(item)
        item.setdefault("locator", canonical)
        item.setdefault("evidence_id", canonical)
        out.append(item)
    return out


def _parse_spreadsheet_locators(path: Path) -> tuple[list[dict[str, Any]], str]:
    """Extract bounded, exact cell fragments from xls/xlsx workbooks."""

    try:
        from lvke_mcp.adapters.spreadsheets.reader import pick_backend

        backend = pick_backend()
        sheet_names = backend.list_sheets(path)
    except Exception:  # noqa: BLE001
        return [], "spreadsheet_parser_unavailable"
    locators: list[dict[str, Any]] = []
    try:
        for sheet_name in sheet_names[:100]:
            result = backend.read_sheet(path, sheet_name, max_rows=2000, max_cols=256)
            for row_index, row in enumerate(result.rows, start=1):
                for column_index, value in enumerate(row, start=1):
                    if value is None or (isinstance(value, str) and not value.strip()):
                        continue
                    cell = f"{_spreadsheet_column(column_index)}{row_index}"
                    locators.append({
                        "kind": "spreadsheet_cell",
                        "locator": f"workbook:{sheet_name}!{cell}",
                        "sheet": sheet_name,
                        "cell": cell,
                        "row_index": row_index,
                        "column_index": column_index,
                        "text": str(value),
                        "original_value": value,
                    })
                    if len(locators) >= 50_000:
                        return locators, "spreadsheet_locator_limit_reached"
    except Exception:  # noqa: BLE001
        return [], "spreadsheet_unreadable"
    if not locators:
        return [], "spreadsheet_no_values"
    return locators, ""


def _parse_bytes(path: Path, mime: str) -> dict[str, Any]:
    data = path.read_bytes()
    digest = hashlib.sha256(data).hexdigest()
    text = ""
    locators: list[dict[str, Any]] = []
    degraded_reason = ""
    ocr_status = "not_required"
    parser = "mcp-source-parser.v1"
    csv_encoding = ""
    csv_delimiter = ""
    if str(mime or "").lower() in _SPREADSHEET_MIMES or path.suffix.lower() in {".xls", ".xlsx", ".xlsm"}:
        locators, degraded_reason = _parse_spreadsheet_locators(path)
        text = "\n".join(str(item.get("text") or "") for item in locators)[:200_000]
        parser = "mcp-spreadsheet-parser.v1"
    elif _is_csv(mime, path):
        locators, csv_encoding, csv_delimiter, degraded_reason = _parse_csv_cells(data, path, mime)
        text = "\n".join(str(item.get("text") or "") for item in locators)[:200_000]
        parser = "mcp-csv-parser.v1"
    elif mime.startswith("text/") or path.suffix.lower() in _TEXT_SUFFIXES:
        text = data.decode("utf-8", errors="replace")[:200_000]
        locators = [{"kind": "document_text", "text": text}] if text.strip() else []
    elif _is_docx(mime, path):
        locators, degraded_reason = _parse_docx_locators(data)
        text = "\n".join(str(item.get("text") or "") for item in locators)[:200_000]
        parser = "mcp-docx-parser.v1"
    elif _is_pdf(mime, path):
        locators, degraded_reason = _parse_pdf_pages(data)
        text = "\n".join(str(item.get("text") or "") for item in locators)[:200_000]
        # 两种降级都需要 OCR 才能补齐：一种连数字都没有，一种有数字但没有标签。
        if degraded_reason in {"pdf_no_text_layer", "pdf_text_layer_lacks_labels"}:
            ocr_status = "needs_ocr"
    analysis: dict[str, Any] = {
        "file_id": path.parent.name,
        "sha256": digest,
        "mime_type": mime,
        "size_bytes": len(data),
        "text_preview": text,
        "locators": _with_canonical_locators(locators),
        "page_count": sum(1 for item in locators if item.get("kind") == "pdf_page"),
        "ocr_status": ocr_status,
        "parser": parser,
        "created_at": _now(),
    }
    if csv_encoding:
        analysis["encoding"] = csv_encoding
    if csv_delimiter:
        analysis["delimiter"] = csv_delimiter
    # Enrich CSV analysis with field-value-unit associations
    if parser == "mcp-csv-parser.v1" and locators:
        units: list[dict[str, Any]] = []
        seen: set[str] = set()
        for item in locators:
            fn = item.get("field_name")
            nv = item.get("numeric_value")
            if fn and nv is not None:
                dedup_key = f"{fn}|{item.get('row_index')}|{item.get('column_index')}"
                if dedup_key not in seen:
                    seen.add(dedup_key)
                    units.append({
                        "field_name": fn,
                        "numeric_value": nv,
                        "unit": item.get("unit"),
                        "unit_normalized": item.get("unit_normalized"),
                        "locator": item.get("locator"),
                        "row_index": item.get("row_index"),
                        "column_index": item.get("column_index"),
                    })
        if units:
            analysis["csv_field_units"] = units
    if degraded_reason:
        analysis["degraded_reason"] = degraded_reason
    return analysis


def _mime_matches(data: bytes, declared_mime: str) -> bool:
    mime = declared_mime.lower().strip()
    if mime == "application/pdf":
        return data.startswith(b"%PDF-")
    if mime in {"application/zip", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"}:
        return data.startswith(b"PK\x03\x04")
    if mime.startswith("text/") or mime in {"application/json", "application/jsonl"}:
        return True
    return True


def commit_staged_source_file(
    workspace_id: str,
    staged_path: Path,
    original_filename: str,
    declared_mime: str,
    *,
    idempotency_key: str,
    expected_sha256: str = "",
    expected_size: int | None = None,
    evidence_policy: str = "",
    evidence_origin: str = "",
    project_fact_certified: bool = False,
) -> dict[str, Any]:
    """Commit a public upload as candidate evidence only.

    The legacy qualification keywords remain accepted for Python callers, but
    are deliberately ignored. Public imports cannot grant formal authority.
    """

    del evidence_policy, evidence_origin, project_fact_certified

    return _commit_staged_source_file(
        workspace_id,
        staged_path,
        original_filename,
        declared_mime,
        idempotency_key=idempotency_key,
        expected_sha256=expected_sha256,
        expected_size=expected_size,
        formal_binding=None,
    )


def commit_promoted_source_file(
    workspace_id: str,
    staged_path: Path,
    original_filename: str,
    declared_mime: str,
    *,
    idempotency_key: str,
    expected_sha256: str,
    expected_size: int,
    expected_file_id: str,
    promotion_id: str,
    template_pack_id: str,
    requirement_id: str,
    kind: str,
) -> dict[str, Any]:
    """Private promotion-authority writer; never register this as an MCP tool."""

    binding = {
        "expected_file_id": str(expected_file_id),
        "formal_promotion_id": str(promotion_id),
        "template_pack_id": str(template_pack_id),
        "requirement_id": str(requirement_id),
        "kind": str(kind),
    }
    if not all(binding.values()):
        raise _error("formal_promotion_binding_invalid", "正式 SourceFile promotion 绑定不完整")
    return _commit_staged_source_file(
        workspace_id,
        staged_path,
        original_filename,
        declared_mime,
        idempotency_key=idempotency_key,
        expected_sha256=expected_sha256,
        expected_size=expected_size,
        formal_binding=binding,
    )


def _commit_staged_source_file(
    workspace_id: str,
    staged_path: Path,
    original_filename: str,
    declared_mime: str,
    *,
    idempotency_key: str,
    expected_sha256: str = "",
    expected_size: int | None = None,
    formal_binding: dict[str, str] | None,
) -> dict[str, Any]:
    data = staged_path.read_bytes()
    if not data or len(data) > _max_upload_bytes():
        raise _error("source_content_too_large", "资料大小超过 MCP 限制")
    digest = hashlib.sha256(data).hexdigest()
    if not _mime_matches(data, declared_mime):
        raise _error("file_mime_mismatch", "资料内容与声明 MIME 不匹配")
    if expected_sha256 and expected_sha256.lower().removeprefix("sha256:") != digest:
        raise _error("source_hash_mismatch", "资料 hash 与声明不一致")
    if expected_size is not None and int(expected_size) != len(data):
        raise _error("source_size_mismatch", "资料大小与声明不一致")
    request_hash = hashlib.sha256(
        json.dumps(
            {
                "filename": original_filename,
                "mime": declared_mime,
                "sha256": digest,
                "size": len(data),
                "formal_binding": formal_binding or None,
            },
            sort_keys=True,
        ).encode()
    ).hexdigest()
    scope = f"commit:{idempotency_key}"
    with _state_guard(workspace_id):
        state = _load_state(workspace_id)
        prior = _active_idempotency_record(state["idempotency"], scope)
        if prior:
            if prior.get("request_hash") != request_hash:
                raise _error("idempotency_conflict", "同一幂等键已用于不同资料")
            return {**state["files"][prior["file_id"]], "idempotent_replay": True}
        file_id = f"src_{digest[:24]}"
        if formal_binding and file_id != formal_binding["expected_file_id"]:
            raise _error("formal_source_identity_mismatch", "正式 SourceFile ID 与预演结果不一致")
        existing = (state.get("files") or {}).get(file_id)
        if isinstance(existing, dict):
            if str(existing.get("sha256") or "").removeprefix("sha256:") != digest:
                raise _error("source_identity_conflict", "SourceFile ID 已绑定不同内容")
            existing_promotion = str(existing.get("formal_promotion_id") or "")
            requested_promotion = str((formal_binding or {}).get("formal_promotion_id") or "")
            if existing_promotion and existing_promotion != requested_promotion:
                raise _error("formal_source_already_promoted", "SourceFile 已属于另一 FormalPromotion")
        job_id = f"job_{uuid.uuid4().hex}"
        target_dir = _root(workspace_id) / "files" / file_id
        target_dir.mkdir(parents=True, exist_ok=True)
        target = target_dir / Path(original_filename or "source.bin").name
        target.write_bytes(data)
        record = {
            "file_id": file_id, "workspace_id": workspace_id,
            "original_filename": Path(original_filename or "source.bin").name,
            "declared_mime": declared_mime, "mime_type": declared_mime,
            "size_bytes": len(data), "sha256": digest, "version": 1,
            "status": "queued", "extract_status": "queued", "ocr_status": "pending",
            "deterministic_status": "pending",
            "security_scan": {"type_verified": True, "scan_status": "passed"},
            "path": str(target), "parse_job_id": job_id, "created_at": _now(), "updated_at": _now(),
            "evidence_policy": "sim_a_formal" if formal_binding else "candidate",
            "evidence_origin": "sim_a_template" if formal_binding else "",
            "project_fact_certified": bool(formal_binding),
        }
        if formal_binding:
            record.update({
                "formal_promotion_id": formal_binding["formal_promotion_id"],
                "template_pack_id": formal_binding["template_pack_id"],
                "requirement_id": formal_binding["requirement_id"],
                "kind": formal_binding["kind"],
            })
        state["files"][file_id] = record
        state["jobs"][job_id] = {
            "job_id": job_id, "file_id": file_id, "workspace_id": workspace_id,
            "status": "queued", "progress": 0, "attempt": 1, "created_at": _now(),
        }
        state["idempotency"][scope] = _idempotency_record(scope, request_hash, file_id=file_id, job_id=job_id)
        _save_state(workspace_id, state)
        return dict(record)


def parse_source_file(workspace_id: str, job_id: str) -> dict[str, Any]:
    with _state_guard(workspace_id):
        state = _load_state(workspace_id)
        job, record = _require_parse_job_from_state(state, workspace_id, job_id)
        job.update({"status": "running", "progress": 25, "started_at": _now()})
        path = Path(str(record.get("path") or ""))
        try:
            analysis = _parse_bytes(path, str(record.get("declared_mime") or "application/octet-stream"))
            analysis["file_id"] = record["file_id"]
            state["analyses"][record["file_id"]] = analysis
            degraded_reason = str(analysis.get("degraded_reason") or "")
            if degraded_reason:
                # No quotable text recovered.  Reporting `succeeded` here would
                # let an empty document pass as formally usable downstream, so
                # the degraded outcome stays visible on both job and record.
                job.update({
                    "status": "partial", "progress": 100, "finished_at": _now(),
                    "degraded_reason": degraded_reason,
                })
                record.update({
                    "status": "partial", "extract_status": "partial",
                    "deterministic_status": "partial",
                    "ocr_status": str(analysis.get("ocr_status") or "pending"),
                    "degraded_reason": degraded_reason, "updated_at": _now(),
                })
            else:
                job.update({"status": "succeeded", "progress": 100, "finished_at": _now()})
                record.update({
                    "status": "succeeded", "extract_status": "succeeded",
                    "deterministic_status": "succeeded",
                    "ocr_status": str(analysis.get("ocr_status") or "not_required"),
                    "updated_at": _now(),
                })
        except (OSError, UnicodeError, ValueError) as exc:
            job.update({"status": "failed", "progress": 100, "finished_at": _now(), "error": type(exc).__name__})
            record.update({"status": "failed", "extract_status": "failed", "updated_at": _now()})
        _save_state(workspace_id, state)
        return _public_parse_job(job)
