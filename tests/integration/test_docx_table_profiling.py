from __future__ import annotations

import base64
import io
import os
import tempfile
import unittest

from docx import Document

import lvke_mcp.servers.lvke_data_analysis.service as analysis_service
import lvke_mcp.servers.lvke_source_files.service as source_service

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _docx_with_tables() -> bytes:
    document = Document()
    document.add_paragraph("项目总投资 46968 万元，建设期 24 个月。")
    first = document.add_table(rows=3, cols=3)
    for row_index, row in enumerate(
        [["科目", "金额", "备注"], ["建设投资", "38000", "含设备"], ["流动资金", "1743", "周转"]]
    ):
        for column_index, value in enumerate(row):
            first.rows[row_index].cells[column_index].text = value
    second = document.add_table(rows=2, cols=2)
    for row_index, row in enumerate([["年份", "产量"], ["2028", "1200"]]):
        for column_index, value in enumerate(row):
            second.rows[row_index].cells[column_index].text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class DocxTableProfilingTest(unittest.TestCase):
    """DOCX 里的表格必须能被 analysis_profile_tabular 画像。

    缺陷：DOCX 解析器已把表格提取成 `docx_table_row`（带 cells 数组），但
    `profile_tabular` 只认 `kind in {"cell","spreadsheet_cell"}`，于是含表格的
    DOCX 一律被判 `no_cell_locators` —— 用户视角是「我的 DOCX 里明明有表」。
    修 DOCX 导入时没同步接进画像消费方，属「同一语义两处实现只改一侧」形态。

    解析器保持**行级** locator 不变（citation 复核按 `table:N:row:M` 回指，
    改成单元格级会打断已固化 locator 的可解析性），由画像侧投影成单元格。
    """

    def setUp(self) -> None:
        self.tempdir = tempfile.TemporaryDirectory(prefix="lvke-docx-profile-")
        self.previous = os.environ.get("LVKE_MCP_DATA_DIR")
        os.environ["LVKE_MCP_DATA_DIR"] = self.tempdir.name
        self.workspace = "docx-table-profiling"

    def tearDown(self) -> None:
        if self.previous is None:
            os.environ.pop("LVKE_MCP_DATA_DIR", None)
        else:
            os.environ["LVKE_MCP_DATA_DIR"] = self.previous
        self.tempdir.cleanup()

    def _profiled(self) -> dict:
        imported = source_service.import_content(
            workspace_id=self.workspace,
            original_filename="投资估算.docx",
            declared_mime=_DOCX_MIME,
            content_base64=base64.b64encode(_docx_with_tables()).decode("ascii"),
            idempotency_key="docx-profile-import",
        )
        self.assertTrue(imported.get("success"), imported)
        ingested = analysis_service.ingest(self.workspace, [], [imported["file_id"]])
        return analysis_service.profile_tabular(
            self.workspace, ingested["analysis_task_id"],
        )

    def test_docx_tables_are_profiled_not_skipped(self) -> None:
        result = self._profiled()
        self.assertEqual(result.get("status"), "ok", result)
        self.assertEqual(result.get("skipped"), [], result)
        profiles = result.get("profiles") or []
        # 两张表必须各自成表，不能被合并成一张（否则行列数是错的）。
        self.assertEqual(len(profiles), 2, profiles)
        sheets = sorted(str(item.get("sheet")) for item in profiles)
        self.assertEqual(sheets, ["docx_table_1", "docx_table_2"])

    def test_headers_and_dimensions_match_the_source_table(self) -> None:
        profiles = {str(item.get("sheet")): item for item in self._profiled()["profiles"]}
        first = profiles["docx_table_1"]
        self.assertEqual(first["headers"], ["科目", "金额", "备注"])
        self.assertEqual(first["observed_row_count"], 3)
        self.assertEqual(first["observed_column_count"], 3)
        self.assertEqual(first["observed_cell_count"], 9)
        second = profiles["docx_table_2"]
        self.assertEqual(second["headers"], ["年份", "产量"])
        self.assertEqual(second["observed_column_count"], 2)

    def test_docx_cells_are_not_counted_as_numeric(self) -> None:
        """DOCX 单元格是文本，没有已解析的数值 —— 不得凭字面像数字就当数值。

        这条守住"不造值"边界：`numeric_cell_count` 必须为 0，否则下游会把
        未经解析的文本当作可参与算术的数值候选。
        """

        first = {str(i.get("sheet")): i for i in self._profiled()["profiles"]}["docx_table_1"]
        self.assertEqual(first["numeric_cell_count"], 0)
        self.assertEqual(first["formula_cell_count"], 0)
        self.assertEqual(first["text_cell_count"], 9)


if __name__ == "__main__":
    unittest.main()
