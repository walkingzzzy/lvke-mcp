from __future__ import annotations

import base64
import io
import json
from datetime import datetime, timedelta, timezone

from docx import Document
from openpyxl import Workbook
from openpyxl import load_workbook

from lvke_mcp.servers.lvke_deliverable_review import service
from lvke_mcp.servers.lvke_source_files import service as source_files


DIMENSIONS = (
    "compliance",
    "article_quality",
    "data_quality",
    "source_quality",
    "financial_model",
    "financial_tables",
    "feasibility",
)
SEMANTIC_CHECKS = {
    "compliance": ["COMP.SUBSTANTIVE.REVIEW"],
    "article_quality": ["ARTICLE.LANGUAGE.LOGIC"],
    "data_quality": ["DATA.ANOMALY.RECONCILIATION"],
    "source_quality": ["SOURCE.CLAIM.SUPPORT"],
    "financial_model": ["FINMODEL.ASSUMPTION.LOGIC"],
    "financial_tables": ["FINTABLE.SUBSTANTIVE.REVIEW"],
    "feasibility": ["FEASIBILITY.DECISION.LOGIC"],
}


def _docx() -> bytes:
    document = Document()
    document.add_heading("项目可行性研究报告", level=1)
    document.add_paragraph(
        "本项目位于湖北省。市场需求分析说明服务对象与需求依据。建设方案和技术方案明确。"
        "投资估算、融资方案与资金筹措相互对应。财务评价包括内部收益率和现金流。"
        "风险分析列明主要风险，结论与建议说明决策条件。"
    )
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _xlsx(*, tables: bool) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "财务表" if tables else "财务模型"
    sheet.append(["项目", "第一年", "第二年"])
    sheet.append(["收入", 100, 120])
    sheet.append(["成本", 60, 70])
    sheet.append(["利润", "=B2-B3", "=C2-C3"])
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _import(workspace: str, name: str, mime: str, content: bytes) -> str:
    result = source_files.import_content(
        workspace_id=workspace,
        original_filename=name,
        declared_mime=mime,
        content_base64=base64.b64encode(content).decode("ascii"),
        idempotency_key=f"import-{name}",
        parse_immediately=True,
    )
    assert result["status"] == "ok", result
    return str(result["file_id"])


def _five_files(workspace: str) -> dict[str, str]:
    return {
        "report": _import(
            workspace,
            "项目可行性研究报告.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            _docx(),
        ),
        "source_evidence": _import(
            workspace,
            "政策来源证据.txt",
            "text/plain",
            "湖北省项目依据，发布机关和有效日期已记录。".encode(),
        ),
        "base_data": _import(
            workspace,
            "基础数据.csv",
            "text/csv",
            "metric,year,value\nrevenue,2026,100\ncost,2026,60\n".encode(),
        ),
        "finance_model": _import(
            workspace,
            "财务模型.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            _xlsx(tables=False),
        ),
        "finance_tables": _import(
            workspace,
            "正式财务表.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            _xlsx(tables=True),
        ),
    }


def _prepare_package(workspace: str, roles: dict[str, str], *, suffix: str = "") -> dict:
    draft = service.prepare_package({
        "workspace_id": workspace,
        "idempotency_key": f"package-draft{suffix}",
        "source_file_ids": list(roles.values()),
        "internal_targets": [],
        "review_mode": "external",
        "review_profile": "standard",
        "project_scope": {
            "region": "湖北省",
            "project_type": "generic_feasibility",
            "industry": "general",
        },
    })
    assert draft["status"] == "ok", draft
    confirmed = service.confirm_package({
        "workspace_id": workspace,
        "idempotency_key": f"package-confirm{suffix}",
        "review_package_draft_id": draft["review_package_draft_id"],
        "expected_draft_hash": draft["draft_hash"],
        "component_roles": [
            {"component_id": file_id, "role": role}
            for role, file_id in roles.items()
        ],
        "confirmation_statement": "已逐项核对五类材料角色。",
    })
    assert confirmed["status"] == "ok", confirmed
    return confirmed


def _start_suite(workspace: str, package: dict, *, suffix: str = "") -> str:
    prepared = service.prepare({
        "workspace_id": workspace,
        "idempotency_key": f"review-prepare{suffix}",
        "target": {"target_type": "review_package", "target_id": package["review_package_id"]},
        "review_profile": "standard",
        "review_mode": "external",
        "project_context": {"review_purpose": "process_acceptance"},
    })
    assert prepared["status"] == "ok", prepared
    started = service.start({
        "workspace_id": workspace,
        "idempotency_key": f"review-start{suffix}",
        "review_preparation_id": prepared["review_preparation_id"],
        "mode": "standard",
        "execution": "sync",
        "deployment_mode": "enforced",
    })
    assert started["status"] in {"ok", "partial", "incomplete"}, started
    return str(started["review_id"])


def _submit_dimension(
    workspace: str,
    review_id: str,
    package_id: str,
    dimension: str,
    *,
    status: str = "passed",
    findings: list[dict] | None = None,
    suffix: str = "",
) -> dict:
    return service.submit_assessment({
        "workspace_id": workspace,
        "idempotency_key": f"assessment-{dimension}{suffix}",
        "review_id": review_id,
        "review_package_id": package_id,
        "dimension": dimension,
        "status": status,
        "coverage": {
            "reviewed": "all applicable semantic checks",
            "checked_check_ids": SEMANTIC_CHECKS[dimension],
        },
        "findings": findings or [],
        "limitations": [],
        "skill": f"review-{dimension}",
        "skill_version": "1.0.0",
        "model": "test-reviewer",
        "model_version": "1.0.0",
        "execution_environment": "controlled_current_environment",
        "independent_context": True,
        "reviewer_context_id": f"ctx-{dimension}{suffix}",
    })


def _confirm_all_dimensions(workspace: str, review_id: str, package_id: str, *, suffix: str = "") -> None:
    for dimension in DIMENSIONS:
        assessment = _submit_dimension(
            workspace, review_id, package_id, dimension, suffix=suffix,
        )
        assert assessment["status"] == "ok", assessment
        confirmed = service.confirm_dimension({
            "workspace_id": workspace,
            "idempotency_key": f"confirmation-{dimension}{suffix}",
            "review_id": review_id,
            "dimension": dimension,
            "role_declaration": f"{dimension} reviewer",
            "review_statement": "已独立复核冻结材料及本域结果。",
            "limitations_accepted": [],
        })
        assert confirmed["status"] == "ok", confirmed


def test_full_external_suite_reaches_dossier_but_not_release(tmp_path) -> None:
    workspace = "seven-domain-success"
    package = _prepare_package(workspace, _five_files(workspace))
    assert package["full_suite"] is True

    prepared = service.prepare({
        "workspace_id": workspace,
        "idempotency_key": "review-prepare",
        "target": {
            "target_type": "review_package",
            "target_id": package["review_package_id"],
        },
        "review_profile": "standard",
        "review_mode": "external",
        "project_context": {
            "project_type": "generic_feasibility",
            "transaction_structure": "new_build",
            "asset_type": "general",
            "evidence_track": "real",
            "review_purpose": "process_acceptance",
        },
    })
    assert prepared["status"] == "ok", prepared
    started = service.start({
        "workspace_id": workspace,
        "idempotency_key": "review-start",
        "review_preparation_id": prepared["review_preparation_id"],
        "mode": "standard",
        "execution": "sync",
        "deployment_mode": "enforced",
    })
    review_id = started["review_id"]

    for dimension in DIMENSIONS:
        assessment = service.submit_assessment({
            "workspace_id": workspace,
            "idempotency_key": f"assessment-{dimension}",
            "review_id": review_id,
            "review_package_id": package["review_package_id"],
            "dimension": dimension,
            "status": "passed",
            "coverage": {
                "reviewed": "all applicable semantic checks",
                "checked_check_ids": SEMANTIC_CHECKS[dimension],
            },
            "findings": [],
            "limitations": [],
            "skill": f"review-{dimension}",
            "skill_version": "1.0.0",
            "model": "test-reviewer",
            "model_version": "1.0.0",
            "execution_environment": "controlled_current_environment",
            "independent_context": True,
            "reviewer_context_id": f"ctx-{dimension}",
        })
        assert assessment["status"] == "ok", assessment
        confirmation = service.confirm_dimension({
            "workspace_id": workspace,
            "idempotency_key": f"confirmation-{dimension}",
            "review_id": review_id,
            "dimension": dimension,
            "role_declaration": f"{dimension} reviewer",
            "review_statement": "已独立复核冻结材料及本域结果。",
            "limitations_accepted": [],
        })
        assert confirmation["status"] == "ok", confirmation

    finalized = service.finalize({
        "workspace_id": workspace,
        "idempotency_key": "review-finalize",
        "review_id": review_id,
    })
    assert finalized["overall_verdict"] == "pass", finalized
    assert finalized["formal_suite_review_complete"] is True
    assert len(finalized["dimension_result_ids"]) == 7
    dimension_result_uris = [
        uri for uri in finalized["resource_uris"]
        if "/dimension-results/" in uri
    ]
    assert len(dimension_result_uris) == 7
    for uri in dimension_result_uris:
        resolved = service.resolve_resource(uri, workspace)
        assert resolved is not None
        record = json.loads(resolved[0])
        assert record["schema_version"] == "ReviewDimensionResult.v1"
        assert record["payload"]["schema_version"] == "ReviewDimensionResult.v1"
        assert record["payload"]["review_id"] == review_id
    state = service.get_review({"workspace_id": workspace, "review_id": review_id})
    assert state["review"]["schema_version"] == "ReviewDossier.v2"
    assert "external_review_release_forbidden" in state["review"]["blockers"]
    exported = service.export_review({
        "workspace_id": workspace,
        "idempotency_key": "review-export-suite",
        "review_id": review_id,
    })
    assert exported["status"] == "ok", exported
    assert {row["format"] for row in exported["files"]} == {
        "json", "markdown", "docx", "xlsx", "annotated_docx",
    }
    from lvke_mcp.servers.lvke_deliverable_review._service.export import _export_root

    workbook = load_workbook(
        _export_root(workspace, exported["export_id"]) / "findings.xlsx",
        read_only=True,
    )
    assert {"findings", "dimension_results", "standards_snapshot", "audit_manifest"} <= set(workbook.sheetnames)
    workbook.close()
    pdf_export = service.export_review({
        "workspace_id": workspace,
        "idempotency_key": "review-export-pdf",
        "review_id": review_id,
        "formats": ["pdf"],
    })
    assert pdf_export["status"] == "ok", pdf_export
    assert pdf_export["files"][0]["format"] == "pdf"


def test_missing_role_only_allows_specialized_review() -> None:
    workspace = "seven-domain-missing"
    roles = _five_files(workspace)
    roles.pop("finance_tables")
    package = _prepare_package(workspace, roles)
    assert package["full_suite"] is False
    assert package["missing_required_roles"] == ["finance_tables"]


def test_assessment_requires_independent_context_and_valid_locator() -> None:
    workspace = "seven-domain-negative"
    roles = _five_files(workspace)
    package = _prepare_package(workspace, roles)
    prepared = service.prepare({
        "workspace_id": workspace,
        "idempotency_key": "review-prepare",
        "target": {"target_type": "review_package", "target_id": package["review_package_id"]},
        "review_profile": "standard",
        "review_mode": "external",
        "project_context": {"review_purpose": "process_acceptance"},
    })
    started = service.start({
        "workspace_id": workspace,
        "idempotency_key": "review-start",
        "review_preparation_id": prepared["review_preparation_id"],
        "mode": "standard",
        "execution": "sync",
        "deployment_mode": "enforced",
    })
    base = {
        "workspace_id": workspace,
        "review_id": started["review_id"],
        "review_package_id": package["review_package_id"],
        "dimension": "source_quality",
        "status": "failed",
        "coverage": {"checked_check_ids": SEMANTIC_CHECKS["source_quality"]},
        "limitations": [],
        "skill": "source-review",
        "skill_version": "1",
        "model": "test-reviewer",
        "model_version": "1",
        "execution_environment": "controlled_current_environment",
        "reviewer_context_id": "ctx-source-negative",
    }
    not_independent = service.submit_assessment({
        **base,
        "idempotency_key": "not-independent",
        "independent_context": False,
        "findings": [{
            "check_id": "SOURCE.CLAIM.SUPPORT",
            "severity": "P1",
            "message": "来源不支持结论",
            "target_location": {"claim": 1},
            "missing_evidence_reason": "未提供原始来源",
            "remediation": "补充原始来源",
        }],
    })
    assert not_independent["code"] == "independent_review_context_required"

    evidence_id = roles["source_evidence"]
    state = __import__(
        "lvke_mcp.adapters.source_files_repository", fromlist=["_load_state"]
    )._load_state(workspace)
    source_hash = "sha256:" + str(state["files"][evidence_id]["sha256"]).removeprefix("sha256:")
    forged = service.submit_assessment({
        **base,
        "idempotency_key": "forged-locator",
        "independent_context": True,
        "findings": [{
            "check_id": "SOURCE.CLAIM.SUPPORT",
            "severity": "P1",
            "message": "来源不支持结论",
            "target_location": {"claim": 1},
            "evidence": [{
                "source_id": evidence_id,
                "source_hash": source_hash,
                "locator": "paragraph:999",
            }],
            "remediation": "补充有效来源",
        }],
    })
    assert forged["status"] == "blocked"
    assert forged["code"] == "citation_locator_not_found"


def test_public_server_exposes_suite_tools_without_verdict_input() -> None:
    from lvke_mcp.servers.lvke_deliverable_review.server import build_server

    specs = {item.name: item for item in build_server().tool_specs}
    expected = {
        "review_package_prepare", "review_package_confirm", "review_confirm_extraction",
        "review_submit_assessment", "review_get_dimension", "review_confirm_dimension",
        "review_finalize",
    }
    assert expected <= set(specs)
    finalize_schema = specs["review_finalize"].input_schema
    assert "verdict" not in finalize_schema["properties"]
    assert "professional_signoff" not in finalize_schema["properties"]


def test_internal_package_rejects_unsigned_external_files() -> None:
    workspace = "seven-domain-internal-unsigned"
    roles = _five_files(workspace)
    draft = service.prepare_package({
        "workspace_id": workspace,
        "idempotency_key": "package-draft-internal",
        "source_file_ids": list(roles.values()),
        "internal_targets": [],
        "review_mode": "internal",
        "review_profile": "standard",
        "project_scope": {"region": "湖北省", "project_type": "generic_feasibility"},
    })
    confirmed = service.confirm_package({
        "workspace_id": workspace,
        "idempotency_key": "package-confirm-internal",
        "review_package_draft_id": draft["review_package_draft_id"],
        "expected_draft_hash": draft["draft_hash"],
        "component_roles": [
            {"component_id": file_id, "role": role} for role, file_id in roles.items()
        ],
        "confirmation_statement": "确认材料角色。",
    })
    assert confirmed["status"] == "blocked", confirmed
    assert confirmed["code"] == "formal_source_policy_required"


def test_unconfirmed_and_not_determinable_dimensions_block_finalize() -> None:
    workspace = "seven-domain-incomplete"
    package = _prepare_package(workspace, _five_files(workspace))
    review_id = _start_suite(workspace, package)
    for dimension in DIMENSIONS:
        result = _submit_dimension(
            workspace,
            review_id,
            package["review_package_id"],
            dimension,
            status="not_determinable" if dimension == "source_quality" else "passed",
        )
        assert result["status"] == "ok", result
        if dimension != "financial_tables":
            confirmed = service.confirm_dimension({
                "workspace_id": workspace,
                "idempotency_key": f"confirm-{dimension}",
                "review_id": review_id,
                "dimension": dimension,
                "role_declaration": "independent reviewer",
                "review_statement": "已检查可用材料。",
                "limitations_accepted": [],
            })
            assert confirmed["status"] == "ok", confirmed
    finalized = service.finalize({
        "workspace_id": workspace,
        "idempotency_key": "finalize-incomplete",
        "review_id": review_id,
    })
    assert finalized["overall_verdict"] == "incomplete", finalized
    assert "review_dimension_not_determinable:source_quality" in finalized["blockers"]
    assert "review_dimension_unconfirmed:financial_tables" in finalized["blockers"]


def test_semantic_p0_never_waived_and_p1_requires_complete_conditions() -> None:
    workspace = "seven-domain-waiver"
    roles = _five_files(workspace)
    package = _prepare_package(workspace, roles)
    review_id = _start_suite(workspace, package)
    for severity, suffix in (("P0", "-p0"), ("P1", "-p1")):
        result = _submit_dimension(
            workspace,
            review_id,
            package["review_package_id"],
            "compliance",
            status="failed",
            suffix=suffix,
            findings=[{
                "check_id": "COMP.SUBSTANTIVE.REVIEW",
                "severity": severity,
                "message": f"{severity} 合规问题",
                    "target_location": {"section": f"审批-{severity}"},
                "missing_evidence_reason": "缺少审批依据",
                "remediation": "补充依据并重审",
            }],
        )
        assert result["status"] == "ok", result
    findings = service.list_findings({"workspace_id": workspace, "review_id": review_id})["findings"]
    by_severity = {row["severity"]: row for row in findings if row.get("rule_id") == "COMP.SUBSTANTIVE.REVIEW"}
    evidence = [{
        "source_id": roles["source_evidence"],
        "locator": "paragraph:1",
        "content_hash": "sha256:" + "a" * 64,
    }]
    common = {
        "workspace_id": workspace,
        "review_id": review_id,
        "disposition": "waived",
        "note": "测试豁免边界",
        "waiver_scope": "仅当前版本",
        "waiver_expires_at": (datetime.now(timezone.utc) + timedelta(days=30)).isoformat(),
        "waiver_invalidation_conditions": ["材料变化"],
        "remediation_evidence": evidence,
    }
    p0 = service.disposition_finding({
        **common,
        "idempotency_key": "waive-p0",
        "finding_id": by_severity["P0"]["finding_id"],
        "waiver_impact": "影响发布",
        "waiver_compensating_controls": "人工复核",
        "waiver_responsible_party": "项目负责人",
    })
    assert p0["code"] == "p0_waiver_forbidden"
    p1 = service.disposition_finding({
        **common,
        "idempotency_key": "waive-p1-incomplete",
        "finding_id": by_severity["P1"]["finding_id"],
    })
    assert p1["code"] == "waiver_conditions_required"


def test_suite_retest_waits_for_fresh_assessments() -> None:
    workspace = "seven-domain-retest"
    roles = _five_files(workspace)
    package = _prepare_package(workspace, roles)
    parent_id = _start_suite(workspace, package, suffix="-parent")
    failed = _submit_dimension(
        workspace,
        parent_id,
        package["review_package_id"],
        "article_quality",
        status="failed",
        suffix="-parent",
        findings=[{
            "check_id": "ARTICLE.LANGUAGE.LOGIC",
            "severity": "P1",
            "message": "存在论证跳跃",
            "target_location": {"section": "市场分析"},
            "missing_evidence_reason": "论证链缺少中间依据",
            "remediation": "补齐论证链",
        }],
    )
    assert failed["status"] == "ok", failed
    for dimension in DIMENSIONS:
        if dimension == "article_quality":
            continue
        assert _submit_dimension(
            workspace, parent_id, package["review_package_id"], dimension, suffix="-parent",
        )["status"] == "ok"

    revised_document = Document()
    revised_document.add_heading("项目可行性研究报告（整改版）", level=1)
    revised_document.add_paragraph(
        "本项目位于湖北省。市场、需求、建设方案、投资融资、财务评价、风险及结论的论证链已补齐。"
    )
    revised_buffer = io.BytesIO()
    revised_document.save(revised_buffer)
    revised_roles = dict(roles)
    revised_roles["report"] = _import(
        workspace,
        "项目可行性研究报告整改版.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        revised_buffer.getvalue(),
    )
    revised_package = _prepare_package(workspace, revised_roles, suffix="-revised")
    result = service.retest({
        "workspace_id": workspace,
        "idempotency_key": "suite-retest",
        "review_id": parent_id,
        "target": {"target_type": "review_package", "target_id": revised_package["review_package_id"]},
        "mode": "standard",
        "remediation_evidence": [{
            "source_id": revised_roles["report"],
            "locator": "paragraph:1",
            "content_hash": "sha256:" + "b" * 64,
        }],
    })
    assert result["status"] == "accepted", result
    assert result["code"] == "retest_assessment_required"
    parent = service.get_review({"workspace_id": workspace, "review_id": parent_id})["review"]
    assert parent["pending_retest_operation_ids"]
    assert result["closed_finding_ids"] == []

    child_id = result["retest_review_id"]
    _confirm_all_dimensions(
        workspace,
        child_id,
        revised_package["review_package_id"],
        suffix="-child",
    )
    finalized = service.finalize({
        "workspace_id": workspace,
        "idempotency_key": "finalize-retest-child",
        "review_id": child_id,
    })
    assert finalized["overall_verdict"] == "pass", finalized
    assert finalized["retest_result"]["status"] == "completed", finalized
    assert finalized["retest_result"]["closed_finding_ids"]
    parent = service.get_review({"workspace_id": workspace, "review_id": parent_id})["review"]
    assert parent["pending_retest_operation_ids"] == []
